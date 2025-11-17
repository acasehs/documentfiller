#!/usr/bin/env python3
"""
Document Content Generator - OpenWebUI Integration
Automatically fills Word document sections using OpenWebUI/Ollama with RAG support
Features: Markdown conversion, prompt learning, auto-features, model comparison
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, colorchooser
from docx import Document  # pip install python-docx
from docx.shared import Inches, Pt, RGBColor # pip install python-docx
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL  # pip install python-docx
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml 
import json
import requests  # pip install requests
import threading
import os
import re
from datetime import datetime
import shutil
import hashlib
from typing import Dict, List, Tuple
import sqlite3
import uuid

# dependency pips:  python -m pip install --break-system-packages cryptography textstat nltk tiktoken scikit-learn numpy requests


# Add after existing imports, before class definitions

# ============================================================================
# ENHANCED MODULE IMPORTS - Optional for graceful degradation
# ============================================================================
ENHANCED_FEATURES_AVAILABLE = {
    'encryption': False,
    'advanced_review': False,
    'intelligent_processing': False
}

try:
    from credential_manager import CredentialManager
    ENHANCED_FEATURES_AVAILABLE['encryption'] = True
    print("✓ Encrypted credentials module loaded")
except ImportError:
    print("⚠ credential_manager not available - using standard config")
    CredentialManager = None

try:
    from document_reviewer import TechnicalDocumentReviewer
    ENHANCED_FEATURES_AVAILABLE['advanced_review'] = True
    print("✓ Advanced document review module loaded")
except ImportError:
    print("⚠ document_reviewer not available - using standard review")
    TechnicalDocumentReviewer = None

try:
    from content_processor import IntelligentContentProcessor
    ENHANCED_FEATURES_AVAILABLE['intelligent_processing'] = True
    print("✓ Intelligent content processing module loaded")
except ImportError:
    print("⚠ content_processor not available - using standard processing")
    IntelligentContentProcessor = None

class DocumentSection:
    """Represents a hierarchical document section"""
    def __init__(self, level, text, paragraph, full_path=""):
        self.level = level
        self.text = text
        self.paragraph = paragraph
        self.full_path = full_path
        self.children = []
        self.parent = None
        self.content_paragraphs = []
        
    def add_child(self, child):
        child.parent = self
        self.children.append(child)
        
    def get_full_path(self):
        """Get full hierarchical path"""
        if self.parent:
            parent_path = self.parent.get_full_path()
            return f"{parent_path} > {self.text}" if parent_path else self.text
        return self.text
        
    def get_existing_content(self):
        """Get existing text content in this section"""
        content = []
        for para in self.content_paragraphs:
            if para.text.strip():
                content.append(para.text)
        return "\n".join(content)
        
    def has_content(self):
        """Check if section has existing content"""
        return len(self.get_existing_content().strip()) > 0
    
    def get_section_hash(self):
        """Get unique hash for this section based on path"""
        return hashlib.md5(self.get_full_path().encode()).hexdigest()

class DocumentReview:
    """Represents a technical writing review of a document section"""
    def __init__(self, section_path: str):
        self.section_path = section_path
        self.timestamp = datetime.now()
        self.cohesion_score = 0
        self.clarity_score = 0
        self.technical_accuracy_score = 0
        self.factual_veracity_score = 0
        self.completeness_score = 0
        self.overall_score = 0
        self.feedback = {
            'cohesion': [],
            'clarity': [],
            'technical_accuracy': [],
            'factual_veracity': [],
            'completeness': [],
            'general': []
        }
        self.recommendations = []
        self.content_suggestions = {}

        # Initialize credential manager if available
        self.credential_manager = None
        if ENHANCED_FEATURES_AVAILABLE['encryption'] and CredentialManager:
            try:
                self.credential_manager = CredentialManager()
                if self.load_encrypted_credentials():
                    self.log_message("✓ Encrypted credentials loaded successfully")
            except Exception as e:
                self.log_message(f"Credential manager init failed: {e}")
        
        # Initialize advanced document reviewer if available
        self.advanced_reviewer = None
        if ENHANCED_FEATURES_AVAILABLE['advanced_review'] and TechnicalDocumentReviewer:
            try:
                reviewer_config = {
                    'base_url': self.openwebui_base_url,
                    'api_key': self.openwebui_api_key,
                    'review_model': self.selected_model.get() or 'llama3.1:latest',
                    'comment_model': self.selected_model.get() or 'llama3.1:latest'
                }
                self.advanced_reviewer = TechnicalDocumentReviewer(reviewer_config)
                self.log_message("✓ Advanced document reviewer initialized")
            except Exception as e:
                self.log_message(f"Advanced reviewer init failed: {e}")
        
        # Initialize intelligent content processor if available
        self.content_processor = None
        if ENHANCED_FEATURES_AVAILABLE['intelligent_processing'] and IntelligentContentProcessor:
            try:
                processor_config = {
                    'rag_threshold': 10000,
                    'max_prompt_tokens': self.max_tokens.get(),
                    'chunk_size': 1000,
                    'chunk_overlap': 100,
                    'document_store_path': os.path.join(os.path.dirname(__file__), 'document_store.db')
                }
                self.content_processor = IntelligentContentProcessor(processor_config)
                self.log_message("✓ Intelligent content processor initialized")
            except Exception as e:
                self.log_message(f"Content processor init failed: {e}")
        
    def calculate_overall_score(self):
        """Calculate overall review score"""
        scores = [
            self.cohesion_score,
            self.clarity_score,
            self.technical_accuracy_score,
            self.factual_veracity_score,
            self.completeness_score
        ]
        self.overall_score = sum(scores) / len(scores) if scores else 0
        return self.overall_score
    
    def get_summary(self) -> str:
        """Get formatted review summary"""
        self.calculate_overall_score()
        summary = f"=== REVIEW SUMMARY ===\n"
        summary += f"Section: {self.section_path}\n"
        summary += f"Timestamp: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        summary += f"Overall Score: {self.overall_score:.1f}/10\n\n"
        summary += f"Individual Scores:\n"
        summary += f"  • Cohesion: {self.cohesion_score}/10\n"
        summary += f"  • Clarity: {self.clarity_score}/10\n"
        summary += f"  • Technical Accuracy: {self.technical_accuracy_score}/10\n"
        summary += f"  • Factual Veracity: {self.factual_veracity_score}/10\n"
        summary += f"  • Completeness: {self.completeness_score}/10\n\n"
        
        if self.recommendations:
            summary += "Key Recommendations:\n"
            for i, rec in enumerate(self.recommendations, 1):
                summary += f"  {i}. {rec}\n"
        
        return summary


class DocxDocumentFiller:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Document Content Generator - OpenWebUI Integration")
        self.root.geometry("1400x900")
        self.root.configure(bg="#2b2b2b")
        
        # Configure dark theme
        self.configure_dark_theme()
        
        # Data storage
        self.document = None
        self.document_path = None
        self.last_document_path = None
        self.sections = []
        self.selected_section = None
        self.generated_content = ""
        self.last_sent_prompt = ""

        # External RAG content storage
        self.external_content_db = None
        self.init_external_content_db()

        # Section chat storage
        self.section_chat_history = {}  # {section_hash: [(role, message), ...]}
        self.current_chat_section = None
        
        # Document review configuration
        self.review_mode_active = False
        self.current_review = None
        self.review_history = {}
        
        # Section tracking
        self.section_tracking = {}
        self.tracking_file = None
        
        # OpenWebUI configuration
        self.openwebui_base_url = "http://172.16.27.122:3000"
        self.openwebui_api_key = ""
        self.available_models = []
        self.selected_model = tk.StringVar()
        self.available_knowledge_collections = []
        self.selected_knowledge_collections = []
        self.temperature = tk.DoubleVar(value=0.1)
        self.max_tokens = tk.IntVar(value=8000)
        
        # Formatting configuration
        self.format_config = {
            'highlight_enabled': tk.BooleanVar(value=True),
            'highlight_color': tk.StringVar(value='YELLOW'),
            'bold_enabled': tk.BooleanVar(value=False),
            'italic_enabled': tk.BooleanVar(value=False),
            'underline_enabled': tk.BooleanVar(value=False),
            'font_color': tk.StringVar(value='000000'),
            'font_size': tk.IntVar(value=11)
        }
        
        # Auto features configuration
        self.auto_config = {
            'auto_backup': tk.BooleanVar(value=True),
            'backup_interval': tk.IntVar(value=5),
            'auto_save': tk.BooleanVar(value=False),
            'auto_reload': tk.BooleanVar(value=True),
            'ask_backup': tk.BooleanVar(value=True)
        }
        
        # Operation mode
        self.operation_mode = tk.StringVar(value="replace")

        # Default master prompt
        self.master_prompt = tk.StringVar()
        self.set_default_prompt()

        # Initialize credential manager if available
        self.credential_manager = None
        self.content_processor = None
        self.advanced_reviewer = None

        if ENHANCED_FEATURES_AVAILABLE['encryption'] and CredentialManager:
            try:
                self.credential_manager = CredentialManager()
                print("✓ Credential manager initialized")
            except Exception as e:
                print(f"⚠ Credential manager init failed: {e}")

        # Load settings
        self.load_settings()

        # Create GUI
        self.create_gui()
        self.setup_text_widgets_colors()

        # Initialize advanced modules after settings are loaded
        if ENHANCED_FEATURES_AVAILABLE['intelligent_processing'] and IntelligentContentProcessor:
            try:
                self.content_processor = IntelligentContentProcessor()
                print("✓ Intelligent content processor initialized")
            except Exception as e:
                print(f"⚠ Content processor init failed: {e}")

        if ENHANCED_FEATURES_AVAILABLE['advanced_review'] and TechnicalDocumentReviewer:
            try:
                reviewer_config = {
                    'base_url': self.openwebui_base_url,
                    'api_key': self.openwebui_api_key,
                    'review_model': self.selected_model.get() or 'llama3.1:latest',
                    'comment_model': self.selected_model.get() or 'llama3.1:latest'
                }
                self.advanced_reviewer = TechnicalDocumentReviewer(reviewer_config)
                print("✓ Advanced document reviewer initialized")
            except Exception as e:
                print(f"⚠ Advanced reviewer init failed: {e}")

        # Load last document if available
        if self.last_document_path and os.path.exists(self.last_document_path):
            self.load_document(self.last_document_path)

        # Start auto-backup timer if enabled
        if self.auto_config['auto_backup'].get():
            self.schedule_auto_backup()
        
    def configure_dark_theme(self):
        """Configure dark theme"""
        try:
            style = ttk.Style()
            available_themes = style.theme_names()
            
            if 'clam' in available_themes:
                style.theme_use('clam')
            
            style.configure('TLabel', background='#2b2b2b', foreground='#ffffff')
            style.configure('TButton', background='#404040', foreground='#ffffff')
            style.configure('TFrame', background='#2b2b2b')
            style.configure('TLabelframe', background='#2b2b2b', foreground='#ffffff')
            style.configure('TLabelframe.Label', background='#2b2b2b', foreground='#ffffff')
            style.configure('TRadiobutton', background='#2b2b2b', foreground='#ffffff')
            style.configure('TCheckbutton', background='#2b2b2b', foreground='#ffffff')
            
            style.configure('Treeview', 
                          background='#1e1e1e',
                          foreground='#ffffff',
                          fieldbackground='#1e1e1e',
                          bordercolor='#666666')
            style.map('Treeview', background=[('selected', '#404040')])
            
        except Exception as e:
            print(f"Dark theme configuration failed: {e}")

    def init_external_content_db(self):
        """Initialize SQLite database for external RAG content"""
        try:
            db_path = os.path.join(os.path.expanduser("~"), ".documentfiller_external_rag.db")
            self.external_content_db = sqlite3.connect(db_path, check_same_thread=False)
            cursor = self.external_content_db.cursor()

            cursor.execute('''
                CREATE TABLE IF NOT EXISTS external_content (
                    id TEXT PRIMARY KEY,
                    title TEXT NOT NULL,
                    content TEXT NOT NULL,
                    category TEXT,
                    tags TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            self.external_content_db.commit()
            print("✓ External RAG content database initialized")
        except Exception as e:
            print(f"⚠ External content DB init failed: {e}")
            self.external_content_db = None

    def set_default_prompt(self):
        """Set default master prompt"""
        default = """You are a Department of Defense cybersecurity expert tasked with creating comprehensive, technically accurate documentation for a DoD system.

CONTEXT:
- This is for a Zero Trust Architecture implementation
- Must comply with DoD standards including RMF, NIST SP 800-53, FedRAMP High, and DoD IL5/IL6
- System is a COTS ERP solution in commercial cloud hosting
- Integrates with 250+ legacy systems
- Handles PII and requires strict data protection

REQUIREMENTS:
- Use specific technical terminology appropriate for DoD cybersecurity documentation
- Reference applicable NIST controls, DoD directives, and compliance frameworks
- Be detailed and comprehensive - this is for ATO documentation
- Use proper markdown formatting: **bold**, *italic*, `code`, bullet points, numbered lists
- Focus on implementation details, not just policy statements
- Include specific requirements for Zero Trust pillars where applicable

ZERO TRUST TOOLS:
- CyberArk PAM
- OKTA Secure Authentication
- Palo Alto Secure Browser 
- Palo Alto NGFWs
- Tenable Cloud Security
- Tenable SC
- Zscaler Private Access
- Zscaler Internet Access
- Trellix ePolicy Orchestrator
- Splunk
- Fortify
- SonarQube

CYBERSECURITY TEAM COMPOSITION
- Information Systems Security Officer (ISSOs) - Governance and Compliance, RMF Documentation
- Information Systems Security Engineers (ISSEs) - Security Engineering, STIG Integration, Security Impact Assessments
- Security Operations (SecOps) - System Scans, Policy Enforcement, Continuous Monitoring
- Identity, Credential, and Access Management (ICAM) - Credential Management, Access Auditing and Monitoring 
- Application Security (AppSec) - Application Vulnerability Testing, SAST, DAST, Penetration Testing

SECTION TO FILL: {section_name}
PARENT CONTEXT: {parent_context}
OPERATION: {operation_mode}"""
        
        self.master_prompt.set(default)
        
    def set_review_prompt(self):
        """Set technical writer review prompt"""
        review_prompt = """You are an expert technical writer and documentation specialist conducting a comprehensive review of technical documentation.

REVIEW PERSPECTIVE:
You are reviewing this content from the perspective of a senior technical writer with expertise in:
- Department of Defense cybersecurity documentation
- RMF, NIST, and FedRAMP compliance documentation
- Zero Trust Architecture implementations
- System Security Plans and ATO documentation

REVIEW CRITERIA:

1. COHESION (Score 1-10):
   - Does the content flow logically?
   - Are transitions between ideas smooth?
   - Does it connect well with preceding and following sections?
   - Is the narrative consistent with document purpose?

2. CLARITY (Score 1-10):
   - Is the writing clear and unambiguous?
   - Are technical terms properly defined when first used?
   - Is the reading level appropriate for the audience?
   - Are sentences concise without sacrificing accuracy?

3. TECHNICAL ACCURACY (Score 1-10):
   - Are technical details correct?
   - Are control references (NIST, STIG) accurate?
   - Are tool names and capabilities correctly stated?
   - Are implementation details technically sound?

4. FACTUAL VERACITY (Score 1-10):
   - Are all statements factually correct?
   - Are statistics and metrics accurate?
   - Are compliance requirements correctly cited?
   - Are claims properly supported?

5. COMPLETENESS (Score 1-10):
   - Does it fully address the section's purpose?
   - Are all required elements present?
   - Are there gaps in coverage?
   - Does it answer likely reader questions?

REVIEW OUTPUT FORMAT:
Provide your review in the following JSON structure:
```json
{
  "scores": {
    "cohesion": <1-10>,
    "clarity": <1-10>,
    "technical_accuracy": <1-10>,
    "factual_veracity": <1-10>,
    "completeness": <1-10>
  },
  "feedback": {
    "cohesion": ["specific issue 1", "specific issue 2"],
    "clarity": ["specific issue 1", "specific issue 2"],
    "technical_accuracy": ["specific issue 1", "specific issue 2"],
    "factual_veracity": ["specific issue 1", "specific issue 2"],
    "completeness": ["specific issue 1", "specific issue 2"],
    "general": ["overall observations"]
  },
  "recommendations": [
    "Priority recommendation 1",
    "Priority recommendation 2",
    "Priority recommendation 3"
  ],
  "content_suggestions": {
    "additions": "Suggested content to add...",
    "revisions": "Suggested revisions to existing content...",
    "deletions": "Suggested content to remove or consolidate..."
  }
}
```

Be specific and actionable in your feedback. Cite specific sentences or phrases when identifying issues.
"""
        return review_prompt

    def conduct_section_review(self):
            """Conduct technical writer review of selected section"""
            if not self.selected_section:
                messagebox.showwarning("No Selection", "Please select a section to review")
                return
            
            if not self.selected_section.has_content():
                messagebox.showinfo("No Content", "Selected section has no content to review")
                return
            
            # Confirm review
            result = messagebox.askyesno(
                "Conduct Review",
                f"Conduct technical writing review of:\n\n{self.selected_section.get_full_path()}\n\n"
                "This will analyze the section for:\n"
                "• Cohesion\n"
                "• Clarity\n"
                "• Technical Accuracy\n"
                "• Factual Veracity\n"
                "• Completeness\n\n"
                "Continue?"
            )
            
            if not result:
                return
            
            # Disable review button and start review
            self.review_btn.config(state=tk.DISABLED)
            self.log_message(f"Starting technical review of: {self.selected_section.get_full_path()}")
            
            # Run review in thread
            thread = threading.Thread(target=self._execute_review)
            thread.daemon = True
            thread.start()
        
    def _execute_review(self):
        """Execute the review process in background thread"""
        try:
            section_content = self.selected_section.get_existing_content()
            section_path = self.selected_section.get_full_path()
            
            # Build review prompt
            review_prompt = self.set_review_prompt()
            
            # Get context from parent sections
            context = self._get_section_context()
            
            # Build full prompt
            full_prompt = f"{review_prompt}\n\n"
            full_prompt += f"SECTION BEING REVIEWED:\n"
            full_prompt += f"Path: {section_path}\n\n"
            
            if context:
                full_prompt += f"DOCUMENT CONTEXT:\n{context}\n\n"
            
            full_prompt += f"SECTION CONTENT TO REVIEW:\n{section_content}\n\n"
            full_prompt += "Provide your comprehensive technical writing review in the specified JSON format."
            
            # Call OpenWebUI API
            review_response = self.query_openwebui(full_prompt)

            if review_response:
                review = self._parse_review_response(review_response, section_path)
                
                if review and self.advanced_reviewer:
                    try:
                        tense_analysis = self.advanced_reviewer.analyze_tense_consistency(section_content)
                        review.feedback['tense_consistency'] = [
                            f"Dominant tense: {tense_analysis.dominant_tense}",
                            f"Consistency score: {tense_analysis.consistency_score:.1f}/10"
                        ]
                        if tense_analysis.inconsistent_sentences:
                            review.feedback['tense_consistency'].append(
                                f"Found {len(tense_analysis.inconsistent_sentences)} inconsistent sentences"
                            )
                        
                        readability = self.advanced_reviewer.calculate_readability_metrics(section_content)
                        review.feedback['readability'] = [
                            f"Flesch Reading Ease: {readability.flesch_score:.1f}",
                            f"Grade Level: {readability.grade_level:.1f}",
                            f"Avg Sentence Length: {readability.avg_sentence_length:.1f} words"
                        ]
                        
                        self.log_message("✓ Enhanced with advanced analysis")
                    except Exception as e:
                        self.log_message(f"Note: Advanced analysis not applied: {e}")

                if review:
                    self.current_review = review
                    self.review_history[section_path] = review
                    
                    # Update UI on main thread
                    self.root.after(0, self._display_review_results, review)
                else:
                    self.root.after(0, self._show_review_parse_error)
            else:
                self.root.after(0, self._show_review_api_error)
                
        except Exception as ex:
            error_msg = str(ex)
            self.log_message(f"Review error: {error_msg}")
            self.root.after(0, self._show_review_exception, error_msg)
        finally:
            # Re-enable button
            self.root.after(0, self._enable_review_button)

    def _show_review_parse_error(self):
        """Show parse error message"""
        messagebox.showerror("Review Error", "Could not parse review response")

    def _show_review_api_error(self):
        """Show API error message"""
        messagebox.showerror("Review Error", "Failed to get review from API")

    def _show_review_exception(self, error_msg):
        """Show exception error message"""
        messagebox.showerror("Review Error", f"Review failed: {error_msg}")

    def _enable_review_button(self):
        """Re-enable review button"""
        self.review_btn.config(state=tk.NORMAL)

    def _get_section_context(self) -> str:
        """Get context from parent and sibling sections"""
        context = ""
        
        if self.selected_section.parent:
            parent_content = self.selected_section.parent.get_existing_content()
            if parent_content:
                context += f"Parent Section: {self.selected_section.parent.text}\n"
                context += f"{parent_content[:500]}...\n\n"
        
        return context
    
    def _parse_review_response(self, response: str, section_path: str) -> DocumentReview:
        """Parse JSON review response into DocumentReview object"""
        try:
            # Extract JSON from response
            json_match = re.search(r'```json\s*(\{.*?\})\s*```', response, re.DOTALL)
            if not json_match:
                # Try without code blocks
                json_match = re.search(r'(\{.*\})', response, re.DOTALL)
            
            if not json_match:
                self.log_message("Could not find JSON in review response")
                return None
            
            review_data = json.loads(json_match.group(1))
            
            # Create review object
            review = DocumentReview(section_path)
            
            # Set scores
            scores = review_data.get('scores', {})
            review.cohesion_score = scores.get('cohesion', 0)
            review.clarity_score = scores.get('clarity', 0)
            review.technical_accuracy_score = scores.get('technical_accuracy', 0)
            review.factual_veracity_score = scores.get('factual_veracity', 0)
            review.completeness_score = scores.get('completeness', 0)
            
            # Set feedback
            review.feedback = review_data.get('feedback', {})
            
            # Set recommendations
            review.recommendations = review_data.get('recommendations', [])
            
            # Set content suggestions
            review.content_suggestions = review_data.get('content_suggestions', {})
            
            review.calculate_overall_score()
            
            return review
            
        except json.JSONDecodeError as e:
            self.log_message(f"JSON parse error: {str(e)}")
            return None
        except Exception as e:
            self.log_message(f"Review parse error: {str(e)}")
            return None
    
    def _display_review_results(self, review: DocumentReview):
        """Display review results in the UI"""
        # Clear existing content
        self.generated_text.delete(1.0, tk.END)
        
        # Display summary
        summary = review.get_summary()
        self.generated_text.insert(tk.END, summary)
        self.generated_text.insert(tk.END, "\n" + "="*50 + "\n\n")
        
        # Display detailed feedback
        self.generated_text.insert(tk.END, "=== DETAILED FEEDBACK ===\n\n")
        
        for category, items in review.feedback.items():
            if items:
                category_title = category.replace('_', ' ').title()
                self.generated_text.insert(tk.END, f"{category_title}:\n")
                for item in items:
                    self.generated_text.insert(tk.END, f"  • {item}\n")
                self.generated_text.insert(tk.END, "\n")
        
        # Display content suggestions
        if review.content_suggestions:
            self.generated_text.insert(tk.END, "\n=== CONTENT SUGGESTIONS ===\n\n")
            
            if review.content_suggestions.get('additions'):
                self.generated_text.insert(tk.END, "Suggested Additions:\n")
                self.generated_text.insert(tk.END, review.content_suggestions['additions'])
                self.generated_text.insert(tk.END, "\n\n")
            
            if review.content_suggestions.get('revisions'):
                self.generated_text.insert(tk.END, "Suggested Revisions:\n")
                self.generated_text.insert(tk.END, review.content_suggestions['revisions'])
                self.generated_text.insert(tk.END, "\n\n")
            
            if review.content_suggestions.get('deletions'):
                self.generated_text.insert(tk.END, "Suggested Deletions:\n")
                self.generated_text.insert(tk.END, review.content_suggestions['deletions'])
                self.generated_text.insert(tk.END, "\n\n")
        
        # Log completion
        self.log_message(f"Review completed. Overall score: {review.overall_score:.1f}/10")
        
        # Enable action buttons
        self.enable_review_actions()
    
    def enable_review_actions(self):
        """Enable review-based action buttons"""
        self.apply_suggestions_btn.config(state=tk.NORMAL)
        self.regenerate_from_review_btn.config(state=tk.NORMAL)

    def apply_review_suggestions(self):
        """Apply review suggestions by generating improved content"""
        if not self.current_review:
            messagebox.showwarning("No Review", "No active review to apply")
            return
        
        result = messagebox.askyesno(
            "Apply Suggestions",
            "Generate improved content based on review feedback?\n\n"
            "This will:\n"
            "1. Use the review recommendations\n"
            "2. Generate new content addressing issues\n"
            "3. Show you the result for approval\n\n"
            "Continue?"
        )
        
        if not result:
            return
        
        self.apply_suggestions_btn.config(state=tk.DISABLED)
        self.log_message("Generating improved content based on review...")
        
        thread = threading.Thread(target=self._generate_improved_content)
        thread.daemon = True
        thread.start()

    def _generate_improved_content(self):
        """Generate improved content based on review feedback"""
        try:
            section_content = self.selected_section.get_existing_content()
            review = self.current_review
            
            # Build improvement prompt
            prompt = f"{self.master_prompt.get()}\n\n"
            prompt += "TASK: Improve the following content based on technical writing review feedback.\n\n"
            prompt += f"SECTION: {self.selected_section.get_full_path()}\n\n"
            prompt += f"CURRENT CONTENT:\n{section_content}\n\n"
            prompt += f"REVIEW SCORES:\n"
            prompt += f"  • Cohesion: {review.cohesion_score}/10\n"
            prompt += f"  • Clarity: {review.clarity_score}/10\n"
            prompt += f"  • Technical Accuracy: {review.technical_accuracy_score}/10\n"
            prompt += f"  • Factual Veracity: {review.factual_veracity_score}/10\n"
            prompt += f"  • Completeness: {review.completeness_score}/10\n\n"
            
            prompt += "KEY ISSUES TO ADDRESS:\n"
            for category, items in review.feedback.items():
                if items:
                    prompt += f"\n{category.replace('_', ' ').title()}:\n"
                    for item in items:
                        prompt += f"  • {item}\n"
            
            prompt += "\nRECOMMENDATIONS:\n"
            for i, rec in enumerate(review.recommendations, 1):
                prompt += f"  {i}. {rec}\n"
            
            if review.content_suggestions:
                prompt += "\nCONTENT SUGGESTIONS:\n"
                if review.content_suggestions.get('additions'):
                    prompt += f"\nAdditions:\n{review.content_suggestions['additions']}\n"
                if review.content_suggestions.get('revisions'):
                    prompt += f"\nRevisions:\n{review.content_suggestions['revisions']}\n"
                if review.content_suggestions.get('deletions'):
                    prompt += f"\nDeletions:\n{review.content_suggestions['deletions']}\n"
            
            prompt += "\n\nGenerate improved content that addresses all review feedback while maintaining technical accuracy and appropriate detail level."
            
            # Generate improved content
            improved_content = self.query_openwebui(prompt)
            
            if improved_content:
                self.generated_content = improved_content
                
                # Update UI
                self.root.after(0, self._show_improved_content, improved_content)
            else:
                self.root.after(0, self._show_generation_error)
                
        except Exception as ex:
            error_msg = str(ex)
            self.log_message(f"Content improvement error: {error_msg}")
            self.root.after(0, self._show_improvement_exception, error_msg)
        finally:
            self.root.after(0, self._enable_apply_suggestions_button)

    def _show_generation_error(self):
        """Show generation error message"""
        messagebox.showerror("Generation Error", "Failed to generate improved content")

    def _show_improvement_exception(self, error_msg):
        """Show improvement exception message"""
        messagebox.showerror("Error", f"Failed to improve content: {error_msg}")

    def _enable_apply_suggestions_button(self):
        """Re-enable apply suggestions button"""
        self.apply_suggestions_btn.config(state=tk.NORMAL)    

    def _show_improved_content(self, content: str):
        """Display improved content"""
        self.generated_text.delete(1.0, tk.END)
        self.generated_text.insert(tk.END, "=== IMPROVED CONTENT ===\n\n")
        self.generated_text.insert(tk.END, content)
        
        self.log_message("Improved content generated. Review and insert if acceptable.")
        
        # Enable commit button
        self.commit_btn.config(state=tk.NORMAL)
    
    def regenerate_from_review(self):
        """Regenerate content completely based on review requirements"""
        if not self.current_review:
            messagebox.showwarning("No Review", "No active review available")
            return
        
        result = messagebox.askyesno(
            "Regenerate Content",
            "Completely regenerate this section based on review feedback?\n\n"
            "This will create entirely new content that addresses the identified issues.\n\n"
            "Continue?"
        )
        
        if not result:
            return
        
        self.regenerate_from_review_btn.config(state=tk.DISABLED)
        self.log_message("Regenerating content from scratch based on review...")
        
        thread = threading.Thread(target=self._regenerate_content)
        thread.daemon = True
        thread.start()

    def _regenerate_content(self):
        """Regenerate content from scratch based on review"""
        try:
            review = self.current_review
            
            # Build regeneration prompt
            prompt = f"{self.master_prompt.get()}\n\n"
            prompt += f"TASK: Write comprehensive content for the following section, addressing specific quality requirements.\n\n"
            prompt += f"SECTION: {self.selected_section.get_full_path()}\n\n"
            
            prompt += "QUALITY REQUIREMENTS (based on previous review):\n\n"
            
            if review.cohesion_score < 8:
                prompt += "COHESION: Ensure strong logical flow and smooth transitions. "
                if review.feedback.get('cohesion'):
                    prompt += "Address: " + "; ".join(review.feedback['cohesion'][:3]) + "\n"
            
            if review.clarity_score < 8:
                prompt += "CLARITY: Write clearly with well-defined terms and appropriate reading level. "
                if review.feedback.get('clarity'):
                    prompt += "Address: " + "; ".join(review.feedback['clarity'][:3]) + "\n"
            
            if review.technical_accuracy_score < 8:
                prompt += "TECHNICAL ACCURACY: Ensure all technical details and references are correct. "
                if review.feedback.get('technical_accuracy'):
                    prompt += "Address: " + "; ".join(review.feedback['technical_accuracy'][:3]) + "\n"
            
            if review.factual_veracity_score < 8:
                prompt += "FACTUAL VERACITY: Verify all statements and properly cite requirements. "
                if review.feedback.get('factual_veracity'):
                    prompt += "Address: " + "; ".join(review.feedback['factual_veracity'][:3]) + "\n"
            
            if review.completeness_score < 8:
                prompt += "COMPLETENESS: Provide comprehensive coverage of all required elements. "
                if review.feedback.get('completeness'):
                    prompt += "Address: " + "; ".join(review.feedback['completeness'][:3]) + "\n"
            
            prompt += "\nKEY RECOMMENDATIONS TO IMPLEMENT:\n"
            for i, rec in enumerate(review.recommendations, 1):
                prompt += f"  {i}. {rec}\n"
            
            if review.content_suggestions.get('additions'):
                prompt += f"\nINCLUDE THESE ELEMENTS:\n{review.content_suggestions['additions']}\n"
            
            prompt += "\n\nWrite comprehensive, high-quality content for this section that achieves 9-10 scores in all review criteria."
            
            # Generate new content
            new_content = self.query_openwebui(prompt)
            
            if new_content:
                self.generated_content = new_content
                
                # Update UI
                self.root.after(0, self._show_regenerated_content, new_content)
            else:
                self.root.after(0, self._show_regeneration_error)
                
        except Exception as ex:
            error_msg = str(ex)
            self.log_message(f"Content regeneration error: {error_msg}")
            self.root.after(0, self._show_regeneration_exception, error_msg)
        finally:
            self.root.after(0, self._enable_regenerate_button)

    def _show_regeneration_error(self):
        """Show regeneration error message"""
        messagebox.showerror("Generation Error", "Failed to regenerate content")

    def _show_regeneration_exception(self, error_msg):
        """Show regeneration exception message"""
        messagebox.showerror("Error", f"Failed to regenerate content: {error_msg}")

    def _enable_regenerate_button(self):
        """Re-enable regenerate button"""
        self.regenerate_from_review_btn.config(state=tk.NORMAL)    

    def _show_regenerated_content(self, content: str):
        """Display regenerated content"""
        self.generated_text.delete(1.0, tk.END)
        self.generated_text.insert(tk.END, "=== REGENERATED CONTENT ===\n\n")
        self.generated_text.insert(tk.END, content)
        
        self.log_message("Content regenerated. Review and insert if acceptable.")
        
        # Enable commit button
        self.commit_btn.config(state=tk.NORMAL)

    def load_settings(self, config_file=None):
        """Load saved settings"""
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            
            # Use provided config file or default
            if not config_file:
                config_file = os.path.join(script_dir, "openwebui_config.json")
            
            if os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    config = json.load(f)
                    self.openwebui_base_url = config.get('base_url', 'http://172.16.27.122:3000')
                    self.openwebui_api_key = config.get('api_key', '')
                    self.selected_model.set(config.get('model', ''))
                    self.temperature.set(config.get('temperature', 0.1))
                    self.max_tokens.set(config.get('max_tokens', 8000))
                    
                    # Load RAG knowledge collections
                    if 'knowledge_collections' in config:
                        self.selected_knowledge_collections = config['knowledge_collections']
                    
                    # Load master prompt
                    if 'master_prompt' in config:
                        self.master_prompt.set(config.get('master_prompt'))
                    
                    if 'format_config' in config:
                        fmt = config['format_config']
                        self.format_config['highlight_enabled'].set(fmt.get('highlight_enabled', True))
                        self.format_config['highlight_color'].set(fmt.get('highlight_color', 'YELLOW'))
                        self.format_config['bold_enabled'].set(fmt.get('bold_enabled', False))
                        self.format_config['italic_enabled'].set(fmt.get('italic_enabled', False))
                        self.format_config['underline_enabled'].set(fmt.get('underline_enabled', False))
                        self.format_config['font_color'].set(fmt.get('font_color', '000000'))
                        self.format_config['font_size'].set(fmt.get('font_size', 11))
                    
                    if 'auto_config' in config:
                        auto = config['auto_config']
                        self.auto_config['auto_backup'].set(auto.get('auto_backup', True))
                        self.auto_config['backup_interval'].set(auto.get('backup_interval', 5))
                        self.auto_config['auto_save'].set(auto.get('auto_save', False))
                        self.auto_config['auto_reload'].set(auto.get('auto_reload', True))
                        self.auto_config['ask_backup'].set(auto.get('ask_backup', True))
            
            last_doc_file = os.path.join(script_dir, "last_document.txt")
            if os.path.exists(last_doc_file):
                with open(last_doc_file, 'r') as f:
                    self.last_document_path = f.read().strip()
                    
        except Exception as e:
            print(f"Note: Could not load settings: {str(e)}")

    def save_settings(self):
        """Save current settings"""
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            config_file = os.path.join(script_dir, "openwebui_config.json")
            
            # Create knowledge collections list with minimal info
            saved_collections = []
            for col in self.selected_knowledge_collections:
                saved_collections.append({
                    'id': col['id'],
                    'name': col['name']
                })
            
            config = {
                'base_url': self.openwebui_base_url,
                'api_key': self.openwebui_api_key,
                'model': self.selected_model.get(),
                'temperature': self.temperature.get(),
                'max_tokens': self.max_tokens.get(),
                'knowledge_collections': saved_collections,
                'master_prompt': self.master_prompt.get(),
                'format_config': {
                    'highlight_enabled': self.format_config['highlight_enabled'].get(),
                    'highlight_color': self.format_config['highlight_color'].get(),
                    'bold_enabled': self.format_config['bold_enabled'].get(),
                    'italic_enabled': self.format_config['italic_enabled'].get(),
                    'underline_enabled': self.format_config['underline_enabled'].get(),
                    'font_color': self.format_config['font_color'].get(),
                    'font_size': self.format_config['font_size'].get()
                },
                'auto_config': {
                    'auto_backup': self.auto_config['auto_backup'].get(),
                    'backup_interval': self.auto_config['backup_interval'].get(),
                    'auto_save': self.auto_config['auto_save'].get(),
                    'auto_reload': self.auto_config['auto_reload'].get(),
                    'ask_backup': self.auto_config['ask_backup'].get()
                }
            }
            
            with open(config_file, 'w') as f:
                json.dump(config, f, indent=2)
                
            if self.document_path:
                last_doc_file = os.path.join(script_dir, "last_document.txt")
                with open(last_doc_file, 'w') as f:
                    f.write(self.document_path)
                    
            self.log_message("Settings saved")
            
            if self.credential_manager and self.credential_manager.is_encrypted:
                try:
                    self.save_to_encrypted_credentials()
                    self.log_message("✓ Also backed up to encrypted credentials")
                except Exception as e:
                    self.log_message(f"Note: Encrypted backup not saved: {e}")
        

        except Exception as e:
            self.log_message(f"Error saving settings: {str(e)}")

    def load_section_tracking(self):
        """Load section tracking from file"""
        if not self.document_path:
            return
        
        try:
            doc_dir = os.path.dirname(self.document_path)
            doc_name = os.path.splitext(os.path.basename(self.document_path))[0]
            self.tracking_file = os.path.join(doc_dir, f".{doc_name}_tracking.json")
            
            if os.path.exists(self.tracking_file):
                with open(self.tracking_file, 'r') as f:
                    self.section_tracking = json.load(f)
                self.log_message(f"Loaded tracking for {len(self.section_tracking)} sections")
        except Exception as e:
            self.log_message(f"Note: Could not load section tracking: {str(e)}")
    
    def save_section_tracking(self):
        """Save section tracking to file"""
        if not self.tracking_file:
            return
        
        try:
            with open(self.tracking_file, 'w') as f:
                json.dump(self.section_tracking, f, indent=2)
        except Exception as e:
            self.log_message(f"Error saving section tracking: {str(e)}")
    
    def mark_section_edited(self, section):
        """Mark a section as edited in tracking"""
        section_hash = section.get_section_hash()
        self.section_tracking[section_hash] = {
            'edited': True,
            'last_modified': datetime.now().isoformat(),
            'section_path': section.get_full_path()
        }
        self.save_section_tracking()
    
    def is_section_edited(self, section):
        """Check if section has been edited"""
        section_hash = section.get_section_hash()
        return self.section_tracking.get(section_hash, {}).get('edited', False)
    
    def schedule_auto_backup(self):
        """Schedule automatic backup"""
        if self.auto_config['auto_backup'].get() and self.document_path:
            interval = self.auto_config['backup_interval'].get() * 60000  # Convert to ms
            self.root.after(interval, self.perform_auto_backup)
    
    def perform_auto_backup(self):
        """Perform automatic backup"""
        if self.document and self.document_path:
            try:
                self.create_backup()
                self.log_message("Auto-backup completed")
            except Exception as e:
                self.log_message(f"Auto-backup failed: {str(e)}")
        
        # Reschedule
        if self.auto_config['auto_backup'].get():
            self.schedule_auto_backup()

    def browse_config_file(self):
        """Browse for a different configuration file"""
        filename = filedialog.askopenfilename(
            title="Select Configuration File",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        if filename:
            try:
                self.load_settings(filename)
                self.update_config_status()
                self.log_message(f"Loaded configuration from: {filename}")
                messagebox.showinfo("Success", "Configuration loaded successfully")
            except Exception as e:
                self.log_message(f"Error loading configuration: {str(e)}")
                messagebox.showerror("Error", f"Failed to load configuration: {str(e)}")

    def save_config_as(self):
        """Save configuration to a new file"""
        filename = filedialog.asksaveasfilename(
            title="Save Configuration As",
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        if filename:
            try:
                script_dir = os.path.dirname(os.path.abspath(__file__))
                
                # Create knowledge collections list with minimal info
                saved_collections = []
                for col in self.selected_knowledge_collections:
                    saved_collections.append({
                        'id': col['id'],
                        'name': col['name']
                    })
                
                config = {
                    'base_url': self.openwebui_base_url,
                    'api_key': self.openwebui_api_key,
                    'model': self.selected_model.get(),
                    'temperature': self.temperature.get(),
                    'max_tokens': self.max_tokens.get(),
                    'knowledge_collections': saved_collections,
                    'master_prompt': self.master_prompt.get(),
                    'format_config': {
                        'highlight_enabled': self.format_config['highlight_enabled'].get(),
                        'highlight_color': self.format_config['highlight_color'].get(),
                        'bold_enabled': self.format_config['bold_enabled'].get(),
                        'italic_enabled': self.format_config['italic_enabled'].get(),
                        'underline_enabled': self.format_config['underline_enabled'].get(),
                        'font_color': self.format_config['font_color'].get(),
                        'font_size': self.format_config['font_size'].get()
                    },
                    'auto_config': {
                        'auto_backup': self.auto_config['auto_backup'].get(),
                        'backup_interval': self.auto_config['backup_interval'].get(),
                        'auto_save': self.auto_config['auto_save'].get(),
                        'auto_reload': self.auto_config['auto_reload'].get(),
                        'ask_backup': self.auto_config['ask_backup'].get()
                    }
                }
                
                with open(filename, 'w') as f:
                    json.dump(config, f, indent=2)
                
                self.log_message(f"Configuration saved to: {filename}")
                messagebox.showinfo("Success", "Configuration saved successfully")
            except Exception as e:
                self.log_message(f"Error saving configuration: {str(e)}")
                messagebox.showerror("Error", f"Failed to save configuration: {str(e)}")

    def setup_text_widgets_colors(self):
        """Enhanced text widget styling with additional markdown support"""
        try:
            # Configure existing tags
            self.generated_text.tag_configure('bold', font=('TkDefaultFont', 10, 'bold'))
            self.generated_text.tag_configure('italic', font=('TkDefaultFont', 10, 'italic'))
            self.generated_text.tag_configure('code', 
                                            font=('Consolas', 9),
                                            background='#f0f0f0',
                                            foreground='#008000')
            
            # Add new enhanced tags
            self.generated_text.tag_configure('highlight', 
                                            background='#ffff00',
                                            foreground='#000000')
            
            self.generated_text.tag_configure('strike', 
                                            overstrike=True,
                                            foreground='#808080')
            
            self.generated_text.tag_configure('quote', 
                                            font=('TkDefaultFont', 10, 'italic'),
                                            foreground='#606060',
                                            lmargin1=20,
                                            lmargin2=20)
            
            self.generated_text.tag_configure('hr', 
                                            foreground='#808080',
                                            justify='center')
            
            self.generated_text.tag_configure('code_lang', 
                                            font=('TkDefaultFont', 8),
                                            foreground='#808080')
            
            # Table styling
            self.generated_text.tag_configure('table_header', 
                                            font=('TkDefaultFont', 10, 'bold'),
                                            foreground='#000080')
            
            self.generated_text.tag_configure('table_sep', 
                                            foreground='#808080')
            
            self.generated_text.tag_configure('table_row', 
                                            font=('Consolas', 9),
                                            foreground='#404040')
            
            # Apply dark theme colors if using dark theme
            if hasattr(self, 'is_dark_theme') and self.is_dark_theme:
                self.generated_text.configure(bg='#1e1e1e', fg='#ffffff', insertbackground='#ffffff')
                
                # Adjust colors for dark theme
                self.generated_text.tag_configure('code', 
                                                background='#2d2d2d',
                                                foreground='#90ee90')
                
                self.generated_text.tag_configure('quote', 
                                                foreground='#c0c0c0')
                
                self.generated_text.tag_configure('table_header', 
                                                foreground='#87ceeb')
                
                self.generated_text.tag_configure('table_row', 
                                                foreground='#d0d0d0')
                
                self.generated_text.tag_configure('hr', 
                                                foreground='#c0c0c0')
                
                self.generated_text.tag_configure('code_lang', 
                                                foreground='#c0c0c0')
            
            # Configure prompt text widget
            if hasattr(self, 'prompt_text'):
                self.prompt_text.tag_configure('bold', font=('TkDefaultFont', 10, 'bold'))
                self.prompt_text.tag_configure('italic', font=('TkDefaultFont', 10, 'italic'))
                self.prompt_text.tag_configure('code', 
                                            font=('Consolas', 9),
                                            background='#f0f0f0' if not hasattr(self, 'is_dark_theme') or not self.is_dark_theme else '#2d2d2d',
                                            foreground='#008000' if not hasattr(self, 'is_dark_theme') or not self.is_dark_theme else '#90ee90')
                
        except Exception as e:
            print(f"Error setting up enhanced text widget colors: {e}")

    def create_gui(self):
        """Create main GUI"""
        main_container = ttk.Frame(self.root, padding="10")
        main_container.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_container.columnconfigure(0, weight=1)
        main_container.rowconfigure(2, weight=1)
        
        # Title row with buttons
        title_frame = ttk.Frame(main_container)
        title_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        title_frame.columnconfigure(0, weight=1)

        ttk.Label(title_frame, text="Document Content Generator - OpenWebUI",
                font=("Arial", 16, "bold")).pack(side=tk.LEFT)

        # Configuration buttons on title row
        btn_container = ttk.Frame(title_frame)
        btn_container.pack(side=tk.RIGHT)

        # Reduce button padding and add more buttons to top row
        btn_style = {'padx': 1, 'pady': 2, 'side': tk.LEFT}

        ttk.Button(btn_container, text="⚙ Config Files",
                command=self.open_config_file_manager).pack(**btn_style)
        ttk.Button(btn_container, text="⚙ Formatting",
                command=self.open_formatting_dialog).pack(**btn_style)
        ttk.Button(btn_container, text="🔄 Auto",
                command=self.open_auto_features_dialog).pack(**btn_style)
        ttk.Button(btn_container, text="📎 External RAG",
                command=self.open_external_content_manager).pack(**btn_style)
        ttk.Button(btn_container, text="✏️ Master Prompt",
                command=self.edit_master_prompt).pack(**btn_style)
        ttk.Button(btn_container, text="📚 Prompts",
                command=self.open_prompt_library).pack(**btn_style)
        ttk.Button(btn_container, text="🔐 Credentials",
                command=self.manage_encrypted_credentials_dialog).pack(**btn_style)
        
        # Top controls - Side-by-side layout with reduced padding
        top_frame = ttk.Frame(main_container)
        top_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        top_frame.columnconfigure(0, weight=1)
        top_frame.columnconfigure(1, weight=1)

        # Document selection - Left side with reduced padding
        doc_frame = ttk.LabelFrame(top_frame, text="Document", padding="5")
        doc_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        doc_frame.columnconfigure(2, weight=1)

        ttk.Button(doc_frame, text="Load", command=self.browse_document).grid(row=0, column=0, padx=(0, 3))

        # Reload button - initially disabled until document is loaded
        self.reload_btn = ttk.Button(doc_frame, text="Reload", command=self.reload_document_wrapper, state=tk.DISABLED)
        self.reload_btn.grid(row=0, column=1, padx=(0, 3))

        self.doc_label_var = tk.StringVar(value="No document loaded")
        ttk.Label(doc_frame, textvariable=self.doc_label_var).grid(row=0, column=2, sticky=tk.W, padx=3)

        # OpenWebUI configuration - Right side with reduced padding
        config_frame = ttk.LabelFrame(top_frame, text="OpenWebUI", padding="5")
        config_frame.grid(row=0, column=1, sticky=(tk.W, tk.E))
        config_frame.columnconfigure(1, weight=1)

        ttk.Button(config_frame, text="Configure AI", command=self.open_config_dialog).grid(row=0, column=0, padx=(0, 5))
        self.config_status_var = tk.StringVar(value="Not configured")
        ttk.Label(config_frame, textvariable=self.config_status_var).grid(row=0, column=1, sticky=tk.W)
        
        # Main content area
        content_frame = ttk.Frame(main_container)
        content_frame.grid(row=2, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        content_frame.columnconfigure(0, weight=1)
        content_frame.columnconfigure(1, weight=2)
        content_frame.rowconfigure(0, weight=1)
        
        # Left panel
        left_panel = ttk.Frame(content_frame)
        left_panel.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        left_panel.columnconfigure(0, weight=1)
        left_panel.rowconfigure(1, weight=1)
        
        # Section tree - Reduced padding
        tree_frame = ttk.LabelFrame(left_panel, text="Document Sections", padding="3")
        tree_frame.grid(row=0, column=0, rowspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 5))
        tree_frame.columnconfigure(0, weight=1)
        tree_frame.rowconfigure(0, weight=1)

        self.tree = ttk.Treeview(tree_frame, selectmode='browse')
        tree_scroll = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=tree_scroll.set)

        self.tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        tree_scroll.grid(row=0, column=1, sticky=(tk.N, tk.S))

        self.tree.bind('<<TreeviewSelect>>', self.on_section_select)

        # Operation controls - Reduced padding
        op_frame = ttk.LabelFrame(left_panel, text="Operation Mode", padding="5")
        op_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        
        ttk.Radiobutton(op_frame, text="Replace (write from scratch)", 
                    variable=self.operation_mode, value="replace").pack(anchor=tk.W)
        ttk.Radiobutton(op_frame, text="Rework (rewrite existing)", 
                    variable=self.operation_mode, value="rework").pack(anchor=tk.W)
        ttk.Radiobutton(op_frame, text="Append (add to existing)", 
                    variable=self.operation_mode, value="append").pack(anchor=tk.W)
        
        # Action buttons - Reduced padding
        button_frame = ttk.Frame(left_panel)
        button_frame.grid(row=3, column=0, sticky=(tk.W, tk.E))

        # Reduced padding for all buttons
        btn_pady = (0, 5)

        self.generate_btn = ttk.Button(button_frame, text="Generate Content",
                                    command=self.generate_content, state='disabled')
        self.generate_btn.pack(fill=tk.X, pady=btn_pady)

        # Review button
        self.review_btn = ttk.Button(button_frame, text="📝 Review Section",
                                    command=self.conduct_section_review, state='disabled')
        self.review_btn.pack(fill=tk.X, pady=btn_pady)

        self.tense_btn = ttk.Button(button_frame, text="🎯 Analyze Tenses",
                                    command=self.analyze_document_tenses, state='disabled')
        self.tense_btn.pack(fill=tk.X, pady=btn_pady)

        # Processing Strategy button
        ttk.Button(button_frame, text="🧠 Processing Strategy",
                command=self.show_processing_strategy_dialog).pack(fill=tk.X, pady=btn_pady)

        # Apply review suggestions button
        self.apply_suggestions_btn = ttk.Button(button_frame, text="✅ Apply Suggestions",
                                            command=self.apply_review_suggestions,
                                            state='disabled')
        self.apply_suggestions_btn.pack(fill=tk.X, pady=btn_pady)

        # Regenerate from review button
        self.regenerate_from_review_btn = ttk.Button(button_frame, text="🔄 Regenerate Review",
                                                    command=self.regenerate_from_review,
                                                    state='disabled')
        self.regenerate_from_review_btn.pack(fill=tk.X, pady=btn_pady)

        # Whole Document Review
        ttk.Button(button_frame, text="📋 Review Document",
                command=self.review_whole_document).pack(fill=tk.X, pady=btn_pady)

        ttk.Button(button_frame, text="🚀 Auto Complete",
                command=self.auto_complete_document).pack(fill=tk.X, pady=btn_pady)
        
        # Right panel
        right_panel = ttk.Frame(content_frame)
        right_panel.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S))
        right_panel.columnconfigure(0, weight=1)
        right_panel.rowconfigure(0, weight=1)
        
        # Create notebook for tabs
        self.notebook = ttk.Notebook(right_panel)
        self.notebook.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Bind tab change event
        self.notebook.bind('<<NotebookTabChanged>>', self.on_tab_changed)
        
        # Preview tab
        self.create_preview_tab()
        
        # Prompt tab
        self.create_prompt_tab()
        
        # NEW: Prompt History tab
        self.create_prompt_history_tab()

        # NEW: Section Chat tab
        self.create_section_chat_tab()

        # Console tab
        self.create_console_tab()
        
        # Update config status
        self.update_config_status()   

        self.add_markdown_export_functionality()
        self.add_table_editing_dialog()

    def analyze_document_tenses(self):
        """NEW FEATURE: Dedicated tense consistency analysis"""
        if not self.selected_section:
            messagebox.showwarning("No Selection", "Please select a section to analyze")
            return
        
        if not self.selected_section.has_content():
            messagebox.showinfo("No Content", "Selected section has no content to analyze")
            return
        
        if not self.advanced_reviewer:
            messagebox.showinfo("Feature Unavailable", 
                            "Advanced tense analysis requires document_reviewer module.\n\n"
                            "To enable this feature:\n"
                            "1. Ensure document_reviewer.py is in the same directory\n"
                            "2. Install dependencies:\n"
                            "   pip install textstat nltk --break-system-packages\n"
                            "3. Run NLTK setup:\n"
                            "   python -c \"import nltk; nltk.download('punkt'); "
                            "nltk.download('averaged_perceptron_tagger'); "
                            "nltk.download('stopwords')\"\n"
                            "4. Restart the application")
            return
        
        content = self.selected_section.get_existing_content()
        
        try:
            self.log_message("Starting tense consistency analysis...")
            tense_analysis = self.advanced_reviewer.analyze_tense_consistency(content)
            
            # Display results in generated_text area
            self.generated_text.delete(1.0, tk.END)
            self.generated_text.insert(tk.END, "=== TENSE CONSISTENCY ANALYSIS ===\n\n")
            self.generated_text.insert(tk.END, f"Section: {self.selected_section.get_full_path()}\n\n")
            self.generated_text.insert(tk.END, f"Dominant Tense: {tense_analysis.dominant_tense.upper()}\n")
            self.generated_text.insert(tk.END, f"Consistency Score: {tense_analysis.consistency_score:.1f}/10\n\n")
            self.generated_text.insert(tk.END, f"Tense Distribution:\n")
            self.generated_text.insert(tk.END, f"  • Past: {tense_analysis.past_count} sentences\n")
            self.generated_text.insert(tk.END, f"  • Present: {tense_analysis.present_count} sentences\n")
            self.generated_text.insert(tk.END, f"  • Future: {tense_analysis.future_count} sentences\n\n")
            
            if tense_analysis.inconsistent_sentences:
                self.generated_text.insert(tk.END, 
                    f"Inconsistent Sentences ({len(tense_analysis.inconsistent_sentences)}):\n\n")
                for i, sentence in enumerate(tense_analysis.inconsistent_sentences[:10], 1):
                    self.generated_text.insert(tk.END, f"{i}. {sentence}\n\n")
                
                if len(tense_analysis.inconsistent_sentences) > 10:
                    self.generated_text.insert(tk.END, 
                        f"... and {len(tense_analysis.inconsistent_sentences) - 10} more\n")
            else:
                self.generated_text.insert(tk.END, "✓ No tense inconsistencies detected\n")
            
            self.log_message(f"Tense analysis completed: {tense_analysis.consistency_score:.1f}/10")
            self.notebook.select(0)  # Switch to preview tab
            
        except Exception as e:
            messagebox.showerror("Analysis Error", f"Failed to analyze tenses: {str(e)}")
            self.log_message(f"Tense analysis error: {e}")

    def show_processing_strategy_dialog(self):
        """NEW FEATURE: Display content processing strategy analysis"""
        if not self.content_processor:
            messagebox.showinfo("Feature Unavailable",
                            "Intelligent processing requires content_processor module.\n\n"
                            "To enable this feature:\n"
                            "1. Ensure content_processor.py is in the same directory\n"
                            "2. Install dependencies:\n"
                            "   pip install tiktoken scikit-learn numpy --break-system-packages\n"
                            "3. Restart the application")
            return
        
        if not self.selected_section or not self.selected_section.has_content():
            messagebox.showwarning("No Content", "Please select a section with content")
            return
        
        content = self.selected_section.get_existing_content()
        
        try:
            strategy = self.content_processor.determine_processing_strategy(
                content, [self.selected_section], "Generate comprehensive content"
            )
            
            # Create dialog
            dialog = tk.Toplevel(self.root)
            dialog.title("Processing Strategy Analysis")
            dialog.geometry("600x500")
            dialog.configure(bg="#2b2b2b")
            dialog.grab_set()
            
            main_frame = ttk.Frame(dialog, padding="20")
            main_frame.pack(fill=tk.BOTH, expand=True)
            
            ttk.Label(main_frame, text="Content Processing Strategy", 
                    font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 15))
            
            # Strategy info
            info_frame = ttk.LabelFrame(main_frame, text="Recommended Strategy", padding="10")
            info_frame.pack(fill=tk.X, pady=(0, 10))
            
            ttk.Label(info_frame, text=f"Method: {strategy.method.upper()}",
                    font=("Arial", 10, "bold")).pack(anchor=tk.W)
            ttk.Label(info_frame, text=f"Reason: {strategy.reason}").pack(anchor=tk.W, pady=5)
            ttk.Label(info_frame, text=f"Token Estimate: {strategy.token_estimate}").pack(anchor=tk.W)
            ttk.Label(info_frame, text=f"Confidence: {strategy.confidence:.1%}").pack(anchor=tk.W)
            
            # Metrics
            metrics_frame = ttk.LabelFrame(main_frame, text="Content Metrics", padding="10")
            metrics_frame.pack(fill=tk.X, pady=(0, 10))
            
            metrics = self.content_processor.analyze_content_metrics(content, [self.selected_section])
            ttk.Label(metrics_frame, text=f"Character Count: {metrics.char_count:,}").pack(anchor=tk.W)
            ttk.Label(metrics_frame, text=f"Token Count: ~{metrics.token_count:,}").pack(anchor=tk.W)
            ttk.Label(metrics_frame, text=f"Complexity Score: {metrics.complexity_score:.2f}").pack(anchor=tk.W)
            ttk.Label(metrics_frame, text=f"Technical Density: {metrics.technical_density:.2%}").pack(anchor=tk.W)
            
            # Explanation
            explain_frame = ttk.LabelFrame(main_frame, text="What This Means", padding="10")
            explain_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
            
            explain_text = tk.Text(explain_frame, wrap=tk.WORD, height=8, width=60,
                                bg="#1e1e1e", fg="#ffffff", font=("Consolas", 9))
            explain_text.pack(fill=tk.BOTH, expand=True)
            
            explanation = ""
            if strategy.method == "full_prompt":
                explanation = "FULL PROMPT: Content is small enough to include entirely in the prompt. This provides complete context for the AI."
            elif strategy.method == "rag":
                explanation = "RAG (Retrieval Augmented Generation): Content is large and complex. The system will:\n1. Break content into chunks\n2. Store in vector database\n3. Retrieve only relevant chunks for each query\n4. Provide focused context to AI"
            elif strategy.method == "hybrid":
                explanation = "HYBRID: Uses overview + targeted retrieval. Provides document structure with detailed context for specific queries."
            
            explain_text.insert("1.0", explanation)
            explain_text.config(state=tk.DISABLED)
            
            # Close button
            ttk.Button(main_frame, text="Close", command=dialog.destroy).pack(side=tk.RIGHT)
            
            self.log_message(f"Processing strategy: {strategy.method}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to analyze strategy: {str(e)}")
            self.log_message(f"Strategy analysis error: {e}")

    def create_prompt_history_tab(self):
        """Create prompt history tab with navigation sidebar"""
        history_frame = ttk.Frame(self.notebook)
        self.notebook.add(history_frame, text="Prompt History")
        history_frame.columnconfigure(0, weight=3)  # Main content area gets more space
        history_frame.columnconfigure(1, weight=1)  # Shortcut sidebar
        history_frame.rowconfigure(0, weight=1)
        
        # Create left side - Main history content
        main_frame = ttk.Frame(history_frame)
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 2), pady=5)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(0, weight=1)
        
        # Create text widget with scrollbars for main content
        self.prompt_history_text = scrolledtext.ScrolledText(main_frame, 
                                                        bg="#1e1e1e", fg="#ffffff",
                                                        font=("Consolas", 10), wrap=tk.WORD)
        self.prompt_history_text.pack(fill=tk.BOTH, expand=True)
        
        # Create right side - Shortcuts list
        shortcut_frame = ttk.LabelFrame(history_frame, text="Navigation Shortcuts")
        shortcut_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(2, 5), pady=5)
        shortcut_frame.columnconfigure(0, weight=1)
        shortcut_frame.rowconfigure(0, weight=1)
        
        # Shortcuts list with scrollbar
        shortcuts_container = ttk.Frame(shortcut_frame)
        shortcuts_container.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        shortcut_scrollbar = ttk.Scrollbar(shortcuts_container)
        shortcut_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.shortcut_list = tk.Listbox(shortcuts_container, 
                                    bg="#1e1e1e", fg="#ffffff", 
                                    font=("Consolas", 9),
                                    yscrollcommand=shortcut_scrollbar.set)
        self.shortcut_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        shortcut_scrollbar.config(command=self.shortcut_list.yview)
        
        # Bind selection event to jump to position
        self.shortcut_list.bind('<<ListboxSelect>>', self.jump_to_prompt_position)
        
        # Store prompt positions (for navigation)
        self.prompt_positions = []
        
        # Create tags for formatting
        self.prompt_history_text.tag_configure("sent", foreground="#90EE90", justify="right")
        self.prompt_history_text.tag_configure("received", foreground="#87CEFA", justify="left")
        self.prompt_history_text.tag_configure("timestamp", foreground="#888888", justify="center")
        self.prompt_history_text.tag_configure("header", font=("Consolas", 10, "bold"))
        
        # Add clear button at bottom
        btn_frame = ttk.Frame(shortcut_frame)
        btn_frame.pack(fill=tk.X, side=tk.BOTTOM, padx=5, pady=5)
        
        ttk.Button(btn_frame, text="Clear History", 
                command=self.clear_prompt_history).pack(fill=tk.X)
        ttk.Button(btn_frame, text="Export History", 
                command=self.export_prompt_history).pack(fill=tk.X, pady=(5, 0))

    def jump_to_prompt_position(self, event):
        """Jump to the selected prompt position in history"""
        if not hasattr(self, 'shortcut_list') or not self.shortcut_list.curselection():
            return
            
        index = self.shortcut_list.curselection()[0]
        if index < len(self.prompt_positions):
            position = self.prompt_positions[index]
            
            # Scroll to position
            self.prompt_history_text.see(position)
            
            # Highlight the prompt temporarily
            self.prompt_history_text.tag_remove("highlight", "1.0", tk.END)
            
            # Calculate the end position (approximate)
            line_num = int(position.split('.')[0])
            self.prompt_history_text.tag_add("highlight", position, f"{line_num+2}.0")
            
            # Configure highlight
            self.prompt_history_text.tag_configure("highlight", background="#404040")
            
            # Schedule removal of highlight
            self.root.after(2000, lambda: self.prompt_history_text.tag_remove("highlight", "1.0", tk.END))

    def clear_prompt_history(self):
        """Clear the prompt history"""
        if hasattr(self, 'prompt_history_text'):
            confirm = messagebox.askyesno("Confirm", "Clear prompt history? This cannot be undone.")
            if confirm:
                self.prompt_history_text.delete("1.0", tk.END)
                self.shortcut_list.delete(0, tk.END)
                self.prompt_positions = []
                self.log_message("Prompt history cleared")

    def export_prompt_history(self):
        """Export prompt history to a file"""
        if not hasattr(self, 'prompt_history_text'):
            return
            
        history_text = self.prompt_history_text.get("1.0", tk.END)
        if not history_text.strip():
            messagebox.showinfo("Info", "Prompt history is empty.")
            return
            
        filename = filedialog.asksaveasfilename(
            title="Export Prompt History",
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
        )
        
        if filename:
            try:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(history_text)
                self.log_message(f"Prompt history exported to: {filename}")
                messagebox.showinfo("Success", f"Prompt history exported to: {filename}")
            except Exception as e:
                self.log_message(f"Error exporting history: {str(e)}")
                messagebox.showerror("Error", f"Failed to export history: {str(e)}")

    def log_prompt_history(self, prompt, response=None, is_sent=True):
        """Log prompts and responses to the prompt history tab"""
        if not hasattr(self, 'prompt_history_text') or not hasattr(self, 'shortcut_list'):
            return
        
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Get current position for shortcut navigation
        current_position = self.prompt_history_text.index(tk.END)
        
        # Add to shortcut list
        shortcut_text = f"[{timestamp}] " + ("SENT" if is_sent else "RECEIVED")
        if is_sent and len(prompt) > 30:
            # Add first part of prompt for context in shortcut
            shortcut_text += f": {prompt[:30]}..."
        
        self.shortcut_list.insert(tk.END, shortcut_text)
        self.prompt_positions.append(current_position)
        
        # Add timestamp
        self.prompt_history_text.insert(tk.END, f"\n[{timestamp}]\n", "timestamp")
        
        if is_sent:
            # Add sent prompt header
            self.prompt_history_text.insert(tk.END, "PROMPT SENT:\n", "header")
            
            # Add actual prompt with right justification and green color
            self.prompt_history_text.insert(tk.END, f"{prompt}\n\n", "sent")
        else:
            # Add received response header
            self.prompt_history_text.insert(tk.END, "RESPONSE RECEIVED:\n", "header")
            
            # Add actual response with left justification and blue color
            self.prompt_history_text.insert(tk.END, f"{response}\n\n", "received")
        
        # Scroll to the end
        self.prompt_history_text.see(tk.END)
        
        # Keep shortcut list scrolled to bottom as well
        self.shortcut_list.see(tk.END)

    def reset_config_to_defaults(self):
        """Reset configuration to default values"""
        result = messagebox.askyesno(
            "Reset Configuration",
            "Are you sure you want to reset all settings to defaults?\n\n"
            "This will reset:\n"
            "- AI connection settings\n"
            "- Model settings\n"
            "- Formatting options\n"
            "- Auto features\n\n"
            "This action cannot be undone."
        )
        
        if result:
            # Reset OpenWebUI settings
            self.openwebui_base_url = "http://172.16.27.122:3000"
            self.openwebui_api_key = ""
            self.selected_model.set("")
            self.available_models = []
            self.selected_knowledge_collections = []
            self.available_knowledge_collections = []
            self.temperature.set(0.1)
            self.max_tokens.set(8000)
            
            # Reset master prompt
            self.set_default_prompt()
            
            # Reset formatting
            self.format_config['highlight_enabled'].set(True)
            self.format_config['highlight_color'].set('YELLOW')
            self.format_config['bold_enabled'].set(False)
            self.format_config['italic_enabled'].set(False)
            self.format_config['underline_enabled'].set(False)
            self.format_config['font_color'].set('000000')
            self.format_config['font_size'].set(11)
            
            # Reset auto features
            self.auto_config['auto_backup'].set(True)
            self.auto_config['backup_interval'].set(5)
            self.auto_config['auto_save'].set(False)
            self.auto_config['auto_reload'].set(True)
            self.auto_config['ask_backup'].set(True)
            
            # Save reset configuration
            self.save_settings()
            
            # Update UI
            self.update_config_status()
            
            self.log_message("Configuration reset to defaults")
            messagebox.showinfo("Reset Complete", "Configuration has been reset to defaults")

    def open_config_file_manager(self):
        """Open configuration file manager dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Configuration Files")
        dialog.geometry("500x300")
        dialog.configure(bg="#2b2b2b")
        dialog.grab_set()
        
        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Configuration File Management", 
                font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 15))
        
        # Current file
        script_dir = os.path.dirname(os.path.abspath(__file__))
        current_config = os.path.join(script_dir, "openwebui_config.json")
        
        if os.path.exists(current_config):
            config_size = os.path.getsize(current_config)
            config_date = datetime.fromtimestamp(os.path.getmtime(current_config)).strftime("%Y-%m-%d %H:%M")
            config_info = f"Current: {os.path.basename(current_config)} ({config_size/1024:.1f} KB, {config_date})"
        else:
            config_info = "No configuration file found"
        
        ttk.Label(main_frame, text=config_info).pack(anchor=tk.W, pady=(0, 15))
        
        # Action buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(0, 15))
        
        ttk.Button(btn_frame, text="Load Configuration", 
                command=self.browse_config_file, width=20).pack(pady=5)
        
        ttk.Button(btn_frame, text="Save Configuration As", 
                command=self.save_config_as, width=20).pack(pady=5)
        
        ttk.Button(btn_frame, text="Reset to Defaults", 
                command=self.reset_config_to_defaults, width=20).pack(pady=5)
        
        # Close button
        ttk.Button(main_frame, text="Close", 
                command=dialog.destroy).pack(side=tk.RIGHT, pady=(15, 0))

    # Add new function for formatting dialog
    def open_formatting_dialog(self):
        """Open formatting configuration dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Formatting Options")
        dialog.geometry("500x400")
        dialog.configure(bg="#2b2b2b")
        dialog.grab_set()
        
        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Content Formatting Options", 
                font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 10))
        
        # Formatting options
        format_frame = ttk.LabelFrame(main_frame, text="Text Formatting", padding="10")
        format_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Checkbutton(format_frame, text="Highlight Changes", 
                    variable=self.format_config['highlight_enabled']).pack(anchor=tk.W)
        
        color_frame = ttk.Frame(format_frame)
        color_frame.pack(fill=tk.X, pady=(5, 10))
        ttk.Label(color_frame, text="Highlight Color:").pack(side=tk.LEFT)
        color_options = ['YELLOW', 'GREEN', 'CYAN', 'PINK', 'BRIGHT_GREEN']
        color_combo = ttk.Combobox(color_frame, textvariable=self.format_config['highlight_color'],
                                values=color_options, state="readonly", width=15)
        color_combo.pack(side=tk.LEFT, padx=(5, 0))
        
        ttk.Checkbutton(format_frame, text="Bold", 
                    variable=self.format_config['bold_enabled']).pack(anchor=tk.W)
        ttk.Checkbutton(format_frame, text="Italic", 
                    variable=self.format_config['italic_enabled']).pack(anchor=tk.W)
        ttk.Checkbutton(format_frame, text="Underline", 
                    variable=self.format_config['underline_enabled']).pack(anchor=tk.W)
        
        font_frame = ttk.Frame(format_frame)
        font_frame.pack(fill=tk.X, pady=(5, 0))
        ttk.Label(font_frame, text="Font Size:").pack(side=tk.LEFT)
        ttk.Spinbox(font_frame, from_=8, to=16, textvariable=self.format_config['font_size'],
                width=10).pack(side=tk.LEFT, padx=(5, 0))
        
        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        def save_and_close():
            self.save_settings()
            self.log_message("Formatting options saved")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="Save", command=save_and_close).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT)

    # Add new function for auto features dialog
    def open_auto_features_dialog(self):
        """Open auto features configuration dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Automatic Features")
        dialog.geometry("500x350")
        dialog.configure(bg="#2b2b2b")
        dialog.grab_set()
        
        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Automatic Features Configuration", 
                font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 10))
        
        # Auto features
        auto_frame = ttk.LabelFrame(main_frame, text="Auto Features", padding="10")
        auto_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Checkbutton(auto_frame, text="Auto Backup", 
                    variable=self.auto_config['auto_backup']).pack(anchor=tk.W)
        
        backup_frame = ttk.Frame(auto_frame)
        backup_frame.pack(fill=tk.X, pady=(5, 10))
        ttk.Label(backup_frame, text="Backup Interval (minutes):").pack(side=tk.LEFT)
        ttk.Spinbox(backup_frame, from_=1, to=60, textvariable=self.auto_config['backup_interval'],
                width=10).pack(side=tk.LEFT, padx=(5, 0))
        
        ttk.Checkbutton(auto_frame, text="Auto Save After Commit", 
                    variable=self.auto_config['auto_save']).pack(anchor=tk.W)
        ttk.Checkbutton(auto_frame, text="Auto Reload Sections After Commit", 
                    variable=self.auto_config['auto_reload']).pack(anchor=tk.W)
        ttk.Checkbutton(auto_frame, text="Ask Before Creating Backup", 
                    variable=self.auto_config['ask_backup']).pack(anchor=tk.W)
        
        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        def save_and_close():
            self.save_settings()
            self.log_message("Auto features configuration saved")
            
            # Restart auto-backup if settings changed
            if self.auto_config['auto_backup'].get():
                self.schedule_auto_backup()
            
            dialog.destroy()
        
        ttk.Button(btn_frame, text="Save", command=save_and_close).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT)


    # Add this new function for tab change event
    def on_tab_changed(self, event):
        """Handle notebook tab changes"""
        selected_tab = event.widget.select()
        tab_text = event.widget.tab(selected_tab, "text")
        
        # If switching to Prompt Editor tab, refresh model list
        if tab_text == "Prompt Editor":
            self.update_prompt_model_selector()

    def create_preview_tab(self):
        """Create content preview tab"""
        preview_frame = ttk.Frame(self.notebook)
        self.notebook.add(preview_frame, text="Content Preview")
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.columnconfigure(1, weight=1)
        preview_frame.rowconfigure(1, weight=1)
        
        # Existing content
        ttk.Label(preview_frame, text="Existing Content", font=("Arial", 10, "bold")).grid(
            row=0, column=0, sticky=tk.W, padx=5, pady=5)
        
        self.existing_text = scrolledtext.ScrolledText(preview_frame, height=20, 
                                                      bg="#1e1e1e", fg="#ffffff",
                                                      font=("Consolas", 10), wrap=tk.WORD)
        self.existing_text.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 2))
        self.existing_text.config(state='disabled')
        
        # Generated content
        ttk.Label(preview_frame, text="Generated Content (Editable)", 
                 font=("Arial", 10, "bold")).grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        
        self.generated_text = scrolledtext.ScrolledText(preview_frame, height=20,
                                                       bg="#1e1e1e", fg="#90EE90",
                                                       font=("Consolas", 10), wrap=tk.WORD)
        self.generated_text.grid(row=1, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(2, 5))
        
        # Setup text tags for markdown rendering
        self.setup_text_tags()
        
        # Buttons
        commit_frame = ttk.Frame(preview_frame)
        commit_frame.grid(row=2, column=0, columnspan=2, pady=10)
        
        self.commit_btn = ttk.Button(commit_frame, text="Commit to Document", 
                                    command=self.commit_content, state='disabled')
        self.commit_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(commit_frame, text="Clear Preview", 
                  command=self.clear_preview).pack(side=tk.LEFT, padx=5)
    
    def create_prompt_tab(self):
        """Create prompt editor tab"""
        prompt_frame = ttk.Frame(self.notebook)
        self.notebook.add(prompt_frame, text="Prompt Editor")
        prompt_frame.columnconfigure(0, weight=1)
        prompt_frame.rowconfigure(1, weight=1)
        
        ttk.Label(prompt_frame, text="Last Sent Prompt (Edit and Regenerate)", 
                 font=("Arial", 10, "bold")).pack(anchor=tk.W, padx=5, pady=5)
        
        self.prompt_text = scrolledtext.ScrolledText(prompt_frame, height=20,
                                                     bg="#1e1e1e", fg="#ffffff",
                                                     font=("Consolas", 10), wrap=tk.WORD)
        self.prompt_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        prompt_btn_frame = ttk.Frame(prompt_frame)
        prompt_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(prompt_btn_frame, text="Model:").pack(side=tk.LEFT, padx=(0, 5))
        self.prompt_model_var = tk.StringVar()
        self.prompt_model_combo = ttk.Combobox(prompt_btn_frame, textvariable=self.prompt_model_var,
                                               state="readonly", width=30)
        self.prompt_model_combo.pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(prompt_btn_frame, text="Regenerate", 
                  command=self.regenerate_with_prompt).pack(side=tk.LEFT, padx=5)
        ttk.Button(prompt_btn_frame, text="Compare 3 Models", 
                  command=self.compare_three_models).pack(side=tk.LEFT, padx=5)
        ttk.Button(prompt_btn_frame, text="Update Master Prompt", 
                  command=self.offer_prompt_update).pack(side=tk.LEFT, padx=5)
    
    def create_console_tab(self):
        """Create console log tab"""
        console_frame = ttk.Frame(self.notebook)
        self.notebook.add(console_frame, text="Console Log")
        console_frame.columnconfigure(0, weight=1)
        console_frame.rowconfigure(0, weight=1)
        
        self.console = scrolledtext.ScrolledText(console_frame, height=20,
                                                bg="#1e1e1e", fg="#ffffff",
                                                font=("Consolas", 9))
        self.console.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5, pady=5)
    
    def setup_text_tags(self):
        """Setup text widget tags for markdown formatting"""
        self.generated_text.tag_configure("bold", font=("Consolas", 10, "bold"))
        self.generated_text.tag_configure("italic", font=("Consolas", 10, "italic"))
        self.generated_text.tag_configure("bold_italic", font=("Consolas", 10, "bold italic"))
        self.generated_text.tag_configure("code", background="#333333", foreground="#00ff00")
        self.generated_text.tag_configure("heading", font=("Consolas", 12, "bold"), foreground="#00aaff")
        
    def log_message(self, message):
        """Log message to console"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.console.insert(tk.END, f"[{timestamp}] {message}\n")
        self.console.see(tk.END)
        self.root.update()
        
    def browse_document(self):
        """Browse for Word document"""
        filename = filedialog.askopenfilename(
            title="Select Word Document",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            self.load_document(filename)
            
    def load_document(self, filepath):
        """Load Word document and parse sections"""
        try:
            self.log_message(f"Loading document: {filepath}")
            self.document = Document(filepath)
            self.document_path = filepath
            self.doc_label_var.set(os.path.basename(filepath))
            
            # Enable reload button since document is now loaded
            self.reload_btn.config(state=tk.NORMAL)
            
            # Load section tracking
            self.load_section_tracking()
            
            # Parse document structure
            self.parse_document_structure()
            
            # Populate tree
            self.populate_tree()
            
            # Save as last document
            self.save_settings()
            
            self.log_message(f"Document loaded: {len(self.sections)} sections found")
            
        except PermissionError:
            self.log_message("Error: File is open in another application")
            messagebox.showerror("Error", "File is open. Please close it and try again.")
        except Exception as e:
            self.log_message(f"Error loading document: {str(e)}")
            messagebox.showerror("Error", f"Failed to load document: {str(e)}")
    
    def reload_document_wrapper(self):
        """
        Wrapper for reload_document that provides user feedback and button state management.
        
        This method ensures:
        - User confirmation before reload
        - Proper error handling with user notifications
        - Button state consistency (disabled until document reloaded)
        - Logging for audit trail
        
        Security Consideration: CUI/Confidential data is refreshed from disk,
        ensuring no stale data persists in memory.
        """
        if not self.document_path:
            messagebox.showwarning("No Document", "Please load a document first")
            return
        
        # Confirm reload to prevent accidental loss of unsaved changes
        result = messagebox.askyesno(
            "Reload Document",
            "Reload the document from disk?\n\n"
            "WARNING: Any unsaved changes will be discarded.\n"
            "Selected section will be preserved if it still exists.",
            icon=messagebox.WARNING
        )
        
        if not result:
            self.log_message("Document reload cancelled by user")
            return
        
        try:
            # Disable button during reload to prevent multiple simultaneous reloads
            self.reload_btn.config(state=tk.DISABLED)
            self.reload_document()
            
            # Re-enable button after successful reload
            self.reload_btn.config(state=tk.NORMAL)
            messagebox.showinfo("Success", "Document reloaded successfully")
            self.log_message(f"Document reloaded from: {self.document_path}")
            
        except Exception as e:
            # Ensure button is re-enabled even on error
            self.reload_btn.config(state=tk.NORMAL)
            self.log_message(f"Error during reload: {str(e)}")
            messagebox.showerror("Reload Failed", f"Failed to reload document: {str(e)}")
    
    def parse_document_structure(self):
        """Parse document into hierarchical structure"""
        self.sections = []
        section_stack = []
        current_section = None
        
        for para in self.document.paragraphs:
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                    if level <= 4:
                        section = DocumentSection(level, para.text.strip(), para)
                        
                        while section_stack and section_stack[-1].level >= level:
                            section_stack.pop()
                        
                        if section_stack:
                            section_stack[-1].add_child(section)
                        else:
                            self.sections.append(section)
                        
                        section_stack.append(section)
                        current_section = section
                        
                except ValueError:
                    pass
            else:
                if current_section and para.text.strip():
                    if not para.text.strip().startswith('-'):
                        current_section.content_paragraphs.append(para)
                        
    def populate_tree(self):
        """Populate treeview with document sections"""
        self.tree.delete(*self.tree.get_children())
        
        def add_to_tree(section, parent=''):
            is_edited = self.is_section_edited(section)
            
            display_text = section.text
            if is_edited:
                display_text += " ✓"
            
            item_id = self.tree.insert(parent, 'end', text=display_text, 
                                      values=(id(section),))
            
            for child in section.children:
                add_to_tree(child, item_id)
        
        for section in self.sections:
            add_to_tree(section)
            
    def on_section_select(self, event):
        """Handle section selection"""
        selection = self.tree.selection()
        if selection:
            item = selection[0]
            section_id = self.tree.item(item, 'values')[0]
            
            self.selected_section = self.find_section_by_id(int(section_id))
            
            if self.selected_section:
                self.generate_btn.config(state='normal')
                # Enable review button if section has content
                if self.selected_section.has_content():
                    self.review_btn.config(state='normal')
                    if hasattr(self, 'tense_btn'):
                        self.tense_btn.config(state='normal')
                else:
                    self.review_btn.config(state='disabled')
                    if hasattr(self, 'tense_btn'):
                        self.tense_btn.config(state='disabled')

                self.show_existing_content()
                self.log_message(f"Selected: {self.selected_section.get_full_path()}")
                
    def find_section_by_id(self, section_id):
        """Find section by its ID"""
        def search(sections):
            for section in sections:
                if id(section) == section_id:
                    return section
                result = search(section.children)
                if result:
                    return result
            return None
        
        return search(self.sections)
        
    def show_existing_content(self):
        """Show existing content in preview"""
        if self.selected_section:
            self.existing_text.config(state='normal')
            self.existing_text.delete('1.0', tk.END)
            
            existing = self.selected_section.get_existing_content()
            if existing:
                self.existing_text.insert('1.0', existing)
            else:
                self.existing_text.insert('1.0', "[No existing content]")
            
            self.existing_text.config(state='disabled')
            
    # Replace the open_config_dialog to only show AI settings

    def save_master_prompt(self, filename=None):
        """Save master prompt to a file"""
        if not filename:
            filename = filedialog.asksaveasfilename(
                title="Save Master Prompt",
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
            )
        
        if not filename:
            return False
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(self.master_prompt.get())
            
            self.log_message(f"Master prompt saved to: {filename}")
            return True
        except Exception as e:
            self.log_message(f"Error saving master prompt: {str(e)}")
            messagebox.showerror("Error", f"Failed to save master prompt: {str(e)}")
            return False

    def load_master_prompt(self):
        """Load master prompt from a file"""
        filename = filedialog.askopenfilename(
            title="Load Master Prompt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
        )
        
        if not filename:
            return False
        
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                new_prompt = f.read()
            
            if not new_prompt:
                messagebox.showwarning("Warning", "The selected file is empty.")
                return False
            
            result = messagebox.askyesno(
                "Load Master Prompt",
                "Are you sure you want to replace the current master prompt?\n\n"
                "This action cannot be undone."
            )
            
            if result:
                self.master_prompt.set(new_prompt)
                self.log_message(f"Master prompt loaded from: {filename}")
                return True
            
            return False
        except Exception as e:
            self.log_message(f"Error loading master prompt: {str(e)}")
            messagebox.showerror("Error", f"Failed to load master prompt: {str(e)}")
            return False

    def refresh_all_ai_resources(self, url, key, model_combo, knowledge_list, status_label):
        """Refresh all AI resources (models and collections)"""
        if not url or not key:
            status_label.config(text="⚠ Please enter URL and API key")
            return
        
        status_label.config(text="🔄 Refreshing all resources...")
        
        # Use threading to avoid freezing the UI
        def refresh_thread():
            try:
                # Test connection first
                success = self.test_connection(url, key, status_label)
                
                if success:
                    # Refresh models
                    self.root.after(0, lambda: status_label.config(text="🔄 Loading models..."))
                    headers = {'Authorization': f'Bearer {key}'}
                    response = requests.get(f"{url}/api/models", headers=headers, timeout=10)
                    
                    if response.status_code == 200:
                        data = response.json()
                        self.available_models = []
                        
                        if isinstance(data, dict) and 'data' in data:
                            models_list = data['data']
                        elif isinstance(data, list):
                            models_list = data
                        else:
                            models_list = [data] if data else []
                        
                        for model in models_list:
                            if isinstance(model, dict):
                                model_id = model.get('id', model.get('name', 'unknown'))
                                self.available_models.append(model_id)
                            else:
                                self.available_models.append(str(model))
                        
                        self.root.after(0, lambda: model_combo.config(values=self.available_models))
                        self.log_message(f"Refreshed models: {len(self.available_models)} found")
                        
                        # Update selection if possible
                        if self.selected_model.get() in self.available_models:
                            self.root.after(0, lambda: model_combo.set(self.selected_model.get()))
                        elif self.available_models:
                            self.root.after(0, lambda: model_combo.set(self.available_models[0]))
                    
                    # Refresh collections
                    self.root.after(0, lambda: status_label.config(text="🔄 Loading collections..."))
                    response = requests.get(f"{url}/api/v1/knowledge/", headers=headers, timeout=10)
                    
                    if response.status_code == 200:
                        data = response.json()
                        self.available_knowledge_collections = []
                        
                        if isinstance(data, list):
                            for collection in data:
                                if isinstance(collection, dict):
                                    self.available_knowledge_collections.append({
                                        'id': collection.get('id', 'unknown'),
                                        'name': collection.get('name', 'Unknown'),
                                        'description': collection.get('description', '')
                                    })
                        
                        # Update list
                        self.root.after(0, lambda: knowledge_list.delete(0, tk.END))
                        for col in self.available_knowledge_collections:
                            display = f"{col['name']} ({col['id']})"
                            self.root.after(0, lambda d=display: knowledge_list.insert(tk.END, d))
                        
                        # Restore selections
                        for i, col in enumerate(self.available_knowledge_collections):
                            for sel in self.selected_knowledge_collections:
                                if col['id'] == sel.get('id', ''):
                                    self.root.after(0, lambda idx=i: knowledge_list.selection_set(idx))
                        
                        self.log_message(f"Refreshed collections: {len(self.available_knowledge_collections)} found")
                    
                    self.root.after(0, lambda: status_label.config(text="✓ All resources refreshed"))
                else:
                    self.root.after(0, lambda: status_label.config(text="✗ Connection failed"))
                    
            except Exception as e:
                self.log_message(f"Error refreshing resources: {str(e)}")
                self.root.after(0, lambda err=str(e): status_label.config(text=f"✗ Error: {err[:20]}..."))
        
        thread = threading.Thread(target=refresh_thread)
        thread.daemon = True
        thread.start()

    def auto_refresh_resources(self, url, key, model_combo, knowledge_list, status_label):
        """Automatically refresh models and knowledge collections on dialog open"""
        if not url or not key:
            status_label.config(text="⚠ Please enter URL and API key")
            return
            
        # Test connection first
        success = self.test_connection(url, key, status_label)
        
        if success:
            # Refresh models and collections in sequence
            self.refresh_models(url, key, model_combo)
            self.refresh_knowledge(url, key, knowledge_list)
            status_label.config(text="✓ Resources loaded")

    def open_config_dialog(self):
        """Open OpenWebUI AI configuration dialog"""
        config_window = tk.Toplevel(self.root)
        config_window.title("OpenWebUI AI Configuration")
        config_window.geometry("700x1000")
        config_window.configure(bg="#2b2b2b")
        config_window.grab_set()
        
        main_frame = ttk.Frame(config_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Configuration Files section
        config_file_frame = ttk.LabelFrame(main_frame, text="Configuration Files", padding="10")
        config_file_frame.pack(fill=tk.X, pady=(0, 10))
        
        config_btn_frame = ttk.Frame(config_file_frame)
        config_btn_frame.pack(fill=tk.X)
        
        ttk.Button(config_btn_frame, text="Load Config", 
                command=self.browse_config_file).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(config_btn_frame, text="Save Config As", 
                command=self.save_config_as).pack(side=tk.LEFT)
        
        # Connection settings
        conn_frame = ttk.LabelFrame(main_frame, text="Connection", padding="10")
        conn_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(conn_frame, text="Base URL:").pack(anchor=tk.W)
        url_var = tk.StringVar(value=self.openwebui_base_url)
        ttk.Entry(conn_frame, textvariable=url_var, width=50).pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(conn_frame, text="API Key:").pack(anchor=tk.W)
        key_var = tk.StringVar(value=self.openwebui_api_key)
        ttk.Entry(conn_frame, textvariable=key_var, width=50, show="*").pack(fill=tk.X, pady=(0, 10))
        
        test_frame = ttk.Frame(conn_frame)
        test_frame.pack(fill=tk.X)
        
        status_label = ttk.Label(test_frame, text="")
        
        # Add Refresh All button
        ttk.Button(test_frame, text="Refresh All", 
                command=lambda: self.refresh_all_ai_resources(url_var.get(), key_var.get(), model_combo, knowledge_list, status_label)).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(test_frame, text="Test Connection", 
                command=lambda: self.test_connection(url_var.get(), key_var.get(), status_label)).pack(side=tk.LEFT)
        
        status_label.pack(side=tk.LEFT, padx=(10, 0))
        
        # Model selection
        model_frame = ttk.LabelFrame(main_frame, text="Model", padding="10")
        model_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(model_frame, text="Refresh Models", 
                command=lambda: self.refresh_models(url_var.get(), key_var.get(), model_combo)).pack(anchor=tk.W, pady=(0, 5))
        
        ttk.Label(model_frame, text="Select Default Model:").pack(anchor=tk.W)
        model_combo = ttk.Combobox(model_frame, textvariable=self.selected_model, state="readonly")
        model_combo['values'] = self.available_models
        model_combo.pack(fill=tk.X)
        
        # Knowledge collections
        knowledge_frame = ttk.LabelFrame(main_frame, text="RAG Knowledge Collections", padding="10")
        knowledge_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        ttk.Button(knowledge_frame, text="Refresh Collections", 
                command=lambda: self.refresh_knowledge(url_var.get(), key_var.get(), knowledge_list)).pack(anchor=tk.W, pady=(0, 5))
        
        ttk.Label(knowledge_frame, text="Select Collections (Ctrl for multiple):").pack(anchor=tk.W)
        
        knowledge_scroll = ttk.Scrollbar(knowledge_frame)
        knowledge_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        knowledge_list = tk.Listbox(knowledge_frame, selectmode=tk.MULTIPLE, 
                                    yscrollcommand=knowledge_scroll.set,
                                    bg="#1e1e1e", fg="#ffffff", height=6)
        knowledge_list.pack(fill=tk.BOTH, expand=True)
        knowledge_scroll.config(command=knowledge_list.yview)
        
        # Model parameters
        params_frame = ttk.LabelFrame(main_frame, text="Parameters", padding="10")
        params_frame.pack(fill=tk.X, pady=(0, 10))
        
        temp_label_var = tk.StringVar(value=f"Temperature: {self.temperature.get()}")
        ttk.Label(params_frame, textvariable=temp_label_var).pack(anchor=tk.W)
        temp_scale = ttk.Scale(params_frame, from_=0.0, to=1.0, variable=self.temperature,
                            command=lambda v: temp_label_var.set(f"Temperature: {float(v):.2f}"))
        temp_scale.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(params_frame, text="Max Tokens:").pack(anchor=tk.W)
        ttk.Spinbox(params_frame, from_=500, to=16000, textvariable=self.max_tokens, 
                increment=500).pack(fill=tk.X)
        
        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        def save_and_close():
            self.openwebui_base_url = url_var.get()
            self.openwebui_api_key = key_var.get()
            
            # Get selected knowledge collections
            self.selected_knowledge_collections = []
            for idx in knowledge_list.curselection():
                self.selected_knowledge_collections.append(
                    self.available_knowledge_collections[idx])
            
            self.save_settings()
            self.update_config_status()
            config_window.destroy()
        
        ttk.Button(btn_frame, text="Save", command=save_and_close).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(btn_frame, text="Cancel", command=config_window.destroy).pack(side=tk.RIGHT)
        
        # Auto-refresh on load
        # Use after to let the dialog render first
        config_window.after(100, lambda: self.auto_refresh_resources(url_var.get(), key_var.get(), model_combo, knowledge_list, status_label))

    def test_connection(self, url, api_key, status_label):
        """Test OpenWebUI connection"""
        try:
            status_label.config(text="⏳ Testing connection...")
            headers = {'Authorization': f'Bearer {api_key}'}
            response = requests.get(f"{url}/api/models", headers=headers, timeout=10)
            
            if response.status_code == 200:
                status_label.config(text="✓ Connection successful")
                return True
            else:
                status_label.config(text=f"✗ HTTP {response.status_code}")
        except Exception as e:
            status_label.config(text=f"✗ {str(e)}")
        return False

    def refresh_models(self, url, api_key, combo):
        """Refresh available models"""
        try:
            combo.set("Loading models...")
            headers = {'Authorization': f'Bearer {api_key}'}
            response = requests.get(f"{url}/api/models", headers=headers, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                self.available_models = []
                
                if isinstance(data, dict) and 'data' in data:
                    models_list = data['data']
                elif isinstance(data, list):
                    models_list = data
                else:
                    models_list = [data] if data else []
                
                for model in models_list:
                    if isinstance(model, dict):
                        model_id = model.get('id', model.get('name', 'unknown'))
                        self.available_models.append(model_id)
                    else:
                        self.available_models.append(str(model))
                
                combo['values'] = self.available_models
                
                # Update selection if possible
                if self.selected_model.get() in self.available_models:
                    combo.set(self.selected_model.get())
                elif self.available_models:
                    combo.set(self.available_models[0])
                else:
                    combo.set("")
                    
                self.log_message(f"Refreshed models: {len(self.available_models)} found")
                return True
            else:
                self.log_message(f"Error refreshing models: HTTP {response.status_code}")
                combo.set("Error loading models")
                return False
        except Exception as e:
            self.log_message(f"Error refreshing models: {str(e)}")
            combo.set("Error loading models")
            return False
            
    def refresh_knowledge(self, url, api_key, listbox):
        """Refresh knowledge collections"""
        try:
            headers = {'Authorization': f'Bearer {api_key}'}
            response = requests.get(f"{url}/api/v1/knowledge/", headers=headers, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                self.available_knowledge_collections = []
                
                if isinstance(data, list):
                    for collection in data:
                        if isinstance(collection, dict):
                            self.available_knowledge_collections.append({
                                'id': collection.get('id', 'unknown'),
                                'name': collection.get('name', 'Unknown'),
                                'description': collection.get('description', '')
                            })
                
                listbox.delete(0, tk.END)
                for col in self.available_knowledge_collections:
                    display = f"{col['name']} ({col['id']})"
                    listbox.insert(tk.END, display)
                
                # Restore selections
                for i, col in enumerate(self.available_knowledge_collections):
                    for sel in self.selected_knowledge_collections:
                        if col['id'] == sel.get('id', ''):
                            listbox.selection_set(i)
                
                self.log_message(f"Refreshed collections: {len(self.available_knowledge_collections)} found")
                return True
            else:
                self.log_message(f"Error refreshing collections: HTTP {response.status_code}")
                return False
        except Exception as e:
            self.log_message(f"Error refreshing collections: {str(e)}")
            return False        

    def update_config_status(self):
        """Update configuration status display"""
        if self.selected_model.get() and self.openwebui_api_key:
            status = f"✓ Model: {self.selected_model.get()}"
            if self.selected_knowledge_collections:
                status += f" | RAG: {len(self.selected_knowledge_collections)} collections"
            self.config_status_var.set(status)
        else:
            self.config_status_var.set("⚠ Not configured - click Configure")

    def load_prompt_library(self, listbox, preview):
        """Load saved prompts into the library list"""
        listbox.delete(0, tk.END)
        
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            library_dir = os.path.join(script_dir, "prompt_library")
            
            # Create directory if it doesn't exist
            if not os.path.exists(library_dir):
                os.makedirs(library_dir)
                self.log_message("Created prompt library directory")
                
                # Create a default prompt
                default_path = os.path.join(library_dir, "default_prompt.txt")
                self.save_master_prompt(default_path)
            
            # Add all prompt files to the list without extensions
            for filename in os.listdir(library_dir):
                if filename.endswith('.txt'):
                    # Display without .txt extension
                    display_name = os.path.splitext(filename)[0]
                    listbox.insert(tk.END, display_name)
                    
                    # Store the full filename in the listbox items dictionary
                    self.prompt_library_filenames = getattr(self, 'prompt_library_filenames', {})
                    self.prompt_library_filenames[display_name] = filename
            
            if listbox.size() > 0:
                listbox.selection_set(0)
                self.on_prompt_select(None, listbox, preview)
                
        except Exception as e:
            self.log_message(f"Error loading prompt library: {str(e)}")

    def on_prompt_select(self, event, listbox, preview):
        """Handle selection in prompt library"""
        try:
            if listbox.curselection():
                index = listbox.curselection()[0]
                display_name = listbox.get(index)
                
                # Get the actual filename from our dictionary
                self.prompt_library_filenames = getattr(self, 'prompt_library_filenames', {})
                filename = self.prompt_library_filenames.get(display_name, display_name + ".txt")
                    
                script_dir = os.path.dirname(os.path.abspath(__file__))
                library_dir = os.path.join(script_dir, "prompt_library")
                prompt_path = os.path.join(library_dir, filename)
                
                with open(prompt_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                preview.delete('1.0', tk.END)
                preview.insert('1.0', content)
                
        except Exception as e:
            self.log_message(f"Error loading prompt: {str(e)}")

    def import_prompt_to_library(self, listbox, preview):
        """Import a prompt file to the library"""
        try:
            filename = filedialog.askopenfilename(
                title="Import Prompt to Library",
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
            )
            
            if not filename:
                return
                
            script_dir = os.path.dirname(os.path.abspath(__file__))
            library_dir = os.path.join(script_dir, "prompt_library")
            
            if not os.path.exists(library_dir):
                os.makedirs(library_dir)
            
            # Get basename for new file
            basename = os.path.basename(filename)
            name_input = tk.simpledialog.askstring(
                "Prompt Name", 
                "Enter a name for this prompt:",
                initialvalue=os.path.splitext(basename)[0]
            )
            
            if not name_input:
                return
                
            # Ensure .txt extension for file storage but display without extension
            if not name_input.endswith('.txt'):
                file_name = name_input + '.txt'
                display_name = name_input
            else:
                file_name = name_input
                display_name = os.path.splitext(name_input)[0]  # Remove extension for display
                
            # Copy file to library
            new_path = os.path.join(library_dir, file_name)
            shutil.copy2(filename, new_path)
            
            # Update our filename dictionary
            self.prompt_library_filenames = getattr(self, 'prompt_library_filenames', {})
            self.prompt_library_filenames[display_name] = file_name
            
            # Refresh list
            self.load_prompt_library(listbox, preview)
            
            # Select the new item
            for i in range(listbox.size()):
                if listbox.get(i) == display_name:
                    listbox.selection_clear(0, tk.END)
                    listbox.selection_set(i)
                    self.on_prompt_select(None, listbox, preview)
                    break
                    
            self.log_message(f"Imported prompt to library: {display_name}")
            
        except Exception as e:
            self.log_message(f"Error importing prompt: {str(e)}")

    def remove_prompt_from_library(self, listbox, preview):
        """Remove a prompt from the library"""
        try:
            if not listbox.curselection():
                messagebox.showinfo("Info", "Please select a prompt to remove.")
                return
                
            index = listbox.curselection()[0]
            display_name = listbox.get(index)
            
            # Get the actual filename from our dictionary
            self.prompt_library_filenames = getattr(self, 'prompt_library_filenames', {})
            filename = self.prompt_library_filenames.get(display_name, display_name + ".txt")
            
            result = messagebox.askyesno(
                "Remove Prompt",
                f"Are you sure you want to remove '{display_name}' from the library?\n\n"
                "This action cannot be undone."
            )
            
            if not result:
                return
                
            script_dir = os.path.dirname(os.path.abspath(__file__))
            library_dir = os.path.join(script_dir, "prompt_library")
            prompt_path = os.path.join(library_dir, filename)
            
            os.remove(prompt_path)
            
            # Remove from our filename dictionary
            if display_name in self.prompt_library_filenames:
                del self.prompt_library_filenames[display_name]
            
            # Refresh list
            self.load_prompt_library(listbox, preview)
            preview.delete('1.0', tk.END)
            
            self.log_message(f"Removed prompt from library: {display_name}")
            
        except Exception as e:
            self.log_message(f"Error removing prompt: {str(e)}")           
 
    def use_selected_prompt(self, listbox, preview, window):
        """Use the selected prompt as master prompt"""
        try:
            if not listbox.curselection():
                messagebox.showinfo("Info", "Please select a prompt to use.")
                return
                
            content = preview.get('1.0', tk.END).strip()
            display_name = listbox.get(listbox.curselection()[0])
            
            if not content:
                messagebox.showwarning("Warning", "Selected prompt is empty.")
                return
                
            result = messagebox.askyesno(
                "Use Prompt",
                f"Are you sure you want to replace the current master prompt with '{display_name}'?\n\n"
                "This action cannot be undone."
            )
            
            if result:
                self.master_prompt.set(content)
                self.log_message(f"Master prompt updated from library: {display_name}")
                window.destroy()
                
        except Exception as e:
            self.log_message(f"Error using prompt: {str(e)}")

    def open_prompt_library(self):
        """Open prompt library manager"""
        library_window = tk.Toplevel(self.root)
        library_window.title("Prompt Library")
        library_window.geometry("900x600")
        library_window.configure(bg="#2b2b2b")
        library_window.grab_set()
        
        main_frame = ttk.Frame(library_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Prompt Library Manager", 
                font=("Arial", 14, "bold")).pack(anchor=tk.W, pady=(0, 15))
        
        # Create split view
        pane = ttk.PanedWindow(main_frame, orient=tk.HORIZONTAL)
        pane.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        # Left panel - list of saved prompts
        left_frame = ttk.Frame(pane)
        pane.add(left_frame, weight=1)
        
        ttk.Label(left_frame, text="Saved Prompts:").pack(anchor=tk.W, pady=(0, 5))
        
        # Prompt list with scrollbar
        list_frame = ttk.Frame(left_frame)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        list_scrollbar = ttk.Scrollbar(list_frame)
        list_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        prompt_list = tk.Listbox(list_frame, yscrollcommand=list_scrollbar.set,
                                bg="#1e1e1e", fg="#ffffff", font=("Consolas", 10),
                                width=30)  # Added width parameter to prevent collapsing
        prompt_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        list_scrollbar.config(command=prompt_list.yview)
        
        # Buttons for list operations
        list_btn_frame = ttk.Frame(left_frame)
        list_btn_frame.pack(fill=tk.X, pady=(5, 0))
        
        ttk.Button(list_btn_frame, text="Import", 
                command=lambda: self.import_prompt_to_library(prompt_list, prompt_preview)).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(list_btn_frame, text="Remove", 
                command=lambda: self.remove_prompt_from_library(prompt_list, prompt_preview)).pack(side=tk.LEFT)
        
        # Right panel - prompt preview
        right_frame = ttk.Frame(pane)
        pane.add(right_frame, weight=2)
        
        ttk.Label(right_frame, text="Prompt Preview:").pack(anchor=tk.W, pady=(0, 5))
        
        prompt_preview = scrolledtext.ScrolledText(right_frame, height=20,
                                                bg="#1e1e1e", fg="#ffffff",
                                                font=("Consolas", 10), wrap=tk.WORD)
        prompt_preview.pack(fill=tk.BOTH, expand=True, pady=(0, 5))

        prompt_preview.config(insertbackground="#ffffff")  # White cursor
        prompt_preview.config(selectbackground="#3a5f8f")  # Blue selection background
        prompt_preview.config(selectforeground="#ffffff")  # White selection text
        
        # Buttons for preview operations
        preview_btn_frame = ttk.Frame(right_frame)
        preview_btn_frame.pack(fill=tk.X)
        
        ttk.Button(preview_btn_frame, text="Use This Prompt", 
                command=lambda: self.use_selected_prompt(prompt_list, prompt_preview, library_window)).pack(side=tk.RIGHT)
        
        # Bottom buttons
        bottom_frame = ttk.Frame(main_frame)
        bottom_frame.pack(fill=tk.X, pady=(0, 5))
        
        ttk.Button(bottom_frame, text="Close", 
                command=library_window.destroy).pack(side=tk.RIGHT)
        
        # Load saved prompts from library
        self.load_prompt_library(prompt_list, prompt_preview)
        
        # Bind selection event
        prompt_list.bind('<<ListboxSelect>>', 
                        lambda e: self.on_prompt_select(e, prompt_list, prompt_preview))
        
        # Ensure the window is properly sized
        library_window.update_idletasks()
        pane.sashpos(0, 300)

    def edit_master_prompt(self):
        """Open master prompt editor"""
        editor_window = tk.Toplevel(self.root)
        editor_window.title("Edit Master Prompt")
        editor_window.geometry("800x600")
        editor_window.configure(bg="#2b2b2b")
        editor_window.grab_set()
        
        main_frame = ttk.Frame(editor_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Master Prompt Template:", 
                font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 5))
        
        info_label = ttk.Label(main_frame, 
                            text="Variables: {section_name}, {parent_context}, {operation_mode}",
                            foreground="#888888")
        info_label.pack(anchor=tk.W, pady=(0, 10))
        
        # Add Save/Load buttons at top
        action_frame = ttk.Frame(main_frame)
        action_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(action_frame, text="Load from File", 
                command=self.load_master_prompt).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(action_frame, text="Save to File", 
                command=self.save_master_prompt).pack(side=tk.LEFT)
        
        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        prompt_text = scrolledtext.ScrolledText(text_frame, height=20, width=80,
                                            bg="#1e1e1e", fg="#ffffff",
                                            font=("Consolas", 10), wrap=tk.WORD)
        prompt_text.pack(fill=tk.BOTH, expand=True)
        prompt_text.insert('1.0', self.master_prompt.get())
        
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        prompt_text.config(insertbackground="#ffffff")  # White cursor
        prompt_text.config(selectbackground="#3a5f8f")  # Blue selection background
        prompt_text.config(selectforeground="#ffffff")  # White selection text
        
        def save_prompt():
            self.master_prompt.set(prompt_text.get('1.0', tk.END).strip())
            self.log_message("Master prompt updated")
            editor_window.destroy()
        
        def reset_prompt():
            if messagebox.askyesno("Reset", "Reset to default prompt?"):
                self.set_default_prompt()
                prompt_text.delete('1.0', tk.END)
                prompt_text.insert('1.0', self.master_prompt.get())
        
        ttk.Button(btn_frame, text="Reset to Default", command=reset_prompt).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Cancel", command=editor_window.destroy).pack(side=tk.RIGHT)
        ttk.Button(btn_frame, text="Save", command=save_prompt).pack(side=tk.RIGHT, padx=(0, 5))            

    def generate_content(self):
        """Generate content for selected section"""
        if not self.selected_section:
            messagebox.showwarning("Warning", "No section selected")
            return
            
        if not self.selected_model.get():
            messagebox.showwarning("Warning", "No model configured")
            return
            
        self.notebook.select(2)  # Switch to console
        
        thread = threading.Thread(target=self.generate_content_thread)
        thread.daemon = True
        thread.start()
        
    def generate_content_thread(self):
        """Generate content in background thread"""
        try:
            self.generate_btn.config(state='disabled')
            
            section = self.selected_section
            mode = self.operation_mode.get()
            existing_content = section.get_existing_content()
            
            self.log_message(f"Generating content for: {section.get_full_path()}")
            self.log_message(f"Mode: {mode}")
            
            # Build parent context
            parent_context = ""
            if section.parent:
                parent_path = []
                current = section.parent
                while current:
                    parent_path.insert(0, current.text)
                    current = current.parent
                parent_context = " > ".join(parent_path)
            
            # Build prompt
            prompt = self.master_prompt.get()
            prompt = prompt.replace("{section_name}", section.text)
            prompt = prompt.replace("{parent_context}", parent_context if parent_context else "Root level")
            prompt = prompt.replace("{operation_mode}", mode.upper())
            
            # Add mode-specific instructions
            if mode == "replace":
                prompt += "\n\nYour task: Write comprehensive content for this section from scratch."
            elif mode == "rework":
                prompt += f"\n\nEXISTING CONTENT TO REWORK:\n{existing_content}\n\n"
                prompt += "Your task: Rewrite and enhance the existing content."
            elif mode == "append":
                prompt += f"\n\nEXISTING CONTENT:\n{existing_content}\n\n"
                prompt += "Your task: Add additional relevant content."
            
            # Add knowledge base instruction
            if self.selected_knowledge_collections:
                collection_names = [col['name'] for col in self.selected_knowledge_collections]
                prompt += f"\n\nIMPORTANT: Reference knowledge base: {', '.join(collection_names)}"

            # NEW: Add external RAG content if available
            external_content = self.get_external_rag_content()
            if external_content:
                prompt += "\n\n" + "="*60 + "\n"
                prompt += "ADDITIONAL REFERENCE CONTENT (External RAG):\n"
                prompt += "="*60 + "\n"
                for title, content, category, tags in external_content[:10]:  # Limit to top 10 for token management
                    prompt += f"\n[{title}]"
                    if category:
                        prompt += f" (Category: {category})"
                    prompt += f"\n{content}\n"
                    prompt += "-"*40 + "\n"
                self.log_message(f"✓ Included {len(external_content[:10])} external RAG items")

            # Store the prompt
            self.last_sent_prompt = prompt
            
            self.log_message("Sending request to OpenWebUI...")
            
            # Query OpenWebUI
            response = self.query_openwebui(prompt)
            
            if response and not response.startswith("Error:"):
                self.generated_content = response
                self.root.after(0, self.show_generated_content)
                self.log_message("Content generated successfully")
            else:
                self.log_message(f"Generation failed: {response}")
                self.root.after(0, self.handle_generation_error)
                
        except Exception as e:
            self.log_message(f"Error generating content: {str(e)}")
            self.root.after(0, self.handle_generation_error)
        finally:
            self.root.after(0, lambda: self.generate_btn.config(state='normal'))
            
    def query_openwebui(self, prompt):
        """Query OpenWebUI API"""
        try:
            if self.content_processor and self.document and self.selected_section:
                try:
                    full_content = self.selected_section.get_existing_content()
                    if len(full_content) > 100:
                        strategy_result = self.content_processor.determine_processing_strategy(
                            full_content, [self.selected_section], prompt
                        )
                        
                        self.log_message(f"📊 Processing Strategy: {strategy_result.method} - {strategy_result.reason}")
                        
                        if strategy_result.method == "rag" and self.document_path:
                            context, chunks = self.content_processor.build_rag_context(
                                prompt, self.document_path, [self.selected_section]
                            )
                            if context and len(chunks) > 0:
                                prompt = f"RELEVANT CONTEXT FROM DOCUMENT:\n{context}\n\n===\n\nORIGINAL QUERY:\n{prompt}"
                                self.log_message(f"✓ Enhanced with RAG context ({len(chunks)} chunks)")
                except Exception as e:
                    self.log_message(f"Note: Could not apply intelligent processing: {e}")
            
            # Log the prompt to history first
            self.log_prompt_history(prompt, is_sent=True)
            
            headers = {
                'Authorization': f'Bearer {self.openwebui_api_key}',
                'Content-Type': 'application/json'
            }
            
            payload = {
                "model": self.selected_model.get(),
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "temperature": self.temperature.get(),
                "max_tokens": self.max_tokens.get()
            }
            
            if self.selected_knowledge_collections:
                payload["files"] = [
                    {"type": "collection", "id": col['id']} 
                    for col in self.selected_knowledge_collections
                ]
            
            response = requests.post(
                f"{self.openwebui_base_url}/api/chat/completions",
                headers=headers, json=payload, timeout=300
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and len(result['choices']) > 0:
                    response_content = result['choices'][0]['message']['content']
                    # Log the response to history
                    self.log_prompt_history(response_content, response_content, is_sent=False)
                    return response_content
                elif 'response' in result:
                    response_content = result['response']
                    # Log the response to history
                    self.log_prompt_history(response_content, response_content, is_sent=False)
                    return response_content
                else:
                    error_msg = "Error: No content in response"
                    self.log_prompt_history(error_msg, error_msg, is_sent=False)
                    return error_msg
            else:
                error_msg = f"Error: HTTP {response.status_code} - {response.text}"
                self.log_prompt_history(error_msg, error_msg, is_sent=False)
                return error_msg
                
        except Exception as e:
            error_msg = f"Error: {str(e)}"
            self.log_prompt_history(error_msg, error_msg, is_sent=False)
            return error_msg

    def query_openwebui_with_model(self, prompt, model):
        """Query OpenWebUI API with specific model"""
        try:
            # Log the prompt to history
            self.log_prompt_history(f"[Model: {model}]\n{prompt}", is_sent=True)
            
            headers = {
                'Authorization': f'Bearer {self.openwebui_api_key}',
                'Content-Type': 'application/json'
            }
            
            payload = {
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "temperature": self.temperature.get(),
                "max_tokens": self.max_tokens.get()
            }
            
            if self.selected_knowledge_collections:
                payload["files"] = [
                    {"type": "collection", "id": col['id']} 
                    for col in self.selected_knowledge_collections
                ]
            
            response = requests.post(
                f"{self.openwebui_base_url}/api/chat/completions",
                headers=headers, json=payload, timeout=300
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and len(result['choices']) > 0:
                    response_content = result['choices'][0]['message']['content']
                    # Log the response to history
                    response_with_model = f"[Model: {model}]\n{response_content}"
                    self.log_prompt_history(response_with_model, response_with_model, is_sent=False)
                    return response_content
                elif 'response' in result:
                    response_content = result['response']
                    # Log the response to history
                    response_with_model = f"[Model: {model}]\n{response_content}"
                    self.log_prompt_history(response_with_model, response_with_model, is_sent=False)
                    return response_content
                else:
                    error_msg = f"[Model: {model}]\nError: No content in response"
                    self.log_prompt_history(error_msg, error_msg, is_sent=False)
                    return "Error: No content in response"
            else:
                error_msg = f"[Model: {model}]\nError: HTTP {response.status_code} - {response.text}"
                self.log_prompt_history(error_msg, error_msg, is_sent=False)
                return f"Error: HTTP {response.status_code} - {response.text}"
                
        except Exception as e:
            error_msg = f"[Model: {model}]\nError: {str(e)}"
            self.log_prompt_history(error_msg, error_msg, is_sent=False)
            return f"Error: {str(e)}"
  
    def show_generated_content(self):
        """Show generated content in preview with markdown rendering"""
        self.notebook.select(0)  # Switch to preview
        
        # Render with markdown formatting
        self.render_markdown_preview(self.generated_content)
        
        # Update prompt tab
        self.prompt_text.delete('1.0', tk.END)
        self.prompt_text.insert('1.0', self.last_sent_prompt)
        
        # Update model selector in prompt tab
        self.update_prompt_model_selector()
        
        self.commit_btn.config(state='normal')
    
    def update_prompt_model_selector(self):
        """Update the model selector in prompt tab"""
        # Try to load models if empty
        if not self.available_models:
            self.ensure_models_loaded()
        
        if self.available_models:
            self.prompt_model_combo['values'] = self.available_models
            if not self.prompt_model_var.get():
                self.prompt_model_var.set(self.selected_model.get())
        else:
            self.prompt_model_combo['values'] = ['No models loaded - Configure AI Settings']
            self.log_message("No models available. Please configure AI settings and refresh models.")

    def debug_model_status(self):
        """Debug function to check model status"""
        print(f"Available models count: {len(self.available_models)}")
        print(f"Models list: {self.available_models}")
        print(f"API Key set: {bool(self.openwebui_api_key)}")
        print(f"Base URL: {self.openwebui_base_url}")
        
        self.log_message(f"Debug: {len(self.available_models)} models loaded")
        if self.available_models:
            self.log_message(f"Models: {', '.join(self.available_models[:3])}...")

    def get_document_outline(self):
        """Get complete document outline for context"""
        outline = []
        
        def build_outline(sections, level=0):
            for section in sections:
                indent = "  " * level
                outline.append(f"{indent}{section.text}")
                build_outline(section.children, level + 1)
        
        build_outline(self.sections)
        return "\n".join(outline)

    def auto_complete_document(self):
        """Automatically complete all sections in the document"""
        if not self.document or not self.sections:
            messagebox.showwarning("Warning", "No document loaded")
            return
        
        if not self.selected_model.get():
            messagebox.showwarning("Warning", "No model configured")
            return
        
        # Get all sections that need completion
        all_sections = self.get_all_sections_flat()
        
        # Filter based on user preference
        result = messagebox.askyesnocancel(
            "Auto Complete Document",
            f"Found {len(all_sections)} sections.\n\n"
            "Yes: Complete ALL sections (overwrites existing)\n"
            "No: Complete only EMPTY sections\n"
            "Cancel: Don't proceed"
        )
        
        if result is None:  # Cancel
            return
        elif result is True:  # All sections
            sections_to_complete = all_sections
        else:  # Only empty
            sections_to_complete = [s for s in all_sections if not s.has_content()]
        
        if not sections_to_complete:
            messagebox.showinfo("Info", "No sections to complete!")
            return
        
        # Confirm
        confirm = messagebox.askyesno(
            "Confirm Auto Complete",
            f"Ready to auto-complete {len(sections_to_complete)} sections.\n\n"
            f"This will:\n"
            f"• Process sections one by one\n"
            f"• Use current operation mode: {self.operation_mode.get().upper()}\n"
            f"• Take several minutes to complete\n"
            f"• Auto-save after each section (if enabled)\n\n"
            f"Continue?"
        )
        
        if not confirm:
            return
        
        # Create progress window
        self.create_auto_complete_window(sections_to_complete)

    def get_all_sections_flat(self):
        """Get flat list of all sections"""
        all_sections = []
        
        def collect_sections(sections):
            for section in sections:
                all_sections.append(section)
                collect_sections(section.children)
        
        collect_sections(self.sections)
        return all_sections

    def create_auto_complete_window(self, sections_to_complete):
        """Create progress window for auto-complete"""
        progress_window = tk.Toplevel(self.root)
        progress_window.title("Auto Complete Document")
        progress_window.geometry("800x600")
        progress_window.configure(bg="#2b2b2b")
        progress_window.protocol("WM_DELETE_WINDOW", lambda: None)  # Prevent closing during process
        
        main_frame = ttk.Frame(progress_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Auto Completing Document", 
                font=("Arial", 14, "bold")).pack(pady=(0, 10))
        
        # Progress info
        progress_frame = ttk.Frame(main_frame)
        progress_frame.pack(fill=tk.X, pady=(0, 10))
        
        progress_label = ttk.Label(progress_frame, text="Starting...", font=("Arial", 10))
        progress_label.pack(anchor=tk.W)
        
        # Progress bar
        progress_bar = ttk.Progressbar(progress_frame, mode='determinate', length=400)
        progress_bar.pack(fill=tk.X, pady=(5, 0))
        progress_bar['maximum'] = len(sections_to_complete)
        
        # Current section info
        current_frame = ttk.LabelFrame(main_frame, text="Current Section", padding="10")
        current_frame.pack(fill=tk.X, pady=(0, 10))
        
        current_section_label = ttk.Label(current_frame, text="", wraplength=700)
        current_section_label.pack(anchor=tk.W)
        
        status_label = ttk.Label(current_frame, text="", foreground="#00aaff")
        status_label.pack(anchor=tk.W, pady=(5, 0))
        
        # Log area
        log_frame = ttk.LabelFrame(main_frame, text="Progress Log", padding="5")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        log_text = scrolledtext.ScrolledText(log_frame, height=15, width=80,
                                            bg="#1e1e1e", fg="#ffffff",
                                            font=("Consolas", 9))
        log_text.pack(fill=tk.BOTH, expand=True)
        
        log_text.config(insertbackground="#ffffff")  # White cursor
        log_text.config(selectbackground="#3a5f8f")  # Blue selection background
        log_text.config(selectforeground="#ffffff")  # White selection text
        

        # Control buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X)
        
        pause_var = tk.BooleanVar(value=False)
        stop_var = tk.BooleanVar(value=False)
        
        pause_btn = ttk.Button(button_frame, text="Pause", 
                            command=lambda: pause_var.set(not pause_var.get()))
        pause_btn.pack(side=tk.LEFT, padx=5)
        
        stop_btn = ttk.Button(button_frame, text="Stop", 
                            command=lambda: stop_var.set(True))
        stop_btn.pack(side=tk.LEFT, padx=5)
        
        close_btn = ttk.Button(button_frame, text="Close", state='disabled',
                            command=progress_window.destroy)
        close_btn.pack(side=tk.RIGHT, padx=5)
        
        # Store references
        progress_window.log_text = log_text
        progress_window.progress_bar = progress_bar
        progress_window.progress_label = progress_label
        progress_window.current_section_label = current_section_label
        progress_window.status_label = status_label
        progress_window.pause_btn = pause_btn
        progress_window.stop_btn = stop_btn
        progress_window.close_btn = close_btn
        progress_window.pause_var = pause_var
        progress_window.stop_var = stop_var
        
        # Start processing in thread
        thread = threading.Thread(
            target=self.auto_complete_process,
            args=(sections_to_complete, progress_window)
        )
        thread.daemon = True
        thread.start()

    def auto_complete_process(self, sections, progress_window):
        """Process auto-complete with parent context and subsection awareness"""
        try:
            # Get document outline for structure context
            document_outline = self.get_document_outline()
            
            # Build content map as we go (for parent context)
            generated_content_map = {}  # section_hash -> content
            
            completed = 0
            failed = 0
            
            for idx, section in enumerate(sections):
                # Check if stopped
                if progress_window.stop_var.get():
                    self.update_progress_log(progress_window, "Stopped by user")
                    break
                
                # Check if paused
                while progress_window.pause_var.get():
                    if not hasattr(progress_window, 'pause_notified'):
                        self.update_progress_log(progress_window, "⏸ Paused - Click Pause again to resume")
                        self.root.after(0, lambda: progress_window.pause_btn.config(text="Resume"))
                        progress_window.pause_notified = True
                    self.root.update()
                    import time
                    time.sleep(0.5)
                    if progress_window.stop_var.get():
                        break
                
                # Reset pause notification
                if hasattr(progress_window, 'pause_notified'):
                    delattr(progress_window, 'pause_notified')
                    self.root.after(0, lambda: progress_window.pause_btn.config(text="Pause"))
                
                # Update progress
                self.root.after(0, lambda i=idx: progress_window.progress_bar.configure(value=i))
                self.root.after(0, lambda i=idx, t=len(sections): 
                            progress_window.progress_label.config(
                                text=f"Processing {i+1} of {t}..."))
                
                # Update current section
                section_path = section.get_full_path()
                self.root.after(0, lambda p=section_path: 
                            progress_window.current_section_label.config(text=p))
                self.root.after(0, lambda: 
                            progress_window.status_label.config(text="🔄 Generating content..."))
                
                self.update_progress_log(progress_window, f"\n{'='*60}")
                self.update_progress_log(progress_window, f"[{idx+1}/{len(sections)}] {section_path}")
                
                try:
                    # Build enhanced prompt with parent context and subsection awareness
                    mode = self.operation_mode.get()
                    existing_content = section.get_existing_content()
                    
                    # Get parent context PATH
                    parent_context_path = ""
                    if section.parent:
                        parent_path = []
                        current = section.parent
                        while current:
                            parent_path.insert(0, current.text)
                            current = current.parent
                        parent_context_path = " > ".join(parent_path)
                    
                    # *** NEW: Get parent's ACTUAL CONTENT ***
                    parent_actual_content = ""
                    if section.parent:
                        parent_hash = section.parent.get_section_hash()
                        # First check if we just generated it
                        if parent_hash in generated_content_map:
                            parent_actual_content = generated_content_map[parent_hash]
                        else:
                            # Check if it exists in document
                            parent_actual_content = section.parent.get_existing_content()
                    
                    # *** NEW: Get sibling sections for context ***
                    sibling_info = ""
                    if section.parent:
                        siblings = [s.text for s in section.parent.children if s != section]
                        if siblings:
                            sibling_info = "Sibling sections (for context): " + ", ".join(siblings)
                    
                    # Build enhanced prompt
                    prompt = self.master_prompt.get()
                    prompt = prompt.replace("{section_name}", section.text)
                    prompt = prompt.replace("{parent_context}", 
                                        parent_context_path if parent_context_path else "Root level")
                    prompt = prompt.replace("{operation_mode}", mode.upper())
                    
                    # *** NEW: Add document outline ***
                    prompt += f"\n\nDOCUMENT STRUCTURE (for overall context):\n{document_outline}\n"
                    
                    # *** NEW: Add parent content if this is a subsection ***
                    if section.parent and parent_actual_content:
                        prompt += f"\n\nPARENT SECTION CONTENT (you MUST align with this):\n"
                        prompt += f"Parent Section: {section.parent.text}\n"
                        prompt += f"{'-'*60}\n"
                        prompt += f"{parent_actual_content}\n"
                        prompt += f"{'-'*60}\n"
                        prompt += f"\nIMPORTANT: Your content for '{section.text}' MUST:\n"
                        prompt += f"1. Directly relate to and expand upon the parent section content above\n"
                        prompt += f"2. Use consistent terminology with the parent section\n"
                        prompt += f"3. Reference parent concepts where appropriate\n"
                        prompt += f"4. Be a logical subdivision of the parent topic\n"
                        prompt += f"5. Not contradict anything in the parent section\n"
                    
                    # *** NEW: Add sibling context ***
                    if sibling_info:
                        prompt += f"\n{sibling_info}\n"
                        prompt += f"Note: Ensure your content is distinct from these sibling sections.\n"
                    
                    # Add mode-specific instructions
                    if mode == "replace":
                        prompt += "\n\nYour task: Write comprehensive content for this section from scratch."
                    elif mode == "rework":
                        prompt += f"\n\nEXISTING CONTENT TO REWORK:\n{existing_content}\n\n"
                        prompt += "Your task: Rewrite and enhance the existing content."
                    elif mode == "append":
                        prompt += f"\n\nEXISTING CONTENT:\n{existing_content}\n\n"
                        prompt += "Your task: Add additional relevant content."
                    
                    # Add knowledge base instruction
                    if self.selected_knowledge_collections:
                        collection_names = [col['name'] for col in self.selected_knowledge_collections]
                        prompt += f"\n\nIMPORTANT: Reference knowledge base: {', '.join(collection_names)}"
                    
                    self.update_progress_log(progress_window, "  Sending request to API...")
                    
                    # Generate content
                    response = self.query_openwebui(prompt)
                    
                    if response and not response.startswith("Error:"):
                        self.update_progress_log(progress_window, "  ✓ Content generated")
                        
                        # *** NEW: Store generated content for future subsections ***
                        section_hash = section.get_section_hash()
                        generated_content_map[section_hash] = response
                        
                        # Apply content to section
                        self.root.after(0, lambda: 
                                    progress_window.status_label.config(text="Committing to document..."))
                        
                        if mode == "replace":
                            self.replace_section_content(section, response)
                        elif mode == "rework":
                            self.replace_section_content(section, response)
                        elif mode == "append":
                            self.append_section_content(section, response)
                        
                        # Mark as edited
                        self.mark_section_edited(section)
                        
                        self.update_progress_log(progress_window, "  ✓ Committed to document")
                        
                        # Auto-save if enabled
                        if self.auto_config['auto_save'].get():
                            try:
                                self.document.save(self.document_path)
                                self.update_progress_log(progress_window, "  ✓ Auto-saved")
                            except Exception as e:
                                self.update_progress_log(progress_window, f"  ⚠ Auto-save failed: {str(e)}")
                        
                        completed += 1
                        
                    else:
                        self.update_progress_log(progress_window, f"  ✗ Generation failed: {response}")
                        failed += 1
                        
                except Exception as e:
                    self.update_progress_log(progress_window, f"  ✗ Error: {str(e)}")
                    failed += 1
            
            # Final update
            self.root.after(0, lambda: progress_window.progress_bar.configure(value=len(sections)))
            self.root.after(0, lambda: 
                        progress_window.progress_label.config(text="Complete!"))
            self.root.after(0, lambda: 
                        progress_window.status_label.config(text=""))
            self.root.after(0, lambda: 
                        progress_window.current_section_label.config(text="All sections processed"))
            
            # Summary
            self.update_progress_log(progress_window, f"\n{'='*60}")
            self.update_progress_log(progress_window, "SUMMARY:")
            self.update_progress_log(progress_window, f"  ✓ Completed: {completed}")
            self.update_progress_log(progress_window, f"  ✗ Failed: {failed}")
            self.update_progress_log(progress_window, f"  Total: {len(sections)}")
            
            # Enable close button and disable others
            self.root.after(0, lambda: progress_window.close_btn.config(state='normal'))
            self.root.after(0, lambda: progress_window.pause_btn.config(state='disabled'))
            self.root.after(0, lambda: progress_window.stop_btn.config(state='disabled'))
            
            # Refresh tree and reload if enabled
            self.root.after(0, self.populate_tree)
            if self.auto_config['auto_reload'].get():
                self.root.after(0, self.reload_document)
            
            # Save final document if not auto-saving
            if not self.auto_config['auto_save'].get():
                self.root.after(0, self.prompt_save_document)
            
            self.log_message(f"Auto-complete finished: {completed} completed, {failed} failed")
            
        except Exception as e:
            self.update_progress_log(progress_window, f"\nFATAL ERROR: {str(e)}")
            self.log_message(f"Auto-complete error: {str(e)}")


    def update_progress_log(self, window, message):
        """Update progress log in auto-complete window"""
        def update():
            window.log_text.insert(tk.END, f"{message}\n")
            window.log_text.see(tk.END)
            window.update()
        
        self.root.after(0, update)

    def regenerate_with_prompt(self):
        """Regenerate content using edited prompt with selected model"""
        if not self.selected_section:
            messagebox.showwarning("Warning", "No section selected")
            return
        
        edited_prompt = self.prompt_text.get('1.0', tk.END).strip()
        if not edited_prompt:
            messagebox.showwarning("Warning", "Prompt is empty")
            return
        
        # Get selected model or use default
        selected_model = self.prompt_model_var.get()
        if not selected_model:
            selected_model = self.selected_model.get()
        
        self.notebook.select(2)  # Switch to console
        
        def regenerate_thread():
            try:
                self.generate_btn.config(state='disabled')
                self.log_message(f"Regenerating with model: {selected_model}...")
                
                response = self.query_openwebui_with_model(edited_prompt, selected_model)
                
                if response and not response.startswith("Error:"):
                    self.generated_content = response
                    self.last_sent_prompt = edited_prompt
                    self.root.after(0, self.show_generated_content)
                    self.log_message("Content regenerated successfully")
                else:
                    self.log_message(f"Regeneration failed: {response}")
                    self.root.after(0, self.handle_generation_error)
                    
            except Exception as e:
                self.log_message(f"Regeneration error: {str(e)}")
                self.root.after(0, self.handle_generation_error)
            finally:
                self.root.after(0, lambda: self.generate_btn.config(state='normal'))
        
        thread = threading.Thread(target=regenerate_thread)
        thread.daemon = True
        thread.start()
    
    def ensure_models_loaded(self):
        """Ensure models are loaded, try to refresh if not"""
        if not self.available_models:
            if not self.openwebui_api_key or not self.openwebui_base_url:
                return False
            
            # Try to refresh models automatically
            try:
                self.log_message("Auto-refreshing models...")
                headers = {'Authorization': f'Bearer {self.openwebui_api_key}'}
                response = requests.get(f"{self.openwebui_base_url}/api/models", 
                                    headers=headers, timeout=10)
                
                if response.status_code == 200:
                    data = response.json()
                    self.available_models = []
                    
                    if isinstance(data, dict) and 'data' in data:
                        models_list = data['data']
                    elif isinstance(data, list):
                        models_list = data
                    else:
                        models_list = [data] if data else []
                    
                    for model in models_list:
                        if isinstance(model, dict):
                            model_id = model.get('id', model.get('name', 'unknown'))
                            self.available_models.append(model_id)
                        else:
                            self.available_models.append(str(model))
                    
                    self.log_message(f"Auto-loaded {len(self.available_models)} models")
                    return len(self.available_models) > 0
                else:
                    self.log_message(f"Failed to auto-load models: HTTP {response.status_code}")
                    return False
            except Exception as e:
                self.log_message(f"Failed to auto-load models: {str(e)}")
                return False
        
        return True

    def compare_three_models(self):
        """Generate content with 3 different models and compare side by side"""
        if not self.selected_section:
            messagebox.showwarning("Warning", "No section selected")
            return
        
        edited_prompt = self.prompt_text.get('1.0', tk.END).strip()
        if not edited_prompt:
            messagebox.showwarning("Warning", "Prompt is empty")
            return
        
        # TRY TO LOAD MODELS IF NOT ALREADY LOADED
        if not self.ensure_models_loaded():
            messagebox.showerror("Error", 
                "Could not load models.\n\n"
                "Please:\n"
                "1. Go to 'Configure AI Settings'\n"
                "2. Enter your API URL and Key\n"
                "3. Click 'Refresh Models'")
            return
        
        if len(self.available_models) < 3:
            messagebox.showwarning("Warning", 
                f"Need at least 3 models available for comparison.\n"
                f"Currently have: {len(self.available_models)} model(s)\n\n"
                "Go to 'Configure AI Settings' and ensure models are available.")
            return
        
        # Create comparison window
        self.create_comparison_window(edited_prompt)

    def create_comparison_window(self, prompt):
        """Create model comparison window with easy model selection"""
        comp_window = tk.Toplevel(self.root)
        comp_window.title("Model Comparison")
        comp_window.geometry("1400x800")
        comp_window.configure(bg="#2b2b2b")
        comp_window.grab_set()
        
        main_frame = ttk.Frame(comp_window, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Compare 3 Models - Select Best Result", 
                font=("Arial", 14, "bold")).pack(pady=(0, 10))
        
        # ========== EASY 3-MODEL SELECTION STARTS HERE ==========
        # Model selection frame with listbox
        model_select_frame = ttk.LabelFrame(main_frame, text="Select 3 Models", padding="10")
        model_select_frame.pack(fill=tk.X, pady=(0, 10))
        
        select_inner = ttk.Frame(model_select_frame)
        select_inner.pack(fill=tk.X)
        
        ttk.Label(select_inner, text="Select exactly 3 models (Ctrl+Click):").pack(side=tk.LEFT, padx=(0, 10))
        
        # THIS IS THE LISTBOX - Multi-select with Ctrl+Click
        model_listbox = tk.Listbox(select_inner, selectmode=tk.MULTIPLE, height=5,
                                bg="#1e1e1e", fg="#ffffff", width=60)
        model_listbox.pack(side=tk.LEFT, padx=(0, 10))
        
        # Populate listbox with all available models
        for model in self.available_models:
            model_listbox.insert(tk.END, model)
        
        # AUTO-SELECT FIRST 3 MODELS BY DEFAULT
        for i in range(min(3, len(self.available_models))):
            model_listbox.selection_set(i)
        
        # SELECTION COUNTER - Shows "Selected: X/3"
        selected_label = ttk.Label(select_inner, text="Selected: 3/3", foreground="#00ff00")
        selected_label.pack(side=tk.LEFT, padx=(10, 0))
        
        # UPDATE SELECTION COUNT AND COLOR
        def update_selection_count(event=None):
            count = len(model_listbox.curselection())
            if count == 3:
                # GREEN when 3 selected - Enable button
                selected_label.config(text=f"Selected: {count}/3", foreground="#00ff00")
                generate_btn.config(state='normal')
            else:
                # RED when not 3 - Disable button
                selected_label.config(text=f"Selected: {count}/3", foreground="#ff0000")
                generate_btn.config(state='disabled')
        
        # Bind selection change event
        model_listbox.bind('<<ListboxSelect>>', update_selection_count)
        
        btn_frame = ttk.Frame(model_select_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        # GENERATE BUTTON - Only enabled when exactly 3 models selected
        generate_btn = ttk.Button(btn_frame, text="Generate All", 
                                command=lambda: self.start_comparison_from_listbox(prompt, model_listbox,
                                                                                    result_frames, status_labels))
        generate_btn.pack(side=tk.LEFT)
        # ========== EASY 3-MODEL SELECTION ENDS HERE ==========
        
        # Results frame - 3 columns side by side
        results_container = ttk.Frame(main_frame)
        results_container.pack(fill=tk.BOTH, expand=True)
        results_container.columnconfigure(0, weight=1)
        results_container.columnconfigure(1, weight=1)
        results_container.columnconfigure(2, weight=1)
        results_container.rowconfigure(0, weight=1)
        
        result_frames = []
        status_labels = []
        text_widgets = []
        select_buttons = []
        
        # CREATE 3 SIDE-BY-SIDE RESULT PANELS
        for i in range(3):
            # Column frame
            col_frame = ttk.Frame(results_container, relief="solid", borderwidth=2)
            col_frame.grid(row=0, column=i, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5, pady=5)
            col_frame.columnconfigure(0, weight=1)
            col_frame.rowconfigure(1, weight=1)
            
            result_frames.append(col_frame)
            
            # Header with model name and status
            header_frame = ttk.Frame(col_frame)
            header_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=5, pady=5)
            
            model_label = ttk.Label(header_frame, text=f"Model {i+1}: [Not Selected]", 
                    font=("Arial", 10, "bold"))
            model_label.pack(side=tk.LEFT)
            
            status_label = ttk.Label(header_frame, text="⏳ Pending", foreground="#ffaa00")
            status_label.pack(side=tk.RIGHT)
            status_labels.append(status_label)
            
            # Store model label for updating with actual model name
            col_frame.model_label = model_label
            
            # Text area for generated content
            text_widget = scrolledtext.ScrolledText(col_frame, height=25, width=40,
                                                bg="#1e1e1e", fg="#ffffff",
                                                font=("Consolas", 9), wrap=tk.WORD)
            text_widget.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5, pady=5)
            text_widgets.append(text_widget)
            
            # Select button for choosing this result
            select_btn = ttk.Button(col_frame, text="Select This Result", 
                                command=lambda idx=i: self.select_comparison_result(idx, text_widgets, 
                                                                                    result_frames, comp_window))
            select_btn.grid(row=2, column=0, pady=5)
            select_btn.config(state='disabled')
            select_buttons.append(select_btn)
        
        # Store widgets for access in generation
        comp_window.text_widgets = text_widgets
        comp_window.select_buttons = select_buttons
        comp_window.result_frames = result_frames

    # Add new function to handle listbox-based model selection
    def start_comparison_from_listbox(self, prompt, model_listbox, result_frames, status_labels):
        """Start comparison from listbox selection"""
        selected_indices = model_listbox.curselection()
        
        if len(selected_indices) != 3:
            messagebox.showwarning("Warning", "Please select exactly 3 models")
            return
        
        models = [model_listbox.get(i) for i in selected_indices]
        
        # Update model labels in result frames
        for i, (model, frame) in enumerate(zip(models, result_frames)):
            frame.model_label.config(text=f"Model {i+1}: {model}")
        
        # Get window references
        comp_window = result_frames[0].master.master.master
        text_widgets = comp_window.text_widgets
        select_buttons = comp_window.select_buttons
        
        def generation_thread():
            for i, model in enumerate(models):
                try:
                    # Update status to processing
                    self.root.after(0, lambda idx=i: status_labels[idx].config(
                        text="⚙️ Processing...", foreground="#00aaff"))
                    
                    self.log_message(f"Generating with {model}...")
                    
                    # Generate content
                    response = self.query_openwebui_with_model(prompt, model)
                    
                    if response and not response.startswith("Error:"):
                        # Update text widget
                        self.root.after(0, lambda idx=i, content=response: (
                            text_widgets[idx].delete('1.0', tk.END),
                            text_widgets[idx].insert('1.0', content),
                            status_labels[idx].config(text="✓ Complete", foreground="#00ff00"),
                            select_buttons[idx].config(state='normal')
                        ))
                        self.log_message(f"{model} completed successfully")
                    else:
                        # Show error
                        self.root.after(0, lambda idx=i, err=response: (
                            text_widgets[idx].delete('1.0', tk.END),
                            text_widgets[idx].insert('1.0', f"Error:\n{err}"),
                            status_labels[idx].config(text="✗ Failed", foreground="#ff0000")
                        ))
                        self.log_message(f"{model} failed: {response}")
                        
                except Exception as e:
                    self.root.after(0, lambda idx=i, err=str(e): (
                        text_widgets[idx].delete('1.0', tk.END),
                        text_widgets[idx].insert('1.0', f"Error:\n{err}"),
                        status_labels[idx].config(text="✗ Error", foreground="#ff0000")
                    ))
                    self.log_message(f"{model} error: {str(e)}")
        
        thread = threading.Thread(target=generation_thread)
        thread.daemon = True
        thread.start()

    def start_comparison_generation(self, prompt, model_vars, result_frames, status_labels):
        """Start sequential generation for 3 models"""
        models = [var.get() for var in model_vars]
        
        if len(set(models)) != 3:
            messagebox.showwarning("Warning", "Please select 3 different models")
            return
        
        # Get window references
        comp_window = result_frames[0].master.master.master
        text_widgets = comp_window.text_widgets
        select_buttons = comp_window.select_buttons
        
        def generation_thread():
            for i, model in enumerate(models):
                try:
                    # Update status to processing
                    self.root.after(0, lambda idx=i: status_labels[idx].config(
                        text="⚙️ Processing...", foreground="#00aaff"))
                    
                    self.log_message(f"Generating with {model}...")
                    
                    # Generate content
                    response = self.query_openwebui_with_model(prompt, model)
                    
                    if response and not response.startswith("Error:"):
                        # Update text widget
                        self.root.after(0, lambda idx=i, content=response: (
                            text_widgets[idx].delete('1.0', tk.END),
                            text_widgets[idx].insert('1.0', content),
                            status_labels[idx].config(text="✓ Complete", foreground="#00ff00"),
                            select_buttons[idx].config(state='normal')
                        ))
                        self.log_message(f"{model} completed successfully")
                    else:
                        # Show error
                        self.root.after(0, lambda idx=i, err=response: (
                            text_widgets[idx].delete('1.0', tk.END),
                            text_widgets[idx].insert('1.0', f"Error:\n{err}"),
                            status_labels[idx].config(text="✗ Failed", foreground="#ff0000")
                        ))
                        self.log_message(f"{model} failed: {response}")
                        
                except Exception as e:
                    self.root.after(0, lambda idx=i, err=str(e): (
                        text_widgets[idx].delete('1.0', tk.END),
                        text_widgets[idx].insert('1.0', f"Error:\n{err}"),
                        status_labels[idx].config(text="✗ Error", foreground="#ff0000")
                    ))
                    self.log_message(f"{model} error: {str(e)}")
        
        thread = threading.Thread(target=generation_thread)
        thread.daemon = True
        thread.start()
    
    def select_comparison_result(self, index, text_widgets, result_frames, window):
        """Select a result from comparison and move to preview"""
        # Get the content
        selected_content = text_widgets[index].get('1.0', tk.END).strip()
        
        if not selected_content or selected_content.startswith("Error:"):
            messagebox.showwarning("Warning", "Cannot select error result")
            return
        
        # Highlight selected frame
        for i, frame in enumerate(result_frames):
            if i == index:
                frame.configure(relief="solid", borderwidth=4)
            else:
                frame.configure(relief="solid", borderwidth=2)
        
        # Move to preview
        self.generated_content = selected_content
        self.show_generated_content()
        
        self.log_message(f"Selected result from Model {index + 1}")
        messagebox.showinfo("Success", f"Result from Model {index + 1} moved to Content Preview")
        
        window.destroy()
    


    def offer_prompt_update(self):
        """Offer to update master prompt with improvements"""
        current_prompt = self.prompt_text.get('1.0', tk.END).strip()
        
        dialog = tk.Toplevel(self.root)
        dialog.title("Update Master Prompt?")
        dialog.geometry("600x400")
        dialog.configure(bg="#2b2b2b")
        dialog.grab_set()
        
        frame = ttk.Frame(dialog, padding="20")
        frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(frame, text="Update Master Prompt with Changes?",
                 font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 10))
        
        ttk.Label(frame, text="Enter specific improvements to add to the master prompt:",
                 wraplength=550).pack(anchor=tk.W, pady=(0, 10))
        
        improvements_text = scrolledtext.ScrolledText(frame, height=10, width=70,
                                                      bg="#1e1e1e", fg="#ffffff",
                                                      font=("Consolas", 10), wrap=tk.WORD)
        improvements_text.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        improvements_text.insert('1.0', "e.g., 'Always include specific tool versions'\n'Reference DoD 8500 series'")
        
        btn_frame = ttk.Frame(frame)
        btn_frame.pack(fill=tk.X)
        
        def update_master():
            improvements = improvements_text.get('1.0', tk.END).strip()
            if improvements and not improvements.startswith('e.g.,'):
                current_master = self.master_prompt.get()
                updated_master = current_master + "\n\nADDITIONAL REQUIREMENTS:\n" + improvements
                self.master_prompt.set(updated_master)
                self.save_settings()
                self.log_message("Master prompt updated with new requirements")
                messagebox.showinfo("Success", "Master prompt updated successfully!")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="Update Master Prompt", 
                  command=update_master).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT)
        
    def handle_generation_error(self):
        """Handle generation error"""
        self.notebook.select(1)  # Switch to prompt tab
        messagebox.showerror("Generation Failed", 
                           "Content generation failed. Edit the prompt in the 'Prompt Editor' tab and try 'Regenerate'.")
        
    def clear_preview(self):
        """Clear preview pane"""
        self.generated_text.delete('1.0', tk.END)
        self.commit_btn.config(state='disabled')
        self.log_message("Preview cleared")
        
    def commit_content(self):
        """Commit generated content to document"""
        if not self.selected_section or not self.document:
            return
            
        # Get edited content from preview (without formatting tags)
        content_to_commit = self.generated_text.get('1.0', tk.END).strip()
        
        if not content_to_commit:
            messagebox.showwarning("Warning", "No content to commit")
            return
        
        # Confirm commit
        result = messagebox.askyesno(
            "Commit Content",
            f"Commit content to section: {self.selected_section.text}\n\n"
            f"Mode: {self.operation_mode.get().upper()}\n\n"
            "This will modify the document."
        )
        
        if not result:
            return
        
        try:
            # Check if backup is needed
            if self.auto_config['ask_backup'].get():
                self.prompt_for_backup()
            else:
                self.create_backup()
            
            # Apply content based on mode
            mode = self.operation_mode.get()
            section = self.selected_section
            
            if mode == "replace":
                self.replace_section_content(section, content_to_commit)
            elif mode == "rework":
                self.replace_section_content(section, content_to_commit)
            elif mode == "append":
                self.append_section_content(section, content_to_commit)
            
            # Mark section as edited
            self.mark_section_edited(section)
            
            # Refresh tree
            self.populate_tree()
            
            # Auto-save if enabled
            if self.auto_config['auto_save'].get():
                self.save_document_auto()
            else:
                self.prompt_save_document()
            
            # Clear preview
            self.clear_preview()
            
            # Auto-reload if enabled
            if self.auto_config['auto_reload'].get():
                self.reload_document()
            else:
                self.show_existing_content()
            
            self.log_message(f"Content committed to: {section.text}")
            
        except Exception as e:
            self.log_message(f"Error committing content: {str(e)}")
            messagebox.showerror("Error", f"Failed to commit content: {str(e)}")
    
    def prompt_for_backup(self):
        """Prompt user for backup with 'don't ask again' option"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Create Backup?")
        dialog.geometry("400x150")
        dialog.configure(bg="#2b2b2b")
        dialog.grab_set()
        
        frame = ttk.Frame(dialog, padding="20")
        frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(frame, text="Create a backup before committing changes?",
                 font=("Arial", 10)).pack(pady=(0, 20))
        
        dont_ask_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(frame, text="Don't ask me again", 
                       variable=dont_ask_var).pack(pady=(0, 20))
        
        btn_frame = ttk.Frame(frame)
        btn_frame.pack()
        
        result = {'backup': False, 'dont_ask': False}
        
        def yes_clicked():
            result['backup'] = True
            result['dont_ask'] = dont_ask_var.get()
            dialog.destroy()
        
        def no_clicked():
            result['backup'] = False
            result['dont_ask'] = dont_ask_var.get()
            dialog.destroy()
        
        ttk.Button(btn_frame, text="Yes", command=yes_clicked).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="No", command=no_clicked).pack(side=tk.LEFT, padx=5)
        
        dialog.wait_window()
        
        if result['dont_ask']:
            self.auto_config['ask_backup'].set(False)
            self.save_settings()
            self.log_message("Backup prompts disabled")
        
        if result['backup']:
            self.create_backup()
            
    def create_backup(self):
        """Create backup of document"""
        if not self.document_path:
            return
            
        try:
            backup_dir = os.path.dirname(self.document_path)
            filename = os.path.basename(self.document_path)
            name, ext = os.path.splitext(filename)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_path = os.path.join(backup_dir, f"{name}_backup_{timestamp}{ext}")
            
            shutil.copy2(self.document_path, backup_path)
            self.log_message(f"Backup created: {os.path.basename(backup_path)}")
            
        except Exception as e:
            self.log_message(f"Warning: Could not create backup: {str(e)}")
    
    def replace_section_content(self, section, content):
        """Replace section content with new content"""
        # Remove existing content
        self.remove_section_content(section)
        
        # Add new content with markdown conversion
        self.add_markdown_content_to_section(section, content)
    
    def add_markdown_export_functionality(self):
        """Add markdown export functionality to the application"""
        try:
            import tkinter as tk
            from tkinter import filedialog, messagebox
            
            def export_to_markdown():
                """Export current content to markdown file"""
                try:
                    if not self.generated_content:
                        messagebox.showwarning("No Content", "No generated content to export.")
                        return
                    
                    # Get save location
                    file_path = filedialog.asksaveasfilename(
                        title="Export to Markdown",
                        defaultextension=".md",
                        filetypes=[("Markdown files", "*.md"), ("Text files", "*.txt"), ("All files", "*.*")]
                    )
                    
                    if file_path:
                        with open(file_path, 'w', encoding='utf-8') as f:
                            f.write(self.generated_content)
                        
                        messagebox.showinfo("Export Complete", f"Content exported to: {file_path}")
                        self.log_message(f"Content exported to markdown: {file_path}")
                        
                except PermissionError:
                    messagebox.showerror("Export Error", "Permission denied. Please choose a different location or close the file if it's open.")
                except Exception as e:
                    messagebox.showerror("Export Error", f"Failed to export: {str(e)}")
            
            def import_from_markdown():
                """Import content from markdown file"""
                try:
                    file_path = filedialog.askopenfilename(
                        title="Import from Markdown",
                        filetypes=[("Markdown files", "*.md"), ("Text files", "*.txt"), ("All files", "*.*")]
                    )
                    
                    if file_path:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                        
                        # Set the content and show preview
                        self.generated_content = content
                        self.show_generated_content()
                        
                        messagebox.showinfo("Import Complete", f"Content imported from: {file_path}")
                        self.log_message(f"Content imported from markdown: {file_path}")
                        
                except PermissionError:
                    messagebox.showerror("Import Error", "Permission denied. Please check file permissions.")
                except Exception as e:
                    messagebox.showerror("Import Error", f"Failed to import: {str(e)}")
            
            # Add export/import buttons to the interface
            if hasattr(self, 'button_frame'):
                export_btn = tk.Button(
                    self.button_frame,
                    text="Export MD",
                    command=export_to_markdown,
                    bg='#404040',
                    fg='#ffffff',
                    padx=10
                )
                export_btn.pack(side=tk.LEFT, padx=2)
                
                import_btn = tk.Button(
                    self.button_frame,
                    text="Import MD",
                    command=import_from_markdown,
                    bg='#404040',
                    fg='#ffffff',
                    padx=10
                )
                import_btn.pack(side=tk.LEFT, padx=2)
                
        except Exception as e:
            print(f"Error adding markdown export functionality: {e}")

    def style_table_cell(self, cell, row_idx, col_idx):
        """Apply styling to table cell"""
        try:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            
            if row_idx == 0:  # Header row
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                self.set_cell_background_color(cell, (52, 73, 94))
            else:
                run.font.size = Pt(10)
                if row_idx % 2 == 0:
                    self.set_cell_background_color(cell, (248, 249, 250))
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self.set_cell_border(cell, (52, 58, 64))
        except Exception as e:
            self.log_message(f"Error styling cell: {e}")

    def set_cell_background_color(self, cell, color_tuple):
        """Set cell background color"""
        try:
            cell_element = cell._element
            cell_properties = cell_element.get_or_add_tcPr()
            
            shading_element = OxmlElement('w:shd')
            color_hex = f'{color_tuple[0]:02x}{color_tuple[1]:02x}{color_tuple[2]:02x}'
            shading_element.set(qn('w:fill'), color_hex)
            shading_element.set(qn('w:val'), 'clear')
            
            cell_properties.append(shading_element)
        except Exception as e:
            self.log_message(f"Error setting cell background: {e}")

    def set_cell_border(self, cell, color_tuple, border_size=4):
        """Set cell borders"""
        try:
            cell_element = cell._element
            cell_properties = cell_element.get_or_add_tcPr()
            
            borders = OxmlElement('w:tcBorders')
            color_hex = f'{color_tuple[0]:02x}{color_tuple[1]:02x}{color_tuple[2]:02x}'
            
            for side in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), str(border_size))
                border.set(qn('w:color'), color_hex)
                borders.append(border)
            
            cell_properties.append(borders)
        except Exception as e:
            self.log_message(f"Error setting cell border: {e}")

    def create_table_from_data(self, table_data):
        """Create a styled table from table data"""
        try:
            if not table_data:
                return None
            
            rows, cols = len(table_data), len(table_data[0])
            table = self.document.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Set column widths
            for col in table.columns:
                col.width = Inches(6.5 / cols)
            
            # Populate and style table
            for row_idx, row_data in enumerate(table_data):
                table_row = table.rows[row_idx]
                for col_idx, cell_data in enumerate(row_data):
                    cell = table_row.cells[col_idx]
                    cell.text = str(cell_data)
                    self.style_table_cell(cell, row_idx, col_idx)
            
            return table
        except Exception as e:
            self.log_message(f"Error creating table: {e}")
            return None

    def add_table_editing_dialog(self):
        """Add table creation and editing dialog"""
        try:
            import tkinter as tk
            from tkinter import ttk, simpledialog, messagebox
            
            def create_table_dialog():
                """Show dialog to create a new table"""
                try:
                    dialog = tk.Toplevel(self.root)
                    dialog.title("Create Table")
                    dialog.geometry("500x400")
                    dialog.configure(bg="#2b2b2b")
                    dialog.transient(self.root)
                    dialog.grab_set()
                    
                    # Table configuration
                    config_frame = ttk.Frame(dialog)
                    config_frame.pack(fill=tk.X, padx=10, pady=5)
                    
                    # Rows and columns
                    ttk.Label(config_frame, text="Rows:").grid(row=0, column=0, sticky='w')
                    rows_var = tk.IntVar(value=3)
                    rows_spin = ttk.Spinbox(config_frame, from_=1, to=20, textvariable=rows_var, width=5)
                    rows_spin.grid(row=0, column=1, padx=(5,20))
                    
                    ttk.Label(config_frame, text="Columns:").grid(row=0, column=2, sticky='w')
                    cols_var = tk.IntVar(value=3)
                    cols_spin = ttk.Spinbox(config_frame, from_=1, to=10, textvariable=cols_var, width=5)
                    cols_spin.grid(row=0, column=3, padx=5)
                    
                    # Has headers checkbox
                    headers_var = tk.BooleanVar(value=True)
                    headers_check = ttk.Checkbutton(config_frame, text="First row as headers", variable=headers_var)
                    headers_check.grid(row=1, column=0, columnspan=4, sticky='w', pady=5)
                    
                    # Table data entry
                    table_frame = ttk.Frame(dialog)
                    table_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
                    
                    # Scrollable frame for table entries
                    canvas = tk.Canvas(table_frame, bg="#2b2b2b")
                    scrollbar = ttk.Scrollbar(table_frame, orient="vertical", command=canvas.yview)
                    scrollable_frame = ttk.Frame(canvas)
                    
                    scrollable_frame.bind(
                        "<Configure>",
                        lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
                    )
                    
                    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
                    canvas.configure(yscrollcommand=scrollbar.set)
                    
                    canvas.pack(side="left", fill="both", expand=True)
                    scrollbar.pack(side="right", fill="y")
                    
                    entry_widgets = []
                    
                    def update_table_grid():
                        """Update the table grid based on current settings"""
                        # Clear existing widgets
                        for widget in scrollable_frame.winfo_children():
                            widget.destroy()
                        entry_widgets.clear()
                        
                        rows = rows_var.get()
                        cols = cols_var.get()
                        
                        for r in range(rows):
                            row_entries = []
                            for c in range(cols):
                                entry = tk.Entry(scrollable_frame, width=15)
                                entry.grid(row=r, column=c, padx=2, pady=2)
                                
                                # Set default values
                                if r == 0 and headers_var.get():
                                    entry.insert(0, f"Header {c+1}")
                                else:
                                    entry.insert(0, f"Cell {r+1}-{c+1}")
                                
                                row_entries.append(entry)
                            entry_widgets.append(row_entries)
                    
                    # Initial table grid
                    update_table_grid()
                    
                    # Update button
                    update_btn = ttk.Button(config_frame, text="Update Grid", command=update_table_grid)
                    update_btn.grid(row=2, column=0, columnspan=4, pady=5)
                    
                    # Action buttons
                    button_frame = ttk.Frame(dialog)
                    button_frame.pack(fill=tk.X, padx=10, pady=5)
                    
                    def generate_table():
                        """Generate markdown table from entries"""
                        try:
                            rows = rows_var.get()
                            cols = cols_var.get()
                            has_headers = headers_var.get()
                            
                            markdown_lines = []
                            
                            # Generate table rows
                            for r in range(rows):
                                row_data = []
                                for c in range(cols):
                                    if r < len(entry_widgets) and c < len(entry_widgets[r]):
                                        value = entry_widgets[r][c].get().strip()
                                        row_data.append(value if value else " ")
                                    else:
                                        row_data.append(" ")
                                
                                markdown_lines.append("| " + " | ".join(row_data) + " |")
                                
                                # Add separator after header row
                                if r == 0 and has_headers:
                                    separator = "| " + " | ".join(["---"] * cols) + " |"
                                    markdown_lines.append(separator)
                            
                            # Insert table into current content
                            table_text = "\n".join(markdown_lines) + "\n\n"
                            
                            # Add to generated content
                            if hasattr(self, 'generated_content'):
                                if self.generated_content:
                                    self.generated_content += "\n\n" + table_text
                                else:
                                    self.generated_content = table_text
                            else:
                                self.generated_content = table_text
                            
                            # Update preview
                            self.show_generated_content()
                            
                            dialog.destroy()
                            messagebox.showinfo("Table Created", "Table added to content successfully!")
                            
                        except Exception as e:
                            messagebox.showerror("Error", f"Failed to create table: {str(e)}")
                    
                    ttk.Button(button_frame, text="Create Table", command=generate_table).pack(side=tk.LEFT, padx=5)
                    ttk.Button(button_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)
                    
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to open table dialog: {str(e)}")
            
            # Add table creation button to interface
            if hasattr(self, 'button_frame'):
                table_btn = tk.Button(
                    self.button_frame,
                    text="Create Table",
                    command=create_table_dialog,
                    bg='#404040',
                    fg='#ffffff',
                    padx=10
                )
                table_btn.pack(side=tk.LEFT, padx=2)
                
        except Exception as e:
            print(f"Error adding table editing functionality: {e}")



    def append_section_content(self, section, content):
        """Append content to section"""
        self.add_markdown_content_to_section(section, content, append=True)
    
    def remove_section_content(self, section):
        """Remove all content paragraphs from section"""
        doc_paragraphs = list(self.document.paragraphs)
        heading_index = -1
        
        # Find heading
        for i, para in enumerate(doc_paragraphs):
            if para.text.strip() == section.paragraph.text.strip():
                heading_index = i
                break
        
        if heading_index == -1:
            raise Exception("Could not find section heading")
        
        # Find next heading
        next_heading_index = len(doc_paragraphs)
        for i in range(heading_index + 1, len(doc_paragraphs)):
            if doc_paragraphs[i].style.name.startswith('Heading'):
                next_heading_index = i
                break
        
        # Remove content paragraphs
        for i in range(heading_index + 1, next_heading_index):
            para = doc_paragraphs[i]
            if not para.style.name.startswith('Heading'):
                self.remove_paragraph(para)
        
        section.content_paragraphs = []
    
    def add_markdown_content_to_section(self, section, content, append=False):
        """Add content to section with markdown to Word formatting conversion"""
        heading_para = section.paragraph
        doc_paragraphs = list(self.document.paragraphs)
        
        # Find heading index
        heading_index = -1
        for i, para in enumerate(doc_paragraphs):
            if para.text.strip() == heading_para.text.strip():
                heading_index = i
                break
        
        if heading_index == -1:
            raise Exception("Could not find section heading")
        
        # Determine insertion point
        if append and section.content_paragraphs:
            next_heading_index = len(doc_paragraphs)
            for i in range(heading_index + 1, len(doc_paragraphs)):
                if doc_paragraphs[i].style.name.startswith('Heading'):
                    next_heading_index = i
                    break
            insertion_index = next_heading_index
        else:
            insertion_index = heading_index + 1
        
        # Split content into paragraphs
        content_paragraphs = content.split('\n')
        
        inserted_paras = []
        for para_text in content_paragraphs:
            if para_text.strip():
                # Create new paragraph
                if insertion_index < len(self.document.paragraphs):
                    new_para = self.document.paragraphs[insertion_index].insert_paragraph_before()
                else:
                    new_para = self.document.add_paragraph()
                
                # Apply markdown formatting to paragraph
                self.apply_markdown_to_paragraph(new_para, para_text.strip())
                
                # Apply configured formatting
                self.apply_configured_formatting(new_para)
                
                inserted_paras.append(new_para)
                insertion_index += 1
        
        # Update section's content paragraphs
        if not append:
            section.content_paragraphs = inserted_paras
        else:
            section.content_paragraphs.extend(inserted_paras)

    def insert_generated_content(self, content, section):
        """Insert generated content with markdown table support"""
        try:
            if not content or not section:
                return False
            
            self.log_message(f"Inserting content into section: {section.get_full_path()}")
            
            # Extract tables from content
            table_pattern = r'((?:\s*│[^│\n]*│[^│\n]*│[^\n]*\n)+)'
            table_matches = list(re.finditer(table_pattern, content, re.MULTILINE))
            
            processed_content = content
            tables_created = []
            
            # Process tables (reverse order to maintain positions)
            for match in reversed(table_matches):
                table_text = match.group(1)
                table_data = self.parse_markdown_table(table_text)
                
                if table_data:
                    table = self.create_table_from_data(table_data)
                    if table:
                        tables_created.insert(0, table)
                        processed_content = (processed_content[:match.start()] + 
                                        f"\n[TABLE_{len(tables_created)}]\n" + 
                                        processed_content[match.end():])
            
            # Convert remaining markdown
            text_content = self.convert_markdown_to_docx(processed_content)
            
            # Handle insertion based on operation mode
            target_paragraph = section.paragraph
            operation_mode = self.operation_mode.get()
            
            if operation_mode == "replace":
                for para in section.content_paragraphs[:]:
                    if para.text.strip():
                        self.remove_paragraph(para)
                section.content_paragraphs.clear()
            elif operation_mode == "append":
                for para in reversed(section.content_paragraphs):
                    if para.text.strip():
                        target_paragraph = para
                        break
            
            # Insert text content
            if text_content.strip():
                new_paragraph = target_paragraph.insert_paragraph_after()
                new_paragraph.text = text_content
                
                for run in new_paragraph.runs:
                    self.apply_formatting_to_run(run)
            
            # Log results
            if tables_created:
                self.log_message(f"Content inserted with {len(tables_created)} tables")
            else:
                self.log_message("Content inserted successfully")
            
            return True
            
        except Exception as e:
            self.log_message(f"Error inserting content: {e}")
            return False

    def parse_markdown_table(self, markdown_text):
        """Parse markdown table from text and extract table data"""
        try:
            lines = markdown_text.strip().split('\n')
            table_lines = [line.strip() for line in lines if '│' in line]
            
            if len(table_lines) < 2:
                return None
            
            table_data = []
            for line in table_lines:
                # Skip separator rows
                if '─' in line and len([c for c in line if c == '─']) > 3:
                    continue
                
                cells = [cell.strip() for cell in line.split('│')]
                while cells and not cells[0]:
                    cells.pop(0)
                while cells and not cells[-1]:
                    cells.pop()
                
                if cells:
                    table_data.append(cells)
            
            # Normalize column count
            if table_data:
                max_cols = max(len(row) for row in table_data)
                for row in table_data:
                    while len(row) < max_cols:
                        row.append('')
            
            return table_data if table_data else None
        except Exception as e:
            self.log_message(f"Error parsing table: {e}")
            return None
        
    def add_markdown_table_to_document(self, section, table_data, insertion_index):
        """Enhanced markdown table creation with styling and alignment"""
        try:
            # Determine number of columns and rows
            num_cols = len(table_data['alignments'])
            num_rows = len(table_data['rows'])
            if table_data['has_headers']:
                num_rows += 1  # Add header row
            
            if num_rows == 0 or num_cols == 0:
                return None
            
            # Create table in document
            if insertion_index < len(self.document.paragraphs):
                para = self.document.paragraphs[insertion_index]
                table = self.document._body.add_tbl_before(para._p, num_rows, num_cols)
            else:
                table = self.document.add_table(rows=num_rows, cols=num_cols)
            
            # Apply enhanced table style
            try:
                table.style = 'Light Grid Accent 1'
            except:
                try:
                    table.style = 'Table Grid'
                except:
                    pass  # Use default styling
            
            current_row = 0
            
            # Add headers with enhanced formatting
            if table_data['has_headers']:
                for col_idx, header in enumerate(table_data['headers']):
                    if col_idx < num_cols:
                        cell = table.cell(current_row, col_idx)
                        cell.text = header
                        
                        # Format header cell
                        for paragraph in cell.paragraphs:
                            # Set alignment based on column alignment
                            alignment = table_data['alignments'][col_idx]
                            if alignment == 'center':
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            elif alignment == 'right':
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                            else:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            
                            # Make header bold
                            for run in paragraph.runs:
                                run.bold = True
                        
                        # Add header background shading (simplified approach)
                        try:
                            self._add_cell_shading(cell, "D9D9D9")  # Light gray
                        except:
                            pass  # Skip shading if it fails
                
                current_row += 1
            
            # Add data rows with formatting
            for row_data in table_data['rows']:
                for col_idx, cell_content in enumerate(row_data):
                    if col_idx < num_cols and current_row < num_rows:
                        cell = table.cell(current_row, col_idx)
                        
                        # Apply markdown formatting to cell content
                        self._format_table_cell_content(cell, cell_content, table_data['alignments'][col_idx])
                
                current_row += 1
            
            # Auto-fit table columns with preferred settings
            table.autofit = True
            
            return table
            
        except Exception as e:
            print(f"Error creating markdown table: {e}")
            return None

    def add_markdown_content_to_section(self, section, content, append=False):
        """Enhanced markdown content conversion with improved table and border handling"""
        try:
            heading_para = section.paragraph
            doc_paragraphs = list(self.document.paragraphs)
            
            # Find heading index
            heading_index = -1
            for i, para in enumerate(doc_paragraphs):
                if para.text.strip() == heading_para.text.strip():
                    heading_index = i
                    break
            
            if heading_index == -1:
                raise Exception("Could not find section heading")
            
            # Determine insertion point
            if append and section.content_paragraphs:
                next_heading_index = len(doc_paragraphs)
                for i in range(heading_index + 1, len(doc_paragraphs)):
                    if doc_paragraphs[i].style.name.startswith('Heading'):
                        next_heading_index = i
                        break
                insertion_index = next_heading_index
            else:
                insertion_index = heading_index + 1
            
            # Process content with enhanced parsing
            lines = content.split('\n')
            i = 0
            inserted_paras = []
            
            while i < len(lines):
                current_line = lines[i].strip()
                
                # Enhanced table detection and processing
                if current_line.startswith('|'):
                    # Collect all lines in the table
                    table_lines = []
                    while i < len(lines) and (lines[i].strip().startswith('|') or lines[i].strip() == ''):
                        if lines[i].strip():  # Skip empty lines within table
                            table_lines.append(lines[i])
                        i += 1
                    
                    # Parse and add table
                    table_data = self.parse_markdown_table(table_lines)
                    if table_data:
                        table = self.add_markdown_table_to_document(section, table_data, insertion_index)
                        if table:
                            insertion_index += 1
                            
                            # Add spacing after table
                            if insertion_index < len(self.document.paragraphs):
                                spacer_para = self.document.paragraphs[insertion_index].insert_paragraph_before()
                            else:
                                spacer_para = self.document.add_paragraph()
                            spacer_para.text = ""
                            inserted_paras.append(spacer_para)
                            insertion_index += 1
                    continue
                
                # Enhanced code block processing
                elif current_line.startswith('```'):
                    code_language = current_line[3:].strip()
                    code_lines = []
                    i += 1  # Skip the opening ```
                    
                    # Collect all lines in code block
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_lines.append(lines[i])
                        i += 1
                    
                    if i < len(lines):  # Skip closing ```
                        i += 1
                    
                    # Add enhanced code block
                    code_para = self._create_code_block(insertion_index, code_lines, code_language)
                    if code_para:
                        inserted_paras.append(code_para)
                        insertion_index += 1
                    continue
                
                # Enhanced horizontal rule processing with different styles
                elif re.match(r'^([-*_])\1{2,}$', current_line):
                    rule_char = current_line[0]
                    rule_para = self._create_horizontal_rule(insertion_index, rule_char)
                    if rule_para:
                        inserted_paras.append(rule_para)
                        insertion_index += 1
                    i += 1
                    continue
                
                # Process blockquotes with enhanced styling
                elif current_line.startswith('> '):
                    quote_lines = []
                    while i < len(lines) and (lines[i].strip().startswith('> ') or lines[i].strip() == '>'):
                        quote_lines.append(lines[i].strip()[1:].strip() if lines[i].strip() != '>' else '')
                        i += 1
                    
                    quote_para = self._create_blockquote(insertion_index, quote_lines)
                    if quote_para:
                        inserted_paras.append(quote_para)
                        insertion_index += 1
                    continue
                
                # Process definition lists (term : definition)
                elif ':' in current_line and not current_line.startswith('http'):
                    if self._is_definition_line(current_line):
                        def_para = self._create_definition_item(insertion_index, current_line)
                        if def_para:
                            inserted_paras.append(def_para)
                            insertion_index += 1
                        i += 1
                        continue
                
                # Regular paragraph processing with enhanced continuation logic
                para_text = current_line
                i += 1
                
                # For normal paragraphs, combine lines with better logic
                if not self._is_special_line(para_text):
                    while i < len(lines) and lines[i].strip() and not self._is_special_line(lines[i].strip()):
                        # Smart line joining
                        if not para_text.endswith((':', '-', '—', '–', '.', '!', '?')):
                            para_text += ' '
                        para_text += lines[i].strip()
                        i += 1
                
                # Skip if this is a duplicate of the section heading
                if para_text.lower() == section.text.lower() or (
                    para_text.startswith('#') and para_text.lstrip('#').strip().lower() == section.text.lower()
                ):
                    if hasattr(self, 'log_message'):
                        self.log_message(f"Skipping duplicate heading: {para_text}")
                    continue
                
                # Create and format paragraph
                if para_text.strip():
                    if insertion_index < len(self.document.paragraphs):
                        new_para = self.document.paragraphs[insertion_index].insert_paragraph_before()
                    else:
                        new_para = self.document.add_paragraph()
                    
                    # Apply enhanced markdown formatting
                    self.apply_markdown_to_paragraph(new_para, para_text)
                    
                    # Apply configured formatting if available
                    if hasattr(self, 'apply_configured_formatting'):
                        self.apply_configured_formatting(new_para)
                    
                    inserted_paras.append(new_para)
                    insertion_index += 1
            
            # Update section's content paragraphs
            if not append:
                section.content_paragraphs = inserted_paras
            else:
                section.content_paragraphs.extend(inserted_paras)
                
        except Exception as e:
            print(f"Error in markdown content processing: {e}")
            raise

    def apply_markdown_to_paragraph(self, paragraph, text):
        """Enhanced markdown to paragraph conversion with additional features"""
        try:
            # Handle heading levels 1-6
            if text.startswith('#'):
                heading_level = 0
                for char in text:
                    if char == '#':
                        heading_level += 1
                    else:
                        break
                
                if 1 <= heading_level <= 6:
                    # Map markdown heading level to Word heading style
                    if heading_level <= 4:
                        try:
                            paragraph.style = f'Heading {heading_level}'
                        except:
                            # If style not found, fallback to normal with bold
                            paragraph.style = 'Normal'
                            # Will be formatted below
                    else:
                        # Fall back to normal text with bold for higher levels
                        paragraph.style = 'Normal'
                        
                    text = text[heading_level:].strip()
            
            # Handle bullet points (*, -, +)
            elif text.strip().startswith(('* ', '- ', '+ ')):
                try:
                    paragraph.style = 'List Bullet'
                except:
                    # Create bullet point manually if style not available
                    paragraph.style = 'Normal'
                    try:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.left_indent = Inches(0.25)
                    except:
                        pass
                    text = "• " + text.strip()[2:].strip()  # Add bullet character
                else:
                    text = text.strip()[2:].strip()  # Remove the bullet marker
            
            # Handle numbered lists
            elif re.match(r'^\d+\.', text.strip()):
                try:
                    paragraph.style = 'List Number'
                except:
                    # Try alternative style names
                    try:
                        paragraph.style = 'List Paragraph'
                    except:
                        # Create numbered format manually
                        paragraph.style = 'Normal'
                        try:
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.left_indent = Inches(0.25)
                        except:
                            pass
                        # Keep the number in text
                
                # Extract number for manual handling if needed
                num_match = re.match(r'^(\d+)\.', text.strip())
                if num_match:
                    num = num_match.group(1)
                    text = text.strip()[len(num)+1:].strip()  # Remove the number and dot
                    
                    # For manual numbering if style failed
                    try:
                        if paragraph.style.name == 'Normal':
                            text = f"{num}. {text}"
                    except:
                        text = f"{num}. {text}"
            
            # Handle blockquotes (handled separately in main function now)
            elif text.strip().startswith('> '):
                text = text.strip()[2:].strip()  # Remove the quote marker
            
            # Clear paragraph text since we'll re-add it with formatting
            paragraph.clear()
            
            # Enhanced inline formatting processing
            self._process_inline_formatting(paragraph, text)
            
        except Exception as e:
            print(f"Error applying markdown to paragraph: {e}")
            # Fallback to plain text
            paragraph.clear()
            paragraph.add_run(text)

    def render_markdown_preview(self, text):
        """Enhanced markdown preview with better table and formatting support"""
        try:
            if not hasattr(self, 'generated_text'):
                return
                
            self.generated_text.delete('1.0', tk.END)
            
            lines = text.split('\n')
            i = 0
            
            while i < len(lines):
                line = lines[i]
                
                # Handle tables in preview
                if line.strip().startswith('|'):
                    # Collect table lines
                    table_lines = []
                    while i < len(lines) and (lines[i].strip().startswith('|') or lines[i].strip() == ''):
                        if lines[i].strip():
                            table_lines.append(lines[i])
                        i += 1
                    
                    # Render table preview
                    self._render_table_preview(table_lines)
                    continue
                
                # Handle code blocks
                elif line.strip().startswith('```'):
                    language = line.strip()[3:]
                    if language:
                        self.generated_text.insert(tk.END, f"[{language.upper()}]\n", 'code_lang')
                    
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        self.generated_text.insert(tk.END, lines[i] + '\n', 'code')
                        i += 1
                    i += 1  # Skip closing ```
                    continue
                
                # Handle horizontal rules
                elif re.match(r'^([-*_])\1{2,}$', line.strip()):
                    self.generated_text.insert(tk.END, '─' * 50 + '\n', 'hr')
                    i += 1
                    continue
                
                # Handle blockquotes
                elif line.strip().startswith('> '):
                    quote_text = line.strip()[2:]
                    self.generated_text.insert(tk.END, f'  "{quote_text}"\n', 'quote')
                    i += 1
                    continue
                
                # Regular line processing with enhanced formatting
                self._render_line_with_formatting(line)
                self.generated_text.insert(tk.END, '\n')
                i += 1
                
        except Exception as e:
            print(f"Error rendering markdown preview: {e}")
            # Fallback to basic rendering
            try:
                self.generated_text.delete('1.0', tk.END)
                self.generated_text.insert('1.0', text)
            except:
                pass

    def _add_cell_shading(self, cell, color_hex):
        """Add background shading to a table cell - simplified version"""
        try:
            # Simplified approach using XML string formatting
            shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
            shading_elm = parse_xml(shading_xml)
            cell._tc.get_or_add_tcPr().append(shading_elm)
        except Exception as e:
            print(f"Error adding cell shading: {e}")
            # Fallback: try to set through properties if available
            try:
                from docx.oxml.shared import qn
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), color_hex)
                tcPr.append(shd)
            except:
                pass  # Skip shading entirely if all methods fail

    def _format_table_cell_content(self, cell, content, alignment):
        """Format table cell content with markdown support"""
        try:
            # Clear existing content
            cell.text = ""
            
            # Add content with markdown formatting
            paragraph = cell.paragraphs[0]
            self.apply_markdown_to_paragraph(paragraph, content)
            
            # Set alignment
            if alignment == 'center':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment == 'right':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
        except Exception as e:
            print(f"Error formatting table cell: {e}")
            # Fallback to plain text
            cell.text = content

    def _is_special_line(self, line):
        """Check if a line has special markdown formatting"""
        if not line:
            return False
        
        return (line.startswith('#') or 
                line.startswith('* ') or 
                line.startswith('- ') or 
                line.startswith('+ ') or
                line.startswith('> ') or
                line.startswith('|') or 
                line.startswith('```') or
                re.match(r'^\d+\.', line) or
                re.match(r'^([-*_])\1{2,}$', line) or
                ':' in line and not line.startswith('http'))

    def _is_definition_line(self, line):
        """Check if line is a definition list item"""
        try:
            parts = line.split(':', 1)
            return len(parts) == 2 and len(parts[0].strip()) > 0 and len(parts[1].strip()) > 0
        except:
            return False

    def _create_code_block(self, insertion_index, code_lines, language=""):
        """Create enhanced code block with syntax highlighting indication"""
        try:
            if insertion_index < len(self.document.paragraphs):
                new_para = self.document.paragraphs[insertion_index].insert_paragraph_before()
            else:
                new_para = self.document.add_paragraph()
            
            # Add language label if provided
            if language:
                lang_run = new_para.add_run(f"[{language.upper()}]\n")
                lang_run.font.size = Pt(8)
                lang_run.font.color.rgb = RGBColor(128, 128, 128)
                lang_run.italic = True
            
            # Format as code
            code_text = '\n'.join(code_lines)
            code_run = new_para.add_run(code_text)
            code_run.font.name = 'Consolas'
            code_run.font.color.rgb = RGBColor(0, 128, 0)
            code_run.font.size = Pt(9)
            
            # Add code block styling
            try:
                new_para.paragraph_format.left_indent = Inches(0.25)
                new_para.paragraph_format.right_indent = Inches(0.25)
            except:
                pass
            
            # Add border around code block (simplified)
            try:
                self._add_paragraph_border(new_para, "E0E0E0")
            except:
                pass
            
            return new_para
            
        except Exception as e:
            print(f"Error creating code block: {e}")
            return None

    def _create_horizontal_rule(self, insertion_index, rule_char):
        """Create enhanced horizontal rule with different styles"""
        try:
            if insertion_index < len(self.document.paragraphs):
                new_para = self.document.paragraphs[insertion_index].insert_paragraph_before()
            else:
                new_para = self.document.add_paragraph()
            
            # Different styles based on character used
            if rule_char == '*':
                # Thick double border
                border_width = 2
                border_color = RGBColor(0, 0, 0)
            elif rule_char == '_':
                # Thin gray border
                border_width = 1
                border_color = RGBColor(128, 128, 128)
            else:  # '-'
                # Medium black border
                border_width = 1
                border_color = RGBColor(0, 0, 0)
            
            # Apply bottom border
            try:
                new_para.paragraph_format.bottom_border.width = border_width
                new_para.paragraph_format.bottom_border.color.rgb = border_color
                
                # Add some spacing
                new_para.paragraph_format.space_after = Pt(6)
                new_para.paragraph_format.space_before = Pt(6)
            except Exception as e:
                print(f"Error setting border properties: {e}")
            
            return new_para
            
        except Exception as e:
            print(f"Error creating horizontal rule: {e}")
            return None

    def _create_blockquote(self, insertion_index, quote_lines):
        """Create enhanced blockquote with styling"""
        try:
            if insertion_index < len(self.document.paragraphs):
                new_para = self.document.paragraphs[insertion_index].insert_paragraph_before()
            else:
                new_para = self.document.add_paragraph()
            
            # Combine quote lines
            quote_text = '\n'.join(quote_lines)
            
            # Apply quote formatting
            quote_run = new_para.add_run(quote_text)
            quote_run.italic = True
            quote_run.font.color.rgb = RGBColor(96, 96, 96)
            
            # Set paragraph formatting
            try:
                new_para.paragraph_format.left_indent = Inches(0.5)
                new_para.paragraph_format.right_indent = Inches(0.5)
            except:
                pass
            
            # Add left border for quote styling (simplified)
            try:
                self._add_paragraph_left_border(new_para, "CCCCCC")
            except:
                pass
            
            return new_para
            
        except Exception as e:
            print(f"Error creating blockquote: {e}")
            return None

    def _create_definition_item(self, insertion_index, def_line):
        """Create definition list item"""
        try:
            if insertion_index < len(self.document.paragraphs):
                new_para = self.document.paragraphs[insertion_index].insert_paragraph_before()
            else:
                new_para = self.document.add_paragraph()
            
            # Split term and definition
            term, definition = def_line.split(':', 1)
            term = term.strip()
            definition = definition.strip()
            
            # Add term in bold
            term_run = new_para.add_run(term)
            term_run.bold = True
            
            # Add separator
            new_para.add_run(': ')
            
            # Add definition
            def_run = new_para.add_run(definition)
            
            # Set indentation
            try:
                new_para.paragraph_format.left_indent = Inches(0.25)
                new_para.paragraph_format.hanging_indent = Inches(-0.25)
            except:
                pass
            
            return new_para
            
        except Exception as e:
            print(f"Error creating definition item: {e}")
            return None

    def _add_paragraph_border(self, paragraph, color_hex):
        """Add border around paragraph - simplified version"""
        try:
            # Convert hex to RGB
            r = int(color_hex[0:2], 16)
            g = int(color_hex[2:4], 16) 
            b = int(color_hex[4:6], 16)
            border_color = RGBColor(r, g, b)
            
            # Add borders (may not work in all docx versions)
            paragraph.paragraph_format.top_border.width = 1
            paragraph.paragraph_format.bottom_border.width = 1
            paragraph.paragraph_format.left_border.width = 1
            paragraph.paragraph_format.right_border.width = 1
            
            paragraph.paragraph_format.top_border.color.rgb = border_color
            paragraph.paragraph_format.bottom_border.color.rgb = border_color
            paragraph.paragraph_format.left_border.color.rgb = border_color
            paragraph.paragraph_format.right_border.color.rgb = border_color
            
        except Exception as e:
            print(f"Error adding paragraph border: {e}")

    def _add_paragraph_left_border(self, paragraph, color_hex):
        """Add left border to paragraph (for quotes) - simplified version"""
        try:
            # Convert hex to RGB
            r = int(color_hex[0:2], 16)
            g = int(color_hex[2:4], 16) 
            b = int(color_hex[4:6], 16)
            border_color = RGBColor(r, g, b)
            
            paragraph.paragraph_format.left_border.width = 3
            paragraph.paragraph_format.left_border.color.rgb = border_color
            
        except Exception as e:
            print(f"Error adding left border: {e}")

    def _process_inline_formatting(self, paragraph, text):
        """Enhanced inline formatting with more markdown features"""
        try:
            i = 0
            in_bold = False
            in_italic = False
            current_text = ""
            
            while i < len(text):
                # Process bold (both ** and __ formats)
                if i+1 < len(text) and (text[i:i+2] == '**' or text[i:i+2] == '__'):
                    # Add accumulated text with current formatting
                    if current_text:
                        run = paragraph.add_run(current_text)
                        if in_bold:
                            run.bold = True
                        if in_italic:
                            run.italic = True
                        current_text = ""
                    
                    # Toggle bold state
                    in_bold = not in_bold
                    i += 2
                    continue
                
                # Process italic (both * and _ formats, but not if part of bold)
                elif (text[i] == '*' or text[i] == '_'):
                    # Check it's not part of bold marker
                    if not (i+1 < len(text) and (text[i:i+2] == '**' or text[i:i+2] == '__')):
                        # Add accumulated text with current formatting
                        if current_text:
                            run = paragraph.add_run(current_text)
                            if in_bold:
                                run.bold = True
                            if in_italic:
                                run.italic = True
                            current_text = ""
                        
                        # Toggle italic state
                        in_italic = not in_italic
                        i += 1
                        continue
                
                # Process inline code
                elif text[i] == '`':
                    # First add accumulated text
                    if current_text:
                        run = paragraph.add_run(current_text)
                        if in_bold:
                            run.bold = True
                        if in_italic:
                            run.italic = True
                        current_text = ""
                    
                    # Find the closing backtick
                    end = text.find('`', i+1)
                    if end != -1:
                        # Add the code text
                        code_text = text[i+1:end]
                        run = paragraph.add_run(code_text)
                        run.font.name = 'Consolas'
                        run.font.color.rgb = RGBColor(0, 128, 0)
                        run.font.size = Pt(10)
                        i = end + 1
                    else:
                        # No closing backtick found, treat as regular character
                        current_text += text[i]
                        i += 1
                    continue
                
                # Process strikethrough (~~text~~)
                elif i+1 < len(text) and text[i:i+2] == '~~':
                    # Add accumulated text
                    if current_text:
                        run = paragraph.add_run(current_text)
                        if in_bold:
                            run.bold = True
                        if in_italic:
                            run.italic = True
                        current_text = ""
                    
                    # Find closing ~~
                    end = text.find('~~', i+2)
                    if end != -1:
                        # Add strikethrough text
                        strike_text = text[i+2:end]
                        run = paragraph.add_run(strike_text)
                        try:
                            run.font.strike = True
                        except:
                            pass  # Strikethrough may not be supported in all versions
                        if in_bold:
                            run.bold = True
                        if in_italic:
                            run.italic = True
                        i = end + 2
                    else:
                        # No closing strike found, treat as regular
                        current_text += text[i:i+2]
                        i += 2
                    continue
                
                # Process superscript (^text^)
                elif text[i] == '^' and i+1 < len(text):
                    end = text.find('^', i+1)
                    if end != -1:
                        # Add accumulated text
                        if current_text:
                            run = paragraph.add_run(current_text)
                            if in_bold:
                                run.bold = True
                            if in_italic:
                                run.italic = True
                            current_text = ""
                        
                        # Add superscript text
                        super_text = text[i+1:end]
                        run = paragraph.add_run(super_text)
                        try:
                            run.font.superscript = True
                            run.font.size = Pt(8)
                        except:
                            pass  # Superscript may not be supported
                        if in_bold:
                            run.bold = True
                        if in_italic:
                            run.italic = True
                        i = end + 1
                        continue
                
                # Process hyperlinks [text](url)
                elif text[i] == '[':
                    link_text_end = text.find(']', i)
                    if link_text_end != -1 and link_text_end + 1 < len(text) and text[link_text_end + 1] == '(':
                        url_end = text.find(')', link_text_end)
                        if url_end != -1:
                            # Add accumulated text first
                            if current_text:
                                run = paragraph.add_run(current_text)
                                if in_bold:
                                    run.bold = True
                                if in_italic:
                                    run.italic = True
                                current_text = ""
                            
                            # Extract link text and URL
                            link_text = text[i+1:link_text_end]
                            url = text[link_text_end+2:url_end]
                            
                            # Add hyperlink (visual styling)
                            run = paragraph.add_run(link_text)
                            run.font.underline = True
                            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
                            if in_bold:
                                run.bold = True
                            if in_italic:
                                run.italic = True
                            
                            i = url_end + 1
                            continue
                
                # Process highlight ==text==
                elif i+1 < len(text) and text[i:i+2] == '==':
                    # Add accumulated text
                    if current_text:
                        run = paragraph.add_run(current_text)
                        if in_bold:
                            run.bold = True
                        if in_italic:
                            run.italic = True
                        current_text = ""
                    
                    # Find closing ==
                    end = text.find('==', i+2)
                    if end != -1:
                        # Add highlighted text
                        highlight_text = text[i+2:end]
                        run = paragraph.add_run(highlight_text)
                        try:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        except:
                            pass  # Highlighting may not be supported
                        if in_bold:
                            run.bold = True
                        if in_italic:
                            run.italic = True
                        i = end + 2
                    else:
                        # No closing highlight found, treat as regular
                        current_text += text[i:i+2]
                        i += 2
                    continue
                
                # Add regular character to current text
                current_text += text[i]
                i += 1
            
            # Add any remaining text
            if current_text:
                run = paragraph.add_run(current_text)
                if in_bold:
                    run.bold = True
                if in_italic:
                    run.italic = True
                    
        except Exception as e:
            print(f"Error processing inline formatting: {e}")
            # Fallback to plain text
            paragraph.add_run(text)

    def _render_table_preview(self, table_lines):
        """Render table in preview with ASCII formatting"""
        try:
            table_data = self.parse_markdown_table(table_lines)
            if not table_data:
                return
            
            # Calculate column widths
            all_rows = [table_data['headers']] + table_data['rows'] if table_data['has_headers'] else table_data['rows']
            col_widths = []
            
            for col_idx in range(len(table_data['alignments'])):
                max_width = 0
                for row in all_rows:
                    if col_idx < len(row):
                        max_width = max(max_width, len(str(row[col_idx])))
                col_widths.append(max(max_width, 8))  # Minimum width of 8
            
            # Render header
            if table_data['has_headers']:
                header_line = "│"
                for col_idx, header in enumerate(table_data['headers']):
                    if col_idx < len(col_widths):
                        padded_header = str(header).ljust(col_widths[col_idx])
                        header_line += f" {padded_header} │"
                
                self.generated_text.insert(tk.END, header_line + '\n', 'table_header')
                
                # Render separator
                sep_line = "│"
                for col_idx, width in enumerate(col_widths):
                    alignment = table_data['alignments'][col_idx] if col_idx < len(table_data['alignments']) else 'left'
                    if alignment == 'center':
                        sep = ':' + '─' * (width-2) + ':'
                    elif alignment == 'right':
                        sep = '─' * (width-1) + ':'
                    else:
                        sep = '─' * width
                    sep_line += f" {sep} │"
                
                self.generated_text.insert(tk.END, sep_line + '\n', 'table_sep')
            
            # Render data rows
            for row in table_data['rows']:
                row_line = "│"
                for col_idx, cell in enumerate(row):
                    if col_idx < len(col_widths):
                        padded_cell = str(cell).ljust(col_widths[col_idx])
                        row_line += f" {padded_cell} │"
                
                self.generated_text.insert(tk.END, row_line + '\n', 'table_row')
            
            self.generated_text.insert(tk.END, '\n')
            
        except Exception as e:
            print(f"Error rendering table preview: {e}")

    def _render_line_with_formatting(self, line):
        """Render line with enhanced markdown formatting in preview"""
        try:
            segments = []
            current_line = line
            
            # Process line for various markdown elements
            i = 0
            while i < len(current_line):
                # Check for ==highlight==
                if current_line[i:i+2] == '==':
                    end = current_line.find('==', i+2)
                    if end != -1:
                        if i > 0:
                            segments.append(('normal', current_line[:i]))
                        segments.append(('highlight', current_line[i+2:end]))
                        current_line = current_line[end+2:]
                        i = 0
                        continue
                
                # Check for **bold**
                elif current_line[i:i+2] == '**':
                    end = current_line.find('**', i+2)
                    if end != -1:
                        if i > 0:
                            segments.append(('normal', current_line[:i]))
                        segments.append(('bold', current_line[i+2:end]))
                        current_line = current_line[end+2:]
                        i = 0
                        continue
                
                # Check for *italic*
                elif current_line[i] == '*':
                    end = current_line.find('*', i+1)
                    if end != -1:
                        if i > 0:
                            segments.append(('normal', current_line[:i]))
                        segments.append(('italic', current_line[i+1:end]))
                        current_line = current_line[end+1:]
                        i = 0
                        continue
                
                # Check for `code`
                elif current_line[i] == '`':
                    end = current_line.find('`', i+1)
                    if end != -1:
                        if i > 0:
                            segments.append(('normal', current_line[:i]))
                        segments.append(('code', current_line[i+1:end]))
                        current_line = current_line[end+1:]
                        i = 0
                        continue
                
                # Check for ~~strikethrough~~
                elif current_line[i:i+2] == '~~':
                    end = current_line.find('~~', i+2)
                    if end != -1:
                        if i > 0:
                            segments.append(('normal', current_line[:i]))
                        segments.append(('strike', current_line[i+2:end]))
                        current_line = current_line[end+2:]
                        i = 0
                        continue
                
                i += 1
            
            # Add remaining text
            if current_line:
                segments.append(('normal', current_line))
            
            # Insert segments with tags
            if segments:
                for tag, text_seg in segments:
                    if tag == 'normal':
                        self.generated_text.insert(tk.END, text_seg)
                    else:
                        start_idx = self.generated_text.index(tk.END)
                        self.generated_text.insert(tk.END, text_seg)
                        end_idx = self.generated_text.index(tk.END)
                        self.generated_text.tag_add(tag, start_idx, end_idx)
            
        except Exception as e:
            print(f"Error rendering line formatting: {e}")
            # Fallback to plain text
            self.generated_text.insert(tk.END, line)

    def validate_markdown_syntax(content):
        """Validate markdown syntax and return warnings"""
        warnings = []
        lines = content.split('\n')
        
        try:
            # Check for unmatched formatting
            bold_count = content.count('**')
            if bold_count % 2 != 0:
                warnings.append("Unmatched bold formatting (**)")
            
            italic_count = content.count('*') - bold_count
            if italic_count % 2 != 0:
                warnings.append("Unmatched italic formatting (*)")
            
            code_count = content.count('`')
            if code_count % 2 != 0:
                warnings.append("Unmatched code formatting (`)")
            
            strike_count = content.count('~~')
            if strike_count % 2 != 0:
                warnings.append("Unmatched strikethrough formatting (~~)")
            
            # Check table formatting
            for i, line in enumerate(lines):
                if line.strip().startswith('|'):
                    pipes = line.count('|')
                    if pipes < 3:  # At least | content |
                        warnings.append(f"Line {i+1}: Incomplete table row (need at least 3 pipe characters)")
            
            return warnings
            
        except Exception as e:
            return [f"Error validating syntax: {str(e)}"]

    def optimize_markdown_for_docx(content):
        """Optimize markdown content for better Word document conversion"""
        try:
            lines = content.split('\n')
            optimized_lines = []
            
            for line in lines:
                # Fix common issues
                line = line.strip()
                
                # Ensure proper spacing around headers
                if line.startswith('#'):
                    if optimized_lines and optimized_lines[-1].strip():
                        optimized_lines.append('')  # Add blank line before header
                
                # Fix table formatting
                elif line.startswith('|'):
                    # Ensure consistent spacing in table cells
                    parts = line.split('|')
                    cleaned_parts = [part.strip() for part in parts]
                    line = '| ' + ' | '.join(cleaned_parts[1:-1]) + ' |' if len(cleaned_parts) > 2 else line
                
                # Fix list formatting
                elif line.startswith(('* ', '- ', '+ ')):
                    # Ensure single space after list marker
                    marker = line[:2]
                    content = line[2:].strip()
                    line = f"{marker}{content}"
                
                optimized_lines.append(line)
            
            return '\n'.join(optimized_lines)
            
        except Exception as e:
            print(f"Error optimizing markdown: {e}")
            return content

    def apply_configured_formatting(self, paragraph):
        """Apply configured formatting options to paragraph"""
        for run in paragraph.runs:
            # Highlight
            if self.format_config['highlight_enabled'].get():
                color_name = self.format_config['highlight_color'].get()
                color_map = {
                    'YELLOW': WD_COLOR_INDEX.YELLOW,
                    'GREEN': WD_COLOR_INDEX.GREEN,
                    'CYAN': WD_COLOR_INDEX.TURQUOISE,
                    'PINK': WD_COLOR_INDEX.PINK,
                    'BRIGHT_GREEN': WD_COLOR_INDEX.BRIGHT_GREEN
                }
                run.font.highlight_color = color_map.get(color_name, WD_COLOR_INDEX.YELLOW)
            
            # Bold (if not already from markdown)
            if self.format_config['bold_enabled'].get() and not run.bold:
                run.bold = True
            
            # Italic (if not already from markdown)
            if self.format_config['italic_enabled'].get() and not run.italic:
                run.italic = True
            
            # Underline
            if self.format_config['underline_enabled'].get():
                run.underline = True
            
            # Font size
            run.font.size = Pt(self.format_config['font_size'].get())
            
            # Font color (optional, if different from default)
            if self.format_config['font_color'].get() != '000000':
                color_hex = self.format_config['font_color'].get()
                try:
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except:
                    pass
    
    def remove_paragraph(self, paragraph):
        """Remove a paragraph from the document"""
        p = paragraph._element
        p.getparent().remove(p)
    
    def save_document_auto(self):
        """Auto-save document with permission error handling"""
        if not self.document or not self.document_path:
            return
        
        try:
            self.document.save(self.document_path)
            self.log_message(f"Document auto-saved: {self.document_path}")
        except PermissionError:
            self.log_message("Auto-save failed: File is open elsewhere")
            
            result = messagebox.askyesno(
                "Auto-Save Failed",
                "Cannot auto-save: File is open in another application.\n\n"
                "Close the file and try again, or save as a different file?"
            )
            
            if result:
                # Try regular save with retry
                self.save_with_retry()
        except Exception as e:
            self.log_message(f"Auto-save failed: {str(e)}")
            messagebox.showerror("Auto-Save Error", f"Failed to auto-save: {str(e)}")
    
    def prompt_save_document(self):
        """Prompt to save document with retry on permission error"""
        if not self.document or not self.document_path:
            return
            
        result = messagebox.askyesnocancel(
            "Save Document",
            "Save changes?\n\n"
            "Yes: Save to current file\n"
            "No: Save As (new file)\n"
            "Cancel: Don't save"
        )
        
        if result is True:
            self.save_with_retry()
                
        elif result is False:
            self.save_as_new_file()
    
    def save_with_retry(self):
        """Save document with retry logic for permission errors"""
        max_retries = 3
        retry_count = 0
        
        while retry_count < max_retries:
            try:
                self.document.save(self.document_path)
                self.log_message(f"Document saved: {self.document_path}")
                messagebox.showinfo("Success", "Document saved successfully")
                return True
                
            except PermissionError:
                retry_count += 1
                self.log_message(f"Save failed (attempt {retry_count}): File is open elsewhere")
                
                if retry_count < max_retries:
                    # Prompt user to close and retry
                    result = messagebox.askyesnocancel(
                        "File In Use",
                        f"Cannot save: File is open in another application.\n\n"
                        f"Please close the file and click:\n"
                        f"• Yes - Retry saving to current file\n"
                        f"• No - Save As a different file\n"
                        f"• Cancel - Don't save\n\n"
                        f"Attempt {retry_count} of {max_retries}"
                    )
                    
                    if result is True:
                        # Retry
                        continue
                    elif result is False:
                        # Save as
                        return self.save_as_new_file()
                    else:
                        # Cancel
                        self.log_message("Save cancelled by user")
                        return False
                else:
                    # Max retries reached
                    result = messagebox.askyesno(
                        "Save Failed",
                        "Could not save after multiple attempts.\n\n"
                        "Save as a different file?"
                    )
                    
                    if result:
                        return self.save_as_new_file()
                    else:
                        self.log_message("Save cancelled after max retries")
                        return False
                        
            except Exception as e:
                self.log_message(f"Save error: {str(e)}")
                messagebox.showerror("Error", f"Failed to save: {str(e)}")
                return False
        
        return False
    
    def save_as_new_file(self):
        """Save document as a new file"""
        new_path = filedialog.asksaveasfilename(
            title="Save Document As",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            initialfile=os.path.basename(self.document_path)
        )
        
        if new_path:
            try:
                self.document.save(new_path)
                self.document_path = new_path
                self.doc_label_var.set(os.path.basename(new_path))
                self.save_settings()
                
                # Update tracking file location
                self.load_section_tracking()
                
                self.log_message(f"Document saved as: {new_path}")
                messagebox.showinfo("Success", "Document saved successfully")
                return True
            except Exception as e:
                self.log_message(f"Save As error: {str(e)}")
                messagebox.showerror("Error", f"Failed to save: {str(e)}")
                return False
        
        return False
    
    def reload_document(self):
        """Reload document to refresh content"""
        if not self.document_path:
            return
        
        try:
            self.log_message("Reloading document...")
            current_selection = self.selected_section
            
            # Reload document
            self.document = Document(self.document_path)
            
            # Re-parse structure
            self.parse_document_structure()
            
            # Refresh tree
            self.populate_tree()
            
            # Try to re-select the same section
            if current_selection:
                reselected = self.find_section_by_path(current_selection.get_full_path())
                if reselected:
                    self.selected_section = reselected
                    self.show_existing_content()
            
            self.log_message("Document reloaded successfully")
            
        except Exception as e:
            self.log_message(f"Error reloading document: {str(e)}")
    
    def find_section_by_path(self, path):
        """Find section by full path"""
        def search(sections):
            for section in sections:
                if section.get_full_path() == path:
                    return section
                result = search(section.children)
                if result:
                    return result
            return None
        
        return search(self.sections)
            
    def run(self):
        """Run the application"""
        self.log_message("Document Content Generator started")
        self.log_message("Load a Word document to begin")
        if self.selected_model.get():
            self.log_message(f"Configured model: {self.selected_model.get()}")
        self.root.mainloop()

    def analyze_document_tenses(self):
        """NEW FEATURE: Dedicated tense consistency analysis"""
        if not self.selected_section:
            messagebox.showwarning("No Selection", "Please select a section to analyze")
            return
        
        if not self.selected_section.has_content():
            messagebox.showinfo("No Content", "Selected section has no content to analyze")
            return
        
        if not self.advanced_reviewer:
            messagebox.showinfo("Feature Unavailable", 
                            "Advanced tense analysis requires document_reviewer module.\n\n"
                            "To enable this feature:\n"
                            "1. Ensure document_reviewer.py is in the same directory\n"
                            "2. Install dependencies:\n"
                            "   pip install textstat nltk --break-system-packages\n"
                            "3. Run NLTK setup:\n"
                            "   python -c \"import nltk; nltk.download('punkt'); "
                            "nltk.download('averaged_perceptron_tagger'); "
                            "nltk.download('stopwords')\"\n"
                            "4. Restart the application")
            return
        
        content = self.selected_section.get_existing_content()
        
        try:
            self.log_message("Starting tense consistency analysis...")
            tense_analysis = self.advanced_reviewer.analyze_tense_consistency(content)
            
            # Display results in generated_text area
            self.generated_text.delete(1.0, tk.END)
            self.generated_text.insert(tk.END, "=== TENSE CONSISTENCY ANALYSIS ===\n\n")
            self.generated_text.insert(tk.END, f"Section: {self.selected_section.get_full_path()}\n\n")
            self.generated_text.insert(tk.END, f"Dominant Tense: {tense_analysis.dominant_tense.upper()}\n")
            self.generated_text.insert(tk.END, f"Consistency Score: {tense_analysis.consistency_score:.1f}/10\n\n")
            self.generated_text.insert(tk.END, f"Tense Distribution:\n")
            self.generated_text.insert(tk.END, f"  • Past: {tense_analysis.past_count} sentences\n")
            self.generated_text.insert(tk.END, f"  • Present: {tense_analysis.present_count} sentences\n")
            self.generated_text.insert(tk.END, f"  • Future: {tense_analysis.future_count} sentences\n\n")
            
            if tense_analysis.inconsistent_sentences:
                self.generated_text.insert(tk.END, 
                    f"Inconsistent Sentences ({len(tense_analysis.inconsistent_sentences)}):\n\n")
                for i, sentence in enumerate(tense_analysis.inconsistent_sentences[:10], 1):
                    self.generated_text.insert(tk.END, f"{i}. {sentence}\n\n")
                
                if len(tense_analysis.inconsistent_sentences) > 10:
                    self.generated_text.insert(tk.END, 
                        f"... and {len(tense_analysis.inconsistent_sentences) - 10} more\n")
            else:
                self.generated_text.insert(tk.END, "✓ No tense inconsistencies detected\n")
            
            self.log_message(f"Tense analysis completed: {tense_analysis.consistency_score:.1f}/10")
            self.notebook.select(0)  # Switch to preview tab
            
        except Exception as e:
            messagebox.showerror("Analysis Error", f"Failed to analyze tenses: {str(e)}")
            self.log_message(f"Tense analysis error: {e}")

    def show_processing_strategy_dialog(self):
        """NEW FEATURE: Display content processing strategy analysis"""
        if not self.content_processor:
            messagebox.showinfo("Feature Unavailable",
                            "Intelligent processing requires content_processor module.\n\n"
                            "To enable this feature:\n"
                            "1. Ensure content_processor.py is in the same directory\n"
                            "2. Install dependencies:\n"
                            "   pip install tiktoken scikit-learn numpy --break-system-packages\n"
                            "3. Restart the application")
            return
        
        if not self.selected_section or not self.selected_section.has_content():
            messagebox.showwarning("No Content", "Please select a section with content")
            return
        
        content = self.selected_section.get_existing_content()
        
        try:
            strategy = self.content_processor.determine_processing_strategy(
                content, [self.selected_section], "Generate comprehensive content"
            )
            
            # Create dialog
            dialog = tk.Toplevel(self.root)
            dialog.title("Processing Strategy Analysis")
            dialog.geometry("600x500")
            dialog.configure(bg="#2b2b2b")
            dialog.grab_set()
            
            main_frame = ttk.Frame(dialog, padding="20")
            main_frame.pack(fill=tk.BOTH, expand=True)
            
            ttk.Label(main_frame, text="Content Processing Strategy", 
                    font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 15))
            
            # Strategy info
            info_frame = ttk.LabelFrame(main_frame, text="Recommended Strategy", padding="10")
            info_frame.pack(fill=tk.X, pady=(0, 10))
            
            ttk.Label(info_frame, text=f"Method: {strategy.method.upper()}",
                    font=("Arial", 10, "bold")).pack(anchor=tk.W)
            ttk.Label(info_frame, text=f"Reason: {strategy.reason}").pack(anchor=tk.W, pady=5)
            ttk.Label(info_frame, text=f"Token Estimate: {strategy.token_estimate}").pack(anchor=tk.W)
            ttk.Label(info_frame, text=f"Confidence: {strategy.confidence:.1%}").pack(anchor=tk.W)
            
            # Metrics
            metrics_frame = ttk.LabelFrame(main_frame, text="Content Metrics", padding="10")
            metrics_frame.pack(fill=tk.X, pady=(0, 10))
            
            metrics = self.content_processor.analyze_content_metrics(content, [self.selected_section])
            ttk.Label(metrics_frame, text=f"Character Count: {metrics.char_count:,}").pack(anchor=tk.W)
            ttk.Label(metrics_frame, text=f"Token Count: ~{metrics.token_count:,}").pack(anchor=tk.W)
            ttk.Label(metrics_frame, text=f"Complexity Score: {metrics.complexity_score:.2f}").pack(anchor=tk.W)
            ttk.Label(metrics_frame, text=f"Technical Density: {metrics.technical_density:.2%}").pack(anchor=tk.W)
            
            # Explanation
            explain_frame = ttk.LabelFrame(main_frame, text="What This Means", padding="10")
            explain_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
            
            explain_text = tk.Text(explain_frame, wrap=tk.WORD, height=8, width=60,
                                bg="#1e1e1e", fg="#ffffff", font=("Consolas", 9))
            explain_text.pack(fill=tk.BOTH, expand=True)
            
            explanation = ""
            if strategy.method == "full_prompt":
                explanation = "FULL PROMPT: Content is small enough to include entirely in the prompt. This provides complete context for the AI."
            elif strategy.method == "rag":
                explanation = "RAG (Retrieval Augmented Generation): Content is large and complex. The system will:\n1. Break content into chunks\n2. Store in vector database\n3. Retrieve only relevant chunks for each query\n4. Provide focused context to AI"
            elif strategy.method == "hybrid":
                explanation = "HYBRID: Uses overview + targeted retrieval. Provides document structure with detailed context for specific queries."
            
            explain_text.insert("1.0", explanation)
            explain_text.config(state=tk.DISABLED)
            
            # Close button
            ttk.Button(main_frame, text="Close", command=dialog.destroy).pack(side=tk.RIGHT)
            
            self.log_message(f"Processing strategy: {strategy.method}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to analyze strategy: {str(e)}")
            self.log_message(f"Strategy analysis error: {e}")

    def manage_encrypted_credentials_dialog(self):
        """Open credential management dialog - FIXED VERSION with visible styling"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Credential Security Management")
        dialog.geometry("550x450")
        dialog.configure(bg="#2b2b2b")
        dialog.grab_set()
        
        # Use regular Frame instead of ttk.Frame for better control
        main_frame = tk.Frame(dialog, bg="#2b2b2b", padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title with better visibility
        title_label = tk.Label(main_frame, 
                            text="Encrypted Credential Management",
                            font=("Arial", 14, "bold"),
                            bg="#2b2b2b",
                            fg="#ffffff")
        title_label.pack(anchor=tk.W, pady=(0, 15))
        
        # Status frame with visible border
        status_frame = tk.LabelFrame(main_frame, 
                                    text="Current Status",
                                    font=("Arial", 10, "bold"),
                                    bg="#2b2b2b",
                                    fg="#ffffff",
                                    padx=10,
                                    pady=10,
                                    relief=tk.GROOVE,
                                    borderwidth=2)
        status_frame.pack(fill=tk.X, pady=(0, 15))
        
        if self.credential_manager:
            encryption_status = "✓ Encrypted" if self.credential_manager.is_encrypted else "⚠ Unencrypted"
            status_color = "#00ff00" if self.credential_manager.is_encrypted else "#ffaa00"
            
            status_label = tk.Label(status_frame,
                                text=f"Credentials: {encryption_status}",
                                font=("Arial", 10),
                                bg="#2b2b2b",
                                fg=status_color)
            status_label.pack(anchor=tk.W, pady=2)
            
            if hasattr(self.credential_manager, 'credentials_file'):
                file_label = tk.Label(status_frame,
                                    text=f"File: {os.path.basename(self.credential_manager.credentials_file)}",
                                    font=("Arial", 9),
                                    bg="#2b2b2b",
                                    fg="#cccccc")
                file_label.pack(anchor=tk.W, pady=2)
            
            # Show current settings
            info_label = tk.Label(status_frame,
                                text=f"API URL: {self.openwebui_base_url}",
                                font=("Arial", 9),
                                bg="#2b2b2b",
                                fg="#cccccc")
            info_label.pack(anchor=tk.W, pady=2)
            
            model_label = tk.Label(status_frame,
                                text=f"Model: {self.selected_model.get() or 'Not set'}",
                                font=("Arial", 9),
                                bg="#2b2b2b",
                                fg="#cccccc")
            model_label.pack(anchor=tk.W, pady=2)
        else:
            error_label = tk.Label(status_frame,
                                text="⚠ Credential encryption not available",
                                font=("Arial", 10),
                                bg="#2b2b2b",
                                fg="#ff5555")
            error_label.pack(anchor=tk.W, pady=2)
            
            help_label = tk.Label(status_frame,
                                text="Install credential_manager.py and cryptography package",
                                font=("Arial", 9),
                                bg="#2b2b2b",
                                fg="#cccccc")
            help_label.pack(anchor=tk.W, pady=2)
        
        # Actions frame with visible border
        action_frame = tk.LabelFrame(main_frame,
                                    text="Available Actions",
                                    font=("Arial", 10, "bold"),
                                    bg="#2b2b2b",
                                    fg="#ffffff",
                                    padx=10,
                                    pady=10,
                                    relief=tk.GROOVE,
                                    borderwidth=2)
        action_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        if self.credential_manager:
            # Create buttons with proper styling
            if self.credential_manager.is_encrypted:
                btn1 = tk.Button(action_frame,
                            text="🔑 Change Password",
                            command=lambda: self.change_credential_password(),
                            bg="#404040",
                            fg="#ffffff",
                            font=("Arial", 10),
                            relief=tk.RAISED,
                            padx=10,
                            pady=5,
                            cursor="hand2")
                btn1.pack(fill=tk.X, pady=3)
                
                btn2 = tk.Button(action_frame,
                            text="💾 Backup Encrypted Credentials",
                            command=lambda: self.backup_credentials(),
                            bg="#404040",
                            fg="#ffffff",
                            font=("Arial", 10),
                            relief=tk.RAISED,
                            padx=10,
                            pady=5,
                            cursor="hand2")
                btn2.pack(fill=tk.X, pady=3)
            
            btn3 = tk.Button(action_frame,
                        text="💾 Save Current Settings to Encrypted Storage",
                        command=lambda: self.save_to_encrypted_credentials(),
                        bg="#404040",
                        fg="#ffffff",
                        font=("Arial", 10),
                        relief=tk.RAISED,
                        padx=10,
                        pady=5,
                        cursor="hand2")
            btn3.pack(fill=tk.X, pady=3)

            # Load from file button
            btn_load = tk.Button(action_frame,
                        text="📂 Load Credentials from File",
                        command=lambda: self.load_from_encrypted_credentials(),
                        bg="#404040",
                        fg="#ffffff",
                        font=("Arial", 10),
                        relief=tk.RAISED,
                        padx=10,
                        pady=5,
                        cursor="hand2")
            btn_load.pack(fill=tk.X, pady=3)

            if not self.credential_manager.is_encrypted:
                btn4 = tk.Button(action_frame,
                            text="🔒 Enable Encryption",
                            command=lambda: self.enable_encryption(),
                            bg="#00aa00",
                            fg="#ffffff",
                            font=("Arial", 10, "bold"),
                            relief=tk.RAISED,
                            padx=10,
                            pady=5,
                            cursor="hand2")
                btn4.pack(fill=tk.X, pady=3)
            
            # Add info text
            info_text = tk.Text(action_frame,
                            height=4,
                            width=50,
                            bg="#1e1e1e",
                            fg="#cccccc",
                            font=("Consolas", 9),
                            relief=tk.SUNKEN,
                            borderwidth=1,
                            wrap=tk.WORD)
            info_text.pack(fill=tk.X, pady=(10, 0))
            
            if self.credential_manager.is_encrypted:
                info_text.insert("1.0", "✓ Your credentials are encrypted with AES-256\n"
                                    "✓ Password required on each app start\n"
                                    "✓ Secure storage of API keys")
            else:
                info_text.insert("1.0", "⚠ Your credentials are stored in plain text\n"
                                    "⚠ Enable encryption for better security\n"
                                    "ℹ You'll need a password on each app start")
            info_text.config(state=tk.DISABLED)
            
        else:
            # Show installation instructions
            instructions = tk.Text(action_frame,
                                height=8,
                                width=50,
                                bg="#1e1e1e",
                                fg="#cccccc",
                                font=("Consolas", 9),
                                relief=tk.SUNKEN,
                                borderwidth=1,
                                wrap=tk.WORD)
            instructions.pack(fill=tk.BOTH, expand=True)
            
            instructions.insert("1.0", 
                "To enable encrypted credentials:\n\n"
                "1. Ensure credential_manager.py is in app directory\n\n"
                "2. Install cryptography package:\n"
                "   pip install cryptography --break-system-packages\n\n"
                "3. Restart application\n\n"
                "Benefits:\n"
                "• Secure API key storage\n"
                "• AES-256 encryption\n"
                "• Password protection")
            instructions.config(state=tk.DISABLED)
        
        # Close button with styling
        btn_frame = tk.Frame(main_frame, bg="#2b2b2b")
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        close_btn = tk.Button(btn_frame,
                            text="Close",
                            command=dialog.destroy,
                            bg="#404040",
                            fg="#ffffff",
                            font=("Arial", 10),
                            relief=tk.RAISED,
                            padx=20,
                            pady=5,
                            cursor="hand2")
        close_btn.pack(side=tk.RIGHT)

    def change_credential_password(self):
        """Change encryption password"""
        if not self.credential_manager:
            messagebox.showinfo("Not Available", "Credential manager not available")
            return
        
        if self.credential_manager.change_password(self.root):
            messagebox.showinfo("Success", "Password changed successfully!")
            self.log_message("Credential password changed")
        else:
            self.log_message("Password change cancelled or failed")

    def backup_credentials(self):
        """Create backup of credentials"""
        if not self.credential_manager:
            messagebox.showinfo("Not Available", "Credential manager not available")
            return
        
        backup_path = filedialog.asksaveasfilename(
            title="Save Credentials Backup",
            defaultextension=".enc",
            filetypes=[("Encrypted files", "*.enc"), ("All files", "*.*")]
        )
        
        if backup_path:
            if self.credential_manager.backup_credentials(backup_path):
                messagebox.showinfo("Success", f"Credentials backed up to:\n{backup_path}")
                self.log_message(f"Credentials backed up to: {backup_path}")
            else:
                messagebox.showerror("Failed", "Failed to backup credentials")

    def enable_encryption(self):
        """Enable encryption for currently unencrypted credentials"""
        if not self.credential_manager:
            return
        
        result = messagebox.askyesno(
            "Enable Encryption",
            "This will encrypt your credentials file.\n\n"
            "You will need to enter a password each time you start the application.\n\n"
            "Continue?"
        )
        
        if result:
            password = tk.simpledialog.askstring(
                "Create Password",
                "Enter a password for encryption:",
                show='*',
                parent=self.root
            )
            
            if password:
                confirm = tk.simpledialog.askstring(
                    "Confirm Password",
                    "Confirm password:",
                    show='*',
                    parent=self.root
                )
                
                if password == confirm:
                    self.credential_manager.is_encrypted = True
                    if self.credential_manager.save_credentials(password):
                        messagebox.showinfo("Success", "Credentials encrypted successfully!")
                        self.log_message("Credentials encrypted")
                    else:
                        messagebox.showerror("Failed", "Failed to encrypt credentials")
                        self.credential_manager.is_encrypted = False
                else:
                    messagebox.showerror("Error", "Passwords do not match")

    def load_encrypted_credentials(self):
        """Load encrypted credentials and sync with standard config"""
        if not self.credential_manager:
            return False
        
        try:
            if self.credential_manager.load_credentials(self.root):
                # Sync credentials with standard config
                self.openwebui_base_url = self.credential_manager.get_credential(
                    'openwebui', 'base_url', self.openwebui_base_url)
                self.openwebui_api_key = self.credential_manager.get_credential(
                    'openwebui', 'api_key', self.openwebui_api_key)
                
                model = self.credential_manager.get_credential(
                    'openwebui', 'default_model', '')
                if model:
                    self.selected_model.set(model)
                
                temp = self.credential_manager.get_credential(
                    'openwebui', 'temperature', 0.1)
                self.temperature.set(temp)
                
                max_tok = self.credential_manager.get_credential(
                    'openwebui', 'max_tokens', 8000)
                self.max_tokens.set(max_tok)
                
                return True
        except Exception as e:
            self.log_message(f"Error loading encrypted credentials: {e}")
        
        return False

    def save_to_encrypted_credentials(self):
        """Save current settings to encrypted credentials as backup"""
        if not self.credential_manager:
            messagebox.showinfo("Not Available", "Credential manager not available")
            return False
        
        try:
            # Update credentials with current settings
            self.credential_manager.set_credential('openwebui', 'base_url', self.openwebui_base_url)
            self.credential_manager.set_credential('openwebui', 'api_key', self.openwebui_api_key)
            self.credential_manager.set_credential('openwebui', 'default_model', self.selected_model.get())
            self.credential_manager.set_credential('openwebui', 'temperature', self.temperature.get())
            self.credential_manager.set_credential('openwebui', 'max_tokens', self.max_tokens.get())
            
            # Save (will prompt for password if needed)
            if self.credential_manager.is_encrypted:
                password = tk.simpledialog.askstring(
                    "Save Encrypted Credentials",
                    "Enter password to save encrypted credentials:",
                    show='*',
                    parent=self.root
                )
                if password:
                    if self.credential_manager.save_credentials(password):
                        messagebox.showinfo("Success", "Settings saved to encrypted credentials")
                        self.log_message("Settings saved to encrypted credentials")
                        return True
                else:
                    self.log_message("Save to encrypted credentials cancelled")
                    return False
            else:
                if self.credential_manager.save_credentials():
                    messagebox.showinfo("Success", "Settings saved to credentials file")
                    self.log_message("Settings saved to credentials file")
                    return True
                
        except Exception as e:
            self.log_message(f"Error saving encrypted credentials: {e}")
            messagebox.showerror("Error", f"Failed to save: {str(e)}")
        
        return False

    def change_credential_password(self):
        """Change encryption password"""
        if not self.credential_manager:
            messagebox.showinfo("Not Available", "Credential manager not available")
            return
        
        if self.credential_manager.change_password(self.root):
            messagebox.showinfo("Success", "Password changed successfully!")
            self.log_message("Credential password changed")
        else:
            self.log_message("Password change cancelled or failed")

    def backup_credentials(self):
        """Create backup of credentials"""
        if not self.credential_manager:
            messagebox.showinfo("Not Available", "Credential manager not available")
            return
        
        backup_path = filedialog.asksaveasfilename(
            title="Save Credentials Backup",
            defaultextension=".enc",
            filetypes=[("Encrypted files", "*.enc"), ("All files", "*.*")]
        )
        
        if backup_path:
            if self.credential_manager.backup_credentials(backup_path):
                messagebox.showinfo("Success", f"Credentials backed up to:\n{backup_path}")
                self.log_message(f"Credentials backed up to: {backup_path}")
            else:
                messagebox.showerror("Failed", "Failed to backup credentials")

    def enable_encryption(self):
        """Enable encryption for currently unencrypted credentials"""
        if not self.credential_manager:
            return
        
        result = messagebox.askyesno(
            "Enable Encryption",
            "This will encrypt your credentials file.\n\n"
            "You will need to enter a password each time you start the application.\n\n"
            "Continue?"
        )
        
        if result:
            password = tk.simpledialog.askstring(
                "Create Password",
                "Enter a password for encryption:",
                show='*',
                parent=self.root
            )
            
            if password:
                confirm = tk.simpledialog.askstring(
                    "Confirm Password",
                    "Confirm password:",
                    show='*',
                    parent=self.root
                )
                
                if password == confirm:
                    self.credential_manager.is_encrypted = True
                    if self.credential_manager.save_credentials(password):
                        messagebox.showinfo("Success", "Credentials encrypted successfully!")
                        self.log_message("Credentials encrypted")
                    else:
                        messagebox.showerror("Failed", "Failed to encrypt credentials")
                        self.credential_manager.is_encrypted = False
                else:
                    messagebox.showerror("Error", "Passwords do not match")

    def load_encrypted_credentials(self):
        """Load encrypted credentials and sync with standard config"""
        if not self.credential_manager:
            return False
        
        try:
            if self.credential_manager.load_credentials(self.root):
                # Sync credentials with standard config
                self.openwebui_base_url = self.credential_manager.get_credential(
                    'openwebui', 'base_url', self.openwebui_base_url)
                self.openwebui_api_key = self.credential_manager.get_credential(
                    'openwebui', 'api_key', self.openwebui_api_key)
                
                model = self.credential_manager.get_credential(
                    'openwebui', 'default_model', '')
                if model:
                    self.selected_model.set(model)
                
                temp = self.credential_manager.get_credential(
                    'openwebui', 'temperature', 0.1)
                self.temperature.set(temp)
                
                max_tok = self.credential_manager.get_credential(
                    'openwebui', 'max_tokens', 8000)
                self.max_tokens.set(max_tok)
                
                return True
        except Exception as e:
            self.log_message(f"Error loading encrypted credentials: {e}")
        
        return False

    def save_to_encrypted_credentials(self):
        """Save current settings to encrypted credentials as backup"""
        if not self.credential_manager:
            messagebox.showinfo("Not Available", "Credential manager not available")
            return False
        
        try:
            # Update credentials with current settings
            self.credential_manager.set_credential('openwebui', 'base_url', self.openwebui_base_url)
            self.credential_manager.set_credential('openwebui', 'api_key', self.openwebui_api_key)
            self.credential_manager.set_credential('openwebui', 'default_model', self.selected_model.get())
            self.credential_manager.set_credential('openwebui', 'temperature', self.temperature.get())
            self.credential_manager.set_credential('openwebui', 'max_tokens', self.max_tokens.get())
            
            # Save (will prompt for password if needed)
            if self.credential_manager.is_encrypted:
                password = tk.simpledialog.askstring(
                    "Save Encrypted Credentials",
                    "Enter password to save encrypted credentials:",
                    show='*',
                    parent=self.root
                )
                if password:
                    if self.credential_manager.save_credentials(password):
                        messagebox.showinfo("Success", "Settings saved to encrypted credentials")
                        self.log_message("Settings saved to encrypted credentials")
                        return True
                else:
                    self.log_message("Save to encrypted credentials cancelled")
                    return False
            else:
                if self.credential_manager.save_credentials():
                    messagebox.showinfo("Success", "Settings saved to credentials file")
                    self.log_message("Settings saved to credentials file")
                    return True
                
        except Exception as e:
            self.log_message(f"Error saving encrypted credentials: {e}")
            messagebox.showerror("Error", f"Failed to save: {str(e)}")

        return False

    def load_from_encrypted_credentials(self):
        """Load credentials from a specific file"""
        if not self.credential_manager:
            messagebox.showinfo("Not Available", "Credential manager not available")
            return False

        # Prompt for file selection
        from tkinter import filedialog
        filepath = filedialog.askopenfilename(
            title="Select Credentials File",
            filetypes=[
                ("Encrypted Credentials", "*.enc"),
                ("JSON Files", "*.json"),
                ("All Files", "*.*")
            ],
            initialdir=os.path.dirname(os.path.abspath(__file__))
        )

        if not filepath:
            self.log_message("Load credentials cancelled")
            return False

        try:
            # Load credentials from the selected file
            if self.credential_manager.load_credentials_from_file(filepath, self.root):
                # Apply loaded credentials to current settings
                creds = self.credential_manager.get_all_credentials()

                if 'openwebui' in creds:
                    self.openwebui_base_url = creds['openwebui'].get('base_url', self.openwebui_base_url)
                    self.openwebui_api_key = creds['openwebui'].get('api_key', '')
                    self.selected_model.set(creds['openwebui'].get('default_model', ''))
                    self.temperature.set(creds['openwebui'].get('temperature', 0.1))
                    self.max_tokens.set(creds['openwebui'].get('max_tokens', 8000))

                # Update UI
                self.update_config_status()

                self.log_message(f"✓ Credentials loaded from: {os.path.basename(filepath)}")
                messagebox.showinfo("Success",
                    f"Credentials loaded successfully!\n\n" +
                    f"From: {os.path.basename(filepath)}\n\n" +
                    "Settings have been applied.")
                return True

        except Exception as e:
            self.log_message(f"Error loading credentials: {e}")
            messagebox.showerror("Error", f"Failed to load credentials: {str(e)}")

        return False

    def analyze_document_tenses(self):
        """NEW FEATURE: Dedicated tense consistency analysis"""
        if not self.selected_section:
            messagebox.showwarning("No Selection", "Please select a section to analyze")
            return
        
        if not self.selected_section.has_content():
            messagebox.showinfo("No Content", "Selected section has no content to analyze")
            return
        
        if not self.advanced_reviewer:
            messagebox.showinfo("Feature Unavailable", 
                            "Advanced tense analysis requires document_reviewer module")
            return
        
        content = self.selected_section.get_existing_content()
        
        try:
            tense_analysis = self.advanced_reviewer.analyze_tense_consistency(content)
            
            # Display results
            self.generated_text.delete(1.0, tk.END)
            self.generated_text.insert(tk.END, "=== TENSE CONSISTENCY ANALYSIS ===\n\n")
            self.generated_text.insert(tk.END, f"Section: {self.selected_section.get_full_path()}\n\n")
            self.generated_text.insert(tk.END, f"Dominant Tense: {tense_analysis.dominant_tense.upper()}\n")
            self.generated_text.insert(tk.END, f"Consistency Score: {tense_analysis.consistency_score:.1f}/10\n\n")
            self.generated_text.insert(tk.END, f"Tense Distribution:\n")
            self.generated_text.insert(tk.END, f"  • Past: {tense_analysis.past_count} sentences\n")
            self.generated_text.insert(tk.END, f"  • Present: {tense_analysis.present_count} sentences\n")
            self.generated_text.insert(tk.END, f"  • Future: {tense_analysis.future_count} sentences\n\n")
            
            if tense_analysis.inconsistent_sentences:
                self.generated_text.insert(tk.END, f"Inconsistent Sentences ({len(tense_analysis.inconsistent_sentences)}):\n\n")
                for i, sentence in enumerate(tense_analysis.inconsistent_sentences[:10], 1):
                    self.generated_text.insert(tk.END, f"{i}. {sentence}\n\n")
            else:
                self.generated_text.insert(tk.END, "✓ No tense inconsistencies detected\n")
            
            self.log_message(f"Tense analysis completed: {tense_analysis.consistency_score:.1f}/10")
            self.notebook.select(0)
            
        except Exception as e:
            messagebox.showerror("Analysis Error", f"Failed to analyze tenses: {str(e)}")

    def show_processing_strategy_dialog(self):
        """NEW FEATURE: Display content processing strategy analysis"""
        if not self.content_processor:
            messagebox.showinfo("Feature Unavailable",
                            "Intelligent processing requires content_processor module")
            return
        
        if not self.selected_section or not self.selected_section.has_content():
            messagebox.showwarning("No Content", "Please select a section with content")
            return
        
        content = self.selected_section.get_existing_content()
        
        try:
            strategy = self.content_processor.determine_processing_strategy(
                content, [self.selected_section], "Generate comprehensive content"
            )
            
            dialog = tk.Toplevel(self.root)
            dialog.title("Processing Strategy Analysis")
            dialog.geometry("600x500")
            dialog.configure(bg="#2b2b2b")
            dialog.grab_set()
            
            main_frame = ttk.Frame(dialog, padding="20")
            main_frame.pack(fill=tk.BOTH, expand=True)
            
            ttk.Label(main_frame, text="Content Processing Strategy", 
                    font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 15))
            
            info_frame = ttk.LabelFrame(main_frame, text="Recommended Strategy", padding="10")
            info_frame.pack(fill=tk.X, pady=(0, 10))
            
            ttk.Label(info_frame, text=f"Method: {strategy.method.upper()}",
                    font=("Arial", 10, "bold")).pack(anchor=tk.W)
            ttk.Label(info_frame, text=f"Reason: {strategy.reason}").pack(anchor=tk.W, pady=5)
            ttk.Label(info_frame, text=f"Token Estimate: {strategy.token_estimate}").pack(anchor=tk.W)
            ttk.Label(info_frame, text=f"Confidence: {strategy.confidence:.1%}").pack(anchor=tk.W)
            
            ttk.Button(main_frame, text="Close", command=dialog.destroy).pack(side=tk.RIGHT)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to analyze strategy: {str(e)}")

    # ============================================================================
    # NEW FEATURES: External RAG Content, Section Chat, Whole Document Review
    # ============================================================================

    def get_external_rag_content(self):
        """Retrieve all external RAG content for inclusion in queries"""
        if not self.external_content_db:
            return []

        try:
            cursor = self.external_content_db.cursor()
            cursor.execute("SELECT title, content, category, tags FROM external_content ORDER BY updated_at DESC")
            return cursor.fetchall()
        except Exception as e:
            self.log_message(f"Error retrieving external content: {e}")
            return []

    def open_external_content_manager(self):
        """Open dialog to manage external RAG content"""
        dialog = tk.Toplevel(self.root)
        dialog.title("External RAG Content Manager")
        dialog.geometry("900x700")
        dialog.configure(bg="#2b2b2b")
        dialog.grab_set()

        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=1)

        # Title
        ttk.Label(main_frame, text="External RAG Content Manager",
                font=("Arial", 12, "bold")).grid(row=0, column=0, sticky=tk.W, pady=(0, 15))

        # Content list
        list_frame = ttk.LabelFrame(main_frame, text="Content Library", padding="10")
        list_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        list_frame.columnconfigure(0, weight=1)
        list_frame.rowconfigure(0, weight=1)

        # Create treeview for content list
        columns = ("title", "category", "tags", "created")
        content_tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=15)
        content_tree.heading("title", text="Title")
        content_tree.heading("category", text="Category")
        content_tree.heading("tags", text="Tags")
        content_tree.heading("created", text="Created")
        content_tree.column("title", width=300)
        content_tree.column("category", width=150)
        content_tree.column("tags", width=200)
        content_tree.column("created", width=150)

        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=content_tree.yview)
        content_tree.configure(yscrollcommand=scrollbar.set)

        content_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))

        # Button frame
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=2, column=0, sticky=tk.E, pady=(10, 0))

        def refresh_content_list():
            """Refresh the content list"""
            content_tree.delete(*content_tree.get_children())
            if self.external_content_db:
                try:
                    cursor = self.external_content_db.cursor()
                    cursor.execute("SELECT id, title, category, tags, created_at FROM external_content ORDER BY updated_at DESC")
                    for row in cursor.fetchall():
                        content_id, title, category, tags, created = row
                        content_tree.insert("", tk.END, values=(title, category or "", tags or "", created[:16] if created else ""))
                        content_tree.set(content_tree.get_children()[-1], "#0", content_id)
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to load content: {e}")

        def add_content():
            """Add new external content"""
            add_dialog = tk.Toplevel(dialog)
            add_dialog.title("Add External Content")
            add_dialog.geometry("700x600")
            add_dialog.configure(bg="#2b2b2b")
            add_dialog.grab_set()

            add_frame = ttk.Frame(add_dialog, padding="20")
            add_frame.pack(fill=tk.BOTH, expand=True)

            ttk.Label(add_frame, text="Add External RAG Content", font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 15))

            # Title
            ttk.Label(add_frame, text="Title:").pack(anchor=tk.W)
            title_entry = ttk.Entry(add_frame, width=60)
            title_entry.pack(fill=tk.X, pady=(0, 10))

            # Category
            ttk.Label(add_frame, text="Category (optional):").pack(anchor=tk.W)
            category_entry = ttk.Entry(add_frame, width=60)
            category_entry.pack(fill=tk.X, pady=(0, 10))

            # Tags
            ttk.Label(add_frame, text="Tags (comma-separated, optional):").pack(anchor=tk.W)
            tags_entry = ttk.Entry(add_frame, width=60)
            tags_entry.pack(fill=tk.X, pady=(0, 10))

            # Content
            ttk.Label(add_frame, text="Content:").pack(anchor=tk.W)
            content_text = scrolledtext.ScrolledText(add_frame, width=60, height=15, bg='#1e1e1e', fg='#ffffff', insertbackground='#ffffff')
            content_text.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

            def save_content():
                title = title_entry.get().strip()
                content = content_text.get('1.0', tk.END).strip()

                if not title or not content:
                    messagebox.showwarning("Validation", "Title and content are required")
                    return

                try:
                    content_id = str(uuid.uuid4())
                    cursor = self.external_content_db.cursor()
                    cursor.execute("""
                        INSERT INTO external_content (id, title, content, category, tags)
                        VALUES (?, ?, ?, ?, ?)
                    """, (content_id, title, content, category_entry.get().strip(), tags_entry.get().strip()))
                    self.external_content_db.commit()
                    self.log_message(f"✓ Added external content: {title}")
                    messagebox.showinfo("Success", "Content added successfully")
                    refresh_content_list()
                    add_dialog.destroy()
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to save content: {e}")

            def load_from_file():
                """Load content from a file"""
                file_path = filedialog.askopenfilename(
                    title="Select File",
                    filetypes=[("Text files", "*.txt"), ("Markdown", "*.md"), ("All files", "*.*")]
                )
                if file_path:
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content_text.delete('1.0', tk.END)
                            content_text.insert('1.0', f.read())
                        if not title_entry.get():
                            title_entry.insert(0, os.path.basename(file_path))
                    except Exception as e:
                        messagebox.showerror("Error", f"Failed to load file: {e}")

            save_btn_frame = ttk.Frame(add_frame)
            save_btn_frame.pack(fill=tk.X)

            ttk.Button(save_btn_frame, text="Load from File", command=load_from_file).pack(side=tk.LEFT)
            ttk.Button(save_btn_frame, text="Cancel", command=add_dialog.destroy).pack(side=tk.RIGHT)
            ttk.Button(save_btn_frame, text="Save", command=save_content).pack(side=tk.RIGHT, padx=(0, 5))

        def delete_content():
            """Delete selected content"""
            selected = content_tree.selection()
            if not selected:
                messagebox.showwarning("No Selection", "Please select content to delete")
                return

            if messagebox.askyesno("Confirm Delete", "Delete selected content?"):
                try:
                    for item in selected:
                        content_id = content_tree.set(item, "#0")
                        cursor = self.external_content_db.cursor()
                        cursor.execute("DELETE FROM external_content WHERE id = ?", (content_id,))
                    self.external_content_db.commit()
                    refresh_content_list()
                    self.log_message("✓ Content deleted")
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to delete content: {e}")

        def view_content():
            """View selected content"""
            selected = content_tree.selection()
            if not selected:
                messagebox.showwarning("No Selection", "Please select content to view")
                return

            try:
                content_id = content_tree.set(selected[0], "#0")
                cursor = self.external_content_db.cursor()
                cursor.execute("SELECT title, content, category, tags FROM external_content WHERE id = ?", (content_id,))
                row = cursor.fetchone()

                if row:
                    view_dialog = tk.Toplevel(dialog)
                    view_dialog.title(f"View: {row[0]}")
                    view_dialog.geometry("700x500")
                    view_dialog.configure(bg="#2b2b2b")

                    view_frame = ttk.Frame(view_dialog, padding="20")
                    view_frame.pack(fill=tk.BOTH, expand=True)

                    ttk.Label(view_frame, text=f"Title: {row[0]}", font=("Arial", 10, "bold")).pack(anchor=tk.W)
                    if row[2]:
                        ttk.Label(view_frame, text=f"Category: {row[2]}").pack(anchor=tk.W)
                    if row[3]:
                        ttk.Label(view_frame, text=f"Tags: {row[3]}").pack(anchor=tk.W)
                    ttk.Label(view_frame, text="Content:", font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(10, 5))

                    content_display = scrolledtext.ScrolledText(view_frame, width=60, height=20, bg='#1e1e1e', fg='#ffffff', wrap=tk.WORD)
                    content_display.pack(fill=tk.BOTH, expand=True)
                    content_display.insert('1.0', row[1])
                    content_display.config(state=tk.DISABLED)

                    ttk.Button(view_frame, text="Close", command=view_dialog.destroy).pack(pady=(10, 0))
            except Exception as e:
                messagebox.showerror("Error", f"Failed to view content: {e}")

        ttk.Button(btn_frame, text="Add Content", command=add_content).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(btn_frame, text="View", command=view_content).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(btn_frame, text="Delete", command=delete_content).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(btn_frame, text="Refresh", command=refresh_content_list).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(btn_frame, text="Close", command=dialog.destroy).pack(side=tk.RIGHT)

        # Initial load
        refresh_content_list()

        # Show count
        count = len(content_tree.get_children())
        self.log_message(f"External RAG content manager opened ({count} items)")

    def create_section_chat_tab(self):
        """Create the Section Chat tab for iterative refinement"""
        chat_frame = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(chat_frame, text="Section Chat")

        chat_frame.columnconfigure(0, weight=1)
        chat_frame.rowconfigure(1, weight=1)

        # Info label
        info_frame = ttk.Frame(chat_frame)
        info_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 10))

        self.chat_section_label = tk.StringVar(value="No section selected for chat")
        ttk.Label(info_frame, textvariable=self.chat_section_label, font=("Arial", 10, "bold")).pack(side=tk.LEFT)

        ttk.Button(info_frame, text="Start Chat with Selected Section",
                command=self.start_section_chat).pack(side=tk.RIGHT)
        ttk.Button(info_frame, text="Clear Chat",
                command=self.clear_section_chat).pack(side=tk.RIGHT, padx=(0, 5))

        # Chat history display
        history_frame = ttk.LabelFrame(chat_frame, text="Conversation", padding="10")
        history_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        history_frame.columnconfigure(0, weight=1)
        history_frame.rowconfigure(0, weight=1)

        self.chat_history_text = scrolledtext.ScrolledText(history_frame, width=80, height=20,
                                                          bg='#1e1e1e', fg='#ffffff',
                                                          insertbackground='#ffffff', wrap=tk.WORD)
        self.chat_history_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.chat_history_text.config(state=tk.DISABLED)

        # Message input
        input_frame = ttk.LabelFrame(chat_frame, text="Your Message", padding="10")
        input_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        input_frame.columnconfigure(0, weight=1)

        self.chat_input_text = scrolledtext.ScrolledText(input_frame, width=80, height=5,
                                                        bg='#1e1e1e', fg='#ffffff',
                                                        insertbackground='#ffffff', wrap=tk.WORD)
        self.chat_input_text.grid(row=0, column=0, sticky=(tk.W, tk.E))

        # Buttons
        btn_frame = ttk.Frame(chat_frame)
        btn_frame.grid(row=3, column=0, sticky=tk.E)

        self.chat_send_btn = ttk.Button(btn_frame, text="Send Message", command=self.send_chat_message, state=tk.DISABLED)
        self.chat_send_btn.pack(side=tk.LEFT, padx=(0, 5))

        self.chat_apply_btn = ttk.Button(btn_frame, text="Apply Last Response to Section",
                                        command=self.apply_chat_to_section, state=tk.DISABLED)
        self.chat_apply_btn.pack(side=tk.LEFT)

    def start_section_chat(self):
        """Start a chat session with the selected section"""
        if not self.selected_section:
            messagebox.showwarning("No Selection", "Please select a section first")
            return

        if not self.selected_model.get():
            messagebox.showwarning("No Model", "Please configure a model first")
            return

        self.current_chat_section = self.selected_section
        section_hash = self.selected_section.get_section_hash()

        # Initialize chat history for this section if needed
        if section_hash not in self.section_chat_history:
            self.section_chat_history[section_hash] = []

        # Update UI
        self.chat_section_label.set(f"Chatting about: {self.selected_section.get_full_path()}")
        self.chat_send_btn.config(state=tk.NORMAL)

        # Display chat history
        self.refresh_chat_display()

        # Add initial context message if this is a new chat
        if len(self.section_chat_history[section_hash]) == 0:
            context_msg = f"Started chat for section: {self.selected_section.get_full_path()}\n"
            if self.selected_section.has_content():
                context_msg += f"\nExisting content:\n{self.selected_section.get_existing_content()}\n"
            else:
                context_msg += "\n(Section currently has no content)\n"

            self.chat_history_text.config(state=tk.NORMAL)
            self.chat_history_text.insert(tk.END, f"{'='*60}\n{context_msg}{'='*60}\n\n")
            self.chat_history_text.config(state=tk.DISABLED)

        self.log_message(f"Started chat for section: {self.selected_section.get_full_path()}")
        self.notebook.select(3)  # Switch to chat tab

    def clear_section_chat(self):
        """Clear the current section chat"""
        if not self.current_chat_section:
            return

        if messagebox.askyesno("Clear Chat", "Clear the chat history for this section?"):
            section_hash = self.current_chat_section.get_section_hash()
            self.section_chat_history[section_hash] = []
            self.refresh_chat_display()
            self.chat_apply_btn.config(state=tk.DISABLED)
            self.log_message("Chat history cleared")

    def refresh_chat_display(self):
        """Refresh the chat history display"""
        if not self.current_chat_section:
            return

        section_hash = self.current_chat_section.get_section_hash()
        history = self.section_chat_history.get(section_hash, [])

        self.chat_history_text.config(state=tk.NORMAL)
        self.chat_history_text.delete('1.0', tk.END)

        for role, message in history:
            if role == "user":
                self.chat_history_text.insert(tk.END, f"YOU:\n{message}\n\n", "user")
            else:
                self.chat_history_text.insert(tk.END, f"AI:\n{message}\n\n", "assistant")
                self.chat_history_text.insert(tk.END, "-" * 60 + "\n\n")

        self.chat_history_text.tag_config("user", foreground="#4CAF50", font=("Arial", 10, "bold"))
        self.chat_history_text.tag_config("assistant", foreground="#2196F3", font=("Arial", 10, "bold"))

        self.chat_history_text.config(state=tk.DISABLED)
        self.chat_history_text.see(tk.END)

    def send_chat_message(self):
        """Send a message in the section chat"""
        if not self.current_chat_section:
            messagebox.showwarning("No Chat", "Please start a chat session first")
            return

        user_message = self.chat_input_text.get('1.0', tk.END).strip()
        if not user_message:
            messagebox.showwarning("Empty Message", "Please enter a message")
            return

        # Disable send button during processing
        self.chat_send_btn.config(state=tk.DISABLED)
        self.chat_input_text.delete('1.0', tk.END)

        # Run in thread
        thread = threading.Thread(target=self.process_chat_message, args=(user_message,))
        thread.daemon = True
        thread.start()

    def process_chat_message(self, user_message):
        """Process chat message in background thread"""
        try:
            section_hash = self.current_chat_section.get_section_hash()

            # Add user message to history
            self.section_chat_history[section_hash].append(("user", user_message))
            self.root.after(0, self.refresh_chat_display)

            # Build context for AI
            section = self.current_chat_section
            context = f"""You are helping refine content for a document section.

Section: {section.get_full_path()}

Current content:
{section.get_existing_content() if section.has_content() else "(No content yet)"}

Previous conversation:
"""
            # Include last 4 messages for context
            history = self.section_chat_history[section_hash]
            for role, msg in history[-5:-1]:  # Exclude the current user message
                context += f"\n{role.upper()}: {msg}\n"

            context += f"\nUser's current question/request: {user_message}\n\n"
            context += "Respond helpfully. If the user asks you to write or modify content, provide the complete updated content in your response."

            # Query AI
            self.log_message("Sending chat message to AI...")
            response = self.query_openwebui(context)

            if response and not response.startswith("Error:"):
                # Add AI response to history
                self.section_chat_history[section_hash].append(("assistant", response))
                self.root.after(0, self.refresh_chat_display)
                self.root.after(0, lambda: self.chat_apply_btn.config(state=tk.NORMAL))
                self.log_message("AI response received")
            else:
                self.log_message(f"Chat error: {response}")
                messagebox.showerror("Error", f"Failed to get response: {response}")

        except Exception as e:
            self.log_message(f"Chat error: {str(e)}")
            self.root.after(0, lambda: messagebox.showerror("Error", f"Chat failed: {str(e)}"))
        finally:
            self.root.after(0, lambda: self.chat_send_btn.config(state=tk.NORMAL))

    def apply_chat_to_section(self):
        """Apply the last AI response to the current section"""
        if not self.current_chat_section:
            messagebox.showwarning("No Chat", "No active chat session")
            return

        section_hash = self.current_chat_section.get_section_hash()
        history = self.section_chat_history.get(section_hash, [])

        if not history or history[-1][0] != "assistant":
            messagebox.showwarning("No Response", "No AI response to apply")
            return

        # Get the last AI response
        last_response = history[-1][1]

        # Set it as generated content
        self.generated_content = last_response

        # Show it in preview
        self.root.after(0, self.show_generated_content)

        # Enable commit
        self.commit_btn.config(state=tk.NORMAL)

        self.log_message("Applied chat response to section - ready to commit")
        messagebox.showinfo("Applied", "Chat response applied to section. Review in Preview tab and click Commit when ready.")

    def review_whole_document(self):
        """Review the entire document for cohesion, clarity, etc."""
        if not self.document:
            messagebox.showwarning("No Document", "Please load a document first")
            return

        if not self.selected_model.get():
            messagebox.showwarning("No Model", "Please configure a model first")
            return

        if not self.sections:
            messagebox.showwarning("No Sections", "Document has no sections to review")
            return

        # Confirm action
        section_count = len([s for s in self.sections if s.has_content()])
        if not messagebox.askyesno("Confirm Review",
                                  f"Review entire document?\n\n"
                                  f"This will analyze {section_count} sections with content.\n"
                                  f"This may take several minutes."):
            return

        self.log_message("Starting whole document review...")
        self.notebook.select(4)  # Switch to console

        # Run in thread
        thread = threading.Thread(target=self.perform_whole_document_review)
        thread.daemon = True
        thread.start()

    def perform_whole_document_review(self):
        """Perform whole document review in background thread"""
        try:
            # Collect all sections with content
            sections_with_content = [s for s in self.sections if s.has_content()]

            if not sections_with_content:
                self.root.after(0, lambda: messagebox.showinfo("No Content", "No sections have content to review"))
                return

            # Build comprehensive document content
            full_document = ""
            for section in sections_with_content:
                full_document += f"\n\n{'='*60}\n"
                full_document += f"SECTION: {section.get_full_path()}\n"
                full_document += f"{'='*60}\n\n"
                full_document += section.get_existing_content()

            # Build review prompt
            review_prompt = """You are an expert technical writer reviewing a complete document for quality, cohesion, and clarity.

REVIEW THE ENTIRE DOCUMENT BELOW and provide a comprehensive analysis.

Evaluate:
1. **Overall Cohesion** (1-10): Does the document flow logically? Are sections well-connected?
2. **Clarity** (1-10): Is the writing clear and understandable throughout?
3. **Technical Accuracy** (1-10): Are technical details accurate and appropriate?
4. **Completeness** (1-10): Does the document cover all necessary topics comprehensively?
5. **Consistency** (1-10): Is terminology, style, and tone consistent across sections?

For each criterion:
- Provide a score (1-10)
- List specific issues or concerns
- Provide recommendations for improvement

Also identify:
- Sections that are particularly strong
- Sections that need significant improvement
- Any gaps or missing information
- Any redundancy or unnecessary repetition
- Transitions that could be improved

DOCUMENT TO REVIEW:
""" + full_document

            self.log_message("Sending document to AI for comprehensive review...")

            # Query AI
            response = self.query_openwebui(review_prompt)

            if response and not response.startswith("Error:"):
                # Display results
                self.root.after(0, lambda: self.show_document_review_results(response, sections_with_content))
                self.log_message("✓ Whole document review completed")
            else:
                self.log_message(f"Review failed: {response}")
                self.root.after(0, lambda: messagebox.showerror("Error", f"Review failed: {response}"))

        except Exception as e:
            self.log_message(f"Review error: {str(e)}")
            self.root.after(0, lambda: messagebox.showerror("Error", f"Review failed: {str(e)}"))

    def show_document_review_results(self, review_content, sections):
        """Display whole document review results"""
        # Create results dialog
        dialog = tk.Toplevel(self.root)
        dialog.title("Whole Document Review Results")
        dialog.geometry("1000x800")
        dialog.configure(bg="#2b2b2b")

        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=1)

        # Header
        header_frame = ttk.Frame(main_frame)
        header_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 15))

        ttk.Label(header_frame, text="Document Review Results",
                font=("Arial", 14, "bold")).pack(side=tk.LEFT)
        ttk.Label(header_frame, text=f"Reviewed {len(sections)} sections",
                font=("Arial", 10)).pack(side=tk.LEFT, padx=(20, 0))

        # Results display
        results_frame = ttk.LabelFrame(main_frame, text="Review Analysis", padding="10")
        results_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(0, weight=1)

        results_text = scrolledtext.ScrolledText(results_frame, width=100, height=35,
                                                bg='#1e1e1e', fg='#ffffff',
                                                insertbackground='#ffffff', wrap=tk.WORD)
        results_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # Insert content
        results_text.insert('1.0', f"{'='*80}\n")
        results_text.insert(tk.END, f"WHOLE DOCUMENT REVIEW\n")
        results_text.insert(tk.END, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        results_text.insert(tk.END, f"Sections Analyzed: {len(sections)}\n")
        results_text.insert(tk.END, f"{'='*80}\n\n")
        results_text.insert(tk.END, review_content)

        results_text.config(state=tk.DISABLED)

        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=2, column=0, sticky=tk.E, pady=(10, 0))

        def export_review():
            """Export review to file"""
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt"), ("Markdown", "*.md"), ("All files", "*.*")]
            )
            if file_path:
                try:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(results_text.get('1.0', tk.END))
                    messagebox.showinfo("Success", f"Review exported to:\n{file_path}")
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to export: {e}")

        ttk.Button(btn_frame, text="Export Review", command=export_review).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(btn_frame, text="Close", command=dialog.destroy).pack(side=tk.LEFT)

        self.log_message("Review results displayed")


if __name__ == "__main__":
    try:
        print("="*60)
        print("Document Content Generator")
        print("="*60)
        print(f"Encrypted Credentials: {'✓' if ENHANCED_FEATURES_AVAILABLE['encryption'] else '✗'}")
        print(f"Advanced Review: {'✓' if ENHANCED_FEATURES_AVAILABLE['advanced_review'] else '✗'}")
        print(f"Intelligent Processing: {'✓' if ENHANCED_FEATURES_AVAILABLE['intelligent_processing'] else '✗'}")
        print("="*60)
        
        app = DocxDocumentFiller()
        app.run()
    except Exception as e:
        print(f"Application error: {str(e)}")
        import traceback
        traceback.print_exc()
        input("Press Enter to exit...")