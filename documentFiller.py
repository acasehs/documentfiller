#!/usr/bin/env python3
"""
DoD Document Filler - OpenWebUI Integration
Automatically fills Word document sections using OpenWebUI/Ollama with RAG support
Supports hierarchical section selection, multiple operation modes, and live preview
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from docx import Document  # pip install python-docx
from docx.shared import RGBColor, Pt  # pip install python-docx
from docx.enum.text import WD_COLOR_INDEX
import json
import requests  # pip install requests
import threading
import os
import re
from datetime import datetime
import shutil


class DocumentSection:
    """Represents a hierarchical document section"""
    def __init__(self, level, text, paragraph, full_path=""):
        self.level = level
        self.text = text
        self.paragraph = paragraph
        self.full_path = full_path
        self.children = []
        self.parent = None
        self.content_paragraphs = []  # Paragraphs belonging to this section
        
    def add_child(self, child):
        child.parent = self
        self.children.append(child)
        
    def get_full_path(self):
        """Get full hierarchical path"""
        if self.parent:
            parent_path = self.parent.get_full_path()
            return f"{parent_path} > {self.text}" if parent_path else self.text
        return self.text
        
    def get_existing_content(self):
        """Get existing text content in this section"""
        content = []
        for para in self.content_paragraphs:
            if para.text.strip():
                content.append(para.text)
        return "\n".join(content)
        
    def has_content(self):
        """Check if section has existing content"""
        return len(self.get_existing_content().strip()) > 0


class DocxDocumentFiller:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("DoD Document Filler - OpenWebUI Integration")
        self.root.geometry("1400x900")
        self.root.configure(bg="#2b2b2b")
        
        # Configure dark theme
        self.configure_dark_theme()
        
        # Data storage
        self.document = None
        self.document_path = None
        self.last_document_path = None
        self.sections = []
        self.selected_section = None
        self.generated_content = ""
        
        # OpenWebUI configuration
        self.openwebui_base_url = "http://172.16.27.122:3000"
        self.openwebui_api_key = ""
        self.available_models = []
        self.selected_model = tk.StringVar()
        self.available_knowledge_collections = []
        self.selected_knowledge_collections = []  # Multiple selections
        self.temperature = tk.DoubleVar(value=0.1)
        self.max_tokens = tk.IntVar(value=8000)
        
        # Operation mode
        self.operation_mode = tk.StringVar(value="replace")
        
        # Default master prompt
        self.master_prompt = tk.StringVar()
        self.set_default_prompt()
        
        # Load settings
        self.load_settings()
        
        # Create GUI
        self.create_gui()
        
        # Load last document if available
        if self.last_document_path and os.path.exists(self.last_document_path):
            self.load_document(self.last_document_path)
        
    def configure_dark_theme(self):
        """Configure dark theme"""
        try:
            style = ttk.Style()
            available_themes = style.theme_names()
            
            if 'clam' in available_themes:
                style.theme_use('clam')
            
            # Configure colors
            style.configure('TLabel', background='#2b2b2b', foreground='#ffffff')
            style.configure('TButton', background='#404040', foreground='#ffffff')
            style.configure('TFrame', background='#2b2b2b')
            style.configure('TLabelframe', background='#2b2b2b', foreground='#ffffff')
            style.configure('TLabelframe.Label', background='#2b2b2b', foreground='#ffffff')
            style.configure('TRadiobutton', background='#2b2b2b', foreground='#ffffff')
            
            # Treeview styling
            style.configure('Treeview', 
                          background='#1e1e1e',
                          foreground='#ffffff',
                          fieldbackground='#1e1e1e',
                          bordercolor='#666666')
            style.map('Treeview', background=[('selected', '#404040')])
            
        except Exception as e:
            print(f"Dark theme configuration failed: {e}")
            
    def set_default_prompt(self):
        """Set default master prompt"""
        default = """You are a Department of Defense cybersecurity expert tasked with creating comprehensive, technically accurate documentation for a DoD system.

CONTEXT:
- This is for a Zero Trust Architecture implementation
- Must comply with DoD standards including RMF, NIST SP 800-53, FedRAMP High, and DoD IL5/IL6
- System is a COTS ERP solution in commercial cloud hosting
- Integrates with 250+ legacy systems
- Handles PII and requires strict data protection

REQUIREMENTS:
- Use specific technical terminology appropriate for DoD cybersecurity documentation
- Reference applicable NIST controls, DoD directives, and compliance frameworks
- Be detailed and comprehensive - this is for ATO documentation
- Use proper formatting with bullet points and structured content where appropriate
- Focus on implementation details, not just policy statements
- Include specific requirements for Zero Trust pillars where applicable

ZERO TRUST TOOLS:
- CyberArk PAM
- OKTA Secure Authenttication
- Palo Alto Secure Browser 
- Palo Alto NGFWs
- Tenable Cloud Security
- Tenable SC
- Zscaler Private Access
- Zscaler Internet Access
- Trellix ePolicy Orchestrator
- Splunk
- Fortify
- SonarQube

CYBERSECURITY TEAM COMPOSITION
- Information Systems Security Officer (ISSOs) - Governance and Compliance, RMF Documentation
- Information Systems Security Engineers (ISSEs) - Security Engineering, STIG Integration, Security Impact Assessments
- Security Operations (SecOps) - System Scans, Policy Enforcement, Continuous Monitoring
- Identity, Credential, and Access Management (ICAM) - Credential Management, Access Auditing and Monitoring 
- Application Security (AppSec) - Application Vulnerability Testing, SAST, DAST, Penetration Testing

SECTION TO FILL: {section_name}
PARENT CONTEXT: {parent_context}
OPERATION: {operation_mode}"""
        
        self.master_prompt.set(default)
        
    def load_settings(self):
        """Load saved settings"""
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            config_file = os.path.join(script_dir, "openwebui_config.json")
            
            if os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    config = json.load(f)
                    self.openwebui_base_url = config.get('base_url', 'http://172.16.27.122:3000')
                    self.openwebui_api_key = config.get('api_key', '')
                    self.selected_model.set(config.get('model', ''))
                    self.temperature.set(config.get('temperature', 0.1))
                    self.max_tokens.set(config.get('max_tokens', 8000))
                    
            # Load last document path
            last_doc_file = os.path.join(script_dir, "last_document.txt")
            if os.path.exists(last_doc_file):
                with open(last_doc_file, 'r') as f:
                    self.last_document_path = f.read().strip()
                    
        except Exception as e:
            print(f"Note: Could not load settings: {str(e)}")
            
    def save_settings(self):
        """Save current settings"""
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            config_file = os.path.join(script_dir, "openwebui_config.json")
            
            config = {
                'base_url': self.openwebui_base_url,
                'api_key': self.openwebui_api_key,
                'model': self.selected_model.get(),
                'temperature': self.temperature.get(),
                'max_tokens': self.max_tokens.get()
            }
            
            with open(config_file, 'w') as f:
                json.dump(config, f, indent=2)
                
            # Save last document path
            if self.document_path:
                last_doc_file = os.path.join(script_dir, "last_document.txt")
                with open(last_doc_file, 'w') as f:
                    f.write(self.document_path)
                    
            self.log_message("Settings saved")
            
        except Exception as e:
            self.log_message(f"Error saving settings: {str(e)}")
            
    def create_gui(self):
        """Create main GUI"""
        # Main container with grid layout
        main_container = ttk.Frame(self.root, padding="10")
        main_container.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_container.columnconfigure(0, weight=1)
        main_container.rowconfigure(2, weight=1)
        
        # Title
        title_label = ttk.Label(main_container, text="DoD Document Filler - OpenWebUI", 
                               font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, pady=(0, 10), sticky=tk.W)
        
        # Top controls frame
        top_frame = ttk.Frame(main_container)
        top_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        top_frame.columnconfigure(1, weight=1)
        
        # Document selection
        doc_frame = ttk.LabelFrame(top_frame, text="Document", padding="10")
        doc_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        doc_frame.columnconfigure(1, weight=1)
        
        ttk.Button(doc_frame, text="Load Document", command=self.browse_document).grid(row=0, column=0, padx=(0, 5))
        self.doc_label_var = tk.StringVar(value="No document loaded")
        ttk.Label(doc_frame, textvariable=self.doc_label_var).grid(row=0, column=1, sticky=tk.W, padx=5)
        
        # OpenWebUI configuration
        config_frame = ttk.LabelFrame(top_frame, text="OpenWebUI Configuration", padding="10")
        config_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        ttk.Button(config_frame, text="Configure", command=self.open_config_dialog).grid(row=0, column=0, padx=(0, 10))
        self.config_status_var = tk.StringVar(value="Not configured")
        ttk.Label(config_frame, textvariable=self.config_status_var).grid(row=0, column=1, sticky=tk.W)
        
        # Main content area - split into left and right
        content_frame = ttk.Frame(main_container)
        content_frame.grid(row=2, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        content_frame.columnconfigure(0, weight=1)
        content_frame.columnconfigure(1, weight=2)
        content_frame.rowconfigure(0, weight=1)
        
        # Left panel - Section tree and controls
        left_panel = ttk.Frame(content_frame)
        left_panel.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        left_panel.columnconfigure(0, weight=1)
        left_panel.rowconfigure(1, weight=1)
        
        # Section tree
        tree_frame = ttk.LabelFrame(left_panel, text="Document Sections", padding="5")
        tree_frame.grid(row=0, column=0, rowspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        tree_frame.columnconfigure(0, weight=1)
        tree_frame.rowconfigure(0, weight=1)
        
        self.tree = ttk.Treeview(tree_frame, selectmode='browse')
        tree_scroll = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=tree_scroll.set)
        
        self.tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        tree_scroll.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        self.tree.bind('<<TreeviewSelect>>', self.on_section_select)
        
        # Operation controls
        op_frame = ttk.LabelFrame(left_panel, text="Operation Mode", padding="10")
        op_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        
        ttk.Radiobutton(op_frame, text="Replace (write from scratch)", 
                       variable=self.operation_mode, value="replace").pack(anchor=tk.W)
        ttk.Radiobutton(op_frame, text="Rework (rewrite existing)", 
                       variable=self.operation_mode, value="rework").pack(anchor=tk.W)
        ttk.Radiobutton(op_frame, text="Append (add to existing)", 
                       variable=self.operation_mode, value="append").pack(anchor=tk.W)
        
        # Action buttons
        button_frame = ttk.Frame(left_panel)
        button_frame.grid(row=3, column=0, sticky=(tk.W, tk.E))
        
        self.generate_btn = ttk.Button(button_frame, text="Generate Content", 
                                      command=self.generate_content, state='disabled')
        self.generate_btn.pack(fill=tk.X, pady=(0, 5))
        
        ttk.Button(button_frame, text="Edit Master Prompt", 
                  command=self.edit_master_prompt).pack(fill=tk.X)
        
        # Right panel - Preview and content
        right_panel = ttk.Frame(content_frame)
        right_panel.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S))
        right_panel.columnconfigure(0, weight=1)
        right_panel.rowconfigure(0, weight=1)
        
        # Create notebook for tabs
        self.notebook = ttk.Notebook(right_panel)
        self.notebook.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Preview tab
        preview_frame = ttk.Frame(self.notebook)
        self.notebook.add(preview_frame, text="Content Preview")
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.columnconfigure(1, weight=1)
        preview_frame.rowconfigure(1, weight=1)
        
        # Existing content (left side)
        ttk.Label(preview_frame, text="Existing Content", font=("Arial", 10, "bold")).grid(
            row=0, column=0, sticky=tk.W, padx=5, pady=5)
        
        self.existing_text = scrolledtext.ScrolledText(preview_frame, height=20, 
                                                      bg="#1e1e1e", fg="#ffffff",
                                                      font=("Consolas", 10), wrap=tk.WORD)
        self.existing_text.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 2))
        self.existing_text.config(state='disabled')
        
        # Generated content (right side - editable)
        ttk.Label(preview_frame, text="Generated Content (Editable)", 
                 font=("Arial", 10, "bold")).grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        
        self.generated_text = scrolledtext.ScrolledText(preview_frame, height=20,
                                                       bg="#1e1e1e", fg="#90EE90",
                                                       font=("Consolas", 10), wrap=tk.WORD)
        self.generated_text.grid(row=1, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(2, 5))
        
        # Commit button
        commit_frame = ttk.Frame(preview_frame)
        commit_frame.grid(row=2, column=0, columnspan=2, pady=10)
        
        self.commit_btn = ttk.Button(commit_frame, text="Commit to Document", 
                                    command=self.commit_content, state='disabled')
        self.commit_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(commit_frame, text="Clear Preview", 
                  command=self.clear_preview).pack(side=tk.LEFT, padx=5)
        
        # Console tab
        console_frame = ttk.Frame(self.notebook)
        self.notebook.add(console_frame, text="Console Log")
        console_frame.columnconfigure(0, weight=1)
        console_frame.rowconfigure(0, weight=1)
        
        self.console = scrolledtext.ScrolledText(console_frame, height=20,
                                                bg="#1e1e1e", fg="#ffffff",
                                                font=("Consolas", 9))
        self.console.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5, pady=5)
        
        # Update config status
        self.update_config_status()
        
    def log_message(self, message):
        """Log message to console"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.console.insert(tk.END, f"[{timestamp}] {message}\n")
        self.console.see(tk.END)
        self.root.update()
        
    def browse_document(self):
        """Browse for Word document"""
        filename = filedialog.askopenfilename(
            title="Select Word Document",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            self.load_document(filename)
            
    def load_document(self, filepath):
        """Load Word document and parse sections"""
        try:
            self.log_message(f"Loading document: {filepath}")
            self.document = Document(filepath)
            self.document_path = filepath
            self.doc_label_var.set(os.path.basename(filepath))
            
            # Parse document structure
            self.parse_document_structure()
            
            # Populate tree
            self.populate_tree()
            
            # Save as last document
            self.save_settings()
            
            self.log_message(f"Document loaded successfully: {len(self.sections)} sections found")
            
        except Exception as e:
            self.log_message(f"Error loading document: {str(e)}")
            messagebox.showerror("Error", f"Failed to load document: {str(e)}")
            
    def parse_document_structure(self):
        """Parse document into hierarchical structure"""
        self.sections = []
        section_stack = []  # Stack to track parent sections
        current_section = None
        
        for para in self.document.paragraphs:
            # Check if this is a heading
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                    if level <= 4:  # Only process up to Heading 4
                        # Create new section
                        section = DocumentSection(level, para.text.strip(), para)
                        
                        # Pop stack to find appropriate parent
                        while section_stack and section_stack[-1].level >= level:
                            section_stack.pop()
                        
                        # Add to parent or root
                        if section_stack:
                            section_stack[-1].add_child(section)
                        else:
                            self.sections.append(section)
                        
                        # Push to stack
                        section_stack.append(section)
                        current_section = section
                        
                except ValueError:
                    pass  # Not a numbered heading
            else:
                # Add paragraph to current section's content
                if current_section and para.text.strip():
                    # Don't add if it's a bullet point header (those are just guidance)
                    if not para.text.strip().startswith('-'):
                        current_section.content_paragraphs.append(para)
                        
    def populate_tree(self):
        """Populate treeview with document sections"""
        self.tree.delete(*self.tree.get_children())
        
        def add_to_tree(section, parent=''):
            # Determine if section has content
            has_content = section.has_content()
            display_text = section.text
            if has_content:
                display_text += " ✓"
            
            item_id = self.tree.insert(parent, 'end', text=display_text, 
                                      values=(id(section),))
            
            for child in section.children:
                add_to_tree(child, item_id)
        
        for section in self.sections:
            add_to_tree(section)
            
    def on_section_select(self, event):
        """Handle section selection"""
        selection = self.tree.selection()
        if selection:
            item = selection[0]
            section_id = self.tree.item(item, 'values')[0]
            
            # Find section by ID
            self.selected_section = self.find_section_by_id(int(section_id))
            
            if self.selected_section:
                # Enable generate button
                self.generate_btn.config(state='normal')
                
                # Show existing content
                self.show_existing_content()
                
                self.log_message(f"Selected section: {self.selected_section.get_full_path()}")
                
    def find_section_by_id(self, section_id):
        """Find section by its ID"""
        def search(sections):
            for section in sections:
                if id(section) == section_id:
                    return section
                result = search(section.children)
                if result:
                    return result
            return None
        
        return search(self.sections)
        
    def show_existing_content(self):
        """Show existing content in preview"""
        if self.selected_section:
            self.existing_text.config(state='normal')
            self.existing_text.delete('1.0', tk.END)
            
            existing = self.selected_section.get_existing_content()
            if existing:
                self.existing_text.insert('1.0', existing)
            else:
                self.existing_text.insert('1.0', "[No existing content]")
            
            self.existing_text.config(state='disabled')
            
    def open_config_dialog(self):
        """Open OpenWebUI configuration dialog"""
        config_window = tk.Toplevel(self.root)
        config_window.title("OpenWebUI Configuration")
        config_window.geometry("700x700")
        config_window.configure(bg="#2b2b2b")
        config_window.grab_set()
        
        main_frame = ttk.Frame(config_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Connection settings
        conn_frame = ttk.LabelFrame(main_frame, text="Connection", padding="10")
        conn_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(conn_frame, text="Base URL:").pack(anchor=tk.W)
        url_var = tk.StringVar(value=self.openwebui_base_url)
        ttk.Entry(conn_frame, textvariable=url_var, width=50).pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(conn_frame, text="API Key:").pack(anchor=tk.W)
        key_var = tk.StringVar(value=self.openwebui_api_key)
        ttk.Entry(conn_frame, textvariable=key_var, width=50, show="*").pack(fill=tk.X, pady=(0, 10))
        
        test_frame = ttk.Frame(conn_frame)
        test_frame.pack(fill=tk.X)
        
        ttk.Button(test_frame, text="Test Connection", 
                  command=lambda: self.test_connection(url_var.get(), key_var.get(), status_label)).pack(side=tk.LEFT)
        status_label = ttk.Label(test_frame, text="")
        status_label.pack(side=tk.LEFT, padx=(10, 0))
        
        # Model selection
        model_frame = ttk.LabelFrame(main_frame, text="Model", padding="10")
        model_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(model_frame, text="Refresh Models", 
                  command=lambda: self.refresh_models(url_var.get(), key_var.get(), model_combo)).pack(anchor=tk.W, pady=(0, 5))
        
        ttk.Label(model_frame, text="Select Model:").pack(anchor=tk.W)
        model_combo = ttk.Combobox(model_frame, textvariable=self.selected_model, state="readonly")
        model_combo['values'] = self.available_models
        model_combo.pack(fill=tk.X)
        
        # Knowledge collections (multiple selection)
        knowledge_frame = ttk.LabelFrame(main_frame, text="RAG Knowledge Collections", padding="10")
        knowledge_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        ttk.Button(knowledge_frame, text="Refresh Collections", 
                  command=lambda: self.refresh_knowledge(url_var.get(), key_var.get(), knowledge_list)).pack(anchor=tk.W, pady=(0, 5))
        
        ttk.Label(knowledge_frame, text="Select Collections (hold Ctrl for multiple):").pack(anchor=tk.W)
        
        knowledge_scroll = ttk.Scrollbar(knowledge_frame)
        knowledge_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        knowledge_list = tk.Listbox(knowledge_frame, selectmode=tk.MULTIPLE, 
                                    yscrollcommand=knowledge_scroll.set,
                                    bg="#1e1e1e", fg="#ffffff", height=8)
        knowledge_list.pack(fill=tk.BOTH, expand=True)
        knowledge_scroll.config(command=knowledge_list.yview)
        
        # Pre-populate if collections exist
        if self.available_knowledge_collections:
            for i, col in enumerate(self.available_knowledge_collections):
                display = f"{col['name']} ({col['id']})"
                knowledge_list.insert(tk.END, display)
                # Select previously selected ones
                if col['id'] in [c['id'] for c in self.selected_knowledge_collections]:
                    knowledge_list.selection_set(i)
        
        # Model parameters
        params_frame = ttk.LabelFrame(main_frame, text="Parameters", padding="10")
        params_frame.pack(fill=tk.X, pady=(0, 10))
        
        temp_label_var = tk.StringVar(value=f"Temperature: {self.temperature.get()}")
        ttk.Label(params_frame, textvariable=temp_label_var).pack(anchor=tk.W)
        temp_scale = ttk.Scale(params_frame, from_=0.0, to=1.0, variable=self.temperature,
                              command=lambda v: temp_label_var.set(f"Temperature: {float(v):.2f}"))
        temp_scale.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(params_frame, text="Max Tokens:").pack(anchor=tk.W)
        ttk.Spinbox(params_frame, from_=500, to=16000, textvariable=self.max_tokens, 
                   increment=500).pack(fill=tk.X)
        
        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        def save_and_close():
            self.openwebui_base_url = url_var.get()
            self.openwebui_api_key = key_var.get()
            
            # Get selected knowledge collections
            self.selected_knowledge_collections = []
            for idx in knowledge_list.curselection():
                self.selected_knowledge_collections.append(
                    self.available_knowledge_collections[idx])
            
            self.save_settings()
            self.update_config_status()
            config_window.destroy()
        
        ttk.Button(btn_frame, text="Save", command=save_and_close).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(btn_frame, text="Cancel", command=config_window.destroy).pack(side=tk.RIGHT)
        
    def test_connection(self, url, api_key, status_label):
        """Test OpenWebUI connection"""
        try:
            headers = {'Authorization': f'Bearer {api_key}'}
            response = requests.get(f"{url}/api/models", headers=headers, timeout=10)
            
            if response.status_code == 200:
                status_label.config(text="✓ Connection successful")
                return True
            else:
                status_label.config(text=f"✗ HTTP {response.status_code}")
        except Exception as e:
            status_label.config(text=f"✗ {str(e)}")
        return False
        
    def refresh_models(self, url, api_key, combo):
        """Refresh available models"""
        try:
            headers = {'Authorization': f'Bearer {api_key}'}
            response = requests.get(f"{url}/api/models", headers=headers, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                self.available_models = []
                
                # Handle different response formats
                if isinstance(data, dict) and 'data' in data:
                    models_list = data['data']
                elif isinstance(data, list):
                    models_list = data
                else:
                    models_list = [data] if data else []
                
                for model in models_list:
                    if isinstance(model, dict):
                        model_id = model.get('id', model.get('name', 'unknown'))
                        self.available_models.append(model_id)
                    else:
                        self.available_models.append(str(model))
                
                combo['values'] = self.available_models
                self.log_message(f"Refreshed models: {len(self.available_models)} found")
            else:
                self.log_message(f"Error refreshing models: HTTP {response.status_code}")
        except Exception as e:
            self.log_message(f"Error refreshing models: {str(e)}")
            
    def refresh_knowledge(self, url, api_key, listbox):
        """Refresh knowledge collections"""
        try:
            headers = {'Authorization': f'Bearer {api_key}'}
            response = requests.get(f"{url}/api/v1/knowledge/", headers=headers, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                self.available_knowledge_collections = []
                
                if isinstance(data, list):
                    for collection in data:
                        if isinstance(collection, dict):
                            self.available_knowledge_collections.append({
                                'id': collection.get('id', 'unknown'),
                                'name': collection.get('name', 'Unknown'),
                                'description': collection.get('description', '')
                            })
                
                # Update listbox
                listbox.delete(0, tk.END)
                for col in self.available_knowledge_collections:
                    display = f"{col['name']} ({col['id']})"
                    listbox.insert(tk.END, display)
                
                self.log_message(f"Refreshed collections: {len(self.available_knowledge_collections)} found")
            else:
                self.log_message(f"Error refreshing collections: HTTP {response.status_code}")
        except Exception as e:
            self.log_message(f"Error refreshing collections: {str(e)}")
            
    def update_config_status(self):
        """Update configuration status display"""
        if self.selected_model.get() and self.openwebui_api_key:
            status = f"✓ Model: {self.selected_model.get()}"
            if self.selected_knowledge_collections:
                status += f" | RAG: {len(self.selected_knowledge_collections)} collections"
            self.config_status_var.set(status)
        else:
            self.config_status_var.set("⚠ Not configured - click Configure")
            
    def edit_master_prompt(self):
        """Open master prompt editor"""
        editor_window = tk.Toplevel(self.root)
        editor_window.title("Edit Master Prompt")
        editor_window.geometry("800x600")
        editor_window.configure(bg="#2b2b2b")
        editor_window.grab_set()
        
        main_frame = ttk.Frame(editor_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Master Prompt Template:", 
                 font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 10))
        
        info_label = ttk.Label(main_frame, 
                              text="Variables: {section_name}, {parent_context}, {operation_mode}, {existing_content}",
                              foreground="#888888")
        info_label.pack(anchor=tk.W, pady=(0, 10))
        
        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        prompt_text = scrolledtext.ScrolledText(text_frame, height=20, width=80,
                                               bg="#1e1e1e", fg="#ffffff",
                                               font=("Consolas", 10), wrap=tk.WORD)
        prompt_text.pack(fill=tk.BOTH, expand=True)
        prompt_text.insert('1.0', self.master_prompt.get())
        
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        def save_prompt():
            self.master_prompt.set(prompt_text.get('1.0', tk.END).strip())
            self.log_message("Master prompt updated")
            editor_window.destroy()
        
        def reset_prompt():
            if messagebox.askyesno("Reset", "Reset to default prompt?"):
                self.set_default_prompt()
                prompt_text.delete('1.0', tk.END)
                prompt_text.insert('1.0', self.master_prompt.get())
        
        ttk.Button(btn_frame, text="Reset to Default", command=reset_prompt).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Cancel", command=editor_window.destroy).pack(side=tk.RIGHT)
        ttk.Button(btn_frame, text="Save", command=save_prompt).pack(side=tk.RIGHT, padx=(0, 5))
        
    def generate_content(self):
        """Generate content for selected section"""
        if not self.selected_section:
            messagebox.showwarning("Warning", "No section selected")
            return
            
        if not self.selected_model.get():
            messagebox.showwarning("Warning", "No model configured")
            return
            
        # Switch to console tab
        self.notebook.select(1)
        
        # Run in thread
        thread = threading.Thread(target=self.generate_content_thread)
        thread.daemon = True
        thread.start()
        
    def generate_content_thread(self):
        """Generate content in background thread"""
        try:
            self.generate_btn.config(state='disabled')
            
            section = self.selected_section
            mode = self.operation_mode.get()
            existing_content = section.get_existing_content()
            
            self.log_message(f"Generating content for: {section.get_full_path()}")
            self.log_message(f"Mode: {mode}")
            
            # Build parent context
            parent_context = ""
            if section.parent:
                parent_path = []
                current = section.parent
                while current:
                    parent_path.insert(0, current.text)
                    current = current.parent
                parent_context = " > ".join(parent_path)
            
            # Build prompt
            prompt = self.master_prompt.get()
            prompt = prompt.replace("{section_name}", section.text)
            prompt = prompt.replace("{parent_context}", parent_context if parent_context else "Root level")
            prompt = prompt.replace("{operation_mode}", mode.upper())
            
            # Add mode-specific instructions
            if mode == "replace":
                prompt += "\n\nYour task: Write comprehensive content for this section from scratch."
            elif mode == "rework":
                prompt += f"\n\nEXISTING CONTENT TO REWORK:\n{existing_content}\n\n"
                prompt += "Your task: Rewrite and enhance the existing content. Keep what's good, improve what needs work, add anything missing."
            elif mode == "append":
                prompt += f"\n\nEXISTING CONTENT:\n{existing_content}\n\n"
                prompt += "Your task: Add additional relevant content that complements and extends the existing content."
            
            # Add knowledge base instruction if collections selected
            if self.selected_knowledge_collections:
                collection_names = [col['name'] for col in self.selected_knowledge_collections]
                prompt += f"\n\nIMPORTANT: Reference the following knowledge base collections: {', '.join(collection_names)}"
            
            self.log_message("Sending request to OpenWebUI...")
            
            # Query OpenWebUI
            response = self.query_openwebui(prompt)
            
            if response and not response.startswith("Error:"):
                self.generated_content = response
                self.root.after(0, self.show_generated_content)
                self.log_message("Content generated successfully")
            else:
                self.log_message(f"Generation failed: {response}")
                self.root.after(0, lambda: self.handle_generation_error(prompt))
                
        except Exception as e:
            self.log_message(f"Error generating content: {str(e)}")
            self.root.after(0, lambda: self.handle_generation_error(""))
        finally:
            self.root.after(0, lambda: self.generate_btn.config(state='normal'))
            
    def query_openwebui(self, prompt):
        """Query OpenWebUI API"""
        try:
            headers = {
                'Authorization': f'Bearer {self.openwebui_api_key}',
                'Content-Type': 'application/json'
            }
            
            payload = {
                "model": self.selected_model.get(),
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "temperature": self.temperature.get(),
                "max_tokens": self.max_tokens.get()
            }
            
            # Add knowledge collections if selected
            if self.selected_knowledge_collections:
                payload["files"] = [
                    {"type": "collection", "id": col['id']} 
                    for col in self.selected_knowledge_collections
                ]
            
            response = requests.post(
                f"{self.openwebui_base_url}/api/chat/completions",
                headers=headers, json=payload, timeout=300
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and len(result['choices']) > 0:
                    return result['choices'][0]['message']['content']
                elif 'response' in result:
                    return result['response']
                else:
                    return "Error: No content in response"
            else:
                return f"Error: HTTP {response.status_code} - {response.text}"
                
        except Exception as e:
            return f"Error: {str(e)}"
            
    def show_generated_content(self):
        """Show generated content in preview"""
        # Switch to preview tab
        self.notebook.select(0)
        
        # Clear and populate generated text
        self.generated_text.delete('1.0', tk.END)
        self.generated_text.insert('1.0', self.generated_content)
        
        # Enable commit button
        self.commit_btn.config(state='normal')
        
    def handle_generation_error(self, original_prompt):
        """Handle generation error with options"""
        result = messagebox.askyesnocancel(
            "Generation Failed",
            "Content generation failed.\n\n"
            "Yes: Edit prompt and retry\n"
            "No: Enter content manually\n"
            "Cancel: Go back"
        )
        
        if result is True:  # Edit prompt
            self.edit_and_retry_prompt(original_prompt)
        elif result is False:  # Manual entry
            self.manual_content_entry()
            
    def edit_and_retry_prompt(self, original_prompt):
        """Edit prompt and retry generation"""
        editor_window = tk.Toplevel(self.root)
        editor_window.title("Edit Prompt and Retry")
        editor_window.geometry("900x700")
        editor_window.configure(bg="#2b2b2b")
        editor_window.grab_set()
        
        main_frame = ttk.Frame(editor_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Edit the prompt and try again:", 
                 font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 10))
        
        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        prompt_text = scrolledtext.ScrolledText(text_frame, height=25, width=90,
                                               bg="#1e1e1e", fg="#ffffff",
                                               font=("Consolas", 10), wrap=tk.WORD)
        prompt_text.pack(fill=tk.BOTH, expand=True)
        
        if original_prompt:
            prompt_text.insert('1.0', original_prompt)
        else:
            # Reconstruct the prompt
            section = self.selected_section
            parent_context = ""
            if section.parent:
                parent_path = []
                current = section.parent
                while current:
                    parent_path.insert(0, current.text)
                    current = current.parent
                parent_context = " > ".join(parent_path)
            
            prompt = self.master_prompt.get()
            prompt = prompt.replace("{section_name}", section.text)
            prompt = prompt.replace("{parent_context}", parent_context if parent_context else "Root level")
            prompt = prompt.replace("{operation_mode}", self.operation_mode.get().upper())
            prompt_text.insert('1.0', prompt)
        
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        def retry_with_edited():
            edited_prompt = prompt_text.get('1.0', tk.END).strip()
            editor_window.destroy()
            
            # Run generation with edited prompt
            def retry_thread():
                try:
                    self.generate_btn.config(state='disabled')
                    self.log_message("Retrying with edited prompt...")
                    response = self.query_openwebui(edited_prompt)
                    
                    if response and not response.startswith("Error:"):
                        self.generated_content = response
                        self.root.after(0, self.show_generated_content)
                        self.log_message("Content generated successfully on retry")
                    else:
                        self.log_message(f"Retry failed: {response}")
                        self.root.after(0, lambda: self.handle_generation_error(edited_prompt))
                except Exception as e:
                    self.log_message(f"Retry error: {str(e)}")
                    self.root.after(0, lambda: self.handle_generation_error(edited_prompt))
                finally:
                    self.root.after(0, lambda: self.generate_btn.config(state='normal'))
            
            thread = threading.Thread(target=retry_thread)
            thread.daemon = True
            thread.start()
        
        ttk.Button(btn_frame, text="Cancel", command=editor_window.destroy).pack(side=tk.RIGHT)
        ttk.Button(btn_frame, text="Retry Generation", command=retry_with_edited).pack(side=tk.RIGHT, padx=(0, 5))
        
    def manual_content_entry(self):
        """Allow manual content entry"""
        # Switch to preview tab
        self.notebook.select(0)
        
        # Clear and enable editing
        self.generated_text.delete('1.0', tk.END)
        self.generated_text.insert('1.0', "[Enter content manually here]")
        
        # Enable commit button
        self.commit_btn.config(state='normal')
        
        self.log_message("Manual content entry mode - edit in preview pane")
        
    def clear_preview(self):
        """Clear preview pane"""
        self.generated_text.delete('1.0', tk.END)
        self.commit_btn.config(state='disabled')
        self.log_message("Preview cleared")
        
    def commit_content(self):
        """Commit generated content to document"""
        if not self.selected_section or not self.document:
            return
            
        # Get edited content from preview
        content_to_commit = self.generated_text.get('1.0', tk.END).strip()
        
        if not content_to_commit:
            messagebox.showwarning("Warning", "No content to commit")
            return
        
        # Confirm commit
        result = messagebox.askyesno(
            "Commit Content",
            f"Commit content to section: {self.selected_section.text}\n\n"
            f"Mode: {self.operation_mode.get().upper()}\n\n"
            "This will modify the document."
        )
        
        if not result:
            return
        
        try:
            # Create backup first
            self.create_backup()
            
            # Apply content based on mode
            mode = self.operation_mode.get()
            section = self.selected_section
            
            if mode == "replace":
                # Remove existing content paragraphs
                for para in section.content_paragraphs:
                    para.clear()
                section.content_paragraphs = []
                
                # Add new content
                self.add_content_to_section(section, content_to_commit)
                
            elif mode == "rework":
                # Remove existing content
                for para in section.content_paragraphs:
                    para.clear()
                section.content_paragraphs = []
                
                # Add reworked content
                self.add_content_to_section(section, content_to_commit)
                
            elif mode == "append":
                # Add to existing content
                self.add_content_to_section(section, content_to_commit, append=True)
            
            # Refresh tree to show section has content
            self.populate_tree()
            
            # Prompt to save
            self.save_document()
            
            # Clear preview
            self.clear_preview()
            
            # Refresh existing content display
            self.show_existing_content()
            
            self.log_message(f"Content committed successfully to: {section.text}")
            
        except Exception as e:
            self.log_message(f"Error committing content: {str(e)}")
            messagebox.showerror("Error", f"Failed to commit content: {str(e)}")
            
    def create_backup(self):
        """Create backup of document"""
        if not self.document_path:
            return
            
        try:
            # Create backup filename
            backup_dir = os.path.dirname(self.document_path)
            filename = os.path.basename(self.document_path)
            name, ext = os.path.splitext(filename)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_path = os.path.join(backup_dir, f"{name}_backup_{timestamp}{ext}")
            
            # Copy file
            shutil.copy2(self.document_path, backup_path)
            self.log_message(f"Backup created: {os.path.basename(backup_path)}")
            
        except Exception as e:
            self.log_message(f"Warning: Could not create backup: {str(e)}")
            
    def add_content_to_section(self, section, content, append=False):
        """Add content to section in document"""
        # Find the paragraph after the section heading
        heading_para = section.paragraph
        
        # Split content into paragraphs (handle both \n\n and single \n)
        content_paragraphs = content.split('\n')
        
        # Find insertion point by matching paragraph text
        doc_paragraphs = list(self.document.paragraphs)
        heading_index = -1
        
        # Find heading by text match (more reliable than object comparison)
        for i, para in enumerate(doc_paragraphs):
            if para.text.strip() == heading_para.text.strip():
                heading_index = i
                break
        
        if heading_index == -1:
            raise Exception("Could not find section heading in document")
        
        # If appending, find the last content paragraph of this section
        if append and section.content_paragraphs:
            # Find the next heading to know where this section ends
            next_heading_index = len(doc_paragraphs)
            for i in range(heading_index + 1, len(doc_paragraphs)):
                if doc_paragraphs[i].style.name.startswith('Heading'):
                    next_heading_index = i
                    break
            insertion_index = next_heading_index
        else:
            # Clear existing content first if not appending
            if section.content_paragraphs:
                # Find next heading to determine section boundary
                next_heading_index = len(doc_paragraphs)
                for i in range(heading_index + 1, len(doc_paragraphs)):
                    if doc_paragraphs[i].style.name.startswith('Heading'):
                        next_heading_index = i
                        break
                
                # Remove paragraphs between heading and next heading
                paragraphs_to_remove = []
                for i in range(heading_index + 1, next_heading_index):
                    para = doc_paragraphs[i]
                    if not para.style.name.startswith('Heading'):
                        paragraphs_to_remove.append(para)
                
                for para in paragraphs_to_remove:
                    self.remove_paragraph(para)
            
            # Refresh paragraph list after removals
            doc_paragraphs = list(self.document.paragraphs)
            
            # Find heading again
            heading_index = -1
            for i, para in enumerate(doc_paragraphs):
                if para.text.strip() == heading_para.text.strip():
                    heading_index = i
                    break
            
            insertion_index = heading_index + 1
        
        # Insert new paragraphs with color highlighting
        inserted_paras = []
        for para_text in content_paragraphs:
            if para_text.strip():
                # Create new paragraph at insertion point
                if insertion_index < len(self.document.paragraphs):
                    new_para = self.document.paragraphs[insertion_index].insert_paragraph_before(para_text.strip())
                else:
                    new_para = self.document.add_paragraph(para_text.strip())
                
                # Apply yellow highlight to new content for review
                for run in new_para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
                inserted_paras.append(new_para)
                insertion_index += 1
        
        # Update section's content paragraphs
        if not append:
            section.content_paragraphs = inserted_paras
        else:
            section.content_paragraphs.extend(inserted_paras)
    
    def remove_paragraph(self, paragraph):
        """Remove a paragraph from the document"""
        # Get the paragraph's XML element
        p = paragraph._element
        # Remove from parent
        p.getparent().remove(p)
            
    def save_document(self):
        """Save document with options"""
        if not self.document or not self.document_path:
            return
            
        result = messagebox.askyesnocancel(
            "Save Document",
            "Save changes?\n\n"
            "Yes: Save to current file\n"
            "No: Save As (new file)\n"
            "Cancel: Don't save"
        )
        
        if result is True:  # Save to current
            try:
                self.document.save(self.document_path)
                self.log_message(f"Document saved: {self.document_path}")
                messagebox.showinfo("Success", "Document saved successfully")
            except PermissionError:
                messagebox.showerror("Error", 
                    "File is open in another application. Please close it and try again.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save: {str(e)}")
                
        elif result is False:  # Save As
            new_path = filedialog.asksaveasfilename(
                title="Save Document As",
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
                initialfile=os.path.basename(self.document_path)
            )
            
            if new_path:
                try:
                    self.document.save(new_path)
                    self.document_path = new_path
                    self.doc_label_var.set(os.path.basename(new_path))
                    self.save_settings()
                    self.log_message(f"Document saved as: {new_path}")
                    messagebox.showinfo("Success", "Document saved successfully")
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to save: {str(e)}")
                    
    def run(self):
        """Run the application"""
        self.log_message("DoD Document Filler started")
        self.log_message("Load a Word document to begin")
        if self.selected_model.get():
            self.log_message(f"Configured model: {self.selected_model.get()}")
        self.root.mainloop()


if __name__ == "__main__": 
    try:
        app = DocxDocumentFiller()
        app.run()
    except Exception as e:
        print(f"Application error: {str(e)}")
        import traceback
        traceback.print_exc()
        input("Press Enter to exit...")