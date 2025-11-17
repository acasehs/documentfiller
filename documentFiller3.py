#!/usr/bin/env python3
"""
Document Content Generator - OpenWebUI Integration
Automatically fills Word document sections using OpenWebUI/Ollama with RAG support
Features: Markdown conversion, prompt learning, auto-features, model comparison
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, colorchooser
from docx import Document  # pip install python-docx
from docx.shared import Pt, RGBColor # pip install python-docx
from docx.enum.text import WD_COLOR_INDEX
import json
import requests  # pip install requests
import threading
import os
import re
from datetime import datetime
import shutil
import hashlib


class DocumentSection:
    """Represents a hierarchical document section"""
    def __init__(self, level, text, paragraph, full_path=""):
        self.level = level
        self.text = text
        self.paragraph = paragraph
        self.full_path = full_path
        self.children = []
        self.parent = None
        self.content_paragraphs = []
        
    def add_child(self, child):
        child.parent = self
        self.children.append(child)
        
    def get_full_path(self):
        """Get full hierarchical path"""
        if self.parent:
            parent_path = self.parent.get_full_path()
            return f"{parent_path} > {self.text}" if parent_path else self.text
        return self.text
        
    def get_existing_content(self):
        """Get existing text content in this section"""
        content = []
        for para in self.content_paragraphs:
            if para.text.strip():
                content.append(para.text)
        return "\n".join(content)
        
    def has_content(self):
        """Check if section has existing content"""
        return len(self.get_existing_content().strip()) > 0
    
    def get_section_hash(self):
        """Get unique hash for this section based on path"""
        return hashlib.md5(self.get_full_path().encode()).hexdigest()


class DocxDocumentFiller:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Document Content Generator - OpenWebUI Integration")
        self.root.geometry("1400x900")
        self.root.configure(bg="#2b2b2b")
        
        # Configure dark theme
        self.configure_dark_theme()
        
        # Data storage
        self.document = None
        self.document_path = None
        self.last_document_path = None
        self.sections = []
        self.selected_section = None
        self.generated_content = ""
        self.last_sent_prompt = ""
        
        # Section tracking
        self.section_tracking = {}
        self.tracking_file = None
        
        # OpenWebUI configuration
        self.openwebui_base_url = "http://172.16.27.122:3000"
        self.openwebui_api_key = ""
        self.available_models = []
        self.selected_model = tk.StringVar()
        self.available_knowledge_collections = []
        self.selected_knowledge_collections = []
        self.temperature = tk.DoubleVar(value=0.1)
        self.max_tokens = tk.IntVar(value=8000)
        
        # Formatting configuration
        self.format_config = {
            'highlight_enabled': tk.BooleanVar(value=True),
            'highlight_color': tk.StringVar(value='YELLOW'),
            'bold_enabled': tk.BooleanVar(value=False),
            'italic_enabled': tk.BooleanVar(value=False),
            'underline_enabled': tk.BooleanVar(value=False),
            'font_color': tk.StringVar(value='000000'),
            'font_size': tk.IntVar(value=11)
        }
        
        # Auto features configuration
        self.auto_config = {
            'auto_backup': tk.BooleanVar(value=True),
            'backup_interval': tk.IntVar(value=5),
            'auto_save': tk.BooleanVar(value=False),
            'auto_reload': tk.BooleanVar(value=True),
            'ask_backup': tk.BooleanVar(value=True)
        }
        
        # Operation mode
        self.operation_mode = tk.StringVar(value="replace")
        
        # Default master prompt
        self.master_prompt = tk.StringVar()
        self.set_default_prompt()
        
        # Load settings
        self.load_settings()
        
        # Create GUI
        self.create_gui()
        self.setup_text_widgets_colors()
        
        # Load last document if available
        if self.last_document_path and os.path.exists(self.last_document_path):
            self.load_document(self.last_document_path)
        
        # Start auto-backup timer if enabled
        if self.auto_config['auto_backup'].get():
            self.schedule_auto_backup()
        
    def configure_dark_theme(self):
        """Configure dark theme"""
        try:
            style = ttk.Style()
            available_themes = style.theme_names()
            
            if 'clam' in available_themes:
                style.theme_use('clam')
            
            style.configure('TLabel', background='#2b2b2b', foreground='#ffffff')
            style.configure('TButton', background='#404040', foreground='#ffffff')
            style.configure('TFrame', background='#2b2b2b')
            style.configure('TLabelframe', background='#2b2b2b', foreground='#ffffff')
            style.configure('TLabelframe.Label', background='#2b2b2b', foreground='#ffffff')
            style.configure('TRadiobutton', background='#2b2b2b', foreground='#ffffff')
            style.configure('TCheckbutton', background='#2b2b2b', foreground='#ffffff')
            
            style.configure('Treeview', 
                          background='#1e1e1e',
                          foreground='#ffffff',
                          fieldbackground='#1e1e1e',
                          bordercolor='#666666')
            style.map('Treeview', background=[('selected', '#404040')])
            
        except Exception as e:
            print(f"Dark theme configuration failed: {e}")
            
    def set_default_prompt(self):
        """Set default master prompt"""
        default = """You are a Department of Defense cybersecurity expert tasked with creating comprehensive, technically accurate documentation for a DoD system.

CONTEXT:
- This is for a Zero Trust Architecture implementation
- Must comply with DoD standards including RMF, NIST SP 800-53, FedRAMP High, and DoD IL5/IL6
- System is a COTS ERP solution in commercial cloud hosting
- Integrates with 250+ legacy systems
- Handles PII and requires strict data protection

REQUIREMENTS:
- Use specific technical terminology appropriate for DoD cybersecurity documentation
- Reference applicable NIST controls, DoD directives, and compliance frameworks
- Be detailed and comprehensive - this is for ATO documentation
- Use proper markdown formatting: **bold**, *italic*, `code`, bullet points, numbered lists
- Focus on implementation details, not just policy statements
- Include specific requirements for Zero Trust pillars where applicable

ZERO TRUST TOOLS:
- CyberArk PAM
- OKTA Secure Authentication
- Palo Alto Secure Browser 
- Palo Alto NGFWs
- Tenable Cloud Security
- Tenable SC
- Zscaler Private Access
- Zscaler Internet Access
- Trellix ePolicy Orchestrator
- Splunk
- Fortify
- SonarQube

CYBERSECURITY TEAM COMPOSITION
- Information Systems Security Officer (ISSOs) - Governance and Compliance, RMF Documentation
- Information Systems Security Engineers (ISSEs) - Security Engineering, STIG Integration, Security Impact Assessments
- Security Operations (SecOps) - System Scans, Policy Enforcement, Continuous Monitoring
- Identity, Credential, and Access Management (ICAM) - Credential Management, Access Auditing and Monitoring 
- Application Security (AppSec) - Application Vulnerability Testing, SAST, DAST, Penetration Testing

SECTION TO FILL: {section_name}
PARENT CONTEXT: {parent_context}
OPERATION: {operation_mode}"""
        
        self.master_prompt.set(default)
        
    def load_settings(self, config_file=None):
        """Load saved settings"""
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            
            # Use provided config file or default
            if not config_file:
                config_file = os.path.join(script_dir, "openwebui_config.json")
            
            if os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    config = json.load(f)
                    self.openwebui_base_url = config.get('base_url', 'http://172.16.27.122:3000')
                    self.openwebui_api_key = config.get('api_key', '')
                    self.selected_model.set(config.get('model', ''))
                    self.temperature.set(config.get('temperature', 0.1))
                    self.max_tokens.set(config.get('max_tokens', 8000))
                    
                    # Load RAG knowledge collections
                    if 'knowledge_collections' in config:
                        self.selected_knowledge_collections = config['knowledge_collections']
                    
                    # Load master prompt
                    if 'master_prompt' in config:
                        self.master_prompt.set(config.get('master_prompt'))
                    
                    if 'format_config' in config:
                        fmt = config['format_config']
                        self.format_config['highlight_enabled'].set(fmt.get('highlight_enabled', True))
                        self.format_config['highlight_color'].set(fmt.get('highlight_color', 'YELLOW'))
                        self.format_config['bold_enabled'].set(fmt.get('bold_enabled', False))
                        self.format_config['italic_enabled'].set(fmt.get('italic_enabled', False))
                        self.format_config['underline_enabled'].set(fmt.get('underline_enabled', False))
                        self.format_config['font_color'].set(fmt.get('font_color', '000000'))
                        self.format_config['font_size'].set(fmt.get('font_size', 11))
                    
                    if 'auto_config' in config:
                        auto = config['auto_config']
                        self.auto_config['auto_backup'].set(auto.get('auto_backup', True))
                        self.auto_config['backup_interval'].set(auto.get('backup_interval', 5))
                        self.auto_config['auto_save'].set(auto.get('auto_save', False))
                        self.auto_config['auto_reload'].set(auto.get('auto_reload', True))
                        self.auto_config['ask_backup'].set(auto.get('ask_backup', True))
            
            last_doc_file = os.path.join(script_dir, "last_document.txt")
            if os.path.exists(last_doc_file):
                with open(last_doc_file, 'r') as f:
                    self.last_document_path = f.read().strip()
                    
        except Exception as e:
            print(f"Note: Could not load settings: {str(e)}")

    def save_settings(self):
        """Save current settings"""
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            config_file = os.path.join(script_dir, "openwebui_config.json")
            
            # Create knowledge collections list with minimal info
            saved_collections = []
            for col in self.selected_knowledge_collections:
                saved_collections.append({
                    'id': col['id'],
                    'name': col['name']
                })
            
            config = {
                'base_url': self.openwebui_base_url,
                'api_key': self.openwebui_api_key,
                'model': self.selected_model.get(),
                'temperature': self.temperature.get(),
                'max_tokens': self.max_tokens.get(),
                'knowledge_collections': saved_collections,
                'master_prompt': self.master_prompt.get(),
                'format_config': {
                    'highlight_enabled': self.format_config['highlight_enabled'].get(),
                    'highlight_color': self.format_config['highlight_color'].get(),
                    'bold_enabled': self.format_config['bold_enabled'].get(),
                    'italic_enabled': self.format_config['italic_enabled'].get(),
                    'underline_enabled': self.format_config['underline_enabled'].get(),
                    'font_color': self.format_config['font_color'].get(),
                    'font_size': self.format_config['font_size'].get()
                },
                'auto_config': {
                    'auto_backup': self.auto_config['auto_backup'].get(),
                    'backup_interval': self.auto_config['backup_interval'].get(),
                    'auto_save': self.auto_config['auto_save'].get(),
                    'auto_reload': self.auto_config['auto_reload'].get(),
                    'ask_backup': self.auto_config['ask_backup'].get()
                }
            }
            
            with open(config_file, 'w') as f:
                json.dump(config, f, indent=2)
                
            if self.document_path:
                last_doc_file = os.path.join(script_dir, "last_document.txt")
                with open(last_doc_file, 'w') as f:
                    f.write(self.document_path)
                    
            self.log_message("Settings saved")
            
        except Exception as e:
            self.log_message(f"Error saving settings: {str(e)}")

    def load_section_tracking(self):
        """Load section tracking from file"""
        if not self.document_path:
            return
        
        try:
            doc_dir = os.path.dirname(self.document_path)
            doc_name = os.path.splitext(os.path.basename(self.document_path))[0]
            self.tracking_file = os.path.join(doc_dir, f".{doc_name}_tracking.json")
            
            if os.path.exists(self.tracking_file):
                with open(self.tracking_file, 'r') as f:
                    self.section_tracking = json.load(f)
                self.log_message(f"Loaded tracking for {len(self.section_tracking)} sections")
        except Exception as e:
            self.log_message(f"Note: Could not load section tracking: {str(e)}")
    
    def save_section_tracking(self):
        """Save section tracking to file"""
        if not self.tracking_file:
            return
        
        try:
            with open(self.tracking_file, 'w') as f:
                json.dump(self.section_tracking, f, indent=2)
        except Exception as e:
            self.log_message(f"Error saving section tracking: {str(e)}")
    
    def mark_section_edited(self, section):
        """Mark a section as edited in tracking"""
        section_hash = section.get_section_hash()
        self.section_tracking[section_hash] = {
            'edited': True,
            'last_modified': datetime.now().isoformat(),
            'section_path': section.get_full_path()
        }
        self.save_section_tracking()
    
    def is_section_edited(self, section):
        """Check if section has been edited"""
        section_hash = section.get_section_hash()
        return self.section_tracking.get(section_hash, {}).get('edited', False)
    
    def schedule_auto_backup(self):
        """Schedule automatic backup"""
        if self.auto_config['auto_backup'].get() and self.document_path:
            interval = self.auto_config['backup_interval'].get() * 60000  # Convert to ms
            self.root.after(interval, self.perform_auto_backup)
    
    def perform_auto_backup(self):
        """Perform automatic backup"""
        if self.document and self.document_path:
            try:
                self.create_backup()
                self.log_message("Auto-backup completed")
            except Exception as e:
                self.log_message(f"Auto-backup failed: {str(e)}")
        
        # Reschedule
        if self.auto_config['auto_backup'].get():
            self.schedule_auto_backup()

    def browse_config_file(self):
        """Browse for a different configuration file"""
        filename = filedialog.askopenfilename(
            title="Select Configuration File",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        if filename:
            try:
                self.load_settings(filename)
                self.update_config_status()
                self.log_message(f"Loaded configuration from: {filename}")
                messagebox.showinfo("Success", "Configuration loaded successfully")
            except Exception as e:
                self.log_message(f"Error loading configuration: {str(e)}")
                messagebox.showerror("Error", f"Failed to load configuration: {str(e)}")

    def save_config_as(self):
        """Save configuration to a new file"""
        filename = filedialog.asksaveasfilename(
            title="Save Configuration As",
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        if filename:
            try:
                script_dir = os.path.dirname(os.path.abspath(__file__))
                
                # Create knowledge collections list with minimal info
                saved_collections = []
                for col in self.selected_knowledge_collections:
                    saved_collections.append({
                        'id': col['id'],
                        'name': col['name']
                    })
                
                config = {
                    'base_url': self.openwebui_base_url,
                    'api_key': self.openwebui_api_key,
                    'model': self.selected_model.get(),
                    'temperature': self.temperature.get(),
                    'max_tokens': self.max_tokens.get(),
                    'knowledge_collections': saved_collections,
                    'master_prompt': self.master_prompt.get(),
                    'format_config': {
                        'highlight_enabled': self.format_config['highlight_enabled'].get(),
                        'highlight_color': self.format_config['highlight_color'].get(),
                        'bold_enabled': self.format_config['bold_enabled'].get(),
                        'italic_enabled': self.format_config['italic_enabled'].get(),
                        'underline_enabled': self.format_config['underline_enabled'].get(),
                        'font_color': self.format_config['font_color'].get(),
                        'font_size': self.format_config['font_size'].get()
                    },
                    'auto_config': {
                        'auto_backup': self.auto_config['auto_backup'].get(),
                        'backup_interval': self.auto_config['backup_interval'].get(),
                        'auto_save': self.auto_config['auto_save'].get(),
                        'auto_reload': self.auto_config['auto_reload'].get(),
                        'ask_backup': self.auto_config['ask_backup'].get()
                    }
                }
                
                with open(filename, 'w') as f:
                    json.dump(config, f, indent=2)
                
                self.log_message(f"Configuration saved to: {filename}")
                messagebox.showinfo("Success", "Configuration saved successfully")
            except Exception as e:
                self.log_message(f"Error saving configuration: {str(e)}")
                messagebox.showerror("Error", f"Failed to save configuration: {str(e)}")

    def setup_text_widgets_colors(self):
        """Configure cursor and selection colors for text widgets"""
        # Configure all ScrolledText widgets to have visible cursors
        for widget in [self.existing_text, self.generated_text, self.prompt_text, 
                    self.console, self.prompt_history_text]:
            widget.config(insertbackground="#ffffff")  # White cursor
            widget.config(selectbackground="#3a5f8f")  # Blue selection background
            widget.config(selectforeground="#ffffff")  # White selection text

    def create_gui(self):
        """Create main GUI"""
        main_container = ttk.Frame(self.root, padding="10")
        main_container.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_container.columnconfigure(0, weight=1)
        main_container.rowconfigure(2, weight=1)
        
        # Title row with buttons
        title_frame = ttk.Frame(main_container)
        title_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        title_frame.columnconfigure(0, weight=1)
        
        ttk.Label(title_frame, text="Document Content Generator - OpenWebUI", 
                font=("Arial", 16, "bold")).pack(side=tk.LEFT)
        
        # Configuration buttons on title row
        btn_container = ttk.Frame(title_frame)
        btn_container.pack(side=tk.RIGHT)

        ttk.Button(btn_container, text="⚙ Config Files", 
                command=self.open_config_file_manager).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_container, text="⚙ Formatting", 
                command=self.open_formatting_dialog).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_container, text="🔄 Auto Features", 
                command=self.open_auto_features_dialog).pack(side=tk.LEFT, padx=2)
        
        # Top controls
        top_frame = ttk.Frame(main_container)
        top_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        top_frame.columnconfigure(1, weight=1)
        
        # Document selection
        doc_frame = ttk.LabelFrame(top_frame, text="Document", padding="10")
        doc_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        doc_frame.columnconfigure(1, weight=1)
        
        ttk.Button(doc_frame, text="Load Document", command=self.browse_document).grid(row=0, column=0, padx=(0, 5))
        self.doc_label_var = tk.StringVar(value="No document loaded")
        ttk.Label(doc_frame, textvariable=self.doc_label_var).grid(row=0, column=1, sticky=tk.W, padx=5)
        
        # OpenWebUI configuration (AI only)
        config_frame = ttk.LabelFrame(top_frame, text="OpenWebUI Configuration", padding="10")
        config_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        ttk.Button(config_frame, text="Configure AI Settings", command=self.open_config_dialog).grid(row=0, column=0, padx=(0, 10))
        self.config_status_var = tk.StringVar(value="Not configured")
        ttk.Label(config_frame, textvariable=self.config_status_var).grid(row=0, column=1, sticky=tk.W)
        
        # Main content area
        content_frame = ttk.Frame(main_container)
        content_frame.grid(row=2, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        content_frame.columnconfigure(0, weight=1)
        content_frame.columnconfigure(1, weight=2)
        content_frame.rowconfigure(0, weight=1)
        
        # Left panel
        left_panel = ttk.Frame(content_frame)
        left_panel.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        left_panel.columnconfigure(0, weight=1)
        left_panel.rowconfigure(1, weight=1)
        
        # Section tree
        tree_frame = ttk.LabelFrame(left_panel, text="Document Sections", padding="5")
        tree_frame.grid(row=0, column=0, rowspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        tree_frame.columnconfigure(0, weight=1)
        tree_frame.rowconfigure(0, weight=1)
        
        self.tree = ttk.Treeview(tree_frame, selectmode='browse')
        tree_scroll = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=tree_scroll.set)
        
        self.tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        tree_scroll.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        self.tree.bind('<<TreeviewSelect>>', self.on_section_select)
        
        # Operation controls
        op_frame = ttk.LabelFrame(left_panel, text="Operation Mode", padding="10")
        op_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        
        ttk.Radiobutton(op_frame, text="Replace (write from scratch)", 
                    variable=self.operation_mode, value="replace").pack(anchor=tk.W)
        ttk.Radiobutton(op_frame, text="Rework (rewrite existing)", 
                    variable=self.operation_mode, value="rework").pack(anchor=tk.W)
        ttk.Radiobutton(op_frame, text="Append (add to existing)", 
                    variable=self.operation_mode, value="append").pack(anchor=tk.W)
        
        # Action buttons
        button_frame = ttk.Frame(left_panel)
        button_frame.grid(row=3, column=0, sticky=(tk.W, tk.E))
        
        self.generate_btn = ttk.Button(button_frame, text="Generate Content", 
                                    command=self.generate_content, state='disabled')
        self.generate_btn.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(button_frame, text="🚀 Auto Complete Document", 
                command=self.auto_complete_document).pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(button_frame, text="Edit Master Prompt", 
                command=self.edit_master_prompt).pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(button_frame, text="📚 Prompt Library", 
        command=self.open_prompt_library).pack(fill=tk.X, pady=(0, 10))
        
        # Right panel
        right_panel = ttk.Frame(content_frame)
        right_panel.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S))
        right_panel.columnconfigure(0, weight=1)
        right_panel.rowconfigure(0, weight=1)
        
        # Create notebook for tabs
        self.notebook = ttk.Notebook(right_panel)
        self.notebook.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Bind tab change event
        self.notebook.bind('<<NotebookTabChanged>>', self.on_tab_changed)
        
        # Preview tab
        self.create_preview_tab()
        
        # Prompt tab
        self.create_prompt_tab()
        
        # NEW: Prompt History tab
        self.create_prompt_history_tab()
        
        # Console tab
        self.create_console_tab()
        
        # Update config status
        self.update_config_status()   

    def create_prompt_history_tab(self):
        """Create prompt history tab with navigation sidebar"""
        history_frame = ttk.Frame(self.notebook)
        self.notebook.add(history_frame, text="Prompt History")
        history_frame.columnconfigure(0, weight=3)  # Main content area gets more space
        history_frame.columnconfigure(1, weight=1)  # Shortcut sidebar
        history_frame.rowconfigure(0, weight=1)
        
        # Create left side - Main history content
        main_frame = ttk.Frame(history_frame)
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 2), pady=5)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(0, weight=1)
        
        # Create text widget with scrollbars for main content
        self.prompt_history_text = scrolledtext.ScrolledText(main_frame, 
                                                        bg="#1e1e1e", fg="#ffffff",
                                                        font=("Consolas", 10), wrap=tk.WORD)
        self.prompt_history_text.pack(fill=tk.BOTH, expand=True)
        
        # Create right side - Shortcuts list
        shortcut_frame = ttk.LabelFrame(history_frame, text="Navigation Shortcuts")
        shortcut_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(2, 5), pady=5)
        shortcut_frame.columnconfigure(0, weight=1)
        shortcut_frame.rowconfigure(0, weight=1)
        
        # Shortcuts list with scrollbar
        shortcuts_container = ttk.Frame(shortcut_frame)
        shortcuts_container.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        shortcut_scrollbar = ttk.Scrollbar(shortcuts_container)
        shortcut_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.shortcut_list = tk.Listbox(shortcuts_container, 
                                    bg="#1e1e1e", fg="#ffffff", 
                                    font=("Consolas", 9),
                                    yscrollcommand=shortcut_scrollbar.set)
        self.shortcut_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        shortcut_scrollbar.config(command=self.shortcut_list.yview)
        
        # Bind selection event to jump to position
        self.shortcut_list.bind('<<ListboxSelect>>', self.jump_to_prompt_position)
        
        # Store prompt positions (for navigation)
        self.prompt_positions = []
        
        # Create tags for formatting
        self.prompt_history_text.tag_configure("sent", foreground="#90EE90", justify="right")
        self.prompt_history_text.tag_configure("received", foreground="#87CEFA", justify="left")
        self.prompt_history_text.tag_configure("timestamp", foreground="#888888", justify="center")
        self.prompt_history_text.tag_configure("header", font=("Consolas", 10, "bold"))
        
        # Add clear button at bottom
        btn_frame = ttk.Frame(shortcut_frame)
        btn_frame.pack(fill=tk.X, side=tk.BOTTOM, padx=5, pady=5)
        
        ttk.Button(btn_frame, text="Clear History", 
                command=self.clear_prompt_history).pack(fill=tk.X)
        ttk.Button(btn_frame, text="Export History", 
                command=self.export_prompt_history).pack(fill=tk.X, pady=(5, 0))

    def jump_to_prompt_position(self, event):
        """Jump to the selected prompt position in history"""
        if not hasattr(self, 'shortcut_list') or not self.shortcut_list.curselection():
            return
            
        index = self.shortcut_list.curselection()[0]
        if index < len(self.prompt_positions):
            position = self.prompt_positions[index]
            
            # Scroll to position
            self.prompt_history_text.see(position)
            
            # Highlight the prompt temporarily
            self.prompt_history_text.tag_remove("highlight", "1.0", tk.END)
            
            # Calculate the end position (approximate)
            line_num = int(position.split('.')[0])
            self.prompt_history_text.tag_add("highlight", position, f"{line_num+2}.0")
            
            # Configure highlight
            self.prompt_history_text.tag_configure("highlight", background="#404040")
            
            # Schedule removal of highlight
            self.root.after(2000, lambda: self.prompt_history_text.tag_remove("highlight", "1.0", tk.END))

    def clear_prompt_history(self):
        """Clear the prompt history"""
        if hasattr(self, 'prompt_history_text'):
            confirm = messagebox.askyesno("Confirm", "Clear prompt history? This cannot be undone.")
            if confirm:
                self.prompt_history_text.delete("1.0", tk.END)
                self.shortcut_list.delete(0, tk.END)
                self.prompt_positions = []
                self.log_message("Prompt history cleared")

    def export_prompt_history(self):
        """Export prompt history to a file"""
        if not hasattr(self, 'prompt_history_text'):
            return
            
        history_text = self.prompt_history_text.get("1.0", tk.END)
        if not history_text.strip():
            messagebox.showinfo("Info", "Prompt history is empty.")
            return
            
        filename = filedialog.asksaveasfilename(
            title="Export Prompt History",
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
        )
        
        if filename:
            try:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(history_text)
                self.log_message(f"Prompt history exported to: {filename}")
                messagebox.showinfo("Success", f"Prompt history exported to: {filename}")
            except Exception as e:
                self.log_message(f"Error exporting history: {str(e)}")
                messagebox.showerror("Error", f"Failed to export history: {str(e)}")

    def log_prompt_history(self, prompt, response=None, is_sent=True):
        """Log prompts and responses to the prompt history tab"""
        if not hasattr(self, 'prompt_history_text') or not hasattr(self, 'shortcut_list'):
            return
        
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Get current position for shortcut navigation
        current_position = self.prompt_history_text.index(tk.END)
        
        # Add to shortcut list
        shortcut_text = f"[{timestamp}] " + ("SENT" if is_sent else "RECEIVED")
        if is_sent and len(prompt) > 30:
            # Add first part of prompt for context in shortcut
            shortcut_text += f": {prompt[:30]}..."
        
        self.shortcut_list.insert(tk.END, shortcut_text)
        self.prompt_positions.append(current_position)
        
        # Add timestamp
        self.prompt_history_text.insert(tk.END, f"\n[{timestamp}]\n", "timestamp")
        
        if is_sent:
            # Add sent prompt header
            self.prompt_history_text.insert(tk.END, "PROMPT SENT:\n", "header")
            
            # Add actual prompt with right justification and green color
            self.prompt_history_text.insert(tk.END, f"{prompt}\n\n", "sent")
        else:
            # Add received response header
            self.prompt_history_text.insert(tk.END, "RESPONSE RECEIVED:\n", "header")
            
            # Add actual response with left justification and blue color
            self.prompt_history_text.insert(tk.END, f"{response}\n\n", "received")
        
        # Scroll to the end
        self.prompt_history_text.see(tk.END)
        
        # Keep shortcut list scrolled to bottom as well
        self.shortcut_list.see(tk.END)

    def reset_config_to_defaults(self):
        """Reset configuration to default values"""
        result = messagebox.askyesno(
            "Reset Configuration",
            "Are you sure you want to reset all settings to defaults?\n\n"
            "This will reset:\n"
            "- AI connection settings\n"
            "- Model settings\n"
            "- Formatting options\n"
            "- Auto features\n\n"
            "This action cannot be undone."
        )
        
        if result:
            # Reset OpenWebUI settings
            self.openwebui_base_url = "http://172.16.27.122:3000"
            self.openwebui_api_key = ""
            self.selected_model.set("")
            self.available_models = []
            self.selected_knowledge_collections = []
            self.available_knowledge_collections = []
            self.temperature.set(0.1)
            self.max_tokens.set(8000)
            
            # Reset master prompt
            self.set_default_prompt()
            
            # Reset formatting
            self.format_config['highlight_enabled'].set(True)
            self.format_config['highlight_color'].set('YELLOW')
            self.format_config['bold_enabled'].set(False)
            self.format_config['italic_enabled'].set(False)
            self.format_config['underline_enabled'].set(False)
            self.format_config['font_color'].set('000000')
            self.format_config['font_size'].set(11)
            
            # Reset auto features
            self.auto_config['auto_backup'].set(True)
            self.auto_config['backup_interval'].set(5)
            self.auto_config['auto_save'].set(False)
            self.auto_config['auto_reload'].set(True)
            self.auto_config['ask_backup'].set(True)
            
            # Save reset configuration
            self.save_settings()
            
            # Update UI
            self.update_config_status()
            
            self.log_message("Configuration reset to defaults")
            messagebox.showinfo("Reset Complete", "Configuration has been reset to defaults")

    def open_config_file_manager(self):
        """Open configuration file manager dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Configuration Files")
        dialog.geometry("500x300")
        dialog.configure(bg="#2b2b2b")
        dialog.grab_set()
        
        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Configuration File Management", 
                font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 15))
        
        # Current file
        script_dir = os.path.dirname(os.path.abspath(__file__))
        current_config = os.path.join(script_dir, "openwebui_config.json")
        
        if os.path.exists(current_config):
            config_size = os.path.getsize(current_config)
            config_date = datetime.fromtimestamp(os.path.getmtime(current_config)).strftime("%Y-%m-%d %H:%M")
            config_info = f"Current: {os.path.basename(current_config)} ({config_size/1024:.1f} KB, {config_date})"
        else:
            config_info = "No configuration file found"
        
        ttk.Label(main_frame, text=config_info).pack(anchor=tk.W, pady=(0, 15))
        
        # Action buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(0, 15))
        
        ttk.Button(btn_frame, text="Load Configuration", 
                command=self.browse_config_file, width=20).pack(pady=5)
        
        ttk.Button(btn_frame, text="Save Configuration As", 
                command=self.save_config_as, width=20).pack(pady=5)
        
        ttk.Button(btn_frame, text="Reset to Defaults", 
                command=self.reset_config_to_defaults, width=20).pack(pady=5)
        
        # Close button
        ttk.Button(main_frame, text="Close", 
                command=dialog.destroy).pack(side=tk.RIGHT, pady=(15, 0))

    # Add new function for formatting dialog
    def open_formatting_dialog(self):
        """Open formatting configuration dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Formatting Options")
        dialog.geometry("500x400")
        dialog.configure(bg="#2b2b2b")
        dialog.grab_set()
        
        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Content Formatting Options", 
                font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 10))
        
        # Formatting options
        format_frame = ttk.LabelFrame(main_frame, text="Text Formatting", padding="10")
        format_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Checkbutton(format_frame, text="Highlight Changes", 
                    variable=self.format_config['highlight_enabled']).pack(anchor=tk.W)
        
        color_frame = ttk.Frame(format_frame)
        color_frame.pack(fill=tk.X, pady=(5, 10))
        ttk.Label(color_frame, text="Highlight Color:").pack(side=tk.LEFT)
        color_options = ['YELLOW', 'GREEN', 'CYAN', 'PINK', 'BRIGHT_GREEN']
        color_combo = ttk.Combobox(color_frame, textvariable=self.format_config['highlight_color'],
                                values=color_options, state="readonly", width=15)
        color_combo.pack(side=tk.LEFT, padx=(5, 0))
        
        ttk.Checkbutton(format_frame, text="Bold", 
                    variable=self.format_config['bold_enabled']).pack(anchor=tk.W)
        ttk.Checkbutton(format_frame, text="Italic", 
                    variable=self.format_config['italic_enabled']).pack(anchor=tk.W)
        ttk.Checkbutton(format_frame, text="Underline", 
                    variable=self.format_config['underline_enabled']).pack(anchor=tk.W)
        
        font_frame = ttk.Frame(format_frame)
        font_frame.pack(fill=tk.X, pady=(5, 0))
        ttk.Label(font_frame, text="Font Size:").pack(side=tk.LEFT)
        ttk.Spinbox(font_frame, from_=8, to=16, textvariable=self.format_config['font_size'],
                width=10).pack(side=tk.LEFT, padx=(5, 0))
        
        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        def save_and_close():
            self.save_settings()
            self.log_message("Formatting options saved")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="Save", command=save_and_close).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT)

    # Add new function for auto features dialog
    def open_auto_features_dialog(self):
        """Open auto features configuration dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Automatic Features")
        dialog.geometry("500x350")
        dialog.configure(bg="#2b2b2b")
        dialog.grab_set()
        
        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Automatic Features Configuration", 
                font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 10))
        
        # Auto features
        auto_frame = ttk.LabelFrame(main_frame, text="Auto Features", padding="10")
        auto_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Checkbutton(auto_frame, text="Auto Backup", 
                    variable=self.auto_config['auto_backup']).pack(anchor=tk.W)
        
        backup_frame = ttk.Frame(auto_frame)
        backup_frame.pack(fill=tk.X, pady=(5, 10))
        ttk.Label(backup_frame, text="Backup Interval (minutes):").pack(side=tk.LEFT)
        ttk.Spinbox(backup_frame, from_=1, to=60, textvariable=self.auto_config['backup_interval'],
                width=10).pack(side=tk.LEFT, padx=(5, 0))
        
        ttk.Checkbutton(auto_frame, text="Auto Save After Commit", 
                    variable=self.auto_config['auto_save']).pack(anchor=tk.W)
        ttk.Checkbutton(auto_frame, text="Auto Reload Sections After Commit", 
                    variable=self.auto_config['auto_reload']).pack(anchor=tk.W)
        ttk.Checkbutton(auto_frame, text="Ask Before Creating Backup", 
                    variable=self.auto_config['ask_backup']).pack(anchor=tk.W)
        
        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        def save_and_close():
            self.save_settings()
            self.log_message("Auto features configuration saved")
            
            # Restart auto-backup if settings changed
            if self.auto_config['auto_backup'].get():
                self.schedule_auto_backup()
            
            dialog.destroy()
        
        ttk.Button(btn_frame, text="Save", command=save_and_close).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT)


    # Add this new function for tab change event
    def on_tab_changed(self, event):
        """Handle notebook tab changes"""
        selected_tab = event.widget.select()
        tab_text = event.widget.tab(selected_tab, "text")
        
        # If switching to Prompt Editor tab, refresh model list
        if tab_text == "Prompt Editor":
            self.update_prompt_model_selector()

    def create_preview_tab(self):
        """Create content preview tab"""
        preview_frame = ttk.Frame(self.notebook)
        self.notebook.add(preview_frame, text="Content Preview")
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.columnconfigure(1, weight=1)
        preview_frame.rowconfigure(1, weight=1)
        
        # Existing content
        ttk.Label(preview_frame, text="Existing Content", font=("Arial", 10, "bold")).grid(
            row=0, column=0, sticky=tk.W, padx=5, pady=5)
        
        self.existing_text = scrolledtext.ScrolledText(preview_frame, height=20, 
                                                      bg="#1e1e1e", fg="#ffffff",
                                                      font=("Consolas", 10), wrap=tk.WORD)
        self.existing_text.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 2))
        self.existing_text.config(state='disabled')
        
        # Generated content
        ttk.Label(preview_frame, text="Generated Content (Editable)", 
                 font=("Arial", 10, "bold")).grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        
        self.generated_text = scrolledtext.ScrolledText(preview_frame, height=20,
                                                       bg="#1e1e1e", fg="#90EE90",
                                                       font=("Consolas", 10), wrap=tk.WORD)
        self.generated_text.grid(row=1, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(2, 5))
        
        # Setup text tags for markdown rendering
        self.setup_text_tags()
        
        # Buttons
        commit_frame = ttk.Frame(preview_frame)
        commit_frame.grid(row=2, column=0, columnspan=2, pady=10)
        
        self.commit_btn = ttk.Button(commit_frame, text="Commit to Document", 
                                    command=self.commit_content, state='disabled')
        self.commit_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(commit_frame, text="Clear Preview", 
                  command=self.clear_preview).pack(side=tk.LEFT, padx=5)
    
    def create_prompt_tab(self):
        """Create prompt editor tab"""
        prompt_frame = ttk.Frame(self.notebook)
        self.notebook.add(prompt_frame, text="Prompt Editor")
        prompt_frame.columnconfigure(0, weight=1)
        prompt_frame.rowconfigure(1, weight=1)
        
        ttk.Label(prompt_frame, text="Last Sent Prompt (Edit and Regenerate)", 
                 font=("Arial", 10, "bold")).pack(anchor=tk.W, padx=5, pady=5)
        
        self.prompt_text = scrolledtext.ScrolledText(prompt_frame, height=20,
                                                     bg="#1e1e1e", fg="#ffffff",
                                                     font=("Consolas", 10), wrap=tk.WORD)
        self.prompt_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        prompt_btn_frame = ttk.Frame(prompt_frame)
        prompt_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(prompt_btn_frame, text="Model:").pack(side=tk.LEFT, padx=(0, 5))
        self.prompt_model_var = tk.StringVar()
        self.prompt_model_combo = ttk.Combobox(prompt_btn_frame, textvariable=self.prompt_model_var,
                                               state="readonly", width=30)
        self.prompt_model_combo.pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(prompt_btn_frame, text="Regenerate", 
                  command=self.regenerate_with_prompt).pack(side=tk.LEFT, padx=5)
        ttk.Button(prompt_btn_frame, text="Compare 3 Models", 
                  command=self.compare_three_models).pack(side=tk.LEFT, padx=5)
        ttk.Button(prompt_btn_frame, text="Update Master Prompt", 
                  command=self.offer_prompt_update).pack(side=tk.LEFT, padx=5)
    
    def create_console_tab(self):
        """Create console log tab"""
        console_frame = ttk.Frame(self.notebook)
        self.notebook.add(console_frame, text="Console Log")
        console_frame.columnconfigure(0, weight=1)
        console_frame.rowconfigure(0, weight=1)
        
        self.console = scrolledtext.ScrolledText(console_frame, height=20,
                                                bg="#1e1e1e", fg="#ffffff",
                                                font=("Consolas", 9))
        self.console.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5, pady=5)
    
    def setup_text_tags(self):
        """Setup text widget tags for markdown formatting"""
        self.generated_text.tag_configure("bold", font=("Consolas", 10, "bold"))
        self.generated_text.tag_configure("italic", font=("Consolas", 10, "italic"))
        self.generated_text.tag_configure("bold_italic", font=("Consolas", 10, "bold italic"))
        self.generated_text.tag_configure("code", background="#333333", foreground="#00ff00")
        self.generated_text.tag_configure("heading", font=("Consolas", 12, "bold"), foreground="#00aaff")
        
    def log_message(self, message):
        """Log message to console"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.console.insert(tk.END, f"[{timestamp}] {message}\n")
        self.console.see(tk.END)
        self.root.update()
        
    def browse_document(self):
        """Browse for Word document"""
        filename = filedialog.askopenfilename(
            title="Select Word Document",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            self.load_document(filename)
            
    def load_document(self, filepath):
        """Load Word document and parse sections"""
        try:
            self.log_message(f"Loading document: {filepath}")
            self.document = Document(filepath)
            self.document_path = filepath
            self.doc_label_var.set(os.path.basename(filepath))
            
            # Load section tracking
            self.load_section_tracking()
            
            # Parse document structure
            self.parse_document_structure()
            
            # Populate tree
            self.populate_tree()
            
            # Save as last document
            self.save_settings()
            
            self.log_message(f"Document loaded: {len(self.sections)} sections found")
            
        except PermissionError:
            self.log_message("Error: File is open in another application")
            messagebox.showerror("Error", "File is open. Please close it and try again.")
        except Exception as e:
            self.log_message(f"Error loading document: {str(e)}")
            messagebox.showerror("Error", f"Failed to load document: {str(e)}")
            
    def parse_document_structure(self):
        """Parse document into hierarchical structure"""
        self.sections = []
        section_stack = []
        current_section = None
        
        for para in self.document.paragraphs:
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                    if level <= 4:
                        section = DocumentSection(level, para.text.strip(), para)
                        
                        while section_stack and section_stack[-1].level >= level:
                            section_stack.pop()
                        
                        if section_stack:
                            section_stack[-1].add_child(section)
                        else:
                            self.sections.append(section)
                        
                        section_stack.append(section)
                        current_section = section
                        
                except ValueError:
                    pass
            else:
                if current_section and para.text.strip():
                    if not para.text.strip().startswith('-'):
                        current_section.content_paragraphs.append(para)
                        
    def populate_tree(self):
        """Populate treeview with document sections"""
        self.tree.delete(*self.tree.get_children())
        
        def add_to_tree(section, parent=''):
            is_edited = self.is_section_edited(section)
            
            display_text = section.text
            if is_edited:
                display_text += " ✓"
            
            item_id = self.tree.insert(parent, 'end', text=display_text, 
                                      values=(id(section),))
            
            for child in section.children:
                add_to_tree(child, item_id)
        
        for section in self.sections:
            add_to_tree(section)
            
    def on_section_select(self, event):
        """Handle section selection"""
        selection = self.tree.selection()
        if selection:
            item = selection[0]
            section_id = self.tree.item(item, 'values')[0]
            
            self.selected_section = self.find_section_by_id(int(section_id))
            
            if self.selected_section:
                self.generate_btn.config(state='normal')
                self.show_existing_content()
                self.log_message(f"Selected: {self.selected_section.get_full_path()}")
                
    def find_section_by_id(self, section_id):
        """Find section by its ID"""
        def search(sections):
            for section in sections:
                if id(section) == section_id:
                    return section
                result = search(section.children)
                if result:
                    return result
            return None
        
        return search(self.sections)
        
    def show_existing_content(self):
        """Show existing content in preview"""
        if self.selected_section:
            self.existing_text.config(state='normal')
            self.existing_text.delete('1.0', tk.END)
            
            existing = self.selected_section.get_existing_content()
            if existing:
                self.existing_text.insert('1.0', existing)
            else:
                self.existing_text.insert('1.0', "[No existing content]")
            
            self.existing_text.config(state='disabled')
            
    # Replace the open_config_dialog to only show AI settings

    def save_master_prompt(self, filename=None):
        """Save master prompt to a file"""
        if not filename:
            filename = filedialog.asksaveasfilename(
                title="Save Master Prompt",
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
            )
        
        if not filename:
            return False
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(self.master_prompt.get())
            
            self.log_message(f"Master prompt saved to: {filename}")
            return True
        except Exception as e:
            self.log_message(f"Error saving master prompt: {str(e)}")
            messagebox.showerror("Error", f"Failed to save master prompt: {str(e)}")
            return False

    def load_master_prompt(self):
        """Load master prompt from a file"""
        filename = filedialog.askopenfilename(
            title="Load Master Prompt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
        )
        
        if not filename:
            return False
        
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                new_prompt = f.read()
            
            if not new_prompt:
                messagebox.showwarning("Warning", "The selected file is empty.")
                return False
            
            result = messagebox.askyesno(
                "Load Master Prompt",
                "Are you sure you want to replace the current master prompt?\n\n"
                "This action cannot be undone."
            )
            
            if result:
                self.master_prompt.set(new_prompt)
                self.log_message(f"Master prompt loaded from: {filename}")
                return True
            
            return False
        except Exception as e:
            self.log_message(f"Error loading master prompt: {str(e)}")
            messagebox.showerror("Error", f"Failed to load master prompt: {str(e)}")
            return False

    def refresh_all_ai_resources(self, url, key, model_combo, knowledge_list, status_label):
        """Refresh all AI resources (models and collections)"""
        if not url or not key:
            status_label.config(text="⚠ Please enter URL and API key")
            return
        
        status_label.config(text="🔄 Refreshing all resources...")
        
        # Use threading to avoid freezing the UI
        def refresh_thread():
            try:
                # Test connection first
                success = self.test_connection(url, key, status_label)
                
                if success:
                    # Refresh models
                    self.root.after(0, lambda: status_label.config(text="🔄 Loading models..."))
                    headers = {'Authorization': f'Bearer {key}'}
                    response = requests.get(f"{url}/api/models", headers=headers, timeout=10)
                    
                    if response.status_code == 200:
                        data = response.json()
                        self.available_models = []
                        
                        if isinstance(data, dict) and 'data' in data:
                            models_list = data['data']
                        elif isinstance(data, list):
                            models_list = data
                        else:
                            models_list = [data] if data else []
                        
                        for model in models_list:
                            if isinstance(model, dict):
                                model_id = model.get('id', model.get('name', 'unknown'))
                                self.available_models.append(model_id)
                            else:
                                self.available_models.append(str(model))
                        
                        self.root.after(0, lambda: model_combo.config(values=self.available_models))
                        self.log_message(f"Refreshed models: {len(self.available_models)} found")
                        
                        # Update selection if possible
                        if self.selected_model.get() in self.available_models:
                            self.root.after(0, lambda: model_combo.set(self.selected_model.get()))
                        elif self.available_models:
                            self.root.after(0, lambda: model_combo.set(self.available_models[0]))
                    
                    # Refresh collections
                    self.root.after(0, lambda: status_label.config(text="🔄 Loading collections..."))
                    response = requests.get(f"{url}/api/v1/knowledge/", headers=headers, timeout=10)
                    
                    if response.status_code == 200:
                        data = response.json()
                        self.available_knowledge_collections = []
                        
                        if isinstance(data, list):
                            for collection in data:
                                if isinstance(collection, dict):
                                    self.available_knowledge_collections.append({
                                        'id': collection.get('id', 'unknown'),
                                        'name': collection.get('name', 'Unknown'),
                                        'description': collection.get('description', '')
                                    })
                        
                        # Update list
                        self.root.after(0, lambda: knowledge_list.delete(0, tk.END))
                        for col in self.available_knowledge_collections:
                            display = f"{col['name']} ({col['id']})"
                            self.root.after(0, lambda d=display: knowledge_list.insert(tk.END, d))
                        
                        # Restore selections
                        for i, col in enumerate(self.available_knowledge_collections):
                            for sel in self.selected_knowledge_collections:
                                if col['id'] == sel.get('id', ''):
                                    self.root.after(0, lambda idx=i: knowledge_list.selection_set(idx))
                        
                        self.log_message(f"Refreshed collections: {len(self.available_knowledge_collections)} found")
                    
                    self.root.after(0, lambda: status_label.config(text="✓ All resources refreshed"))
                else:
                    self.root.after(0, lambda: status_label.config(text="✗ Connection failed"))
                    
            except Exception as e:
                self.log_message(f"Error refreshing resources: {str(e)}")
                self.root.after(0, lambda err=str(e): status_label.config(text=f"✗ Error: {err[:20]}..."))
        
        thread = threading.Thread(target=refresh_thread)
        thread.daemon = True
        thread.start()

    def auto_refresh_resources(self, url, key, model_combo, knowledge_list, status_label):
        """Automatically refresh models and knowledge collections on dialog open"""
        if not url or not key:
            status_label.config(text="⚠ Please enter URL and API key")
            return
            
        # Test connection first
        success = self.test_connection(url, key, status_label)
        
        if success:
            # Refresh models and collections in sequence
            self.refresh_models(url, key, model_combo)
            self.refresh_knowledge(url, key, knowledge_list)
            status_label.config(text="✓ Resources loaded")

    def open_config_dialog(self):
        """Open OpenWebUI AI configuration dialog"""
        config_window = tk.Toplevel(self.root)
        config_window.title("OpenWebUI AI Configuration")
        config_window.geometry("700x1000")
        config_window.configure(bg="#2b2b2b")
        config_window.grab_set()
        
        main_frame = ttk.Frame(config_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Configuration Files section
        config_file_frame = ttk.LabelFrame(main_frame, text="Configuration Files", padding="10")
        config_file_frame.pack(fill=tk.X, pady=(0, 10))
        
        config_btn_frame = ttk.Frame(config_file_frame)
        config_btn_frame.pack(fill=tk.X)
        
        ttk.Button(config_btn_frame, text="Load Config", 
                command=self.browse_config_file).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(config_btn_frame, text="Save Config As", 
                command=self.save_config_as).pack(side=tk.LEFT)
        
        # Connection settings
        conn_frame = ttk.LabelFrame(main_frame, text="Connection", padding="10")
        conn_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(conn_frame, text="Base URL:").pack(anchor=tk.W)
        url_var = tk.StringVar(value=self.openwebui_base_url)
        ttk.Entry(conn_frame, textvariable=url_var, width=50).pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(conn_frame, text="API Key:").pack(anchor=tk.W)
        key_var = tk.StringVar(value=self.openwebui_api_key)
        ttk.Entry(conn_frame, textvariable=key_var, width=50, show="*").pack(fill=tk.X, pady=(0, 10))
        
        test_frame = ttk.Frame(conn_frame)
        test_frame.pack(fill=tk.X)
        
        status_label = ttk.Label(test_frame, text="")
        
        # Add Refresh All button
        ttk.Button(test_frame, text="Refresh All", 
                command=lambda: self.refresh_all_ai_resources(url_var.get(), key_var.get(), model_combo, knowledge_list, status_label)).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(test_frame, text="Test Connection", 
                command=lambda: self.test_connection(url_var.get(), key_var.get(), status_label)).pack(side=tk.LEFT)
        
        status_label.pack(side=tk.LEFT, padx=(10, 0))
        
        # Model selection
        model_frame = ttk.LabelFrame(main_frame, text="Model", padding="10")
        model_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(model_frame, text="Refresh Models", 
                command=lambda: self.refresh_models(url_var.get(), key_var.get(), model_combo)).pack(anchor=tk.W, pady=(0, 5))
        
        ttk.Label(model_frame, text="Select Default Model:").pack(anchor=tk.W)
        model_combo = ttk.Combobox(model_frame, textvariable=self.selected_model, state="readonly")
        model_combo['values'] = self.available_models
        model_combo.pack(fill=tk.X)
        
        # Knowledge collections
        knowledge_frame = ttk.LabelFrame(main_frame, text="RAG Knowledge Collections", padding="10")
        knowledge_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        ttk.Button(knowledge_frame, text="Refresh Collections", 
                command=lambda: self.refresh_knowledge(url_var.get(), key_var.get(), knowledge_list)).pack(anchor=tk.W, pady=(0, 5))
        
        ttk.Label(knowledge_frame, text="Select Collections (Ctrl for multiple):").pack(anchor=tk.W)
        
        knowledge_scroll = ttk.Scrollbar(knowledge_frame)
        knowledge_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        knowledge_list = tk.Listbox(knowledge_frame, selectmode=tk.MULTIPLE, 
                                    yscrollcommand=knowledge_scroll.set,
                                    bg="#1e1e1e", fg="#ffffff", height=6)
        knowledge_list.pack(fill=tk.BOTH, expand=True)
        knowledge_scroll.config(command=knowledge_list.yview)
        
        # Model parameters
        params_frame = ttk.LabelFrame(main_frame, text="Parameters", padding="10")
        params_frame.pack(fill=tk.X, pady=(0, 10))
        
        temp_label_var = tk.StringVar(value=f"Temperature: {self.temperature.get()}")
        ttk.Label(params_frame, textvariable=temp_label_var).pack(anchor=tk.W)
        temp_scale = ttk.Scale(params_frame, from_=0.0, to=1.0, variable=self.temperature,
                            command=lambda v: temp_label_var.set(f"Temperature: {float(v):.2f}"))
        temp_scale.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(params_frame, text="Max Tokens:").pack(anchor=tk.W)
        ttk.Spinbox(params_frame, from_=500, to=16000, textvariable=self.max_tokens, 
                increment=500).pack(fill=tk.X)
        
        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        def save_and_close():
            self.openwebui_base_url = url_var.get()
            self.openwebui_api_key = key_var.get()
            
            # Get selected knowledge collections
            self.selected_knowledge_collections = []
            for idx in knowledge_list.curselection():
                self.selected_knowledge_collections.append(
                    self.available_knowledge_collections[idx])
            
            self.save_settings()
            self.update_config_status()
            config_window.destroy()
        
        ttk.Button(btn_frame, text="Save", command=save_and_close).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(btn_frame, text="Cancel", command=config_window.destroy).pack(side=tk.RIGHT)
        
        # Auto-refresh on load
        # Use after to let the dialog render first
        config_window.after(100, lambda: self.auto_refresh_resources(url_var.get(), key_var.get(), model_combo, knowledge_list, status_label))

    def test_connection(self, url, api_key, status_label):
        """Test OpenWebUI connection"""
        try:
            status_label.config(text="⏳ Testing connection...")
            headers = {'Authorization': f'Bearer {api_key}'}
            response = requests.get(f"{url}/api/models", headers=headers, timeout=10)
            
            if response.status_code == 200:
                status_label.config(text="✓ Connection successful")
                return True
            else:
                status_label.config(text=f"✗ HTTP {response.status_code}")
        except Exception as e:
            status_label.config(text=f"✗ {str(e)}")
        return False

    def refresh_models(self, url, api_key, combo):
        """Refresh available models"""
        try:
            combo.set("Loading models...")
            headers = {'Authorization': f'Bearer {api_key}'}
            response = requests.get(f"{url}/api/models", headers=headers, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                self.available_models = []
                
                if isinstance(data, dict) and 'data' in data:
                    models_list = data['data']
                elif isinstance(data, list):
                    models_list = data
                else:
                    models_list = [data] if data else []
                
                for model in models_list:
                    if isinstance(model, dict):
                        model_id = model.get('id', model.get('name', 'unknown'))
                        self.available_models.append(model_id)
                    else:
                        self.available_models.append(str(model))
                
                combo['values'] = self.available_models
                
                # Update selection if possible
                if self.selected_model.get() in self.available_models:
                    combo.set(self.selected_model.get())
                elif self.available_models:
                    combo.set(self.available_models[0])
                else:
                    combo.set("")
                    
                self.log_message(f"Refreshed models: {len(self.available_models)} found")
                return True
            else:
                self.log_message(f"Error refreshing models: HTTP {response.status_code}")
                combo.set("Error loading models")
                return False
        except Exception as e:
            self.log_message(f"Error refreshing models: {str(e)}")
            combo.set("Error loading models")
            return False
            
    def refresh_knowledge(self, url, api_key, listbox):
        """Refresh knowledge collections"""
        try:
            headers = {'Authorization': f'Bearer {api_key}'}
            response = requests.get(f"{url}/api/v1/knowledge/", headers=headers, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                self.available_knowledge_collections = []
                
                if isinstance(data, list):
                    for collection in data:
                        if isinstance(collection, dict):
                            self.available_knowledge_collections.append({
                                'id': collection.get('id', 'unknown'),
                                'name': collection.get('name', 'Unknown'),
                                'description': collection.get('description', '')
                            })
                
                listbox.delete(0, tk.END)
                for col in self.available_knowledge_collections:
                    display = f"{col['name']} ({col['id']})"
                    listbox.insert(tk.END, display)
                
                # Restore selections
                for i, col in enumerate(self.available_knowledge_collections):
                    for sel in self.selected_knowledge_collections:
                        if col['id'] == sel.get('id', ''):
                            listbox.selection_set(i)
                
                self.log_message(f"Refreshed collections: {len(self.available_knowledge_collections)} found")
                return True
            else:
                self.log_message(f"Error refreshing collections: HTTP {response.status_code}")
                return False
        except Exception as e:
            self.log_message(f"Error refreshing collections: {str(e)}")
            return False        

    def update_config_status(self):
        """Update configuration status display"""
        if self.selected_model.get() and self.openwebui_api_key:
            status = f"✓ Model: {self.selected_model.get()}"
            if self.selected_knowledge_collections:
                status += f" | RAG: {len(self.selected_knowledge_collections)} collections"
            self.config_status_var.set(status)
        else:
            self.config_status_var.set("⚠ Not configured - click Configure")

    def load_prompt_library(self, listbox, preview):
        """Load saved prompts into the library list"""
        listbox.delete(0, tk.END)
        
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            library_dir = os.path.join(script_dir, "prompt_library")
            
            # Create directory if it doesn't exist
            if not os.path.exists(library_dir):
                os.makedirs(library_dir)
                self.log_message("Created prompt library directory")
                
                # Create a default prompt
                default_path = os.path.join(library_dir, "default_prompt.txt")
                self.save_master_prompt(default_path)
            
            # Add all prompt files to the list without extensions
            for filename in os.listdir(library_dir):
                if filename.endswith('.txt'):
                    # Display without .txt extension
                    display_name = os.path.splitext(filename)[0]
                    listbox.insert(tk.END, display_name)
                    
                    # Store the full filename in the listbox items dictionary
                    self.prompt_library_filenames = getattr(self, 'prompt_library_filenames', {})
                    self.prompt_library_filenames[display_name] = filename
            
            if listbox.size() > 0:
                listbox.selection_set(0)
                self.on_prompt_select(None, listbox, preview)
                
        except Exception as e:
            self.log_message(f"Error loading prompt library: {str(e)}")

    def on_prompt_select(self, event, listbox, preview):
        """Handle selection in prompt library"""
        try:
            if listbox.curselection():
                index = listbox.curselection()[0]
                display_name = listbox.get(index)
                
                # Get the actual filename from our dictionary
                self.prompt_library_filenames = getattr(self, 'prompt_library_filenames', {})
                filename = self.prompt_library_filenames.get(display_name, display_name + ".txt")
                    
                script_dir = os.path.dirname(os.path.abspath(__file__))
                library_dir = os.path.join(script_dir, "prompt_library")
                prompt_path = os.path.join(library_dir, filename)
                
                with open(prompt_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                preview.delete('1.0', tk.END)
                preview.insert('1.0', content)
                
        except Exception as e:
            self.log_message(f"Error loading prompt: {str(e)}")

    def import_prompt_to_library(self, listbox, preview):
        """Import a prompt file to the library"""
        try:
            filename = filedialog.askopenfilename(
                title="Import Prompt to Library",
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
            )
            
            if not filename:
                return
                
            script_dir = os.path.dirname(os.path.abspath(__file__))
            library_dir = os.path.join(script_dir, "prompt_library")
            
            if not os.path.exists(library_dir):
                os.makedirs(library_dir)
            
            # Get basename for new file
            basename = os.path.basename(filename)
            name_input = tk.simpledialog.askstring(
                "Prompt Name", 
                "Enter a name for this prompt:",
                initialvalue=os.path.splitext(basename)[0]
            )
            
            if not name_input:
                return
                
            # Ensure .txt extension for file storage but display without extension
            if not name_input.endswith('.txt'):
                file_name = name_input + '.txt'
                display_name = name_input
            else:
                file_name = name_input
                display_name = os.path.splitext(name_input)[0]  # Remove extension for display
                
            # Copy file to library
            new_path = os.path.join(library_dir, file_name)
            shutil.copy2(filename, new_path)
            
            # Update our filename dictionary
            self.prompt_library_filenames = getattr(self, 'prompt_library_filenames', {})
            self.prompt_library_filenames[display_name] = file_name
            
            # Refresh list
            self.load_prompt_library(listbox, preview)
            
            # Select the new item
            for i in range(listbox.size()):
                if listbox.get(i) == display_name:
                    listbox.selection_clear(0, tk.END)
                    listbox.selection_set(i)
                    self.on_prompt_select(None, listbox, preview)
                    break
                    
            self.log_message(f"Imported prompt to library: {display_name}")
            
        except Exception as e:
            self.log_message(f"Error importing prompt: {str(e)}")

    def remove_prompt_from_library(self, listbox, preview):
        """Remove a prompt from the library"""
        try:
            if not listbox.curselection():
                messagebox.showinfo("Info", "Please select a prompt to remove.")
                return
                
            index = listbox.curselection()[0]
            display_name = listbox.get(index)
            
            # Get the actual filename from our dictionary
            self.prompt_library_filenames = getattr(self, 'prompt_library_filenames', {})
            filename = self.prompt_library_filenames.get(display_name, display_name + ".txt")
            
            result = messagebox.askyesno(
                "Remove Prompt",
                f"Are you sure you want to remove '{display_name}' from the library?\n\n"
                "This action cannot be undone."
            )
            
            if not result:
                return
                
            script_dir = os.path.dirname(os.path.abspath(__file__))
            library_dir = os.path.join(script_dir, "prompt_library")
            prompt_path = os.path.join(library_dir, filename)
            
            os.remove(prompt_path)
            
            # Remove from our filename dictionary
            if display_name in self.prompt_library_filenames:
                del self.prompt_library_filenames[display_name]
            
            # Refresh list
            self.load_prompt_library(listbox, preview)
            preview.delete('1.0', tk.END)
            
            self.log_message(f"Removed prompt from library: {display_name}")
            
        except Exception as e:
            self.log_message(f"Error removing prompt: {str(e)}")           
 
    def use_selected_prompt(self, listbox, preview, window):
        """Use the selected prompt as master prompt"""
        try:
            if not listbox.curselection():
                messagebox.showinfo("Info", "Please select a prompt to use.")
                return
                
            content = preview.get('1.0', tk.END).strip()
            display_name = listbox.get(listbox.curselection()[0])
            
            if not content:
                messagebox.showwarning("Warning", "Selected prompt is empty.")
                return
                
            result = messagebox.askyesno(
                "Use Prompt",
                f"Are you sure you want to replace the current master prompt with '{display_name}'?\n\n"
                "This action cannot be undone."
            )
            
            if result:
                self.master_prompt.set(content)
                self.log_message(f"Master prompt updated from library: {display_name}")
                window.destroy()
                
        except Exception as e:
            self.log_message(f"Error using prompt: {str(e)}")

    def open_prompt_library(self):
        """Open prompt library manager"""
        library_window = tk.Toplevel(self.root)
        library_window.title("Prompt Library")
        library_window.geometry("900x600")
        library_window.configure(bg="#2b2b2b")
        library_window.grab_set()
        
        main_frame = ttk.Frame(library_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Prompt Library Manager", 
                font=("Arial", 14, "bold")).pack(anchor=tk.W, pady=(0, 15))
        
        # Create split view
        pane = ttk.PanedWindow(main_frame, orient=tk.HORIZONTAL)
        pane.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        # Left panel - list of saved prompts
        left_frame = ttk.Frame(pane)
        pane.add(left_frame, weight=1)
        
        ttk.Label(left_frame, text="Saved Prompts:").pack(anchor=tk.W, pady=(0, 5))
        
        # Prompt list with scrollbar
        list_frame = ttk.Frame(left_frame)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        list_scrollbar = ttk.Scrollbar(list_frame)
        list_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        prompt_list = tk.Listbox(list_frame, yscrollcommand=list_scrollbar.set,
                                bg="#1e1e1e", fg="#ffffff", font=("Consolas", 10),
                                width=30)  # Added width parameter to prevent collapsing
        prompt_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        list_scrollbar.config(command=prompt_list.yview)
        
        # Buttons for list operations
        list_btn_frame = ttk.Frame(left_frame)
        list_btn_frame.pack(fill=tk.X, pady=(5, 0))
        
        ttk.Button(list_btn_frame, text="Import", 
                command=lambda: self.import_prompt_to_library(prompt_list, prompt_preview)).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(list_btn_frame, text="Remove", 
                command=lambda: self.remove_prompt_from_library(prompt_list, prompt_preview)).pack(side=tk.LEFT)
        
        # Right panel - prompt preview
        right_frame = ttk.Frame(pane)
        pane.add(right_frame, weight=2)
        
        ttk.Label(right_frame, text="Prompt Preview:").pack(anchor=tk.W, pady=(0, 5))
        
        prompt_preview = scrolledtext.ScrolledText(right_frame, height=20,
                                                bg="#1e1e1e", fg="#ffffff",
                                                font=("Consolas", 10), wrap=tk.WORD)
        prompt_preview.pack(fill=tk.BOTH, expand=True, pady=(0, 5))

        prompt_preview.config(insertbackground="#ffffff")  # White cursor
        prompt_preview.config(selectbackground="#3a5f8f")  # Blue selection background
        prompt_preview.config(selectforeground="#ffffff")  # White selection text
        
        # Buttons for preview operations
        preview_btn_frame = ttk.Frame(right_frame)
        preview_btn_frame.pack(fill=tk.X)
        
        ttk.Button(preview_btn_frame, text="Use This Prompt", 
                command=lambda: self.use_selected_prompt(prompt_list, prompt_preview, library_window)).pack(side=tk.RIGHT)
        
        # Bottom buttons
        bottom_frame = ttk.Frame(main_frame)
        bottom_frame.pack(fill=tk.X, pady=(0, 5))
        
        ttk.Button(bottom_frame, text="Close", 
                command=library_window.destroy).pack(side=tk.RIGHT)
        
        # Load saved prompts from library
        self.load_prompt_library(prompt_list, prompt_preview)
        
        # Bind selection event
        prompt_list.bind('<<ListboxSelect>>', 
                        lambda e: self.on_prompt_select(e, prompt_list, prompt_preview))
        
        # Ensure the window is properly sized
        library_window.update_idletasks()
        pane.sashpos(0, 300)

    def edit_master_prompt(self):
        """Open master prompt editor"""
        editor_window = tk.Toplevel(self.root)
        editor_window.title("Edit Master Prompt")
        editor_window.geometry("800x600")
        editor_window.configure(bg="#2b2b2b")
        editor_window.grab_set()
        
        main_frame = ttk.Frame(editor_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Master Prompt Template:", 
                font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 5))
        
        info_label = ttk.Label(main_frame, 
                            text="Variables: {section_name}, {parent_context}, {operation_mode}",
                            foreground="#888888")
        info_label.pack(anchor=tk.W, pady=(0, 10))
        
        # Add Save/Load buttons at top
        action_frame = ttk.Frame(main_frame)
        action_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(action_frame, text="Load from File", 
                command=self.load_master_prompt).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(action_frame, text="Save to File", 
                command=self.save_master_prompt).pack(side=tk.LEFT)
        
        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        prompt_text = scrolledtext.ScrolledText(text_frame, height=20, width=80,
                                            bg="#1e1e1e", fg="#ffffff",
                                            font=("Consolas", 10), wrap=tk.WORD)
        prompt_text.pack(fill=tk.BOTH, expand=True)
        prompt_text.insert('1.0', self.master_prompt.get())
        
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        prompt_text.config(insertbackground="#ffffff")  # White cursor
        prompt_text.config(selectbackground="#3a5f8f")  # Blue selection background
        prompt_text.config(selectforeground="#ffffff")  # White selection text
        
        def save_prompt():
            self.master_prompt.set(prompt_text.get('1.0', tk.END).strip())
            self.log_message("Master prompt updated")
            editor_window.destroy()
        
        def reset_prompt():
            if messagebox.askyesno("Reset", "Reset to default prompt?"):
                self.set_default_prompt()
                prompt_text.delete('1.0', tk.END)
                prompt_text.insert('1.0', self.master_prompt.get())
        
        ttk.Button(btn_frame, text="Reset to Default", command=reset_prompt).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Cancel", command=editor_window.destroy).pack(side=tk.RIGHT)
        ttk.Button(btn_frame, text="Save", command=save_prompt).pack(side=tk.RIGHT, padx=(0, 5))            

    def generate_content(self):
        """Generate content for selected section"""
        if not self.selected_section:
            messagebox.showwarning("Warning", "No section selected")
            return
            
        if not self.selected_model.get():
            messagebox.showwarning("Warning", "No model configured")
            return
            
        self.notebook.select(2)  # Switch to console
        
        thread = threading.Thread(target=self.generate_content_thread)
        thread.daemon = True
        thread.start()
        
    def generate_content_thread(self):
        """Generate content in background thread"""
        try:
            self.generate_btn.config(state='disabled')
            
            section = self.selected_section
            mode = self.operation_mode.get()
            existing_content = section.get_existing_content()
            
            self.log_message(f"Generating content for: {section.get_full_path()}")
            self.log_message(f"Mode: {mode}")
            
            # Build parent context
            parent_context = ""
            if section.parent:
                parent_path = []
                current = section.parent
                while current:
                    parent_path.insert(0, current.text)
                    current = current.parent
                parent_context = " > ".join(parent_path)
            
            # Build prompt
            prompt = self.master_prompt.get()
            prompt = prompt.replace("{section_name}", section.text)
            prompt = prompt.replace("{parent_context}", parent_context if parent_context else "Root level")
            prompt = prompt.replace("{operation_mode}", mode.upper())
            
            # Add mode-specific instructions
            if mode == "replace":
                prompt += "\n\nYour task: Write comprehensive content for this section from scratch."
            elif mode == "rework":
                prompt += f"\n\nEXISTING CONTENT TO REWORK:\n{existing_content}\n\n"
                prompt += "Your task: Rewrite and enhance the existing content."
            elif mode == "append":
                prompt += f"\n\nEXISTING CONTENT:\n{existing_content}\n\n"
                prompt += "Your task: Add additional relevant content."
            
            # Add knowledge base instruction
            if self.selected_knowledge_collections:
                collection_names = [col['name'] for col in self.selected_knowledge_collections]
                prompt += f"\n\nIMPORTANT: Reference knowledge base: {', '.join(collection_names)}"
            
            # Store the prompt
            self.last_sent_prompt = prompt
            
            self.log_message("Sending request to OpenWebUI...")
            
            # Query OpenWebUI
            response = self.query_openwebui(prompt)
            
            if response and not response.startswith("Error:"):
                self.generated_content = response
                self.root.after(0, self.show_generated_content)
                self.log_message("Content generated successfully")
            else:
                self.log_message(f"Generation failed: {response}")
                self.root.after(0, self.handle_generation_error)
                
        except Exception as e:
            self.log_message(f"Error generating content: {str(e)}")
            self.root.after(0, self.handle_generation_error)
        finally:
            self.root.after(0, lambda: self.generate_btn.config(state='normal'))
            
    def query_openwebui(self, prompt):
        """Query OpenWebUI API"""
        try:
            # Log the prompt to history first
            self.log_prompt_history(prompt, is_sent=True)
            
            headers = {
                'Authorization': f'Bearer {self.openwebui_api_key}',
                'Content-Type': 'application/json'
            }
            
            payload = {
                "model": self.selected_model.get(),
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "temperature": self.temperature.get(),
                "max_tokens": self.max_tokens.get()
            }
            
            if self.selected_knowledge_collections:
                payload["files"] = [
                    {"type": "collection", "id": col['id']} 
                    for col in self.selected_knowledge_collections
                ]
            
            response = requests.post(
                f"{self.openwebui_base_url}/api/chat/completions",
                headers=headers, json=payload, timeout=300
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and len(result['choices']) > 0:
                    response_content = result['choices'][0]['message']['content']
                    # Log the response to history
                    self.log_prompt_history(response_content, response_content, is_sent=False)
                    return response_content
                elif 'response' in result:
                    response_content = result['response']
                    # Log the response to history
                    self.log_prompt_history(response_content, response_content, is_sent=False)
                    return response_content
                else:
                    error_msg = "Error: No content in response"
                    self.log_prompt_history(error_msg, error_msg, is_sent=False)
                    return error_msg
            else:
                error_msg = f"Error: HTTP {response.status_code} - {response.text}"
                self.log_prompt_history(error_msg, error_msg, is_sent=False)
                return error_msg
                
        except Exception as e:
            error_msg = f"Error: {str(e)}"
            self.log_prompt_history(error_msg, error_msg, is_sent=False)
            return error_msg

    def query_openwebui_with_model(self, prompt, model):
        """Query OpenWebUI API with specific model"""
        try:
            # Log the prompt to history
            self.log_prompt_history(f"[Model: {model}]\n{prompt}", is_sent=True)
            
            headers = {
                'Authorization': f'Bearer {self.openwebui_api_key}',
                'Content-Type': 'application/json'
            }
            
            payload = {
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "temperature": self.temperature.get(),
                "max_tokens": self.max_tokens.get()
            }
            
            if self.selected_knowledge_collections:
                payload["files"] = [
                    {"type": "collection", "id": col['id']} 
                    for col in self.selected_knowledge_collections
                ]
            
            response = requests.post(
                f"{self.openwebui_base_url}/api/chat/completions",
                headers=headers, json=payload, timeout=300
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and len(result['choices']) > 0:
                    response_content = result['choices'][0]['message']['content']
                    # Log the response to history
                    response_with_model = f"[Model: {model}]\n{response_content}"
                    self.log_prompt_history(response_with_model, response_with_model, is_sent=False)
                    return response_content
                elif 'response' in result:
                    response_content = result['response']
                    # Log the response to history
                    response_with_model = f"[Model: {model}]\n{response_content}"
                    self.log_prompt_history(response_with_model, response_with_model, is_sent=False)
                    return response_content
                else:
                    error_msg = f"[Model: {model}]\nError: No content in response"
                    self.log_prompt_history(error_msg, error_msg, is_sent=False)
                    return "Error: No content in response"
            else:
                error_msg = f"[Model: {model}]\nError: HTTP {response.status_code} - {response.text}"
                self.log_prompt_history(error_msg, error_msg, is_sent=False)
                return f"Error: HTTP {response.status_code} - {response.text}"
                
        except Exception as e:
            error_msg = f"[Model: {model}]\nError: {str(e)}"
            self.log_prompt_history(error_msg, error_msg, is_sent=False)
            return f"Error: {str(e)}"

    def render_markdown_preview(self, text):
        """Render markdown in text widget with formatting"""
        self.generated_text.delete('1.0', tk.END)
        
        lines = text.split('\n')
        for line in lines:
            segments = []
            current_line = line
            
            # Process line character by character
            i = 0
            while i < len(current_line):
                # Check for **bold**
                if current_line[i:i+2] == '**':
                    end = current_line.find('**', i+2)
                    if end != -1:
                        if i > 0:
                            segments.append(('normal', current_line[:i]))
                        segments.append(('bold', current_line[i+2:end]))
                        current_line = current_line[end+2:]
                        i = 0
                        continue
                
                # Check for *italic*
                elif current_line[i] == '*':
                    end = current_line.find('*', i+1)
                    if end != -1:
                        if i > 0:
                            segments.append(('normal', current_line[:i]))
                        segments.append(('italic', current_line[i+1:end]))
                        current_line = current_line[end+1:]
                        i = 0
                        continue
                
                # Check for `code`
                elif current_line[i] == '`':
                    end = current_line.find('`', i+1)
                    if end != -1:
                        if i > 0:
                            segments.append(('normal', current_line[:i]))
                        segments.append(('code', current_line[i+1:end]))
                        current_line = current_line[end+1:]
                        i = 0
                        continue
                
                i += 1
            
            # Add remaining text
            if current_line:
                segments.append(('normal', current_line))
            
            # Insert segments with tags
            if segments:
                for tag, text_seg in segments:
                    if tag == 'normal':
                        self.generated_text.insert(tk.END, text_seg)
                    else:
                        start_idx = self.generated_text.index(tk.END)
                        self.generated_text.insert(tk.END, text_seg)
                        end_idx = self.generated_text.index(tk.END)
                        self.generated_text.tag_add(tag, start_idx, end_idx)
            
            self.generated_text.insert(tk.END, '\n')
            
    def show_generated_content(self):
        """Show generated content in preview with markdown rendering"""
        self.notebook.select(0)  # Switch to preview
        
        # Render with markdown formatting
        self.render_markdown_preview(self.generated_content)
        
        # Update prompt tab
        self.prompt_text.delete('1.0', tk.END)
        self.prompt_text.insert('1.0', self.last_sent_prompt)
        
        # Update model selector in prompt tab
        self.update_prompt_model_selector()
        
        self.commit_btn.config(state='normal')
    
    def update_prompt_model_selector(self):
        """Update the model selector in prompt tab"""
        # Try to load models if empty
        if not self.available_models:
            self.ensure_models_loaded()
        
        if self.available_models:
            self.prompt_model_combo['values'] = self.available_models
            if not self.prompt_model_var.get():
                self.prompt_model_var.set(self.selected_model.get())
        else:
            self.prompt_model_combo['values'] = ['No models loaded - Configure AI Settings']
            self.log_message("No models available. Please configure AI settings and refresh models.")

    def debug_model_status(self):
        """Debug function to check model status"""
        print(f"Available models count: {len(self.available_models)}")
        print(f"Models list: {self.available_models}")
        print(f"API Key set: {bool(self.openwebui_api_key)}")
        print(f"Base URL: {self.openwebui_base_url}")
        
        self.log_message(f"Debug: {len(self.available_models)} models loaded")
        if self.available_models:
            self.log_message(f"Models: {', '.join(self.available_models[:3])}...")

    def get_document_outline(self):
        """Get complete document outline for context"""
        outline = []
        
        def build_outline(sections, level=0):
            for section in sections:
                indent = "  " * level
                outline.append(f"{indent}{section.text}")
                build_outline(section.children, level + 1)
        
        build_outline(self.sections)
        return "\n".join(outline)

    def auto_complete_document(self):
        """Automatically complete all sections in the document"""
        if not self.document or not self.sections:
            messagebox.showwarning("Warning", "No document loaded")
            return
        
        if not self.selected_model.get():
            messagebox.showwarning("Warning", "No model configured")
            return
        
        # Get all sections that need completion
        all_sections = self.get_all_sections_flat()
        
        # Filter based on user preference
        result = messagebox.askyesnocancel(
            "Auto Complete Document",
            f"Found {len(all_sections)} sections.\n\n"
            "Yes: Complete ALL sections (overwrites existing)\n"
            "No: Complete only EMPTY sections\n"
            "Cancel: Don't proceed"
        )
        
        if result is None:  # Cancel
            return
        elif result is True:  # All sections
            sections_to_complete = all_sections
        else:  # Only empty
            sections_to_complete = [s for s in all_sections if not s.has_content()]
        
        if not sections_to_complete:
            messagebox.showinfo("Info", "No sections to complete!")
            return
        
        # Confirm
        confirm = messagebox.askyesno(
            "Confirm Auto Complete",
            f"Ready to auto-complete {len(sections_to_complete)} sections.\n\n"
            f"This will:\n"
            f"• Process sections one by one\n"
            f"• Use current operation mode: {self.operation_mode.get().upper()}\n"
            f"• Take several minutes to complete\n"
            f"• Auto-save after each section (if enabled)\n\n"
            f"Continue?"
        )
        
        if not confirm:
            return
        
        # Create progress window
        self.create_auto_complete_window(sections_to_complete)

    def get_all_sections_flat(self):
        """Get flat list of all sections"""
        all_sections = []
        
        def collect_sections(sections):
            for section in sections:
                all_sections.append(section)
                collect_sections(section.children)
        
        collect_sections(self.sections)
        return all_sections

    def create_auto_complete_window(self, sections_to_complete):
        """Create progress window for auto-complete"""
        progress_window = tk.Toplevel(self.root)
        progress_window.title("Auto Complete Document")
        progress_window.geometry("800x600")
        progress_window.configure(bg="#2b2b2b")
        progress_window.protocol("WM_DELETE_WINDOW", lambda: None)  # Prevent closing during process
        
        main_frame = ttk.Frame(progress_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Auto Completing Document", 
                font=("Arial", 14, "bold")).pack(pady=(0, 10))
        
        # Progress info
        progress_frame = ttk.Frame(main_frame)
        progress_frame.pack(fill=tk.X, pady=(0, 10))
        
        progress_label = ttk.Label(progress_frame, text="Starting...", font=("Arial", 10))
        progress_label.pack(anchor=tk.W)
        
        # Progress bar
        progress_bar = ttk.Progressbar(progress_frame, mode='determinate', length=400)
        progress_bar.pack(fill=tk.X, pady=(5, 0))
        progress_bar['maximum'] = len(sections_to_complete)
        
        # Current section info
        current_frame = ttk.LabelFrame(main_frame, text="Current Section", padding="10")
        current_frame.pack(fill=tk.X, pady=(0, 10))
        
        current_section_label = ttk.Label(current_frame, text="", wraplength=700)
        current_section_label.pack(anchor=tk.W)
        
        status_label = ttk.Label(current_frame, text="", foreground="#00aaff")
        status_label.pack(anchor=tk.W, pady=(5, 0))
        
        # Log area
        log_frame = ttk.LabelFrame(main_frame, text="Progress Log", padding="5")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        log_text = scrolledtext.ScrolledText(log_frame, height=15, width=80,
                                            bg="#1e1e1e", fg="#ffffff",
                                            font=("Consolas", 9))
        log_text.pack(fill=tk.BOTH, expand=True)
        
        log_text.config(insertbackground="#ffffff")  # White cursor
        log_text.config(selectbackground="#3a5f8f")  # Blue selection background
        log_text.config(selectforeground="#ffffff")  # White selection text
        

        # Control buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X)
        
        pause_var = tk.BooleanVar(value=False)
        stop_var = tk.BooleanVar(value=False)
        
        pause_btn = ttk.Button(button_frame, text="Pause", 
                            command=lambda: pause_var.set(not pause_var.get()))
        pause_btn.pack(side=tk.LEFT, padx=5)
        
        stop_btn = ttk.Button(button_frame, text="Stop", 
                            command=lambda: stop_var.set(True))
        stop_btn.pack(side=tk.LEFT, padx=5)
        
        close_btn = ttk.Button(button_frame, text="Close", state='disabled',
                            command=progress_window.destroy)
        close_btn.pack(side=tk.RIGHT, padx=5)
        
        # Store references
        progress_window.log_text = log_text
        progress_window.progress_bar = progress_bar
        progress_window.progress_label = progress_label
        progress_window.current_section_label = current_section_label
        progress_window.status_label = status_label
        progress_window.pause_btn = pause_btn
        progress_window.stop_btn = stop_btn
        progress_window.close_btn = close_btn
        progress_window.pause_var = pause_var
        progress_window.stop_var = stop_var
        
        # Start processing in thread
        thread = threading.Thread(
            target=self.auto_complete_process,
            args=(sections_to_complete, progress_window)
        )
        thread.daemon = True
        thread.start()

    def auto_complete_process(self, sections, progress_window):
        """Process auto-complete with parent context and subsection awareness"""
        try:
            # Get document outline for structure context
            document_outline = self.get_document_outline()
            
            # Build content map as we go (for parent context)
            generated_content_map = {}  # section_hash -> content
            
            completed = 0
            failed = 0
            
            for idx, section in enumerate(sections):
                # Check if stopped
                if progress_window.stop_var.get():
                    self.update_progress_log(progress_window, "Stopped by user")
                    break
                
                # Check if paused
                while progress_window.pause_var.get():
                    if not hasattr(progress_window, 'pause_notified'):
                        self.update_progress_log(progress_window, "⏸ Paused - Click Pause again to resume")
                        self.root.after(0, lambda: progress_window.pause_btn.config(text="Resume"))
                        progress_window.pause_notified = True
                    self.root.update()
                    import time
                    time.sleep(0.5)
                    if progress_window.stop_var.get():
                        break
                
                # Reset pause notification
                if hasattr(progress_window, 'pause_notified'):
                    delattr(progress_window, 'pause_notified')
                    self.root.after(0, lambda: progress_window.pause_btn.config(text="Pause"))
                
                # Update progress
                self.root.after(0, lambda i=idx: progress_window.progress_bar.configure(value=i))
                self.root.after(0, lambda i=idx, t=len(sections): 
                            progress_window.progress_label.config(
                                text=f"Processing {i+1} of {t}..."))
                
                # Update current section
                section_path = section.get_full_path()
                self.root.after(0, lambda p=section_path: 
                            progress_window.current_section_label.config(text=p))
                self.root.after(0, lambda: 
                            progress_window.status_label.config(text="🔄 Generating content..."))
                
                self.update_progress_log(progress_window, f"\n{'='*60}")
                self.update_progress_log(progress_window, f"[{idx+1}/{len(sections)}] {section_path}")
                
                try:
                    # Build enhanced prompt with parent context and subsection awareness
                    mode = self.operation_mode.get()
                    existing_content = section.get_existing_content()
                    
                    # Get parent context PATH
                    parent_context_path = ""
                    if section.parent:
                        parent_path = []
                        current = section.parent
                        while current:
                            parent_path.insert(0, current.text)
                            current = current.parent
                        parent_context_path = " > ".join(parent_path)
                    
                    # *** NEW: Get parent's ACTUAL CONTENT ***
                    parent_actual_content = ""
                    if section.parent:
                        parent_hash = section.parent.get_section_hash()
                        # First check if we just generated it
                        if parent_hash in generated_content_map:
                            parent_actual_content = generated_content_map[parent_hash]
                        else:
                            # Check if it exists in document
                            parent_actual_content = section.parent.get_existing_content()
                    
                    # *** NEW: Get sibling sections for context ***
                    sibling_info = ""
                    if section.parent:
                        siblings = [s.text for s in section.parent.children if s != section]
                        if siblings:
                            sibling_info = "Sibling sections (for context): " + ", ".join(siblings)
                    
                    # Build enhanced prompt
                    prompt = self.master_prompt.get()
                    prompt = prompt.replace("{section_name}", section.text)
                    prompt = prompt.replace("{parent_context}", 
                                        parent_context_path if parent_context_path else "Root level")
                    prompt = prompt.replace("{operation_mode}", mode.upper())
                    
                    # *** NEW: Add document outline ***
                    prompt += f"\n\nDOCUMENT STRUCTURE (for overall context):\n{document_outline}\n"
                    
                    # *** NEW: Add parent content if this is a subsection ***
                    if section.parent and parent_actual_content:
                        prompt += f"\n\nPARENT SECTION CONTENT (you MUST align with this):\n"
                        prompt += f"Parent Section: {section.parent.text}\n"
                        prompt += f"{'-'*60}\n"
                        prompt += f"{parent_actual_content}\n"
                        prompt += f"{'-'*60}\n"
                        prompt += f"\nIMPORTANT: Your content for '{section.text}' MUST:\n"
                        prompt += f"1. Directly relate to and expand upon the parent section content above\n"
                        prompt += f"2. Use consistent terminology with the parent section\n"
                        prompt += f"3. Reference parent concepts where appropriate\n"
                        prompt += f"4. Be a logical subdivision of the parent topic\n"
                        prompt += f"5. Not contradict anything in the parent section\n"
                    
                    # *** NEW: Add sibling context ***
                    if sibling_info:
                        prompt += f"\n{sibling_info}\n"
                        prompt += f"Note: Ensure your content is distinct from these sibling sections.\n"
                    
                    # Add mode-specific instructions
                    if mode == "replace":
                        prompt += "\n\nYour task: Write comprehensive content for this section from scratch."
                    elif mode == "rework":
                        prompt += f"\n\nEXISTING CONTENT TO REWORK:\n{existing_content}\n\n"
                        prompt += "Your task: Rewrite and enhance the existing content."
                    elif mode == "append":
                        prompt += f"\n\nEXISTING CONTENT:\n{existing_content}\n\n"
                        prompt += "Your task: Add additional relevant content."
                    
                    # Add knowledge base instruction
                    if self.selected_knowledge_collections:
                        collection_names = [col['name'] for col in self.selected_knowledge_collections]
                        prompt += f"\n\nIMPORTANT: Reference knowledge base: {', '.join(collection_names)}"
                    
                    self.update_progress_log(progress_window, "  Sending request to API...")
                    
                    # Generate content
                    response = self.query_openwebui(prompt)
                    
                    if response and not response.startswith("Error:"):
                        self.update_progress_log(progress_window, "  ✓ Content generated")
                        
                        # *** NEW: Store generated content for future subsections ***
                        section_hash = section.get_section_hash()
                        generated_content_map[section_hash] = response
                        
                        # Apply content to section
                        self.root.after(0, lambda: 
                                    progress_window.status_label.config(text="Committing to document..."))
                        
                        if mode == "replace":
                            self.replace_section_content(section, response)
                        elif mode == "rework":
                            self.replace_section_content(section, response)
                        elif mode == "append":
                            self.append_section_content(section, response)
                        
                        # Mark as edited
                        self.mark_section_edited(section)
                        
                        self.update_progress_log(progress_window, "  ✓ Committed to document")
                        
                        # Auto-save if enabled
                        if self.auto_config['auto_save'].get():
                            try:
                                self.document.save(self.document_path)
                                self.update_progress_log(progress_window, "  ✓ Auto-saved")
                            except Exception as e:
                                self.update_progress_log(progress_window, f"  ⚠ Auto-save failed: {str(e)}")
                        
                        completed += 1
                        
                    else:
                        self.update_progress_log(progress_window, f"  ✗ Generation failed: {response}")
                        failed += 1
                        
                except Exception as e:
                    self.update_progress_log(progress_window, f"  ✗ Error: {str(e)}")
                    failed += 1
            
            # Final update
            self.root.after(0, lambda: progress_window.progress_bar.configure(value=len(sections)))
            self.root.after(0, lambda: 
                        progress_window.progress_label.config(text="Complete!"))
            self.root.after(0, lambda: 
                        progress_window.status_label.config(text=""))
            self.root.after(0, lambda: 
                        progress_window.current_section_label.config(text="All sections processed"))
            
            # Summary
            self.update_progress_log(progress_window, f"\n{'='*60}")
            self.update_progress_log(progress_window, "SUMMARY:")
            self.update_progress_log(progress_window, f"  ✓ Completed: {completed}")
            self.update_progress_log(progress_window, f"  ✗ Failed: {failed}")
            self.update_progress_log(progress_window, f"  Total: {len(sections)}")
            
            # Enable close button and disable others
            self.root.after(0, lambda: progress_window.close_btn.config(state='normal'))
            self.root.after(0, lambda: progress_window.pause_btn.config(state='disabled'))
            self.root.after(0, lambda: progress_window.stop_btn.config(state='disabled'))
            
            # Refresh tree and reload if enabled
            self.root.after(0, self.populate_tree)
            if self.auto_config['auto_reload'].get():
                self.root.after(0, self.reload_document)
            
            # Save final document if not auto-saving
            if not self.auto_config['auto_save'].get():
                self.root.after(0, self.prompt_save_document)
            
            self.log_message(f"Auto-complete finished: {completed} completed, {failed} failed")
            
        except Exception as e:
            self.update_progress_log(progress_window, f"\nFATAL ERROR: {str(e)}")
            self.log_message(f"Auto-complete error: {str(e)}")


    def update_progress_log(self, window, message):
        """Update progress log in auto-complete window"""
        def update():
            window.log_text.insert(tk.END, f"{message}\n")
            window.log_text.see(tk.END)
            window.update()
        
        self.root.after(0, update)

    def regenerate_with_prompt(self):
        """Regenerate content using edited prompt with selected model"""
        if not self.selected_section:
            messagebox.showwarning("Warning", "No section selected")
            return
        
        edited_prompt = self.prompt_text.get('1.0', tk.END).strip()
        if not edited_prompt:
            messagebox.showwarning("Warning", "Prompt is empty")
            return
        
        # Get selected model or use default
        selected_model = self.prompt_model_var.get()
        if not selected_model:
            selected_model = self.selected_model.get()
        
        self.notebook.select(2)  # Switch to console
        
        def regenerate_thread():
            try:
                self.generate_btn.config(state='disabled')
                self.log_message(f"Regenerating with model: {selected_model}...")
                
                response = self.query_openwebui_with_model(edited_prompt, selected_model)
                
                if response and not response.startswith("Error:"):
                    self.generated_content = response
                    self.last_sent_prompt = edited_prompt
                    self.root.after(0, self.show_generated_content)
                    self.log_message("Content regenerated successfully")
                else:
                    self.log_message(f"Regeneration failed: {response}")
                    self.root.after(0, self.handle_generation_error)
                    
            except Exception as e:
                self.log_message(f"Regeneration error: {str(e)}")
                self.root.after(0, self.handle_generation_error)
            finally:
                self.root.after(0, lambda: self.generate_btn.config(state='normal'))
        
        thread = threading.Thread(target=regenerate_thread)
        thread.daemon = True
        thread.start()
    
    def ensure_models_loaded(self):
        """Ensure models are loaded, try to refresh if not"""
        if not self.available_models:
            if not self.openwebui_api_key or not self.openwebui_base_url:
                return False
            
            # Try to refresh models automatically
            try:
                self.log_message("Auto-refreshing models...")
                headers = {'Authorization': f'Bearer {self.openwebui_api_key}'}
                response = requests.get(f"{self.openwebui_base_url}/api/models", 
                                    headers=headers, timeout=10)
                
                if response.status_code == 200:
                    data = response.json()
                    self.available_models = []
                    
                    if isinstance(data, dict) and 'data' in data:
                        models_list = data['data']
                    elif isinstance(data, list):
                        models_list = data
                    else:
                        models_list = [data] if data else []
                    
                    for model in models_list:
                        if isinstance(model, dict):
                            model_id = model.get('id', model.get('name', 'unknown'))
                            self.available_models.append(model_id)
                        else:
                            self.available_models.append(str(model))
                    
                    self.log_message(f"Auto-loaded {len(self.available_models)} models")
                    return len(self.available_models) > 0
                else:
                    self.log_message(f"Failed to auto-load models: HTTP {response.status_code}")
                    return False
            except Exception as e:
                self.log_message(f"Failed to auto-load models: {str(e)}")
                return False
        
        return True

    def compare_three_models(self):
        """Generate content with 3 different models and compare side by side"""
        if not self.selected_section:
            messagebox.showwarning("Warning", "No section selected")
            return
        
        edited_prompt = self.prompt_text.get('1.0', tk.END).strip()
        if not edited_prompt:
            messagebox.showwarning("Warning", "Prompt is empty")
            return
        
        # TRY TO LOAD MODELS IF NOT ALREADY LOADED
        if not self.ensure_models_loaded():
            messagebox.showerror("Error", 
                "Could not load models.\n\n"
                "Please:\n"
                "1. Go to 'Configure AI Settings'\n"
                "2. Enter your API URL and Key\n"
                "3. Click 'Refresh Models'")
            return
        
        if len(self.available_models) < 3:
            messagebox.showwarning("Warning", 
                f"Need at least 3 models available for comparison.\n"
                f"Currently have: {len(self.available_models)} model(s)\n\n"
                "Go to 'Configure AI Settings' and ensure models are available.")
            return
        
        # Create comparison window
        self.create_comparison_window(edited_prompt)

    def create_comparison_window(self, prompt):
        """Create model comparison window with easy model selection"""
        comp_window = tk.Toplevel(self.root)
        comp_window.title("Model Comparison")
        comp_window.geometry("1400x800")
        comp_window.configure(bg="#2b2b2b")
        comp_window.grab_set()
        
        main_frame = ttk.Frame(comp_window, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Compare 3 Models - Select Best Result", 
                font=("Arial", 14, "bold")).pack(pady=(0, 10))
        
        # ========== EASY 3-MODEL SELECTION STARTS HERE ==========
        # Model selection frame with listbox
        model_select_frame = ttk.LabelFrame(main_frame, text="Select 3 Models", padding="10")
        model_select_frame.pack(fill=tk.X, pady=(0, 10))
        
        select_inner = ttk.Frame(model_select_frame)
        select_inner.pack(fill=tk.X)
        
        ttk.Label(select_inner, text="Select exactly 3 models (Ctrl+Click):").pack(side=tk.LEFT, padx=(0, 10))
        
        # THIS IS THE LISTBOX - Multi-select with Ctrl+Click
        model_listbox = tk.Listbox(select_inner, selectmode=tk.MULTIPLE, height=5,
                                bg="#1e1e1e", fg="#ffffff", width=60)
        model_listbox.pack(side=tk.LEFT, padx=(0, 10))
        
        # Populate listbox with all available models
        for model in self.available_models:
            model_listbox.insert(tk.END, model)
        
        # AUTO-SELECT FIRST 3 MODELS BY DEFAULT
        for i in range(min(3, len(self.available_models))):
            model_listbox.selection_set(i)
        
        # SELECTION COUNTER - Shows "Selected: X/3"
        selected_label = ttk.Label(select_inner, text="Selected: 3/3", foreground="#00ff00")
        selected_label.pack(side=tk.LEFT, padx=(10, 0))
        
        # UPDATE SELECTION COUNT AND COLOR
        def update_selection_count(event=None):
            count = len(model_listbox.curselection())
            if count == 3:
                # GREEN when 3 selected - Enable button
                selected_label.config(text=f"Selected: {count}/3", foreground="#00ff00")
                generate_btn.config(state='normal')
            else:
                # RED when not 3 - Disable button
                selected_label.config(text=f"Selected: {count}/3", foreground="#ff0000")
                generate_btn.config(state='disabled')
        
        # Bind selection change event
        model_listbox.bind('<<ListboxSelect>>', update_selection_count)
        
        btn_frame = ttk.Frame(model_select_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        # GENERATE BUTTON - Only enabled when exactly 3 models selected
        generate_btn = ttk.Button(btn_frame, text="Generate All", 
                                command=lambda: self.start_comparison_from_listbox(prompt, model_listbox,
                                                                                    result_frames, status_labels))
        generate_btn.pack(side=tk.LEFT)
        # ========== EASY 3-MODEL SELECTION ENDS HERE ==========
        
        # Results frame - 3 columns side by side
        results_container = ttk.Frame(main_frame)
        results_container.pack(fill=tk.BOTH, expand=True)
        results_container.columnconfigure(0, weight=1)
        results_container.columnconfigure(1, weight=1)
        results_container.columnconfigure(2, weight=1)
        results_container.rowconfigure(0, weight=1)
        
        result_frames = []
        status_labels = []
        text_widgets = []
        select_buttons = []
        
        # CREATE 3 SIDE-BY-SIDE RESULT PANELS
        for i in range(3):
            # Column frame
            col_frame = ttk.Frame(results_container, relief="solid", borderwidth=2)
            col_frame.grid(row=0, column=i, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5, pady=5)
            col_frame.columnconfigure(0, weight=1)
            col_frame.rowconfigure(1, weight=1)
            
            result_frames.append(col_frame)
            
            # Header with model name and status
            header_frame = ttk.Frame(col_frame)
            header_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=5, pady=5)
            
            model_label = ttk.Label(header_frame, text=f"Model {i+1}: [Not Selected]", 
                    font=("Arial", 10, "bold"))
            model_label.pack(side=tk.LEFT)
            
            status_label = ttk.Label(header_frame, text="⏳ Pending", foreground="#ffaa00")
            status_label.pack(side=tk.RIGHT)
            status_labels.append(status_label)
            
            # Store model label for updating with actual model name
            col_frame.model_label = model_label
            
            # Text area for generated content
            text_widget = scrolledtext.ScrolledText(col_frame, height=25, width=40,
                                                bg="#1e1e1e", fg="#ffffff",
                                                font=("Consolas", 9), wrap=tk.WORD)
            text_widget.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5, pady=5)
            text_widgets.append(text_widget)
            
            # Select button for choosing this result
            select_btn = ttk.Button(col_frame, text="Select This Result", 
                                command=lambda idx=i: self.select_comparison_result(idx, text_widgets, 
                                                                                    result_frames, comp_window))
            select_btn.grid(row=2, column=0, pady=5)
            select_btn.config(state='disabled')
            select_buttons.append(select_btn)
        
        # Store widgets for access in generation
        comp_window.text_widgets = text_widgets
        comp_window.select_buttons = select_buttons
        comp_window.result_frames = result_frames

    # Add new function to handle listbox-based model selection
    def start_comparison_from_listbox(self, prompt, model_listbox, result_frames, status_labels):
        """Start comparison from listbox selection"""
        selected_indices = model_listbox.curselection()
        
        if len(selected_indices) != 3:
            messagebox.showwarning("Warning", "Please select exactly 3 models")
            return
        
        models = [model_listbox.get(i) for i in selected_indices]
        
        # Update model labels in result frames
        for i, (model, frame) in enumerate(zip(models, result_frames)):
            frame.model_label.config(text=f"Model {i+1}: {model}")
        
        # Get window references
        comp_window = result_frames[0].master.master.master
        text_widgets = comp_window.text_widgets
        select_buttons = comp_window.select_buttons
        
        def generation_thread():
            for i, model in enumerate(models):
                try:
                    # Update status to processing
                    self.root.after(0, lambda idx=i: status_labels[idx].config(
                        text="⚙️ Processing...", foreground="#00aaff"))
                    
                    self.log_message(f"Generating with {model}...")
                    
                    # Generate content
                    response = self.query_openwebui_with_model(prompt, model)
                    
                    if response and not response.startswith("Error:"):
                        # Update text widget
                        self.root.after(0, lambda idx=i, content=response: (
                            text_widgets[idx].delete('1.0', tk.END),
                            text_widgets[idx].insert('1.0', content),
                            status_labels[idx].config(text="✓ Complete", foreground="#00ff00"),
                            select_buttons[idx].config(state='normal')
                        ))
                        self.log_message(f"{model} completed successfully")
                    else:
                        # Show error
                        self.root.after(0, lambda idx=i, err=response: (
                            text_widgets[idx].delete('1.0', tk.END),
                            text_widgets[idx].insert('1.0', f"Error:\n{err}"),
                            status_labels[idx].config(text="✗ Failed", foreground="#ff0000")
                        ))
                        self.log_message(f"{model} failed: {response}")
                        
                except Exception as e:
                    self.root.after(0, lambda idx=i, err=str(e): (
                        text_widgets[idx].delete('1.0', tk.END),
                        text_widgets[idx].insert('1.0', f"Error:\n{err}"),
                        status_labels[idx].config(text="✗ Error", foreground="#ff0000")
                    ))
                    self.log_message(f"{model} error: {str(e)}")
        
        thread = threading.Thread(target=generation_thread)
        thread.daemon = True
        thread.start()

    def start_comparison_generation(self, prompt, model_vars, result_frames, status_labels):
        """Start sequential generation for 3 models"""
        models = [var.get() for var in model_vars]
        
        if len(set(models)) != 3:
            messagebox.showwarning("Warning", "Please select 3 different models")
            return
        
        # Get window references
        comp_window = result_frames[0].master.master.master
        text_widgets = comp_window.text_widgets
        select_buttons = comp_window.select_buttons
        
        def generation_thread():
            for i, model in enumerate(models):
                try:
                    # Update status to processing
                    self.root.after(0, lambda idx=i: status_labels[idx].config(
                        text="⚙️ Processing...", foreground="#00aaff"))
                    
                    self.log_message(f"Generating with {model}...")
                    
                    # Generate content
                    response = self.query_openwebui_with_model(prompt, model)
                    
                    if response and not response.startswith("Error:"):
                        # Update text widget
                        self.root.after(0, lambda idx=i, content=response: (
                            text_widgets[idx].delete('1.0', tk.END),
                            text_widgets[idx].insert('1.0', content),
                            status_labels[idx].config(text="✓ Complete", foreground="#00ff00"),
                            select_buttons[idx].config(state='normal')
                        ))
                        self.log_message(f"{model} completed successfully")
                    else:
                        # Show error
                        self.root.after(0, lambda idx=i, err=response: (
                            text_widgets[idx].delete('1.0', tk.END),
                            text_widgets[idx].insert('1.0', f"Error:\n{err}"),
                            status_labels[idx].config(text="✗ Failed", foreground="#ff0000")
                        ))
                        self.log_message(f"{model} failed: {response}")
                        
                except Exception as e:
                    self.root.after(0, lambda idx=i, err=str(e): (
                        text_widgets[idx].delete('1.0', tk.END),
                        text_widgets[idx].insert('1.0', f"Error:\n{err}"),
                        status_labels[idx].config(text="✗ Error", foreground="#ff0000")
                    ))
                    self.log_message(f"{model} error: {str(e)}")
        
        thread = threading.Thread(target=generation_thread)
        thread.daemon = True
        thread.start()
    
    def select_comparison_result(self, index, text_widgets, result_frames, window):
        """Select a result from comparison and move to preview"""
        # Get the content
        selected_content = text_widgets[index].get('1.0', tk.END).strip()
        
        if not selected_content or selected_content.startswith("Error:"):
            messagebox.showwarning("Warning", "Cannot select error result")
            return
        
        # Highlight selected frame
        for i, frame in enumerate(result_frames):
            if i == index:
                frame.configure(relief="solid", borderwidth=4)
            else:
                frame.configure(relief="solid", borderwidth=2)
        
        # Move to preview
        self.generated_content = selected_content
        self.show_generated_content()
        
        self.log_message(f"Selected result from Model {index + 1}")
        messagebox.showinfo("Success", f"Result from Model {index + 1} moved to Content Preview")
        
        window.destroy()
    


    def offer_prompt_update(self):
        """Offer to update master prompt with improvements"""
        current_prompt = self.prompt_text.get('1.0', tk.END).strip()
        
        dialog = tk.Toplevel(self.root)
        dialog.title("Update Master Prompt?")
        dialog.geometry("600x400")
        dialog.configure(bg="#2b2b2b")
        dialog.grab_set()
        
        frame = ttk.Frame(dialog, padding="20")
        frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(frame, text="Update Master Prompt with Changes?",
                 font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=(0, 10))
        
        ttk.Label(frame, text="Enter specific improvements to add to the master prompt:",
                 wraplength=550).pack(anchor=tk.W, pady=(0, 10))
        
        improvements_text = scrolledtext.ScrolledText(frame, height=10, width=70,
                                                      bg="#1e1e1e", fg="#ffffff",
                                                      font=("Consolas", 10), wrap=tk.WORD)
        improvements_text.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        improvements_text.insert('1.0', "e.g., 'Always include specific tool versions'\n'Reference DoD 8500 series'")
        
        btn_frame = ttk.Frame(frame)
        btn_frame.pack(fill=tk.X)
        
        def update_master():
            improvements = improvements_text.get('1.0', tk.END).strip()
            if improvements and not improvements.startswith('e.g.,'):
                current_master = self.master_prompt.get()
                updated_master = current_master + "\n\nADDITIONAL REQUIREMENTS:\n" + improvements
                self.master_prompt.set(updated_master)
                self.save_settings()
                self.log_message("Master prompt updated with new requirements")
                messagebox.showinfo("Success", "Master prompt updated successfully!")
            dialog.destroy()
        
        ttk.Button(btn_frame, text="Update Master Prompt", 
                  command=update_master).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT)
        
    def handle_generation_error(self):
        """Handle generation error"""
        self.notebook.select(1)  # Switch to prompt tab
        messagebox.showerror("Generation Failed", 
                           "Content generation failed. Edit the prompt in the 'Prompt Editor' tab and try 'Regenerate'.")
        
    def clear_preview(self):
        """Clear preview pane"""
        self.generated_text.delete('1.0', tk.END)
        self.commit_btn.config(state='disabled')
        self.log_message("Preview cleared")
        
    def commit_content(self):
        """Commit generated content to document"""
        if not self.selected_section or not self.document:
            return
            
        # Get edited content from preview (without formatting tags)
        content_to_commit = self.generated_text.get('1.0', tk.END).strip()
        
        if not content_to_commit:
            messagebox.showwarning("Warning", "No content to commit")
            return
        
        # Confirm commit
        result = messagebox.askyesno(
            "Commit Content",
            f"Commit content to section: {self.selected_section.text}\n\n"
            f"Mode: {self.operation_mode.get().upper()}\n\n"
            "This will modify the document."
        )
        
        if not result:
            return
        
        try:
            # Check if backup is needed
            if self.auto_config['ask_backup'].get():
                self.prompt_for_backup()
            else:
                self.create_backup()
            
            # Apply content based on mode
            mode = self.operation_mode.get()
            section = self.selected_section
            
            if mode == "replace":
                self.replace_section_content(section, content_to_commit)
            elif mode == "rework":
                self.replace_section_content(section, content_to_commit)
            elif mode == "append":
                self.append_section_content(section, content_to_commit)
            
            # Mark section as edited
            self.mark_section_edited(section)
            
            # Refresh tree
            self.populate_tree()
            
            # Auto-save if enabled
            if self.auto_config['auto_save'].get():
                self.save_document_auto()
            else:
                self.prompt_save_document()
            
            # Clear preview
            self.clear_preview()
            
            # Auto-reload if enabled
            if self.auto_config['auto_reload'].get():
                self.reload_document()
            else:
                self.show_existing_content()
            
            self.log_message(f"Content committed to: {section.text}")
            
        except Exception as e:
            self.log_message(f"Error committing content: {str(e)}")
            messagebox.showerror("Error", f"Failed to commit content: {str(e)}")
    
    def prompt_for_backup(self):
        """Prompt user for backup with 'don't ask again' option"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Create Backup?")
        dialog.geometry("400x150")
        dialog.configure(bg="#2b2b2b")
        dialog.grab_set()
        
        frame = ttk.Frame(dialog, padding="20")
        frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(frame, text="Create a backup before committing changes?",
                 font=("Arial", 10)).pack(pady=(0, 20))
        
        dont_ask_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(frame, text="Don't ask me again", 
                       variable=dont_ask_var).pack(pady=(0, 20))
        
        btn_frame = ttk.Frame(frame)
        btn_frame.pack()
        
        result = {'backup': False, 'dont_ask': False}
        
        def yes_clicked():
            result['backup'] = True
            result['dont_ask'] = dont_ask_var.get()
            dialog.destroy()
        
        def no_clicked():
            result['backup'] = False
            result['dont_ask'] = dont_ask_var.get()
            dialog.destroy()
        
        ttk.Button(btn_frame, text="Yes", command=yes_clicked).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="No", command=no_clicked).pack(side=tk.LEFT, padx=5)
        
        dialog.wait_window()
        
        if result['dont_ask']:
            self.auto_config['ask_backup'].set(False)
            self.save_settings()
            self.log_message("Backup prompts disabled")
        
        if result['backup']:
            self.create_backup()
            
    def create_backup(self):
        """Create backup of document"""
        if not self.document_path:
            return
            
        try:
            backup_dir = os.path.dirname(self.document_path)
            filename = os.path.basename(self.document_path)
            name, ext = os.path.splitext(filename)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_path = os.path.join(backup_dir, f"{name}_backup_{timestamp}{ext}")
            
            shutil.copy2(self.document_path, backup_path)
            self.log_message(f"Backup created: {os.path.basename(backup_path)}")
            
        except Exception as e:
            self.log_message(f"Warning: Could not create backup: {str(e)}")
    
    def replace_section_content(self, section, content):
        """Replace section content with new content"""
        # Remove existing content
        self.remove_section_content(section)
        
        # Add new content with markdown conversion
        self.add_markdown_content_to_section(section, content)
    
    def append_section_content(self, section, content):
        """Append content to section"""
        self.add_markdown_content_to_section(section, content, append=True)
    
    def remove_section_content(self, section):
        """Remove all content paragraphs from section"""
        doc_paragraphs = list(self.document.paragraphs)
        heading_index = -1
        
        # Find heading
        for i, para in enumerate(doc_paragraphs):
            if para.text.strip() == section.paragraph.text.strip():
                heading_index = i
                break
        
        if heading_index == -1:
            raise Exception("Could not find section heading")
        
        # Find next heading
        next_heading_index = len(doc_paragraphs)
        for i in range(heading_index + 1, len(doc_paragraphs)):
            if doc_paragraphs[i].style.name.startswith('Heading'):
                next_heading_index = i
                break
        
        # Remove content paragraphs
        for i in range(heading_index + 1, next_heading_index):
            para = doc_paragraphs[i]
            if not para.style.name.startswith('Heading'):
                self.remove_paragraph(para)
        
        section.content_paragraphs = []
    
    def add_markdown_content_to_section(self, section, content, append=False):
        """Add content to section with markdown to Word formatting conversion"""
        heading_para = section.paragraph
        doc_paragraphs = list(self.document.paragraphs)
        
        # Find heading index
        heading_index = -1
        for i, para in enumerate(doc_paragraphs):
            if para.text.strip() == heading_para.text.strip():
                heading_index = i
                break
        
        if heading_index == -1:
            raise Exception("Could not find section heading")
        
        # Determine insertion point
        if append and section.content_paragraphs:
            next_heading_index = len(doc_paragraphs)
            for i in range(heading_index + 1, len(doc_paragraphs)):
                if doc_paragraphs[i].style.name.startswith('Heading'):
                    next_heading_index = i
                    break
            insertion_index = next_heading_index
        else:
            insertion_index = heading_index + 1
        
        # Split content into paragraphs
        content_paragraphs = content.split('\n')
        
        inserted_paras = []
        for para_text in content_paragraphs:
            if para_text.strip():
                # Create new paragraph
                if insertion_index < len(self.document.paragraphs):
                    new_para = self.document.paragraphs[insertion_index].insert_paragraph_before()
                else:
                    new_para = self.document.add_paragraph()
                
                # Apply markdown formatting to paragraph
                self.apply_markdown_to_paragraph(new_para, para_text.strip())
                
                # Apply configured formatting
                self.apply_configured_formatting(new_para)
                
                inserted_paras.append(new_para)
                insertion_index += 1
        
        # Update section's content paragraphs
        if not append:
            section.content_paragraphs = inserted_paras
        else:
            section.content_paragraphs.extend(inserted_paras)

    def parse_markdown_table(self, table_lines):
        """Parse markdown table into a structured format"""
        if len(table_lines) < 3:  # Need at least header, separator, and one data row
            return None
        
        # Split lines into cells
        rows = []
        for line in table_lines:
            line = line.strip()
            if line.startswith('|'):
                line = line[1:]  # Remove leading pipe
            if line.endswith('|'):
                line = line[:-1]  # Remove trailing pipe
            
            cells = [cell.strip() for cell in line.split('|')]
            rows.append(cells)
        
        # Check if the second row is a separator
        is_separator = all(cell.strip().startswith('-') for cell in rows[1])
        if not is_separator:
            return None
        
        # Extract headers and data rows
        headers = rows[0]
        data_rows = rows[2:]
        
        return {
            'headers': headers,
            'rows': data_rows
        }

    def add_markdown_table_to_document(self, section, table_data, insertion_index):
        """Add a markdown table to the document"""
        # Determine number of columns and rows
        num_cols = len(table_data['headers'])
        num_rows = len(table_data['rows']) + 1  # +1 for header row
        
        # Create table in document
        if insertion_index < len(self.document.paragraphs):
            para = self.document.paragraphs[insertion_index]
            table = self.document._body.add_tbl_before(para._p, num_rows, num_cols)
        else:
            table = self.document.add_table(rows=num_rows, cols=num_cols)
        
        # Apply table style
        table.style = 'Table Grid'
        
        # Add headers with formatting
        for i, header in enumerate(table_data['headers']):
            cell = table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for row_idx, row in enumerate(table_data['rows']):
            for col_idx, cell_content in enumerate(row):
                if col_idx < num_cols:  # Ensure we don't exceed column count
                    cell = table.cell(row_idx + 1, col_idx)  # +1 to account for header
                    cell.text = cell_content
        
        # Auto-fit table columns
        table.autofit = True
        
        # Return the actual table object for reference
        return table

    def add_markdown_content_to_section(self, section, content, append=False):
        """Add content to section with markdown to Word formatting conversion"""
        heading_para = section.paragraph
        doc_paragraphs = list(self.document.paragraphs)
        
        # Find heading index
        heading_index = -1
        for i, para in enumerate(doc_paragraphs):
            if para.text.strip() == heading_para.text.strip():
                heading_index = i
                break
        
        if heading_index == -1:
            raise Exception("Could not find section heading")
        
        # Determine insertion point
        if append and section.content_paragraphs:
            next_heading_index = len(doc_paragraphs)
            for i in range(heading_index + 1, len(doc_paragraphs)):
                if doc_paragraphs[i].style.name.startswith('Heading'):
                    next_heading_index = i
                    break
            insertion_index = next_heading_index
        else:
            insertion_index = heading_index + 1
        
        # Process content
        lines = content.split('\n')
        i = 0
        inserted_paras = []
        
        while i < len(lines):
            current_line = lines[i].strip()
            
            # Check for table (starts with | and has at least 3 lines)
            if current_line.startswith('|') and i + 2 < len(lines) and lines[i+1].strip().startswith('|') and lines[i+2].strip().startswith('|'):
                # Collect all lines in the table
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                
                # Parse and add table
                table_data = self.parse_markdown_table(table_lines)
                if table_data:
                    table = self.add_markdown_table_to_document(section, table_data, insertion_index)
                    insertion_index += 1  # Tables count as one "paragraph" for insertion
                    # We don't add tables to inserted_paras since they're not paragraphs
                continue
                
            # Check for code blocks (```code```)
            elif current_line.startswith('```'):
                code_language = current_line[3:].strip()
                code_lines = []
                i += 1  # Skip the opening ```
                
                # Collect all lines in code block
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                
                if i < len(lines):  # Skip closing ```
                    i += 1
                
                # Add code block as a formatted paragraph
                if insertion_index < len(self.document.paragraphs):
                    new_para = self.document.paragraphs[insertion_index].insert_paragraph_before()
                else:
                    new_para = self.document.add_paragraph()
                
                # Format as code
                code_text = '\n'.join(code_lines)
                run = new_para.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.color.rgb = RGBColor(0, 128, 0)
                run.font.size = Pt(9)  # Smaller size for code
                
                # Add light gray background (this requires more complex XML manipulation)
                # Here we just make it visually distinct
                
                inserted_paras.append(new_para)
                insertion_index += 1
                continue
                
            # Check for horizontal rule (---, ***, ___)
            elif re.match(r'^([-*_])\1{2,}$', current_line):
                if insertion_index < len(self.document.paragraphs):
                    new_para = self.document.paragraphs[insertion_index].insert_paragraph_before()
                else:
                    new_para = self.document.add_paragraph()
                
                # Apply a paragraph border to simulate horizontal rule
                new_para.paragraph_format.bottom_border.width = 1
                new_para.paragraph_format.bottom_border.color.rgb = RGBColor(0, 0, 0)
                
                inserted_paras.append(new_para)
                insertion_index += 1
                i += 1
                continue
            
            # Regular paragraph or list item processing
            # Build paragraph text
            para_text = current_line
            i += 1
            
            # For normal paragraphs, combine lines until we hit an empty line or special line
            if not (para_text.startswith('#') or para_text.startswith('* ') or 
                para_text.startswith('- ') or re.match(r'^\d+\.', para_text)):
                while i < len(lines) and lines[i].strip() and not (
                    lines[i].strip().startswith('#') or 
                    lines[i].strip().startswith('* ') or 
                    lines[i].strip().startswith('- ') or 
                    lines[i].strip().startswith('|') or 
                    lines[i].strip().startswith('```') or
                    re.match(r'^([-*_])\1{2,}$', lines[i].strip()) or
                    re.match(r'^\d+\.', lines[i].strip())
                ):
                    # Add a space before appending unless the line ends with specific characters
                    if not para_text.endswith((':', '-', '—', '–')):
                        para_text += ' '
                    para_text += lines[i].strip()
                    i += 1
            
            # Skip if this is a duplicate of the section heading
            if para_text.lower() == section.text.lower() or (
                para_text.startswith('#') and para_text.lstrip('#').strip().lower() == section.text.lower()
            ):
                self.log_message(f"Skipping duplicate heading: {para_text}")
                continue
            
            # Create and format paragraph
            if para_text.strip():
                if insertion_index < len(self.document.paragraphs):
                    new_para = self.document.paragraphs[insertion_index].insert_paragraph_before()
                else:
                    new_para = self.document.add_paragraph()
                
                # Apply markdown formatting
                self.apply_markdown_to_paragraph(new_para, para_text)
                
                # Apply configured formatting
                self.apply_configured_formatting(new_para)
                
                inserted_paras.append(new_para)
                insertion_index += 1
        
        # Update section's content paragraphs
        if not append:
            section.content_paragraphs = inserted_paras
        else:
            section.content_paragraphs.extend(inserted_paras)

    def apply_markdown_to_paragraph(self, paragraph, text):
        """Convert markdown formatting to Word formatting in paragraph"""
        from docx.shared import Inches, Pt, RGBColor
        
        # Handle heading levels 1-6
        if text.startswith('#'):
            heading_level = 0
            for char in text:
                if char == '#':
                    heading_level += 1
                else:
                    break
            
            if 1 <= heading_level <= 6:
                # Map markdown heading level to Word heading style
                if heading_level <= 4:
                    try:
                        paragraph.style = f'Heading {heading_level}'
                    except:
                        # If style not found, fallback to normal with bold
                        paragraph.style = 'Normal'
                        paragraph.runs[0].bold = True
                        paragraph.runs[0].font.size = Pt(14 - heading_level)  # Decrease size for deeper headings
                else:
                    # Fall back to normal text with bold for higher levels
                    paragraph.style = 'Normal'
                    
                text = text[heading_level:].strip()
        
        # Handle bullet points
        elif text.strip().startswith('* ') or text.strip().startswith('- '):
            try:
                paragraph.style = 'List Bullet'
            except:
                # Create bullet point manually if style not available
                paragraph.style = 'Normal'
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Inches(0.25)
                text = "• " + text.strip()[2:].strip()  # Add bullet character
            else:
                text = text.strip()[2:].strip()  # Remove the bullet marker
        
        # Handle numbered lists
        elif re.match(r'^\d+\.', text.strip()):
            try:
                paragraph.style = 'List Number'
            except:
                # Try alternative style names
                try:
                    paragraph.style = 'List Paragraph'
                except:
                    # Create numbered format manually
                    paragraph.style = 'Normal'
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.left_indent = Inches(0.25)
                    # Keep the number in text
            
            # Extract number for manual handling if needed
            num_match = re.match(r'^(\d+)\.', text.strip())
            if num_match:
                num = num_match.group(1)
                text = text.strip()[len(num)+1:].strip()  # Remove the number and dot
                
                # For manual numbering if style failed
                if paragraph.style.name == 'Normal':
                    text = f"{num}. {text}"
        
        # Handle blockquotes
        elif text.strip().startswith('> '):
            try:
                paragraph.style = 'Quote'
            except:
                # Manual quote formatting
                paragraph.style = 'Normal'
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Inches(0.5)
                paragraph_format.right_indent = Inches(0.5)
            
            text = text.strip()[2:].strip()  # Remove the quote marker
        
        # Clear paragraph text since we'll re-add it with formatting
        paragraph.clear()
        
        # Process for inline formatting
        i = 0
        in_bold = False
        in_italic = False
        current_text = ""
        
        while i < len(text):
            # Process bold (both ** and __ formats)
            if i+1 < len(text) and (text[i:i+2] == '**' or text[i:i+2] == '__'):
                # Add accumulated text with current formatting
                if current_text:
                    run = paragraph.add_run(current_text)
                    if in_bold:
                        run.bold = True
                    if in_italic:
                        run.italic = True
                    current_text = ""
                
                # Toggle bold state
                in_bold = not in_bold
                i += 2
                continue
            
            # Process italic (both * and _ formats, but not if part of bold)
            elif (text[i] == '*' or text[i] == '_'):
                # Check it's not part of bold marker
                if not (i+1 < len(text) and (text[i:i+2] == '**' or text[i:i+2] == '__')):
                    # Add accumulated text with current formatting
                    if current_text:
                        run = paragraph.add_run(current_text)
                        if in_bold:
                            run.bold = True
                        if in_italic:
                            run.italic = True
                        current_text = ""
                    
                    # Toggle italic state
                    in_italic = not in_italic
                    i += 1
                    continue
            
            # Process code
            elif text[i] == '`':
                # First add accumulated text
                if current_text:
                    run = paragraph.add_run(current_text)
                    if in_bold:
                        run.bold = True
                    if in_italic:
                        run.italic = True
                    current_text = ""
                
                # Find the closing backtick
                end = text.find('`', i+1)
                if end != -1:
                    # Add the code text
                    code_text = text[i+1:end]
                    run = paragraph.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.color.rgb = RGBColor(0, 128, 0)
                    i = end + 1
                else:
                    # No closing backtick found, treat as regular character
                    current_text += text[i]
                    i += 1
                continue
            
            # Process strikethrough (~~text~~)
            elif i+1 < len(text) and text[i:i+2] == '~~':
                # Add accumulated text
                if current_text:
                    run = paragraph.add_run(current_text)
                    if in_bold:
                        run.bold = True
                    if in_italic:
                        run.italic = True
                    current_text = ""
                
                # Find closing ~~
                end = text.find('~~', i+2)
                if end != -1:
                    # Add strikethrough text
                    strike_text = text[i+2:end]
                    run = paragraph.add_run(strike_text)
                    run.font.strike = True
                    if in_bold:
                        run.bold = True
                    if in_italic:
                        run.italic = True
                    i = end + 2
                else:
                    # No closing strike found, treat as regular
                    current_text += text[i:i+2]
                    i += 2
                continue
            
            # Process hyperlinks [text](url)
            elif text[i] == '[':
                link_text_end = text.find(']', i)
                if link_text_end != -1 and link_text_end + 1 < len(text) and text[link_text_end + 1] == '(':
                    url_end = text.find(')', link_text_end)
                    if url_end != -1:
                        # Add accumulated text first
                        if current_text:
                            run = paragraph.add_run(current_text)
                            if in_bold:
                                run.bold = True
                            if in_italic:
                                run.italic = True
                            current_text = ""
                        
                        # Extract link text and URL
                        link_text = text[i+1:link_text_end]
                        url = text[link_text_end+2:url_end]
                        
                        # Add hyperlink
                        run = paragraph.add_run(link_text)
                        run.font.underline = True
                        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
                        if in_bold:
                            run.bold = True
                        if in_italic:
                            run.italic = True
                        
                        # Note: True hyperlinks require additional XML manipulation with docx
                        # This styling makes it look like a hyperlink visually
                        
                        i = url_end + 1
                        continue
            
            # Add regular character to current text
            current_text += text[i]
            i += 1
        
        # Add any remaining text
        if current_text:
            run = paragraph.add_run(current_text)
            if in_bold:
                run.bold = True
            if in_italic:
                run.italic = True

    def apply_configured_formatting(self, paragraph):
        """Apply configured formatting options to paragraph"""
        for run in paragraph.runs:
            # Highlight
            if self.format_config['highlight_enabled'].get():
                color_name = self.format_config['highlight_color'].get()
                color_map = {
                    'YELLOW': WD_COLOR_INDEX.YELLOW,
                    'GREEN': WD_COLOR_INDEX.GREEN,
                    'CYAN': WD_COLOR_INDEX.TURQUOISE,
                    'PINK': WD_COLOR_INDEX.PINK,
                    'BRIGHT_GREEN': WD_COLOR_INDEX.BRIGHT_GREEN
                }
                run.font.highlight_color = color_map.get(color_name, WD_COLOR_INDEX.YELLOW)
            
            # Bold (if not already from markdown)
            if self.format_config['bold_enabled'].get() and not run.bold:
                run.bold = True
            
            # Italic (if not already from markdown)
            if self.format_config['italic_enabled'].get() and not run.italic:
                run.italic = True
            
            # Underline
            if self.format_config['underline_enabled'].get():
                run.underline = True
            
            # Font size
            run.font.size = Pt(self.format_config['font_size'].get())
            
            # Font color (optional, if different from default)
            if self.format_config['font_color'].get() != '000000':
                color_hex = self.format_config['font_color'].get()
                try:
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except:
                    pass
    
    def remove_paragraph(self, paragraph):
        """Remove a paragraph from the document"""
        p = paragraph._element
        p.getparent().remove(p)
    
    def save_document_auto(self):
        """Auto-save document with permission error handling"""
        if not self.document or not self.document_path:
            return
        
        try:
            self.document.save(self.document_path)
            self.log_message(f"Document auto-saved: {self.document_path}")
        except PermissionError:
            self.log_message("Auto-save failed: File is open elsewhere")
            
            result = messagebox.askyesno(
                "Auto-Save Failed",
                "Cannot auto-save: File is open in another application.\n\n"
                "Close the file and try again, or save as a different file?"
            )
            
            if result:
                # Try regular save with retry
                self.save_with_retry()
        except Exception as e:
            self.log_message(f"Auto-save failed: {str(e)}")
            messagebox.showerror("Auto-Save Error", f"Failed to auto-save: {str(e)}")
    
    def prompt_save_document(self):
        """Prompt to save document with retry on permission error"""
        if not self.document or not self.document_path:
            return
            
        result = messagebox.askyesnocancel(
            "Save Document",
            "Save changes?\n\n"
            "Yes: Save to current file\n"
            "No: Save As (new file)\n"
            "Cancel: Don't save"
        )
        
        if result is True:
            self.save_with_retry()
                
        elif result is False:
            self.save_as_new_file()
    
    def save_with_retry(self):
        """Save document with retry logic for permission errors"""
        max_retries = 3
        retry_count = 0
        
        while retry_count < max_retries:
            try:
                self.document.save(self.document_path)
                self.log_message(f"Document saved: {self.document_path}")
                messagebox.showinfo("Success", "Document saved successfully")
                return True
                
            except PermissionError:
                retry_count += 1
                self.log_message(f"Save failed (attempt {retry_count}): File is open elsewhere")
                
                if retry_count < max_retries:
                    # Prompt user to close and retry
                    result = messagebox.askyesnocancel(
                        "File In Use",
                        f"Cannot save: File is open in another application.\n\n"
                        f"Please close the file and click:\n"
                        f"• Yes - Retry saving to current file\n"
                        f"• No - Save As a different file\n"
                        f"• Cancel - Don't save\n\n"
                        f"Attempt {retry_count} of {max_retries}"
                    )
                    
                    if result is True:
                        # Retry
                        continue
                    elif result is False:
                        # Save as
                        return self.save_as_new_file()
                    else:
                        # Cancel
                        self.log_message("Save cancelled by user")
                        return False
                else:
                    # Max retries reached
                    result = messagebox.askyesno(
                        "Save Failed",
                        "Could not save after multiple attempts.\n\n"
                        "Save as a different file?"
                    )
                    
                    if result:
                        return self.save_as_new_file()
                    else:
                        self.log_message("Save cancelled after max retries")
                        return False
                        
            except Exception as e:
                self.log_message(f"Save error: {str(e)}")
                messagebox.showerror("Error", f"Failed to save: {str(e)}")
                return False
        
        return False
    
    def save_as_new_file(self):
        """Save document as a new file"""
        new_path = filedialog.asksaveasfilename(
            title="Save Document As",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            initialfile=os.path.basename(self.document_path)
        )
        
        if new_path:
            try:
                self.document.save(new_path)
                self.document_path = new_path
                self.doc_label_var.set(os.path.basename(new_path))
                self.save_settings()
                
                # Update tracking file location
                self.load_section_tracking()
                
                self.log_message(f"Document saved as: {new_path}")
                messagebox.showinfo("Success", "Document saved successfully")
                return True
            except Exception as e:
                self.log_message(f"Save As error: {str(e)}")
                messagebox.showerror("Error", f"Failed to save: {str(e)}")
                return False
        
        return False
    
    def reload_document(self):
        """Reload document to refresh content"""
        if not self.document_path:
            return
        
        try:
            self.log_message("Reloading document...")
            current_selection = self.selected_section
            
            # Reload document
            self.document = Document(self.document_path)
            
            # Re-parse structure
            self.parse_document_structure()
            
            # Refresh tree
            self.populate_tree()
            
            # Try to re-select the same section
            if current_selection:
                reselected = self.find_section_by_path(current_selection.get_full_path())
                if reselected:
                    self.selected_section = reselected
                    self.show_existing_content()
            
            self.log_message("Document reloaded successfully")
            
        except Exception as e:
            self.log_message(f"Error reloading document: {str(e)}")
    
    def find_section_by_path(self, path):
        """Find section by full path"""
        def search(sections):
            for section in sections:
                if section.get_full_path() == path:
                    return section
                result = search(section.children)
                if result:
                    return result
            return None
        
        return search(self.sections)
            
    def run(self):
        """Run the application"""
        self.log_message("Document Content Generator started")
        self.log_message("Load a Word document to begin")
        if self.selected_model.get():
            self.log_message(f"Configured model: {self.selected_model.get()}")
        self.root.mainloop()


if __name__ == "__main__":
    try:
        app = DocxDocumentFiller()
        app.run()
    except Exception as e:
        print(f"Application error: {str(e)}")
        import traceback
        traceback.print_exc()
        input("Press Enter to exit...")